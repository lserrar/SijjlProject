# ─── Transcript Service ─────────────────────────────────────────────────────
# Handles Word document conversion and transcript management

from docx import Document
from io import BytesIO
import re
from typing import List, Dict, Optional
import logging

logger = logging.getLogger(__name__)


def convert_docx_to_markdown(file_content: bytes) -> Dict:
    """
    Convert a Word document to Markdown format.
    
    Returns:
        {
            "title": "Document title (first heading or first line)",
            "content": "Full markdown content",
            "sections": [{"title": "Section", "content": "..."}],
            "word_count": 2375
        }
    """
    try:
        doc = Document(BytesIO(file_content))
        
        markdown_lines = []
        sections = []
        current_section = {"title": "", "content": ""}
        title = ""
        word_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                markdown_lines.append("")
                continue
            
            # Count words
            word_count += len(text.split())
            
            # Check if it's a heading
            style_name = para.style.name.lower() if para.style else ""
            
            if 'heading 1' in style_name or 'titre 1' in style_name:
                # Main title
                if not title:
                    title = text
                markdown_lines.append(f"# {text}")
                markdown_lines.append("")
                
                # Start new section
                if current_section["content"]:
                    sections.append(current_section)
                current_section = {"title": text, "content": ""}
                
            elif 'heading 2' in style_name or 'titre 2' in style_name:
                markdown_lines.append(f"## {text}")
                markdown_lines.append("")
                
                # Start new section
                if current_section["content"]:
                    sections.append(current_section)
                current_section = {"title": text, "content": ""}
                
            elif 'heading 3' in style_name or 'titre 3' in style_name:
                markdown_lines.append(f"### {text}")
                markdown_lines.append("")
                current_section["content"] += f"### {text}\n\n"
                
            else:
                # Check for bold text that might be a section header
                is_bold_header = False
                if para.runs:
                    all_bold = all(run.bold for run in para.runs if run.text.strip())
                    if all_bold and len(text) < 100 and not text.endswith('.'):
                        is_bold_header = True
                
                if is_bold_header:
                    markdown_lines.append(f"## {text}")
                    markdown_lines.append("")
                    
                    if current_section["content"]:
                        sections.append(current_section)
                    current_section = {"title": text, "content": ""}
                else:
                    # Regular paragraph
                    markdown_lines.append(text)
                    markdown_lines.append("")
                    current_section["content"] += text + "\n\n"
        
        # Add last section
        if current_section["content"]:
            sections.append(current_section)
        
        # If no title found, use first line
        if not title and markdown_lines:
            for line in markdown_lines:
                if line.strip():
                    title = line.replace('#', '').strip()
                    break
        
        content = "\n".join(markdown_lines).strip()
        
        return {
            "title": title,
            "content": content,
            "sections": sections,
            "word_count": word_count
        }
        
    except Exception as e:
        logger.error(f"Error converting DOCX: {e}")
        raise ValueError(f"Failed to convert document: {str(e)}")


def extract_text_only(file_content: bytes) -> str:
    """Extract plain text from a Word document."""
    try:
        doc = Document(BytesIO(file_content))
        return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        logger.error(f"Error extracting text: {e}")
        raise ValueError(f"Failed to extract text: {str(e)}")


def estimate_reading_time(word_count: int, words_per_minute: int = 200) -> int:
    """Estimate reading time in minutes."""
    return max(1, round(word_count / words_per_minute))
