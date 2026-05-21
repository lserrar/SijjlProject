from fastapi import FastAPI, APIRouter, HTTPException, Request, Depends, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse, Response
from fastapi.templating import Jinja2Templates
from motor.motor_asyncio import AsyncIOMotorClient
from dotenv import load_dotenv
from pydantic import BaseModel, EmailStr, Field
from typing import List, Optional, Any, Dict
from datetime import datetime, timezone, timedelta
from pathlib import Path
import os, uuid, logging, hashlib, hmac, requests as http_requests, re, io
import asyncio
import boto3
from botocore.config import Config as BotoConfig
from botocore.exceptions import ClientError
from emergentintegrations.payments.stripe.checkout import StripeCheckout, CheckoutSessionResponse, CheckoutStatusResponse, CheckoutSessionRequest
from fastapi.staticfiles import StaticFiles

# Stripe subscriptions (Phase B — 12-month commitment, mode=subscription)
import stripe as stripe_sdk
from utils import stripe_subscriptions as stripe_subs

# Email service for notifications
from utils.email_service import (
    is_email_configured,
    send_email as smtp_send_email,
    send_referral_signup_notification,
    send_referral_conversion_notification,
    send_referee_welcome_notification,
    send_subscription_confirmation,
    get_base_template,
    send_password_reset_email,
    send_welcome_email,
    send_trial_expiration_email
)

# Transcript service for Word document conversion
from utils.transcript_service import (
    convert_docx_to_markdown,
    extract_text_only,
    estimate_reading_time
)

# Apple Sign-In
from utils.apple_auth import (
    is_apple_auth_configured,
    get_apple_auth_url,
    exchange_apple_code_for_tokens,
    validate_apple_identity_token,
    decode_apple_user_data
)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Initialize Jinja2 templates
templates = Jinja2Templates(directory=str(ROOT_DIR / 'admin_templates'))

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

JWT_SECRET = os.environ.get('JWT_SECRET_KEY', 'hikmabyLM_secret')
EMERGENT_AUTH_URL = "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data"

# ─── Utility: Clean title prefixes ───────────────────────────────────────────
def clean_title(title: str) -> str:
    """Remove redundant prefixes like 'Épisode 1 —', 'Cours 1:', etc."""
    if not title:
        return title
    # Remove "Épisode X —" or "Épisode X :" patterns
    title = re.sub(r'^[ÉE]pisode\s+\d+\s*[—:\-–]\s*', '', title, flags=re.IGNORECASE)
    # Remove "Cours X :" or "Cours X —" patterns
    title = re.sub(r'^Cours\s+\d+\s*[—:\-–]\s*', '', title, flags=re.IGNORECASE)
    # Remove "Module X —" patterns
    title = re.sub(r'^Module\s+\d+\s*[—:\-–]\s*', '', title, flags=re.IGNORECASE)
    return title.strip()

# ─── Cloudflare R2 Config ────────────────────────────────────────────────────
R2_ACCOUNT_ID     = os.environ.get('R2_ACCOUNT_ID', '')
R2_ACCESS_KEY_ID  = os.environ.get('R2_ACCESS_KEY_ID', '')
R2_SECRET_KEY     = os.environ.get('R2_SECRET_ACCESS_KEY', '')
R2_BUCKET         = os.environ.get('R2_BUCKET_NAME', 'hikma-audio')
R2_ENDPOINT_URL   = os.environ.get('R2_ENDPOINT_URL', f'https://{R2_ACCOUNT_ID}.r2.cloudflarestorage.com')

r2_client = None
logger_r2 = logging.getLogger(__name__)
logger_r2.info(f"R2 config: ACCOUNT_ID={R2_ACCOUNT_ID[:10] if R2_ACCOUNT_ID else 'None'}..., ACCESS_KEY={R2_ACCESS_KEY_ID[:10] if R2_ACCESS_KEY_ID else 'None'}..., SECRET={'SET' if R2_SECRET_KEY else 'EMPTY'}, BUCKET={R2_BUCKET}")
if R2_ACCOUNT_ID and R2_ACCESS_KEY_ID and R2_SECRET_KEY:
    try:
        r2_client = boto3.client(
            's3',
            endpoint_url=R2_ENDPOINT_URL,
            aws_access_key_id=R2_ACCESS_KEY_ID,
            aws_secret_access_key=R2_SECRET_KEY,
            config=BotoConfig(
                signature_version='s3v4',
                retries={'max_attempts': 3, 'mode': 'standard'}
            ),
            region_name='auto',
        )
        logger_init = logging.getLogger(__name__)
        logger_init.info(f"R2 client initialized for bucket '{R2_BUCKET}'")
    except Exception as e:
        r2_client = None
        logging.getLogger(__name__).error(f"R2 init failed: {e}")

PRESIGNED_URL_EXPIRY = 3600  # 1 hour

# Public base URL for audio proxy (avoids R2 CORS issues)
PUBLIC_URL = os.environ.get('PUBLIC_URL', '')

# ─── Stripe Config ───────────────────────────────────────────────────────────
STRIPE_API_KEY = os.environ.get('STRIPE_API_KEY', '')

# Default subscription plans
DEFAULT_PLANS = {
    'fondateur_mensuel': {'name': 'Fondateur Mensuel', 'price': 7.00, 'duration_days': 30, 'type': 'subscription', 'is_fondateur': True, 'is_active': True},
    'fondateur_annuel': {'name': 'Fondateur Annuel', 'price': 84.00, 'duration_days': 365, 'type': 'subscription', 'is_fondateur': True, 'is_active': True},
    'standard_mensuel': {'name': 'Standard Mensuel', 'price': 12.00, 'duration_days': 30, 'type': 'subscription', 'is_fondateur': False, 'is_active': False},
    'standard_annuel': {'name': 'Standard Annuel', 'price': 120.00, 'duration_days': 365, 'type': 'subscription', 'is_fondateur': False, 'is_active': True},
}

def get_presigned_stream_url(file_key: str) -> Optional[str]:
    """Generate a presigned URL for streaming an audio file from R2."""
    if not r2_client or not file_key:
        return None
    try:
        url = r2_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': R2_BUCKET, 'Key': file_key},
            ExpiresIn=PRESIGNED_URL_EXPIRY,
        )
        return url
    except ClientError as e:
        logging.getLogger(__name__).error(f"Presigned URL error for '{file_key}': {e}")
        return None

def get_presigned_upload_url(file_key: str, content_type: str = 'audio/mpeg') -> Optional[str]:
    """Generate a presigned URL for uploading an audio file to R2."""
    if not r2_client or not file_key:
        return None
    try:
        url = r2_client.generate_presigned_url(
            'put_object',
            Params={'Bucket': R2_BUCKET, 'Key': file_key, 'ContentType': content_type},
            ExpiresIn=3600,
        )
        return url
    except ClientError as e:
        logging.getLogger(__name__).error(f"Upload URL error for '{file_key}': {e}")
        return None

def resolve_audio_url(audio_doc: dict) -> str:
    """Return a streaming URL: proxy via backend (avoids CORS) if possible, else presigned R2 URL."""
    file_key = audio_doc.get('file_key')
    audio_id = audio_doc.get('id', '')
    if file_key and audio_id and PUBLIC_URL:
        return f"{PUBLIC_URL}/api/audios/{audio_id}/stream"
    if file_key:
        presigned = get_presigned_stream_url(file_key)
        if presigned:
            return presigned
    return audio_doc.get('audio_url', '')


# ─── R2 Auto-Detection (Phase 3 generic media sync) ─────────────────────────

_AUDIO_EXTS = {'mp3', 'm4a', 'wav', 'aac', 'ogg', 'flac'}
_VIDEO_EXTS = {'mp4', 'mov', 'webm', 'mkv'}
_DOC_EXTS = {'pdf', 'doc', 'docx'}

_BIBLIO_KEYWORDS = ('biblio',)
_GLOSSAIRE_KEYWORDS = ('glossaire', 'glossary', 'lexique')

# Episode number patterns: episode1, ep1, _1., -1., (1), -01., partie1, etc.
_EPISODE_PATTERNS = [
    # After keyword: "episode1", "ep01", "partie 2", "chapitre-3"
    re.compile(r'(?:episode|épisode|ep|partie|part|chapitre|chap)[\s_\-]*0*(\d{1,2})(?=[\s_\-\.\)]|$)', re.IGNORECASE),
    # Trailing number after separator: "...maimounide-1", "..._02"
    re.compile(r'[\-_](\d{1,2})(?=[\s_\-\.\)]|$)'),
    # Leading number at start: "1_intro", "01-foo"
    re.compile(r'^0*(\d{1,2})[\s_\-\.]'),
]


def _extract_episode_number(filename: str) -> Optional[int]:
    """Best-effort extraction of an episode number from a filename (1..50)."""
    name = filename.lower().rsplit('.', 1)[0]
    for pat in _EPISODE_PATTERNS:
        m = pat.search(name)
        if m:
            try:
                n = int(m.group(1))
                if 1 <= n <= 50:
                    return n
            except (ValueError, IndexError):
                continue
    return None


def _classify_r2_file(key: str, prefix: str = '') -> Optional[dict]:
    """Classify an R2 object key into a media role.
    Returns a dict {role, episode_number, type, label, mime, r2_key, filename, subfolder}
    or None if not classifiable.
    Roles: 'episode_video', 'episode_audio', 'episode_doc', 'course_doc'.
    `subfolder` is the path between the course prefix and the filename (e.g. 'al-kindi/').
    """
    filename = key.rsplit('/', 1)[-1]
    if not filename or filename.startswith('.'):
        return None
    # Extract subfolder (path inside course prefix, before filename)
    subfolder = ''
    if prefix and key.startswith(prefix):
        rel = key[len(prefix):]
        if '/' in rel:
            subfolder = rel.rsplit('/', 1)[0] + '/'
    parts = filename.rsplit('.', 1)
    if len(parts) != 2:
        return None
    ext = parts[1].lower()
    name_l = filename.lower()
    ep = _extract_episode_number(filename)

    if ext in _VIDEO_EXTS:
        if ep is not None:
            return {'role': 'episode_video', 'episode_number': ep, 'type': 'video',
                    'label': f'Vidéo épisode {ep}', 'mime': f'video/{ext if ext != "mov" else "quicktime"}',
                    'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}
        return None  # course-level video not yet supported

    if ext in _AUDIO_EXTS:
        if ep is not None:
            mime_map = {'mp3': 'audio/mpeg', 'm4a': 'audio/mp4', 'wav': 'audio/wav',
                        'ogg': 'audio/ogg', 'aac': 'audio/aac', 'flac': 'audio/flac'}
            return {'role': 'episode_audio', 'episode_number': ep, 'type': 'audio',
                    'label': f'Podcast épisode {ep}', 'mime': mime_map.get(ext, 'audio/mpeg'),
                    'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}
        return None

    if ext in _DOC_EXTS:
        mime_map = {'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'doc': 'application/msword'}
        mime = mime_map[ext]
        is_biblio = any(k in name_l for k in _BIBLIO_KEYWORDS)
        is_gloss = any(k in name_l for k in _GLOSSAIRE_KEYWORDS)
        is_script = name_l.startswith('script-') or name_l.startswith('script_')
        is_slides = name_l.startswith('slide-') or name_l.startswith('slides-') or name_l.startswith('slide_') or name_l.startswith('slides_')
        if is_slides:
            if ep is not None:
                return {'role': 'episode_doc', 'episode_number': ep, 'type': 'slides',
                        'label': f"Slides — Épisode {ep}", 'mime': mime,
                        'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}
            return {'role': 'course_doc', 'episode_number': None, 'type': 'slides',
                    'label': filename.rsplit('.', 1)[0].replace('slide-', '').replace('slides-', '').replace('_', ' ').replace('-', ' ').strip().capitalize() or 'Slides',
                    'mime': mime, 'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}
        if is_script:
            if ep is not None:
                return {'role': 'episode_doc', 'episode_number': ep, 'type': 'script',
                        'label': "Script de l'épisode", 'mime': mime,
                        'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}
            return {'role': 'course_doc', 'episode_number': None, 'type': 'script',
                    'label': "Script", 'mime': mime,
                    'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}
        if is_biblio:
            return {'role': 'course_doc', 'episode_number': None, 'type': 'biblio',
                    'label': 'Bibliographie sélective', 'mime': mime,
                    'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}
        if is_gloss:
            return {'role': 'course_doc', 'episode_number': None, 'type': 'glossaire',
                    'label': 'Glossaire des termes', 'mime': mime,
                    'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}
        if ep is not None:
            return {'role': 'episode_doc', 'episode_number': ep, 'type': 'script',
                    'label': "Script de l'épisode", 'mime': mime,
                    'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}
        return {'role': 'course_doc', 'episode_number': None, 'type': 'document',
                'label': filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').strip().capitalize(),
                'mime': mime, 'r2_key': key, 'filename': filename, 'ext': ext, 'subfolder': subfolder}

    # Unknown ext (jpg/png covers, manuscripts, etc.) — ignore for now
    return None


def _list_r2_keys(prefix: str) -> list:
    """List R2 object keys under the given prefix (recursively)."""
    if not r2_client or not prefix:
        return []
    keys = []
    paginator = r2_client.get_paginator('list_objects_v2')
    try:
        for page in paginator.paginate(Bucket=R2_BUCKET, Prefix=prefix):
            for obj in page.get('Contents', []):
                k = obj.get('Key')
                if k:
                    keys.append(k)
    except Exception as e:
        logging.getLogger(__name__).warning(f"R2 list error for prefix {prefix}: {e}")
    return keys


def _build_r2_detections(prefix: str) -> dict:
    """Scan an R2 prefix and group classified detections by role.
    Multi-intervenant safe: keeps a LIST of detections per episode_number so the
    caller can disambiguate using r2_subprefix on the target audio.
    Returned shape:
      videos[ep] = [detection, ...]   (list, in case multiple subfolders share an episode)
      audios[ep] = [detection, ...]
      episode_docs[ep] = [detection, ...]
      course_docs = [...]
      unclassified = [...]
    """
    out: dict = {
        'prefix': prefix,
        'videos': {},
        'audios': {},
        'episode_docs': {},
        'course_docs': [],
        'unclassified': [],
    }
    for key in _list_r2_keys(prefix):
        cls = _classify_r2_file(key, prefix=prefix)
        if not cls:
            out['unclassified'].append(key)
            continue
        role = cls['role']
        ep = cls.get('episode_number')
        if role == 'episode_video' and ep:
            out['videos'].setdefault(ep, []).append(cls)
        elif role == 'episode_audio' and ep:
            out['audios'].setdefault(ep, []).append(cls)
        elif role == 'episode_doc' and ep:
            out['episode_docs'].setdefault(ep, []).append(cls)
        elif role == 'course_doc':
            if cls['type'] in ('biblio', 'glossaire'):
                if not any(d['type'] == cls['type'] for d in out['course_docs']):
                    out['course_docs'].append(cls)
            else:
                out['course_docs'].append(cls)
    return out


def _match_detection_for_audio(detections_list: list, audio_subprefix: str) -> Optional[dict]:
    """Pick the right detection for an audio based on its r2_subprefix.
    - If audio has r2_subprefix, return the first detection whose subfolder matches.
    - If audio has NO subprefix and only one detection exists, return it (back-compat).
    - Otherwise return None (ambiguous — admin needs to set r2_subprefix).
    """
    if not detections_list:
        return None
    sp = (audio_subprefix or '').strip().strip('/').lower()
    if sp:
        for d in detections_list:
            df = (d.get('subfolder') or '').strip().strip('/').lower()
            if df == sp:
                return d
        return None
    # No subprefix on audio
    if len(detections_list) == 1:
        return detections_list[0]
    # Multiple candidates but no subprefix — prefer one with EMPTY subfolder (flat course)
    flat = [d for d in detections_list if not (d.get('subfolder') or '').strip('/')]
    if len(flat) == 1:
        return flat[0]
    return None  # ambiguous


async def _apply_r2_detections(course_id: str, detections: dict, auto_create_missing: bool = True) -> dict:
    """Apply auto-detection results to the database for this course.
    Multi-intervenant safe: respects each audio's r2_subprefix to pick the right file.
    When auto_create_missing=True, creates new audio docs for episodes/subfolders
    found in R2 that don't yet exist in DB.
    """
    summary = {
        'course_resources_count': 0,
        'episodes_updated': 0,
        'episodes_created': 0,
        'audios_with_video': 0,
        'audios_with_audio': 0,
        'audios_with_script': 0,
        'ambiguous': [],  # list of {ep, candidates} where no audio's subprefix matched
        'unclassified_count': len(detections.get('unclassified') or []),
    }
    course_res = []
    for d in detections.get('course_docs', []):
        course_res.append({
            'type': d['type'], 'label': d['label'],
            'r2_key': d['r2_key'], 'mime': d['mime'],
        })
    if course_res:
        await db.courses.update_one(
            {'id': course_id},
            {'$set': {'course_resources': course_res, 'r2_prefix': detections.get('prefix')}}
        )
        summary['course_resources_count'] = len(course_res)
    else:
        await db.courses.update_one(
            {'id': course_id},
            {'$set': {'r2_prefix': detections.get('prefix')}}
        )

    audios = await db.audios.find(
        {'course_id': course_id},
        {'_id': 0, 'id': 1, 'episode_number': 1, 'r2_subprefix': 1, 'title': 1}
    ).to_list(500)
    # Group existing audios by (ep, subprefix) for precise matching.
    # An audio without r2_subprefix is matched only when no subprefix-bearing audio claims it.

    def _find_audio_for(ep: int, detection: dict):
        """Return id of the audio matching this detection, or None."""
        det_sf = (detection.get('subfolder') or '').strip().strip('/').lower()
        # 1) exact subprefix match
        for a in audios:
            if a.get('episode_number') != ep:
                continue
            asp = (a.get('r2_subprefix') or '').strip().strip('/').lower()
            if asp and asp == det_sf:
                return a['id']
        # 2) no subprefix candidate but exactly one ep without subprefix
        plain = [a for a in audios if a.get('episode_number') == ep and not (a.get('r2_subprefix') or '').strip('/')]
        if det_sf == '' and len(plain) == 1:
            return plain[0]['id']
        # 3) detection flat (no subfolder) + only one audio at this ep total → assume it's the right one
        ep_audios = [a for a in audios if a.get('episode_number') == ep]
        if det_sf == '' and len(ep_audios) == 1:
            return ep_audios[0]['id']
        return None

    async def _create_audio_from_detection(ep: int, audio_det: dict, doc_dets: list) -> str:
        """Create a new audio row for a (subfolder, ep) pair when no match exists."""
        sf = (audio_det.get('subfolder') or '').strip('/')
        sf_slug = sf.replace('/', '-') or 'main'
        new_id = f"aud_{course_id}-{sf_slug or 'main'}-ep{ep:02d}"
        # Derive a human title from subfolder; fall back to a clean default
        if sf_slug and sf_slug != 'main':
            readable = sf_slug.replace('-', ' ').strip().title()
            audio_title = f"{readable} — Épisode {ep}"
        else:
            audio_title = f"Épisode {ep}"
        # Try to find module in course to attach module_id
        course_doc = await db.courses.find_one({'id': course_id}, {'_id': 0, 'modules': 1})
        module_id = None
        if course_doc and (course_doc.get('modules') or []):
            module_id = course_doc['modules'][0].get('id')  # default first module
        ep_res = [
            {'type': d['type'], 'label': d['label'], 'r2_key': d['r2_key'], 'mime': d['mime']}
            for d in (doc_dets or [])
        ]
        new_audio = {
            'id': new_id,
            'course_id': course_id,
            'module_id': module_id,
            'episode_number': ep,
            'title': audio_title,
            'audio_url': '',
            'r2_audio_key': audio_det['r2_key'],
            'has_r2_audio': True,
            'r2_subprefix': sf + ('/' if sf else ''),
            'is_active': True,
            'episode_resources': ep_res,
            'created_at': datetime.now(timezone.utc).isoformat(),
        }
        await db.audios.update_one(
            {'id': new_id},
            {'$set': new_audio},
            upsert=True,
        )
        audios.append({'id': new_id, 'episode_number': ep, 'r2_subprefix': new_audio['r2_subprefix'], 'title': new_audio['title']})
        return new_id

    # 1) Videos
    for ep, vids in (detections.get('videos') or {}).items():
        for v in vids:
            aid = _find_audio_for(ep, v)
            if aid:
                await db.audios.update_one({'id': aid}, {'$set': {'r2_video_key': v['r2_key']}})
                summary['audios_with_video'] += 1
            else:
                summary['ambiguous'].append({'ep': ep, 'type': 'video', 'r2_key': v['r2_key'], 'subfolder': v.get('subfolder')})

    # 2) Audios — for each (ep, audio) try to attach; auto-create if none and allowed
    for ep, auds in (detections.get('audios') or {}).items():
        # Prefer .mp3 over .m4a only when same subfolder
        by_sf: dict = {}
        for a in auds:
            sf = (a.get('subfolder') or '').strip('/')
            existing = by_sf.get(sf)
            if not existing or (a['ext'] == 'mp3' and existing['ext'] != 'mp3'):
                by_sf[sf] = a
        for sf, a in by_sf.items():
            aid = _find_audio_for(ep, a)
            if aid:
                await db.audios.update_one(
                    {'id': aid},
                    {'$set': {'r2_audio_key': a['r2_key'], 'has_r2_audio': True}}
                )
                summary['audios_with_audio'] += 1
            elif auto_create_missing:
                # Pick docs matching this subfolder + ep
                docs_for_this = []
                for d in (detections.get('episode_docs') or {}).get(ep, []):
                    dsf = (d.get('subfolder') or '').strip('/')
                    if dsf == sf:
                        docs_for_this.append(d)
                await _create_audio_from_detection(ep, a, docs_for_this)
                summary['episodes_created'] += 1
                summary['audios_with_audio'] += 1
            else:
                summary['ambiguous'].append({'ep': ep, 'type': 'audio', 'r2_key': a['r2_key'], 'subfolder': a.get('subfolder')})

    # 3) Episode docs (scripts/slides per episode/intervenant)
    # Group by (ep, subfolder) then attach to matching audio
    docs_by_pair: dict = {}
    for ep, docs in (detections.get('episode_docs') or {}).items():
        for d in docs:
            sf = (d.get('subfolder') or '').strip('/')
            docs_by_pair.setdefault((ep, sf), []).append(d)
    for (ep, sf), docs in docs_by_pair.items():
        # Use first doc as proxy for subfolder matching
        proxy = {'subfolder': sf + ('/' if sf else '')}
        aid = _find_audio_for(ep, proxy)
        if not aid:
            summary['ambiguous'].append({'ep': ep, 'type': 'docs', 'subfolder': sf, 'count': len(docs)})
            continue
        ep_resources = [
            {'type': d['type'], 'label': d['label'], 'r2_key': d['r2_key'], 'mime': d['mime']}
            for d in docs
        ]
        await db.audios.update_one(
            {'id': aid},
            {'$set': {'episode_resources': ep_resources}}
        )
        summary['audios_with_script'] += 1

    summary['episodes_updated'] = summary['audios_with_audio'] + summary['audios_with_video']
    return summary



app = FastAPI(title="HikmabyLM API")
api_router = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Simple JWT (HS256 via hmac) ───────────────────────────────────────────

import base64, json as _json

def _b64url(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).rstrip(b'=').decode()

def create_jwt(payload: dict) -> str:
    header = _b64url(_json.dumps({"alg": "HS256", "typ": "JWT"}).encode())
    body   = _b64url(_json.dumps(payload).encode())
    sig    = _b64url(hmac.new(JWT_SECRET.encode(), f"{header}.{body}".encode(), hashlib.sha256).digest())
    return f"{header}.{body}.{sig}"

def verify_jwt(token: str) -> Optional[dict]:
    try:
        parts = token.split('.')
        if len(parts) != 3:
            return None
        header, body, sig = parts
        expected_sig = _b64url(hmac.new(JWT_SECRET.encode(), f"{header}.{body}".encode(), hashlib.sha256).digest())
        if not hmac.compare_digest(sig, expected_sig):
            return None
        pad = len(body) % 4
        if pad:
            body += '=' * (4 - pad)
        payload = _json.loads(base64.urlsafe_b64decode(body))
        if payload.get('exp') and datetime.fromtimestamp(payload['exp'], tz=timezone.utc) < datetime.now(timezone.utc):
            return None
        return payload
    except Exception:
        return None

async def get_current_user(request: Request) -> Optional[dict]:
    token = request.headers.get('Authorization', '').replace('Bearer ', '').strip()
    if not token:
        token = request.cookies.get('session_token', '')
    if not token:
        return None
    payload = verify_jwt(token)
    if not payload:
        return None
    user = await db.users.find_one({'user_id': payload.get('user_id')}, {'_id': 0})
    return user

def hash_password(pw: str) -> str:
    return hashlib.sha256((pw + JWT_SECRET).encode()).hexdigest()

# Only admin role is assignable via code — no API can grant admin
ADMIN_EMAILS = {'loubna.serrar@gmail.com'}

async def require_admin(request: Request) -> dict:
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Accès réservé aux administrateurs")
    return user

async def require_subscriber(request: Request) -> dict:
    """Require authenticated user with active subscription/trial/admin/free_access.
    Used to gate premium content endpoints (frises, contextes, bibliographies)."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    # Inline access check to avoid forward reference to check_user_access
    if user.get('role') == 'admin' or user.get('free_access'):
        return user
    now = datetime.now(timezone.utc)
    trial = user.get('trial')
    if trial and trial.get('expires_at'):
        exp = trial['expires_at']
        if isinstance(exp, str):
            exp = datetime.fromisoformat(exp.replace('Z', '+00:00'))
        if exp > now:
            return user
    sub = user.get('subscription')
    if sub and sub.get('expires_at'):
        exp = sub['expires_at']
        if isinstance(exp, str):
            exp = datetime.fromisoformat(exp.replace('Z', '+00:00'))
        if exp > now:
            return user
    raise HTTPException(403, {"error": "subscription_required", "reason": "no_access"})

async def verify_content_access(request: Request, token: Optional[str] = None) -> dict:
    """Verify access for content endpoints (HTML opened in new tab).
    Accepts either Authorization header (subscribed user) OR a short-lived signed token (?t=...).
    Returns the user dict (or token payload) on success, raises 401/403 otherwise."""
    user = await get_current_user(request)
    if user:
        return await require_subscriber(request)
    if token:
        payload = verify_jwt(token)
        if payload and payload.get('scope') == 'content_access':
            return {'token_user_id': payload.get('sub'), 'token_payload': payload}
    raise HTTPException(401, "Authentification requise")

# ─── Pydantic Models ────────────────────────────────────────────────────────

class RegisterRequest(BaseModel):
    email: EmailStr
    password: str
    name: str
    referral_code: Optional[str] = None  # Code de parrainage du parrain

class LoginRequest(BaseModel):
    email: EmailStr
    password: str

class GoogleSessionRequest(BaseModel):
    session_id: str

class ProgressRequest(BaseModel):
    content_id: str
    content_type: str
    progress: float
    position: Optional[float] = 0

class FavoriteRequest(BaseModel):
    content_id: str
    content_type: str

class LiveRegisterRequest(BaseModel):
    session_id: str

# ─── Cursus Models (was Thematiques) ─────────────────────────────────────────

class CursusCreate(BaseModel):
    name: str
    description: str = ""
    icon: str = "book-open"
    order: int = 0
    is_active: bool = False
    is_featured: bool = False
    hero_title: Optional[str] = None
    hero_description: Optional[str] = None

class CursusUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    icon: Optional[str] = None
    order: Optional[int] = None
    is_active: Optional[bool] = None
    is_featured: Optional[bool] = None
    hero_title: Optional[str] = None
    hero_description: Optional[str] = None

# ─── Module Models ───────────────────────────────────────────────────────────

class ModuleCreate(BaseModel):
    name: str
    description: str = ""
    course_id: str  # Link to course
    scholar_name: Optional[str] = None
    order: int = 0
    episode_count: int = 2  # Default: 1 short + 1 long
    is_active: bool = False  # Default inactive for preparation

class ModuleUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    course_id: Optional[str] = None
    scholar_name: Optional[str] = None
    order: Optional[int] = None
    episode_count: Optional[int] = None
    is_active: Optional[bool] = None

# ─── Bulk Action Models ──────────────────────────────────────────────────────

class BulkToggleRequest(BaseModel):
    ids: List[str]
    is_active: bool

# ─── Audio Category Models ─────────────────────────────────────────────────────

class AudioCategoryCreate(BaseModel):
    name: str
    description: str = ""
    r2_folder: str  # Dossier R2 associé (ex: "hikma-audio/0. Conference/")
    icon: str = "headphones"  # Font Awesome icon name
    is_active: bool = True

class AudioCategoryUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    r2_folder: Optional[str] = None
    icon: Optional[str] = None
    is_active: Optional[bool] = None

class AudioCreate(BaseModel):
    title: str
    description: str = ""
    scholar_id: Optional[str] = None
    scholar_name: Optional[str] = None
    duration: int = 0
    audio_url: Optional[str] = ""
    file_key: Optional[str] = ""
    thumbnail: str = ""
    topic: str = ""
    type: str = "episode"  # episode, short, long
    category_id: Optional[str] = None  # Link to audio_categories collection
    module_id: Optional[str] = None  # Link to a module (NEW - replaces course_id)
    episode_number: Optional[int] = None
    is_active: bool = False
    youtube_url: Optional[str] = None  # Unlisted YouTube video URL
    coming_soon: Optional[bool] = None  # Episode planned but not yet available
    available_date: Optional[str] = None  # e.g. "mai 2026", "sept. 2026"

class AudioUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    scholar_id: Optional[str] = None
    scholar_name: Optional[str] = None
    duration: Optional[int] = None
    thumbnail: Optional[str] = None
    topic: Optional[str] = None
    type: Optional[str] = None
    category_id: Optional[str] = None
    module_id: Optional[str] = None
    episode_number: Optional[int] = None
    is_active: Optional[bool] = None
    file_key: Optional[str] = None
    audio_url: Optional[str] = None
    youtube_url: Optional[str] = None
    coming_soon: Optional[bool] = None
    available_date: Optional[str] = None
    intervenant: Optional[str] = None
    r2_subprefix: Optional[str] = None
    published_at: Optional[str] = None

class ScholarCreate(BaseModel):
    name: str
    university: str
    bio: str
    photo: str = ""
    specializations: List[str] = []

class ScholarUpdate(BaseModel):
    name: Optional[str] = None
    university: Optional[str] = None
    bio: Optional[str] = None
    photo: Optional[str] = None
    specializations: Optional[List[str]] = None

class CourseCreate(BaseModel):
    title: str
    description: str
    topic: str
    level: str
    language: str = "Francais"
    scholar_id: str
    scholar_name: str
    thematique_id: Optional[str] = None
    duration: int = 0
    thumbnail: str = ""
    modules_count: int = 0
    tags: List[str] = []
    is_featured: bool = False
    r2_folder: Optional[str] = None
    hero_title: Optional[str] = None
    hero_description: Optional[str] = None
    youtube_url: Optional[str] = None  # Unlisted YouTube URL (for whole course or single-episode courses)
    is_launch_catalog: Optional[bool] = None  # Show in public launch Catalogue page
    coming_soon: Optional[bool] = None  # Course planned but not yet available
    available_date: Optional[str] = None  # e.g. "mai 2026"

class CourseUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    topic: Optional[str] = None
    level: Optional[str] = None
    language: Optional[str] = None
    scholar_id: Optional[str] = None
    scholar_name: Optional[str] = None
    thematique_id: Optional[str] = None
    duration: Optional[int] = None
    thumbnail: Optional[str] = None
    modules_count: Optional[int] = None
    tags: Optional[List[str]] = None
    is_active: Optional[bool] = None
    is_featured: Optional[bool] = None
    r2_folder: Optional[str] = None
    hero_title: Optional[str] = None
    hero_description: Optional[str] = None
    youtube_url: Optional[str] = None
    is_launch_catalog: Optional[bool] = None
    coming_soon: Optional[bool] = None
    available_date: Optional[str] = None
    summary: Optional[str] = None
    r2_prefix: Optional[str] = None
    co_scholar_ids: Optional[List[str]] = None  # co-intervenants — IDs of additional scholars displayed after the primary on catalogue cards

# ─── Conference Models ─────────────────────────────────────────────────────────

class ConferenceCreate(BaseModel):
    title: str
    description: str
    speaker_name: str  # Peut être un professeur externe
    speaker_bio: str = ""
    thematique_id: Optional[str] = None  # Lien avec un cursus/thématique
    audio_url: str = ""
    video_url: str = ""
    thumbnail: str = ""
    duration: int = 0  # en minutes
    conference_date: Optional[str] = None  # Date de la conférence
    source: str = ""  # Source de la conférence (ex: "Colloque EPHE 2024")
    tags: List[str] = []
    is_public: bool = False  # Conférence gratuite ou payante

class ConferenceUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    speaker_name: Optional[str] = None
    speaker_bio: Optional[str] = None
    thematique_id: Optional[str] = None
    audio_url: Optional[str] = None
    video_url: Optional[str] = None
    thumbnail: Optional[str] = None
    duration: Optional[int] = None
    conference_date: Optional[str] = None
    source: Optional[str] = None
    tags: Optional[List[str]] = None
    is_public: Optional[bool] = None
    is_active: Optional[bool] = None

# ─── Subscription & Payment Models ───────────────────────────────────────────

class PlanCreate(BaseModel):
    plan_id: str
    name: str
    price: float
    duration_days: int
    type: str = "subscription"  # subscription, course_purchase, cursus_purchase
    description: str = ""
    trial_days: int = 0  # Free trial period
    is_active: bool = True

class PlanUpdate(BaseModel):
    name: Optional[str] = None
    price: Optional[float] = None
    duration_days: Optional[int] = None
    description: Optional[str] = None
    is_active: Optional[bool] = None
    trial_days: Optional[int] = None

class CheckoutRequest(BaseModel):
    plan_id: Optional[str] = None
    course_id: Optional[str] = None
    cursus_id: Optional[str] = None
    promo_code: Optional[str] = None
    origin_url: str

class PromoCodeCreate(BaseModel):
    code: str
    discount_percent: Optional[float] = None  # 0-100
    discount_amount: Optional[float] = None   # Fixed amount in EUR
    max_uses: Optional[int] = None
    start_date: Optional[str] = None   # When the promo code becomes valid
    expires_at: Optional[str] = None   # When the promo code expires
    applicable_plans: List[str] = []  # Empty = all plans
    description: str = ""
    is_active: bool = True

class PromoCodeUpdate(BaseModel):
    discount_percent: Optional[float] = None
    discount_amount: Optional[float] = None
    max_uses: Optional[int] = None
    start_date: Optional[str] = None
    expires_at: Optional[str] = None
    applicable_plans: Optional[List[str]] = None
    description: Optional[str] = None
    is_active: Optional[bool] = None

class StartTrialRequest(BaseModel):
    plan_id: str

class Top10UpdateRequest(BaseModel):
    course_ids: List[str]

# ─── Referral Models ─────────────────────────────────────────────────────────

class ApplyReferralRequest(BaseModel):
    referral_code: str

class GrantFreeMonthRequest(BaseModel):
    user_id: str
    months: int = 1
    reason: str = "admin_grant"

def generate_referral_code(user_id: str, name: str) -> str:
    """Generate a unique referral code based on user name and ID."""
    # Clean name: remove accents, spaces, keep first part
    import unicodedata
    clean_name = unicodedata.normalize('NFKD', name).encode('ASCII', 'ignore').decode('ASCII')
    clean_name = clean_name.split()[0].upper()[:6] if clean_name else "USER"
    # Add unique suffix from user_id
    suffix = user_id[-4:].upper()
    return f"SIJILL-{clean_name}{suffix}"

# ─── Auth Routes ────────────────────────────────────────────────────────────

@api_router.post("/auth/register")
async def register(body: RegisterRequest):
    existing = await db.users.find_one({'email': body.email})
    if existing:
        raise HTTPException(400, "Email déjà utilisé")
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)
    role = 'admin' if body.email in ADMIN_EMAILS else 'user'
    
    # Generate unique referral code for this user
    referral_code = generate_referral_code(user_id, body.name)
    
    user_doc = {
        'user_id': user_id,
        'email': body.email,
        'name': body.name,
        'password_hash': hash_password(body.password),
        'picture': f"https://ui-avatars.com/api/?name={body.name.replace(' ','+')}&background=04D182&color=000&bold=true",
        'provider': 'email',
        'role': role,
        'created_at': now,
        'progress': {},
        'favorites': [],
        # Referral fields
        'referral_code': referral_code,
        'referred_by': None,  # user_id of referrer
        'referral_count': 0,  # Number of successful referrals
        'free_months_earned': 0,  # Total free months earned from referrals
        'free_months_remaining': 0,  # Free months not yet used
        'subscription_end_date': None,  # When subscription/free access ends
    }
    
    # Process referral code if provided
    referrer_user = None
    if body.referral_code:
        referrer_user = await db.users.find_one({'referral_code': body.referral_code.upper()})
        if referrer_user:
            user_doc['referred_by'] = referrer_user['user_id']
            user_doc['free_months_remaining'] = 1  # 1 free month for new user (filleul)
            user_doc['subscription_end_date'] = now + timedelta(days=30)
            
            # Create referral record
            await db.referrals.insert_one({
                'id': f"ref_{uuid.uuid4().hex[:12]}",
                'referrer_id': referrer_user['user_id'],
                'referrer_name': referrer_user.get('name', ''),
                'referrer_email': referrer_user.get('email', ''),
                'referee_id': user_id,
                'referee_name': body.name,
                'referee_email': body.email,
                'status': 'pending',  # pending = waiting for referee to subscribe
                'referrer_rewarded': False,
                'created_at': now,
                'converted_at': None,  # When referee subscribes
            })
            logger.info(f"Referral created: {referrer_user['user_id']} -> {user_id}")
            
            # Send notification emails (if configured)
            if is_email_configured():
                # Notify referrer that someone signed up
                send_referral_signup_notification(
                    referrer_email=referrer_user.get('email', ''),
                    referrer_name=referrer_user.get('name', 'Membre'),
                    referee_name=body.name
                )
                # Welcome email to new user with referral bonus info
                send_referee_welcome_notification(
                    referee_email=body.email,
                    referee_name=body.name,
                    referrer_name=referrer_user.get('name', 'Un membre Sijill'),
                    free_months=1
                )
    
    await db.users.insert_one(user_doc)
    token = create_jwt({'user_id': user_id, 'exp': int((now + timedelta(days=7)).timestamp())})
    
    # Send welcome email (if no referral - referral users get a different welcome email above)
    if not referrer_user and is_email_configured():
        send_welcome_email(
            user_email=body.email,
            user_name=body.name
        )
    
    # Return user data (exclude sensitive fields)
    user_response = {k: v for k, v in user_doc.items() if k not in ('_id', 'password_hash')}
    if referrer_user:
        user_response['referrer_name'] = referrer_user.get('name', '')
    
    return {'token': token, 'user': user_response}

@api_router.post("/auth/login")
async def login(body: LoginRequest):
    user = await db.users.find_one({'email': body.email}, {'_id': 0})
    logger.info(f"Login attempt: email={body.email}, user_found={user is not None}")
    if not user:
        raise HTTPException(401, "Email ou mot de passe incorrect")
    
    stored_hash = user.get('password_hash')
    computed_hash = hash_password(body.password)
    logger.info(f"Login: stored_hash={stored_hash[:20]}..., computed_hash={computed_hash[:20]}...")
    
    if stored_hash != computed_hash:
        raise HTTPException(401, "Email ou mot de passe incorrect")
    
    now = datetime.now(timezone.utc)
    token = create_jwt({'user_id': user['user_id'], 'exp': int((now + timedelta(days=7)).timestamp())})
    user.pop('password_hash', None)
    return {'token': token, 'user': user}

class ForgotPasswordRequest(BaseModel):
    email: str

class ResetPasswordRequest(BaseModel):
    token: str
    new_password: str

@api_router.post("/auth/forgot-password")
async def forgot_password(body: ForgotPasswordRequest):
    """Request a password reset. Generates token and sends email."""
    email = body.email.lower().strip()
    now = datetime.now(timezone.utc)
    
    # Check if user exists
    user = await db.users.find_one({'email': email}, {'_id': 0})
    
    if user:
        # Generate secure reset token
        reset_token = f"rst_{uuid.uuid4().hex}"
        expires_at = now + timedelta(hours=1)
        
        # Store reset token in database
        await db.password_reset_tokens.update_one(
            {'email': email},
            {'$set': {
                'email': email,
                'user_id': user['user_id'],
                'token': reset_token,
                'created_at': now,
                'expires_at': expires_at,
                'used': False
            }},
            upsert=True
        )
        
        # Build reset link
        reset_link = f"https://sijillproject.com/reset-password?token={reset_token}"
        
        # Send email if configured
        if is_email_configured():
            result = send_password_reset_email(
                user_email=email,
                user_name=user.get('name', 'Utilisateur'),
                reset_link=reset_link
            )
            logger.info(f"Password reset email sent: email={email}, success={result.get('success')}")
        else:
            logger.warning(f"Email not configured - reset token generated but not sent: {email}")
    
    # Always return success to prevent email enumeration attacks
    return {'message': 'Si un compte existe avec cette adresse, un email de réinitialisation a été envoyé.'}

@api_router.post("/auth/reset-password")
async def reset_password(body: ResetPasswordRequest):
    """Reset password using the token received by email."""
    now = datetime.now(timezone.utc)
    
    # Find valid token
    token_doc = await db.password_reset_tokens.find_one({
        'token': body.token,
        'used': False,
        'expires_at': {'$gt': now}
    })
    
    if not token_doc:
        raise HTTPException(400, "Lien de réinitialisation invalide ou expiré")
    
    # Validate new password
    if len(body.new_password) < 6:
        raise HTTPException(400, "Le mot de passe doit contenir au moins 6 caractères")
    
    # Update user password
    new_hash = hash_password(body.new_password)
    result = await db.users.update_one(
        {'user_id': token_doc['user_id']},
        {'$set': {'password_hash': new_hash}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(404, "Utilisateur non trouvé")
    
    # Mark token as used
    await db.password_reset_tokens.update_one(
        {'token': body.token},
        {'$set': {'used': True, 'used_at': now}}
    )
    
    logger.info(f"Password reset successful: user_id={token_doc['user_id']}")
    
    return {'message': 'Mot de passe réinitialisé avec succès. Vous pouvez maintenant vous connecter.'}

@api_router.get("/auth/reset-password/validate")
async def validate_reset_token(token: str):
    """Validate a reset token before showing the reset form."""
    now = datetime.now(timezone.utc)
    
    token_doc = await db.password_reset_tokens.find_one({
        'token': token,
        'used': False,
        'expires_at': {'$gt': now}
    })
    
    if not token_doc:
        raise HTTPException(400, "Lien de réinitialisation invalide ou expiré")
    
    return {'valid': True, 'email': token_doc.get('email', '')}

@api_router.post("/auth/google/exchange")
async def google_exchange(body: GoogleSessionRequest):
    try:
        resp = http_requests.get(EMERGENT_AUTH_URL, headers={'X-Session-ID': body.session_id}, timeout=10)
        if resp.status_code != 200:
            raise HTTPException(401, "Session Google invalide")
        data = resp.json()
    except Exception as e:
        raise HTTPException(500, f"Erreur d'authentification Google: {str(e)}")

    email = data.get('email')
    name  = data.get('name', email)
    picture = data.get('picture', '')
    now = datetime.now(timezone.utc)

    existing = await db.users.find_one({'email': email}, {'_id': 0})
    if existing:
        user_id = existing['user_id']
        # Update name, picture, and ensure admin role if applicable
        update_data = {'name': name, 'picture': picture}
        if email in ADMIN_EMAILS:
            update_data['role'] = 'admin'
        await db.users.update_one({'user_id': user_id}, {'$set': update_data})
    else:
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        role = 'admin' if email in ADMIN_EMAILS else 'user'
        await db.users.insert_one({
            'user_id': user_id, 'email': email, 'name': name, 'picture': picture,
            'provider': 'google', 'role': role, 'created_at': now, 'password_hash': None
        })

    token = create_jwt({'user_id': user_id, 'exp': int((now + timedelta(days=7)).timestamp())})
    user = await db.users.find_one({'user_id': user_id}, {'_id': 0, 'password_hash': 0})
    return {'token': token, 'user': user}

# ─── Apple Sign-In Routes ────────────────────────────────────────────────────

class AppleAuthRequest(BaseModel):
    code: Optional[str] = None
    id_token: Optional[str] = None
    user: Optional[str] = None  # JSON string with name/email on first auth

@api_router.get("/auth/apple/login")
async def apple_login_redirect():
    """Redirect user to Apple Sign-In page."""
    if not is_apple_auth_configured():
        raise HTTPException(503, "Apple Sign-In n'est pas configuré")
    
    import secrets
    state = secrets.token_urlsafe(32)
    auth_url = get_apple_auth_url(state=state)
    
    return {"auth_url": auth_url, "state": state}

@api_router.post("/auth/apple/callback")
async def apple_callback(
    code: Optional[str] = None,
    id_token: Optional[str] = None,
    user: Optional[str] = None,
    state: Optional[str] = None,
    error: Optional[str] = None
):
    """
    Handle Apple OAuth callback.
    Apple sends data via form POST (response_mode: form_post).
    """
    if error:
        logger.error(f"Apple auth error: {error}")
        raise HTTPException(400, f"Erreur Apple: {error}")
    
    if not is_apple_auth_configured():
        raise HTTPException(503, "Apple Sign-In n'est pas configuré")
    
    try:
        # Exchange code for tokens if we have a code
        if code and not id_token:
            token_response = await exchange_apple_code_for_tokens(code)
            id_token = token_response.get("id_token")
        
        if not id_token:
            raise HTTPException(400, "Token d'identité manquant")
        
        # Validate the identity token
        apple_user_data = await validate_apple_identity_token(id_token)
        apple_user_id = apple_user_data.get("sub")
        
        if not apple_user_id:
            raise HTTPException(400, "Identifiant utilisateur Apple manquant")
        
        # Get email from token or user data
        email = apple_user_data.get("email")
        
        # Parse user info (Apple only sends this on FIRST sign-in)
        user_info = decode_apple_user_data(user) if user else {}
        name_data = user_info.get("name", {})
        first_name = name_data.get("firstName", "")
        last_name = name_data.get("lastName", "")
        full_name = f"{first_name} {last_name}".strip() if first_name or last_name else None
        
        # If no email in token, try from user info
        if not email:
            email = user_info.get("email")
        
        now = datetime.now(timezone.utc)
        
        # Check if user exists (by Apple ID or email)
        existing = await db.users.find_one({
            '$or': [
                {'apple_id': apple_user_id},
                {'email': email} if email else {'_id': None}
            ]
        }, {'_id': 0})
        
        if existing:
            user_id = existing['user_id']
            # Update Apple ID if not set
            update_data = {'apple_id': apple_user_id}
            if full_name and not existing.get('name'):
                update_data['name'] = full_name
            if email and not existing.get('email'):
                update_data['email'] = email
            await db.users.update_one({'user_id': user_id}, {'$set': update_data})
            logger.info(f"Apple Sign-In: existing user {user_id}")
        else:
            # Create new user
            user_id = f"user_{uuid.uuid4().hex[:12]}"
            role = 'admin' if email in ADMIN_EMAILS else 'user'
            
            # Generate referral code
            referral_code = generate_referral_code(user_id, full_name or "User")
            
            user_doc = {
                'user_id': user_id,
                'apple_id': apple_user_id,
                'email': email,
                'name': full_name or (email.split('@')[0] if email else 'Utilisateur Apple'),
                'picture': f"https://ui-avatars.com/api/?name={(full_name or 'A').replace(' ', '+')}&background=04D182&color=000&bold=true",
                'provider': 'apple',
                'role': role,
                'created_at': now,
                'password_hash': None,
                'referral_code': referral_code,
                'referred_by': None,
                'referral_count': 0,
                'free_months_earned': 0,
                'free_months_remaining': 0,
                'subscription_end_date': None,
            }
            await db.users.insert_one(user_doc)
            logger.info(f"Apple Sign-In: new user created {user_id}")
        
        # Create JWT token
        token = create_jwt({'user_id': user_id, 'exp': int((now + timedelta(days=7)).timestamp())})
        user_data = await db.users.find_one({'user_id': user_id}, {'_id': 0, 'password_hash': 0})
        
        return {'token': token, 'user': user_data}
        
    except ValueError as e:
        logger.error(f"Apple auth validation error: {e}")
        raise HTTPException(401, str(e))
    except Exception as e:
        logger.error(f"Apple auth error: {e}")
        raise HTTPException(500, "Erreur d'authentification Apple")

@api_router.get("/auth/me")
async def me(request: Request):
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Non authentifié")
    user.pop('password_hash', None)
    return user

@api_router.post("/auth/logout")
async def logout():
    return {'message': 'Déconnexion réussie'}


# ─── Pre-Registration Routes ────────────────────────────────────────────────

class PreRegistrationRequest(BaseModel):
    email: EmailStr
    prenom: str

@api_router.post("/preregistration")
async def preregister(body: PreRegistrationRequest):
    email = body.email.lower().strip()
    prenom = body.prenom.strip()
    if not prenom:
        raise HTTPException(400, "Le prénom est requis")
    existing = await db.preregistrations.find_one({'email': email})
    if existing:
        return {'message': 'Vous êtes déjà pré-inscrit(e). Merci !', 'already_registered': True}
    await db.preregistrations.insert_one({
        'email': email,
        'prenom': prenom,
        'created_at': datetime.now(timezone.utc).isoformat(),
    })
    count = await db.preregistrations.count_documents({})
    return {'message': 'Pré-inscription confirmée ! Vous serez informé(e) en priorité.', 'already_registered': False, 'total': count}

@api_router.get("/preregistration/count")
async def preregistration_count():
    count = await db.preregistrations.count_documents({})
    return {'count': count}

@api_router.get("/admin/preregistrations")
async def admin_list_preregistrations(request: Request):
    await require_admin(request)
    preinscriptions = await db.preregistrations.find({}, {'_id': 0}).sort('created_at', -1).to_list(1000)
    return preinscriptions


# ─── Scholar Routes ─────────────────────────────────────────────────────────

def _normalize_scholar_photo(s: dict) -> dict:
    """Ensure both 'photo_url' and 'photo' fields are populated from whichever exists.
    Avoids legacy/new field mismatch (admin panel writes 'photo', frontend reads 'photo_url')."""
    pu = s.get('photo_url') or s.get('photo')
    if pu:
        s['photo_url'] = pu
        s['photo'] = pu
    return s

@api_router.get("/scholars")
async def get_scholars():
    scholars = await db.scholars.find({'is_active': {'$ne': False}}, {'_id': 0}).to_list(100)
    return [_normalize_scholar_photo(s) for s in scholars]

@api_router.get("/scholars/{scholar_id}")
async def get_scholar(scholar_id: str):
    s = await db.scholars.find_one({'id': scholar_id}, {'_id': 0})
    if not s:
        raise HTTPException(404, "Érudit non trouvé")
    return _normalize_scholar_photo(s)

# ─── Course Routes ──────────────────────────────────────────────────────────

@api_router.get("/catalogue")
async def get_catalogue():
    """Public catalogue view: returns module-level granularity for multi-module launch courses,
    falls back to course-level for single-module or module-less courses.
    Each item = one 'card' to render on the Catalogue page.
    """
    launch_courses = await db.courses.find(
        {'is_launch_catalog': True, 'is_active': {'$ne': False}},
        {'_id': 0}
    ).sort('order', 1).to_list(200)
    if not launch_courses:
        return []
    course_ids = [c['id'] for c in launch_courses]
    modules = await db.modules.find(
        {'course_id': {'$in': course_ids}, 'is_active': {'$ne': False}},
        {'_id': 0}
    ).sort('order', 1).to_list(1000)
    audios = await db.audios.find(
        {'course_id': {'$in': course_ids}},
        {'_id': 0, 'course_id': 1, 'module_id': 1, 'published_at': 1, 'available_date': 1}
    ).to_list(3000)
    mods_by_course: dict = {}
    for m in modules:
        mods_by_course.setdefault(m['course_id'], []).append(m)
    audios_per_module: dict = {}
    audios_per_course_unassigned: dict = {}
    # Episode-level date aggregation: collect every `published_at` / `available_date`
    # set on at least one episode of the course. Used downstream to compute the
    # effective `available_date` displayed on the catalogue card.
    ep_dates_per_course: dict = {}
    for a in audios:
        cid = a.get('course_id')
        mid = a.get('module_id')
        if mid:
            audios_per_module[mid] = audios_per_module.get(mid, 0) + 1
        else:
            audios_per_course_unassigned[cid] = audios_per_course_unassigned.get(cid, 0) + 1
        # Prefer explicit `available_date` (free-form, ex. "juillet 2026"),
        # fall back to `published_at` (YYYY-MM normalised below).
        d = (a.get('available_date') or a.get('published_at') or '').strip()
        if d:
            ep_dates_per_course.setdefault(cid, []).append(d)

    # ─── Helpers to convert YYYY-MM → "mois YYYY" and compute effective date ──
    _MONTHS_FR = {
        '01': 'janvier', '02': 'février', '03': 'mars', '04': 'avril',
        '05': 'mai', '06': 'juin', '07': 'juillet', '08': 'août',
        '09': 'septembre', '10': 'octobre', '11': 'novembre', '12': 'décembre',
    }

    def _normalise_episode_date(raw: str) -> str:
        """`2026-07` → `juillet 2026`. Other formats passed through unchanged."""
        if not raw:
            return ''
        s = raw.strip()
        # ISO-ish: YYYY-MM or YYYY-MM-DD
        if len(s) >= 7 and s[4] == '-' and s[5:7] in _MONTHS_FR:
            return f"{_MONTHS_FR[s[5:7]]} {s[:4]}"
        return s

    def _effective_available_date(course_doc, episode_dates):
        """Return the date shown on the catalogue card.

        Priority:
        1. Earliest episode-level `published_at` / `available_date` if any episode
           has one (this lets the user edit dates per-episode in the admin).
        2. Course-level `available_date` field (legacy fallback).

        When multiple distinct dates exist, we return a range
        `"<earliest> → <latest>"` sorted CHRONOLOGICALLY when the raw value is
        in ISO format (`YYYY-MM`); otherwise we fall back to a single date
        (the first non-empty) to avoid showing a confusing alphabetic range
        on free-form labels like "juillet 2026" vs "mai 2026".
        """
        # Build (sort_key, display) pairs; sort_key keeps ISO ordering when possible.
        pairs = []
        seen_display = set()
        for d in (episode_dates or []):
            if not d:
                continue
            raw = d.strip()
            display = _normalise_episode_date(raw)
            if display in seen_display:
                continue
            seen_display.add(display)
            # ISO YYYY-MM → use raw for chronological sort; else use display
            sort_key = raw if (len(raw) >= 7 and raw[4] == '-' and raw[5:7] in _MONTHS_FR) else display
            pairs.append((sort_key, display))
        if pairs:
            pairs.sort(key=lambda p: p[0])
            displays = [p[1] for p in pairs]
            return displays[0] if len(displays) == 1 else f"{displays[0]} → {displays[-1]}"
        return course_doc.get('available_date')

    # Pre-load scholars name map (used to join primary + co-intervenant names on each card)
    all_scholars = await db.scholars.find({}, {'_id': 0, 'id': 1, 'name': 1}).to_list(200)
    scholar_name_by_id = {s['id']: s.get('name', '') for s in all_scholars}

    def _scholar_label(c_doc):
        """Return 'Primary · Co1 · Co2' or fallback to scholar_name field."""
        names = []
        primary = scholar_name_by_id.get(c_doc.get('scholar_id'))
        if primary:
            names.append(primary)
        for co_id in (c_doc.get('co_scholar_ids') or []):
            n = scholar_name_by_id.get(co_id)
            if n and n not in names:
                names.append(n)
        return ' · '.join(names) if names else c_doc.get('scholar_name')

    items = []
    for c in launch_courses:
        course_modules = mods_by_course.get(c['id'], [])
        # Compute the date that will appear on this course's cards (used for both
        # module-level rows and the course-level fallback row).
        effective_date = _effective_available_date(c, ep_dates_per_course.get(c['id'], []))
        # Decide whether to split this course into per-module cards.
        # Only do so when at least 2 modules carry actual content (audios)
        # OR are explicitly flagged `is_launch_catalog`. Otherwise show a
        # single course-level card so the editorial summary stays visible.
        modules_with_content = [
            m for m in course_modules
            if audios_per_module.get(m['id'], 0) > 0 or m.get('is_launch_catalog')
        ]
        if len(modules_with_content) >= 2:
            visible_modules = []
            for m in course_modules:
                ep_count = audios_per_module.get(m['id'], 0)
                # Hide modules without any episode (old seed artifacts not part of the launch catalog)
                if ep_count == 0 and not m.get('is_launch_catalog'):
                    continue
                visible_modules.append({
                    'type': 'module',
                    'id': f"{c['id']}::{m['id']}",
                    'course_id': c['id'],
                    'module_id': m['id'],
                    'title': m.get('name') or m.get('title') or '',
                    'description': m.get('description') or '',
                    'cursus_id': c.get('cursus_id') or c.get('thematique_id'),
                    'course_title': c.get('title'),
                    'scholar_name': m.get('scholar_name') or _scholar_label(c),
                    'episode_count': ep_count,
                    'coming_soon': c.get('coming_soon', False) or ep_count == 0,
                    'available_date': effective_date,
                    'recruiting': c.get('recruiting', False),
                    'order': (c.get('order') or 0) * 1000 + (m.get('order') or 0),
                })
            if visible_modules:
                items.extend(visible_modules)
            else:
                # Fallback: course has multiple modules but none visible — show course-level as "Bientôt"
                items.append({
                    'type': 'course',
                    'id': c['id'],
                    'course_id': c['id'],
                    'module_id': None,
                    'title': c.get('title') or c.get('name') or '',
                    'description': c.get('description', ''),
                    'cursus_id': c.get('cursus_id') or c.get('thematique_id'),
                    'course_title': c.get('title'),
                    'scholar_name': _scholar_label(c),
                    'episode_count': 0,
                    'coming_soon': True,
                    'available_date': effective_date,
                    'recruiting': c.get('recruiting', False),
                    'order': (c.get('order') or 0) * 1000,
                })
        else:
            ep_count = sum(audios_per_module.get(m['id'], 0) for m in course_modules)
            ep_count += audios_per_course_unassigned.get(c['id'], 0)
            # Count course-level YouTube URL as +1 episode ONLY when no audio
            # already carries a YT URL (avoids double-counting mono-video courses
            # where the same YT URL is denormalised on both course AND audio).
            if c.get('youtube_url') and ep_count == 0:
                ep_count += 1
            # Prefer the short `summary` (description courte) for the catalogue
            # card. Fall back to the long `description` if no summary set.
            card_blurb = c.get('summary') or c.get('description', '')
            items.append({
                'type': 'course',
                'id': c['id'],
                'course_id': c['id'],
                'module_id': None,
                'title': c.get('title') or c.get('name') or '',
                'description': card_blurb,
                'summary': c.get('summary', ''),
                'cursus_id': c.get('cursus_id') or c.get('thematique_id'),
                'course_title': c.get('title'),
                'scholar_name': _scholar_label(c),
                'episode_count': ep_count,
                'coming_soon': c.get('coming_soon', False),
                'available_date': effective_date,
                'recruiting': c.get('recruiting', False),
                'order': (c.get('order') or 0) * 1000,
            })
    items.sort(key=lambda x: (bool(x['coming_soon']), x['order']))
    return items



@api_router.get("/courses")
async def get_courses(request: Request, topic: Optional[str] = None, level: Optional[str] = None, scholar_id: Optional[str] = None, thematique_id: Optional[str] = None, cursus_id: Optional[str] = None):
    query: dict = {'is_active': {'$ne': False}}  # Only show active courses
    if topic:
        query['topic'] = topic
    if level:
        query['level'] = level
    if scholar_id:
        # Match courses where the scholar is primary OR co-intervenant
        query['$or'] = [{'scholar_id': scholar_id}, {'co_scholar_ids': scholar_id}]
    # Support both old (thematique_id) and new (cursus_id) field names
    filter_id = cursus_id or thematique_id
    if filter_id:
        query['$or'] = [{'cursus_id': filter_id}, {'thematique_id': filter_id}]
    courses = await db.courses.find(query, {'_id': 0}).to_list(100)
    # Determine access per-course to strip youtube_url for non-subscribers
    user = await get_current_user(request)
    for c in courses:
        if c.get('title'):
            c['title'] = clean_title(c['title'])
        if c.get('youtube_url'):
            has_access = False
            if user:
                access = await check_user_access(user['user_id'], content_type='course', content_id=c['id'])
                has_access = access.get('has_access', False)
            if not has_access:
                c.pop('youtube_url', None)
    return courses

@api_router.get("/courses/{course_id}/playlist")
async def get_course_playlist(course_id: str):
    """Return ordered list of audios for a course."""
    # Fetch audios directly by course_id (module_id may not be set)
    audios = await db.audios.find(
        {'course_id': course_id, 'is_active': True}, {'_id': 0}
    ).sort('episode_number', 1).to_list(200)

    # Also fetch modules for naming
    modules = await db.modules.find(
        {'course_id': course_id, 'is_active': True}, {'_id': 0}
    ).sort('order', 1).to_list(200)
    module_map = {mod['id']: mod for mod in modules}

    playlist = []
    for idx, audio in enumerate(audios):
        audio['stream_url'] = resolve_audio_url(audio)
        mod = module_map.get(audio.get('module_id', ''))
        playlist.append({
            'module_id': audio.get('module_id', ''),
            'module_name': clean_title(mod.get('name', '')) if mod else clean_title(audio.get('title', '')),
            'module_order': mod.get('order', idx + 1) if mod else idx + 1,
            'audio_id': audio['id'],
            'audio_title': clean_title(audio.get('title', '')),
            'stream_url': audio.get('stream_url', ''),
            'thumbnail': audio.get('thumbnail', ''),
            'duration': audio.get('duration', 0),
        })
    return playlist

@api_router.get("/courses/featured")
async def get_featured_course():
    """Get the featured course for homepage highlight."""
    course = await db.courses.find_one({'is_featured': True, 'is_active': True}, {'_id': 0})
    if course and course.get('title'):
        course['title'] = clean_title(course['title'])
    return course

@api_router.get("/courses/{course_id}")
async def get_course(course_id: str, request: Request):
    c = await db.courses.find_one({'id': course_id}, {'_id': 0})
    if not c:
        raise HTTPException(404, "Cours non trouvé")
    if c.get('title'):
        c['title'] = clean_title(c['title'])
    # Live-sync scholar_name from scholars collection (single source of truth).
    # The denormalised field on courses can drift when an admin renames a prof.
    # Includes primary + co-intervenants joined by " · ".
    primary_id = c.get('scholar_id')
    co_ids = c.get('co_scholar_ids') or []
    relevant_ids = [sid for sid in [primary_id] + list(co_ids) if sid]
    if relevant_ids:
        scholars_docs = await db.scholars.find(
            {'id': {'$in': relevant_ids}}, {'_id': 0, 'id': 1, 'name': 1}
        ).to_list(20)
        scholar_by_id = {s['id']: s.get('name', '') for s in scholars_docs}
        ordered_names = []
        for sid in relevant_ids:
            n = scholar_by_id.get(sid)
            if not n:
                continue
            # Strip "Dr. " prefix for the display name.
            if n.lower().startswith('dr. '):
                n = n[4:]
            if n and n not in ordered_names:
                ordered_names.append(n)
        if ordered_names:
            c['scholar_name'] = ' · '.join(ordered_names)
    # Live-sync modules_count from the modules collection (avoid stale seed values).
    try:
        c['modules_count'] = await db.modules.count_documents(
            {'course_id': course_id, 'is_active': {'$ne': False}}
        )
    except Exception:
        pass
    # Strip protected fields (youtube_url) if user has no access
    user = await get_current_user(request)
    has_access = False
    if user:
        access = await check_user_access(user['user_id'], content_type='course', content_id=course_id)
        has_access = access.get('has_access', False)
    if not has_access:
        c.pop('youtube_url', None)
        c.pop('course_resources', None)
    return c

@api_router.get("/courses/{course_id}/suggestions")
async def get_course_suggestions(course_id: str):
    """Get 'Pour approfondir' suggestions based on course theme."""
    course = await db.courses.find_one({'id': course_id}, {'_id': 0})
    if not course:
        raise HTTPException(404, "Cours non trouvé")
    
    thematique_id = course.get('thematique_id')
    suggestions = {
        'conferences': [],
        'bibliographies': [],
        'related_courses': []
    }
    
    if thematique_id:
        # Get conferences on same theme
        suggestions['conferences'] = await db.conferences.find(
            {'thematique_id': thematique_id, 'is_active': True}, 
            {'_id': 0}
        ).to_list(5)
        
        # Get bibliographies on same theme
        suggestions['bibliographies'] = await db.bibliographies.find(
            {'thematique_id': thematique_id}, 
            {'_id': 0}
        ).to_list(5)
        
        # Get other courses on same theme (excluding current)
        suggestions['related_courses'] = await db.courses.find(
            {'thematique_id': thematique_id, 'id': {'$ne': course_id}, 'is_active': True}, 
            {'_id': 0}
        ).to_list(3)
    
    return suggestions

# ─── Conference Routes ─────────────────────────────────────────────────────────

@api_router.get("/conferences")
async def get_conferences(thematique_id: Optional[str] = None):
    """Get all conferences, optionally filtered by theme."""
    query: dict = {'is_active': True}
    if thematique_id:
        query['thematique_id'] = thematique_id
    conferences = await db.conferences.find(query, {'_id': 0}).to_list(100)
    return conferences

@api_router.get("/conferences/{conference_id}")
async def get_conference(conference_id: str):
    conf = await db.conferences.find_one({'id': conference_id}, {'_id': 0})
    if not conf:
        raise HTTPException(404, "Conférence non trouvée")
    return conf

# ─── Audio Routes ───────────────────────────────────────────────────────────

@api_router.get("/audios")
async def get_audios(request: Request, topic: Optional[str] = None, audio_type: Optional[str] = None, scholar_id: Optional[str] = None, module_id: Optional[str] = None, course_id: Optional[str] = None):
    query: dict = {'is_active': True}
    if topic:
        query['topic'] = topic
    if audio_type:
        query['type'] = audio_type
    if scholar_id:
        query['scholar_id'] = scholar_id
    if module_id:
        query['module_id'] = module_id
    if course_id:
        query['course_id'] = course_id
    audios = await db.audios.find(query, {'_id': 0}).to_list(100)
    # Determine access once per course to strip youtube_url for non-subscribers
    user = await get_current_user(request)
    has_access_cache: dict = {}
    for a in audios:
        a['stream_url'] = resolve_audio_url(a)
        cid = a.get('course_id', '')
        if cid not in has_access_cache:
            if user:
                access = await check_user_access(user['user_id'], content_type='course', content_id=cid)
                has_access_cache[cid] = access.get('has_access', False)
            else:
                has_access_cache[cid] = False
        if not has_access_cache[cid]:
            a.pop('youtube_url', None)
            a.pop('r2_audio_key', None)
            a.pop('r2_video_key', None)
            a.pop('episode_resources', None)
    return audios

@api_router.get("/audios/{audio_id}")
async def get_audio(audio_id: str, request: Request):
    a = await db.audios.find_one({'id': audio_id, 'is_active': True}, {'_id': 0})
    if not a:
        raise HTTPException(404, "Audio non trouvé")
    a['stream_url'] = resolve_audio_url(a)
    # Gate youtube_url behind access check
    user = await get_current_user(request)
    has_access = False
    if user and a.get('course_id'):
        access = await check_user_access(user['user_id'], content_type='course', content_id=a['course_id'])
        has_access = access.get('has_access', False)
    if not has_access:
        a.pop('youtube_url', None)

    # Enrich with course + cursus data
    course = await db.courses.find_one({'id': a.get('course_id', '')}, {'_id': 0}) if a.get('course_id') else None
    if course:
        a['scholar_name'] = course.get('scholar_name', '')
        a['scholar_id'] = course.get('scholar_id', '')
        a['description'] = a.get('description') or course.get('description', '')
        a['course_title'] = course.get('title', '')
        a['total_episodes'] = course.get('modules_count', 0)

        # Enrich cursus color/letter (7 cursus: A=Histoire, B=Théologie, C=Sciences, D=Arts, E=Falsafa, F=Mystique, G=Pensées non-islamiques)
        CURSUS_LETTERS = ['A', 'B', 'C', 'D', 'E', 'F', 'G']
        CURSUS_COLORS = ['#D97757', '#8B5CF6', '#EAD637', '#EC4899', '#04D182', '#06B6D4', '#F59E0B']
        cursus = await db.cursus.find_one({'id': course.get('cursus_id', '')}, {'_id': 0}) if course.get('cursus_id') else None
        if cursus:
            order = max(0, min(cursus.get('order', 1) - 1, len(CURSUS_LETTERS) - 1))
            a['cursus_id'] = cursus['id']
            a['cursus_name'] = cursus.get('name', '')
            a['cursus_letter'] = CURSUS_LETTERS[order]
            a['cursus_color'] = CURSUS_COLORS[order]
        else:
            a['cursus_letter'] = 'A'
            a['cursus_color'] = '#04D182'
    else:
        a['scholar_name'] = ''
        a['cursus_letter'] = 'A'
        a['cursus_color'] = '#04D182'

    return a

@api_router.get("/search")
async def search_content(q: str, limit: int = 20):
    """Search episodes and courses by keyword."""
    if not q or len(q.strip()) < 2:
        return {'audios': [], 'courses': [], 'total': 0}
    regex = {'$regex': q.strip(), '$options': 'i'}
    audios = await db.audios.find(
        {'$or': [{'title': regex}, {'scholar_name': regex}, {'description': regex}], 'is_active': True},
        {'_id': 0}
    ).limit(limit).to_list(limit)
    for a in audios:
        a['stream_url'] = resolve_audio_url(a)
    courses = await db.courses.find(
        {'$or': [{'title': regex}, {'scholar_name': regex}, {'description': regex}], 'is_active': True},
        {'_id': 0}
    ).limit(10).to_list(10)
    return {'audios': audios, 'courses': courses, 'total': len(audios) + len(courses)}


@api_router.post("/audios/{audio_id}/play")
async def track_play(audio_id: str):
    """Increment play count for an audio and its parent course."""
    await db.audios.update_one({'id': audio_id}, {'$inc': {'play_count': 1}})
    audio = await db.audios.find_one({'id': audio_id}, {'_id': 0, 'course_id': 1})
    if audio and audio.get('course_id'):
        await db.courses.update_one({'id': audio['course_id']}, {'$inc': {'play_count': 1}})
    return {'ok': True}


@api_router.get("/audios/{audio_id}/stream-url")
async def get_audio_stream_url(audio_id: str, request: Request):
    """Return a proxy streaming URL (avoids R2 CORS restrictions). Requires active subscription/trial/admin access."""
    # Authentication + access check
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    access = await check_user_access(user['user_id'], content_type='audio', content_id=audio_id)
    if not access.get('has_access'):
        raise HTTPException(403, {"error": "subscription_required", "reason": access.get('reason', 'no_access')})
    
    a = await db.audios.find_one({'id': audio_id}, {'_id': 0})
    if not a:
        raise HTTPException(404, "Audio non trouvé")
    
    # Sign a short-lived (1h) token for the /stream endpoint (which is hit by <audio src>, no headers)
    stream_token = create_jwt({
        'sub': user['user_id'],
        'audio_id': audio_id,
        'scope': 'audio_stream',
        'exp': int((datetime.now(timezone.utc) + timedelta(hours=1)).timestamp()),
    })
    
    file_key = a.get('file_key')
    if file_key and r2_client:
        scheme = request.headers.get('x-forwarded-proto', 'https')
        host = request.headers.get('x-forwarded-host') or request.headers.get('host', '')
        proxy_url = f"{scheme}://{host}/api/audios/{audio_id}/stream?t={stream_token}"
        return {
            'audio_id': audio_id,
            'stream_url': proxy_url,
            'file_key': file_key,
            'source': 'proxy',
            'expires_in': 3600,
        }
    # Fallback: presigned R2 URL if R2 unavailable
    stream_url = resolve_audio_url(a)
    return {
        'audio_id': audio_id,
        'stream_url': stream_url,
        'file_key': a.get('file_key'),
        'source': 'fallback',
        'expires_in': 3600,
    }

@api_router.api_route("/audios/{audio_id}/stream", methods=["GET", "HEAD"])
async def stream_audio(audio_id: str, request: Request, t: Optional[str] = None):
    """Proxy the audio file from R2 to the client. Requires a valid short-lived stream token (?t=)."""
    # Validate stream token (delivered by /stream-url after auth+access check)
    token = t or request.query_params.get('t')
    if not token:
        # Allow authenticated users via Authorization header as fallback (admin preview)
        user = await get_current_user(request)
        if not user:
            raise HTTPException(401, "Jeton de streaming requis")
        access = await check_user_access(user['user_id'], content_type='audio', content_id=audio_id)
        if not access.get('has_access'):
            raise HTTPException(403, "Abonnement requis")
    else:
        payload = verify_jwt(token)
        if not payload or payload.get('scope') != 'audio_stream' or payload.get('audio_id') != audio_id:
            raise HTTPException(403, "Jeton invalide ou expiré")
    
    import httpx
    a = await db.audios.find_one({'id': audio_id}, {'_id': 0})
    if not a:
        raise HTTPException(404, "Audio non trouvé")
    file_key = a.get('file_key')
    if not file_key or not r2_client:
        # Fallback: redirect to audio_url if set
        fallback = a.get('audio_url')
        if fallback:
            from fastapi.responses import RedirectResponse
            return RedirectResponse(url=fallback)
        raise HTTPException(404, "Fichier audio non disponible")

    range_header = request.headers.get('Range')
    try:
        get_kwargs: dict = {'Bucket': R2_BUCKET, 'Key': file_key}
        if range_header:
            get_kwargs['Range'] = range_header

        resp = r2_client.get_object(**get_kwargs)
        content_type = resp.get('ContentType', 'audio/mp4')
        body = resp['Body'].read()

        headers: dict = {
            'Content-Type': content_type,
            'Accept-Ranges': 'bytes',
            'Cache-Control': 'no-cache',
            'Access-Control-Allow-Origin': '*',
        }
        if 'ContentLength' in resp:
            headers['Content-Length'] = str(resp['ContentLength'])
        if 'ContentRange' in resp:
            headers['Content-Range'] = resp['ContentRange']

        status_code = 206 if range_header else 200
        from fastapi.responses import Response
        return Response(content=body, status_code=status_code, headers=headers, media_type=content_type)
    except ClientError as e:
        code = e.response.get('Error', {}).get('Code', '')
        if code in ('NoSuchKey', '404'):
            raise HTTPException(404, "Fichier audio non disponible")
        logging.getLogger(__name__).error(f"Stream error for {audio_id}: {e}")
        raise HTTPException(500, "Erreur de lecture du fichier audio")
    except Exception as e:
        logging.getLogger(__name__).error(f"Stream error for {audio_id}: {e}")
        raise HTTPException(500, "Erreur de lecture du fichier audio")

@api_router.get("/images/{file_path:path}")
async def serve_image(file_path: str, request: Request):
    """Serve images from R2 bucket (supports both images/ and Prof/ folders)."""
    from fastapi.responses import Response
    
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    # Determine the correct R2 key based on the file path
    if file_path.startswith('Prof_') or file_path.startswith('Prof/'):
        # Professor photos in Prof/ folder
        file_key = f"Prof/{file_path.replace('Prof/', '')}"
    elif file_path.startswith('images/') or '/' in file_path:
        # Already has folder path
        file_key = file_path
    else:
        # Default to images/ folder
        file_key = f"images/{file_path}"
    
    logger.info(f"Serving image: key={file_key}")
    
    try:
        resp = r2_client.get_object(Bucket=R2_BUCKET, Key=file_key)
        content_type = resp.get('ContentType', 'image/jpeg')
        body = resp['Body'].read()
        
        return Response(
            content=body,
            media_type=content_type,
            headers={
                'Cache-Control': 'public, max-age=86400',
                'Access-Control-Allow-Origin': '*',
            }
        )
    except ClientError as e:
        logger.error(f"R2 error for key={file_key}: {e}")
        if e.response['Error']['Code'] == 'NoSuchKey':
            raise HTTPException(404, "Image non trouvée")
        raise HTTPException(500, f"Erreur R2: {str(e)}")

class UploadUrlRequest(BaseModel):
    file_key: str
    content_type: str = 'audio/mpeg'

@api_router.post("/r2/upload-url")
async def get_upload_url(body: UploadUrlRequest, request: Request):
    """Get a presigned upload URL for putting a new audio file into R2."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    upload_url = get_presigned_upload_url(body.file_key, body.content_type)
    if not upload_url:
        raise HTTPException(500, "Impossible de générer l'URL d'upload")
    return {
        'upload_url': upload_url,
        'file_key': body.file_key,
        'bucket': R2_BUCKET,
        'expires_in': 3600,
    }

@api_router.get("/r2/files")
async def list_r2_files(prefix: Optional[str] = None):
    """List files in the R2 bucket (optionally filtered by prefix)."""
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    try:
        kwargs: dict = {'Bucket': R2_BUCKET, 'MaxKeys': 200}
        if prefix:
            kwargs['Prefix'] = prefix
        response = r2_client.list_objects_v2(**kwargs)
        files = [
            {
                'key': obj['Key'],
                'size': obj['Size'],
                'last_modified': obj['LastModified'].isoformat(),
                'stream_url': get_presigned_stream_url(obj['Key']),
            }
            for obj in response.get('Contents', [])
        ]
        return {'files': files, 'count': len(files), 'bucket': R2_BUCKET}
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

@api_router.put("/audios/{audio_id}/file-key")
async def update_audio_file_key(audio_id: str, request: Request):
    """Update the R2 file_key for an audio document."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    body = await request.json()
    file_key = body.get('file_key')
    if not file_key:
        raise HTTPException(400, "file_key requis")
    result = await db.audios.update_one({'id': audio_id}, {'$set': {'file_key': file_key}})
    if result.matched_count == 0:
        raise HTTPException(404, "Audio non trouvé")
    return {'message': 'file_key mis à jour', 'audio_id': audio_id, 'file_key': file_key}

# ─── Transcript Routes ─────────────────────────────────────────────────────────

class TranscriptCreate(BaseModel):
    audio_id: str
    title: Optional[str] = None

@api_router.get("/transcripts/{audio_id}")
async def get_transcript(audio_id: str, request: Request):
    """Get the transcript for an audio episode. Requires auth + active subscription."""
    await require_subscriber(request)
    transcript = await db.transcripts.find_one({'audio_id': audio_id}, {'_id': 0})
    if not transcript:
        raise HTTPException(404, "Transcript non trouvé")
    return transcript

@api_router.get("/audios/{audio_id}/transcript")
async def get_audio_transcript(audio_id: str, request: Request):
    """Get the transcript for an audio episode (alternative endpoint).
    Returns content only for subscribers; non-subscribers see only the existence flag."""
    transcript = await db.transcripts.find_one({'audio_id': audio_id}, {'_id': 0})
    if not transcript:
        # Return empty response instead of 404 (audio may not have transcript yet)
        return {'audio_id': audio_id, 'has_transcript': False}
    # Check subscription before returning content
    user = await get_current_user(request)
    is_subscriber = False
    if user:
        if user.get('role') == 'admin' or user.get('free_access'):
            is_subscriber = True
        else:
            now = datetime.now(timezone.utc)
            for fld in ('trial', 'subscription'):
                v = user.get(fld)
                if v and v.get('expires_at'):
                    exp = v['expires_at']
                    if isinstance(exp, str):
                        exp = datetime.fromisoformat(exp.replace('Z', '+00:00'))
                    if exp > now:
                        is_subscriber = True
                        break
    if not is_subscriber:
        # Strip premium fields, expose only existence + metadata
        return {
            'audio_id': audio_id,
            'has_transcript': True,
            'locked': True,
            'word_count': transcript.get('word_count'),
            'reading_time_minutes': transcript.get('reading_time_minutes'),
            'title': transcript.get('title'),
        }
    transcript['has_transcript'] = True
    return transcript

@api_router.post("/transcripts/upload")
async def upload_transcript(request: Request):
    """Upload a Word document and create/update a transcript."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin access required")
    
    form = await request.form()
    audio_id = form.get('audio_id')
    file = form.get('file')
    
    if not audio_id:
        raise HTTPException(400, "audio_id requis")
    if not file:
        raise HTTPException(400, "Fichier requis")
    
    # Check if audio exists
    audio = await db.audios.find_one({'id': audio_id}, {'_id': 0})
    if not audio:
        raise HTTPException(404, "Audio non trouvé")
    
    # Read and convert the Word document
    content = await file.read()
    try:
        result = convert_docx_to_markdown(content)
    except ValueError as e:
        raise HTTPException(400, str(e))
    
    # Create or update transcript
    transcript_data = {
        'transcript_id': f'tr_{uuid.uuid4().hex[:12]}',
        'audio_id': audio_id,
        'title': result['title'] or audio.get('title', ''),
        'content': result['content'],
        'sections': result['sections'],
        'word_count': result['word_count'],
        'reading_time_minutes': estimate_reading_time(result['word_count']),
        'created_at': datetime.now(timezone.utc).isoformat(),
        'updated_at': datetime.now(timezone.utc).isoformat(),
    }
    
    # Upsert (create or update)
    await db.transcripts.update_one(
        {'audio_id': audio_id},
        {'$set': transcript_data},
        upsert=True
    )
    
    # Update audio to indicate it has a transcript
    await db.audios.update_one(
        {'id': audio_id},
        {'$set': {'has_transcript': True}}
    )
    
    return {
        'success': True,
        'message': 'Transcript uploaded successfully',
        'transcript': transcript_data
    }

@api_router.delete("/transcripts/{audio_id}")
async def delete_transcript(audio_id: str, request: Request):
    """Delete a transcript."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin access required")
    
    result = await db.transcripts.delete_one({'audio_id': audio_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Transcript non trouvé")
    
    # Update audio to indicate it no longer has a transcript
    await db.audios.update_one(
        {'id': audio_id},
        {'$set': {'has_transcript': False}}
    )
    
    return {'success': True, 'message': 'Transcript deleted'}

@api_router.post("/transcripts/from-r2")
async def create_transcript_from_r2(request: Request):
    """Create a transcript from an existing Word file in R2."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin access required")
    
    body = await request.json()
    audio_id = body.get('audio_id')
    file_key = body.get('file_key')  # e.g., "transcripts/episode-01.docx"
    
    if not audio_id or not file_key:
        raise HTTPException(400, "audio_id et file_key requis")
    
    # Check if audio exists
    audio = await db.audios.find_one({'id': audio_id}, {'_id': 0})
    if not audio:
        raise HTTPException(404, "Audio non trouvé")
    
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    # Fetch the file from R2
    try:
        resp = r2_client.get_object(Bucket=R2_BUCKET, Key=file_key)
        content = resp['Body'].read()
    except ClientError as e:
        logger.error(f"R2 error for transcript key={file_key}: {e}")
        raise HTTPException(404, "Fichier non trouvé dans R2")
    
    # Convert the Word document
    try:
        result = convert_docx_to_markdown(content)
    except ValueError as e:
        raise HTTPException(400, str(e))
    
    # Create or update transcript
    transcript_data = {
        'transcript_id': f'tr_{uuid.uuid4().hex[:12]}',
        'audio_id': audio_id,
        'title': result['title'] or audio.get('title', ''),
        'content': result['content'],
        'sections': result['sections'],
        'word_count': result['word_count'],
        'reading_time_minutes': estimate_reading_time(result['word_count']),
        'r2_file_key': file_key,
        'created_at': datetime.now(timezone.utc).isoformat(),
        'updated_at': datetime.now(timezone.utc).isoformat(),
    }
    
    await db.transcripts.update_one(
        {'audio_id': audio_id},
        {'$set': transcript_data},
        upsert=True
    )
    
    await db.audios.update_one(
        {'id': audio_id},
        {'$set': {'has_transcript': True}}
    )
    
    return {
        'success': True,
        'message': 'Transcript created from R2',
        'transcript': transcript_data
    }

@api_router.post("/transcripts/sync-r2")
async def sync_transcripts_from_r2(request: Request):
    """Scan R2 for .docx files matching audio episodes and create transcripts."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin access required")
    
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    # Get all audios from DB
    audios = await db.audios.find({}, {'_id': 0, 'id': 1, 'file_key': 1, 'title': 1}).to_list(None)
    
    results = {
        'synced': [],
        'already_exists': [],
        'no_docx': [],
        'errors': []
    }
    
    for audio in audios:
        audio_id = audio.get('id', '')
        file_key = audio.get('file_key', '')
        if not file_key:
            continue
        
        # Build the expected .docx key by replacing extension
        # file_key can be like "cursus-a-falsafa/01-mouvement-traduction/episode-01.m4a"
        # or with "Audio/" prefix like "Audio/cursus-a-falsafa/..."
        # Try both: the direct path and without "Audio/" prefix
        base_key = file_key.rsplit('.', 1)[0]  # remove extension
        docx_keys_to_try = [f"{base_key}.docx"]
        
        # If key starts with "Audio/", also try without it
        if file_key.startswith('Audio/'):
            stripped = file_key[len('Audio/'):]
            docx_keys_to_try.append(f"{stripped.rsplit('.', 1)[0]}.docx")
        else:
            # Also try with "Audio/" prefix
            docx_keys_to_try.append(f"Audio/{base_key}.docx")
        
        # Check if transcript already exists
        existing = await db.transcripts.find_one({'audio_id': audio_id})
        if existing:
            results['already_exists'].append(audio_id)
            continue
        
        # Try to find .docx in R2
        docx_content = None
        found_key = None
        for docx_key in docx_keys_to_try:
            try:
                resp = r2_client.get_object(Bucket=R2_BUCKET, Key=docx_key)
                docx_content = resp['Body'].read()
                found_key = docx_key
                break
            except ClientError:
                continue
        
        if not docx_content:
            results['no_docx'].append(audio_id)
            continue
        
        # Convert and save
        try:
            result = convert_docx_to_markdown(docx_content)
            transcript_data = {
                'transcript_id': f'tr_{uuid.uuid4().hex[:12]}',
                'audio_id': audio_id,
                'title': result['title'] or audio.get('title', ''),
                'content': result['content'],
                'sections': result['sections'],
                'word_count': result['word_count'],
                'reading_time_minutes': estimate_reading_time(result['word_count']),
                'r2_file_key': found_key,
                'created_at': datetime.now(timezone.utc).isoformat(),
                'updated_at': datetime.now(timezone.utc).isoformat(),
            }
            
            await db.transcripts.update_one(
                {'audio_id': audio_id},
                {'$set': transcript_data},
                upsert=True
            )
            await db.audios.update_one(
                {'id': audio_id},
                {'$set': {'has_transcript': True}}
            )
            results['synced'].append({'audio_id': audio_id, 'r2_key': found_key, 'word_count': result['word_count']})
        except Exception as e:
            logger.error(f"Error syncing transcript for {audio_id}: {e}")
            results['errors'].append({'audio_id': audio_id, 'error': str(e)})
    
    return {
        'success': True,
        'summary': {
            'synced': len(results['synced']),
            'already_exists': len(results['already_exists']),
            'no_docx': len(results['no_docx']),
            'errors': len(results['errors']),
        },
        'details': results
    }

# ─── Article Routes ─────────────────────────────────────────────────────────

@api_router.get("/articles")
async def get_articles(topic: Optional[str] = None):
    query = {'topic': topic} if topic else {}
    articles = await db.articles.find(query, {'_id': 0}).to_list(100)
    return articles

@api_router.get("/articles/{article_id}")
async def get_article(article_id: str):
    a = await db.articles.find_one({'id': article_id}, {'_id': 0})
    if not a:
        raise HTTPException(404, "Article non trouvé")
    return a

# ─── Thematiques Routes ──────────────────────────────────────────────────────

@api_router.get("/thematiques")
async def get_thematiques():
    """Get all cursus (themes) ordered by position."""
    # Use only the cursus collection
    cursus = await db.cursus.find({'is_active': {'$ne': False}}, {'_id': 0}).sort('order', 1).to_list(100)
    
    # Batch count courses per cursus using aggregation (fix N+1 query)
    course_counts = await db.courses.aggregate([
        {'$match': {'is_active': {'$ne': False}}},
        {'$group': {'_id': '$cursus_id', 'count': {'$sum': 1}}}
    ]).to_list(100)
    count_map = {item['_id']: item['count'] for item in course_counts}
    
    # Add course count for each cursus
    for c in cursus:
        c['course_count'] = count_map.get(c['id'], 0)
    return cursus

@api_router.get("/thematiques/{thematique_id}")
async def get_thematique(thematique_id: str):
    t = await db.cursus.find_one({'id': thematique_id}, {'_id': 0})
    if not t:
        raise HTTPException(404, "Cursus non trouvé")
    return t

# ─── Masterclass Routes ──────────────────────────────────────────────────────

@api_router.get("/masterclasses")
async def get_masterclasses():
    """Get all masterclasses for the live tab."""
    masterclasses = await db.masterclasses.find({'is_active': {'$ne': False}}, {'_id': 0}).to_list(100)
    return masterclasses

@api_router.get("/masterclasses/{mc_id}")
async def get_masterclass(mc_id: str):
    mc = await db.masterclasses.find_one({'id': mc_id}, {'_id': 0})
    if not mc:
        raise HTTPException(404, "Masterclass non trouvée")
    return mc

@api_router.post("/masterclasses/{mc_id}/register")
async def register_masterclass(mc_id: str, request: Request):
    """Register current user for a masterclass."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    mc = await db.masterclasses.find_one({'id': mc_id})
    if not mc:
        raise HTTPException(404, "Masterclass non trouvée")
    # Check if already registered
    if user['user_id'] in mc.get('registered_users', []):
        raise HTTPException(400, "Déjà inscrit à cette masterclass")
    # Check capacity
    if mc.get('max_participants') and mc.get('current_participants', 0) >= mc['max_participants']:
        raise HTTPException(400, "Cette masterclass est complète")
    
    await db.masterclasses.update_one(
        {'id': mc_id},
        {'$addToSet': {'registered_users': user['user_id']}, '$inc': {'current_participants': 1}}
    )
    return {'message': 'Inscription confirmée', 'masterclass_id': mc_id}

# ─── Live Session Routes ─────────────────────────────────────────────────────

@api_router.get("/live-sessions")
async def get_live_sessions():
    sessions = await db.live_sessions.find({}, {'_id': 0}).sort('date', 1).to_list(100)
    return sessions

@api_router.get("/live-sessions/{session_id}")
async def get_live_session(session_id: str):
    s = await db.live_sessions.find_one({'id': session_id}, {'_id': 0})
    if not s:
        raise HTTPException(404, "Session non trouvée")
    return s

@api_router.post("/live-sessions/{session_id}/register")
async def register_live_session(session_id: str, request: Request):
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    session = await db.live_sessions.find_one({'id': session_id})
    if not session:
        raise HTTPException(404, "Session non trouvée")
    await db.live_sessions.update_one(
        {'id': session_id},
        {'$addToSet': {'registered_users': user['user_id']}, '$inc': {'registered_count': 1}}
    )
    await db.users.update_one({'user_id': user['user_id']}, {'$addToSet': {'registered_sessions': session_id}})
    return {'message': 'Inscription confirmée'}

@api_router.delete("/live-sessions/{session_id}/register")
async def unregister_live_session(session_id: str, request: Request):
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    await db.live_sessions.update_one(
        {'id': session_id},
        {'$pull': {'registered_users': user['user_id']}, '$inc': {'registered_count': -1}}
    )
    await db.users.update_one({'user_id': user['user_id']}, {'$pull': {'registered_sessions': session_id}})
    return {'message': 'Désinscription effectuée'}

# ─── Home Route ──────────────────────────────────────────────────────────────

@api_router.get("/home")
async def get_home(request: Request):
    user = await get_current_user(request)
    user_id = user['user_id'] if user else None

    # Pre-load all cursus for fast lookup (7 cursus: A=Histoire, B=Théologie, C=Sciences, D=Arts, E=Falsafa, F=Mystique, G=Pensées non-islamiques)
    CURSUS_LETTERS = ['A', 'B', 'C', 'D', 'E', 'F', 'G']
    CURSUS_COLORS = ['#D97757', '#8B5CF6', '#EAD637', '#EC4899', '#04D182', '#06B6D4', '#F59E0B']
    all_cursus_list = await db.cursus.find({'is_active': True}, {'_id': 0}).sort('order', 1).to_list(20)
    cursus_map = {c['id']: c for c in all_cursus_list}

    def enrich_cursus(item: dict) -> dict:
        cid = item.get('cursus_id', '')
        c = cursus_map.get(cid, {})
        order = max(0, min(c.get('order', 1) - 1, len(CURSUS_LETTERS) - 1))
        item['cursus_name'] = c.get('name', '')
        item['cursus_letter'] = CURSUS_LETTERS[order]
        item['cursus_color'] = CURSUS_COLORS[order]
        # Clean title prefixes
        if item.get('title'):
            item['title'] = clean_title(item['title'])
        return item

    # 1. Featured hero: check highlight config first
    featured_hero = None
    
    # Get highlight configuration
    highlight_config = await db.config.find_one({'key': 'highlight_config'}, {'_id': 0})
    highlight_mode = highlight_config.get('mode', 'manual') if highlight_config else 'manual'
    logger.info(f"Highlight mode: {highlight_mode}")
    
    if highlight_mode == 'random':
        # Random mode: pick a random active course or cursus
        import random
        
        # Get all active courses and cursus
        active_courses = await db.courses.find({'is_active': True}, {'_id': 0}).to_list(100)
        active_cursus = await db.cursus.find({'is_active': True}, {'_id': 0}).to_list(50)
        logger.info(f"Random mode: {len(active_courses)} courses, {len(active_cursus)} cursus")
        
        # Combine and pick random
        all_items = []
        for c in active_courses:
            all_items.append(('course', c))
        for c in active_cursus:
            all_items.append(('cursus', c))
        
        if all_items:
            item_type, item = random.choice(all_items)
            logger.info(f"Random picked: {item_type} - {item.get('id')}")
            if item_type == 'cursus':
                order = max(0, min(item.get('order', 1) - 1, len(CURSUS_LETTERS) - 1))
                featured_hero = {
                    'id': item['id'],
                    'hero_type': 'cursus',
                    'title': item.get('hero_title') or item.get('name', ''),
                    'description': item.get('description', ''),
                    'cursus_id': item['id'],
                    'cursus_letter': CURSUS_LETTERS[order],
                    'cursus_color': CURSUS_COLORS[order],
                    'cursus_name': item.get('name', ''),
                }
            else:
                item = enrich_cursus(item)
                item['hero_type'] = 'course'
                if item.get('hero_title'):
                    item['title'] = item['hero_title']
                featured_hero = item
        else:
            logger.info("Random mode: No items found")
    else:
        # Manual mode: check for featured cursus first, then course
        featured_cursus_doc = await db.cursus.find_one({'is_featured': True, 'is_active': True}, {'_id': 0})
        if featured_cursus_doc:
            order = max(0, min(featured_cursus_doc.get('order', 1) - 1, len(CURSUS_LETTERS) - 1))
            featured_hero = {
                'id': featured_cursus_doc['id'],
                'hero_type': 'cursus',
                'title': featured_cursus_doc.get('hero_title') or featured_cursus_doc.get('name', ''),
                'description': featured_cursus_doc.get('description', ''),
                'cursus_id': featured_cursus_doc['id'],
                'cursus_letter': CURSUS_LETTERS[order],
                'cursus_color': CURSUS_COLORS[order],
                'cursus_name': featured_cursus_doc.get('name', ''),
            }
        else:
            featured_course = await db.courses.find_one({'is_featured': True, 'is_active': True}, {'_id': 0})
            if not featured_course:
                featured_course = await db.courses.find_one({'is_active': True}, {'_id': 0})
            if featured_course:
                featured_course = enrich_cursus(featured_course)
                featured_course['hero_type'] = 'course'
                if featured_course.get('hero_title'):
                    featured_course['title'] = featured_course['hero_title']
                featured_hero = featured_course

    # 2. Continue watching (last 3 in-progress audios)
    continue_watching = []
    if user_id:
        progress_items = await db.user_progress.find(
            {'user_id': user_id, 'content_type': 'audio', 'completed': {'$ne': True}, 'progress': {'$gt': 0.01}},
            {'_id': 0}
        ).sort('updated_at', -1).limit(3).to_list(3)
        
        if progress_items:
            # Batch fetch audios (fix N+1)
            audio_ids = [p['content_id'] for p in progress_items]
            audios_list = await db.audios.find({'id': {'$in': audio_ids}}, {'_id': 0}).to_list(10)
            audio_map = {a['id']: a for a in audios_list}
            
            # Batch fetch courses for these audios (fix nested N+1)
            course_ids = [a.get('course_id') for a in audios_list if a.get('course_id')]
            courses_list = await db.courses.find({'id': {'$in': course_ids}}, {'_id': 0, 'title': 1, 'cursus_id': 1, 'id': 1}).to_list(10)
            course_map = {c['id']: c for c in courses_list}
            
            for p in progress_items:
                audio = audio_map.get(p['content_id'])
                if audio:
                    audio = dict(audio)  # Make a copy to avoid modifying cached data
                    audio['stream_url'] = resolve_audio_url(audio)
                    if audio.get('course_id'):
                        course_info = course_map.get(audio['course_id'])
                        if course_info:
                            audio['cursus_id'] = course_info.get('cursus_id', '')
                            audio = enrich_cursus(audio)
                    continue_watching.append({
                        'audio': audio,
                        'progress': p.get('progress', 0),
                        'position': p.get('position', 0),
                    })

    # 3. Recent episodes (last 8 audios, enriched with cursus)
    recent_audios_raw = await db.audios.find({'is_active': True}, {'_id': 0}).sort('published_at', -1).limit(8).to_list(8)
    
    # Batch fetch courses for recent audios (fix N+1)
    recent_course_ids = [a.get('course_id') for a in recent_audios_raw if a.get('course_id')]
    recent_courses = await db.courses.find({'id': {'$in': recent_course_ids}}, {'_id': 0, 'cursus_id': 1, 'id': 1}).to_list(10)
    recent_course_map = {c['id']: c for c in recent_courses}
    
    recent_episodes = []
    for audio in recent_audios_raw:
        audio = dict(audio)  # Make a copy
        audio['stream_url'] = resolve_audio_url(audio)
        if audio.get('course_id'):
            course_info = recent_course_map.get(audio['course_id'])
            if course_info:
                audio['cursus_id'] = course_info.get('cursus_id', '')
        audio = enrich_cursus(audio)
        recent_episodes.append(audio)

    # 4. Recommendations (6 active courses, enriched with cursus)
    recommendations_raw = await db.courses.find({'is_active': True}, {'_id': 0}).limit(6).to_list(6)
    recommendations = [enrich_cursus(c) for c in recommendations_raw]

    # 5. Scholars
    scholars_list = await db.scholars.find({'is_active': True}, {'_id': 0}).to_list(10)

    # 6. Top 5 courses (admin config + auto fill by play_count), enriched with cursus
    top5_courses = []
    config = await db.config.find_one({'key': 'top10_courses'}, {'_id': 0})
    if config and config.get('course_ids'):
        # Batch fetch top courses (fix N+1)
        top_course_ids = config['course_ids'][:5]
        top_courses_raw = await db.courses.find({'id': {'$in': top_course_ids}, 'is_active': True}, {'_id': 0}).to_list(5)
        # Preserve order from config
        top_course_map = {c['id']: c for c in top_courses_raw}
        for cid in top_course_ids:
            if cid in top_course_map:
                top5_courses.append(enrich_cursus(top_course_map[cid]))
    if len(top5_courses) < 5:
        existing_ids = [c['id'] for c in top5_courses]
        extra = await db.courses.find(
            {'is_active': True, 'id': {'$nin': existing_ids}},
            {'_id': 0}
        ).sort('play_count', -1).limit(5 - len(top5_courses)).to_list(5)
        top5_courses.extend([enrich_cursus(c) for c in extra])

    return {
        'featured_hero': featured_hero,
        'featured_course': featured_hero,  # Keep for backwards compatibility
        'continue_watching': continue_watching,
        'recent_episodes': recent_episodes,
        'recommendations': recommendations,
        'scholars': scholars_list,
        'top5_courses': top5_courses,
    }

# ─── User Progress Routes ────────────────────────────────────────────────────

@api_router.get("/user/progress")
async def get_user_progress(request: Request):
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    items = await db.user_progress.find({'user_id': user['user_id']}, {'_id': 0}).to_list(500)
    return items

@api_router.post("/user/progress")
async def update_progress(body: ProgressRequest, request: Request):
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    now = datetime.now(timezone.utc)
    await db.user_progress.update_one(
        {'user_id': user['user_id'], 'content_id': body.content_id},
        {'$set': {
            'user_id': user['user_id'],
            'content_id': body.content_id,
            'content_type': body.content_type,
            'progress': body.progress,
            'position': body.position,
            'updated_at': now,
            'completed': body.progress >= 0.95
        }},
        upsert=True
    )
    return {'message': 'Progression mise à jour'}

# ─── User Favorites Routes ───────────────────────────────────────────────────

@api_router.get("/user/favorites")
async def get_favorites(request: Request):
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    favs = await db.user_favorites.find({'user_id': user['user_id']}, {'_id': 0}).to_list(500)
    
    # Group favorites by content type for batch fetching (fix N+1)
    favs_by_type = {}
    for fav in favs:
        ctype = fav.get('content_type', 'audio')
        if ctype not in favs_by_type:
            favs_by_type[ctype] = []
        favs_by_type[ctype].append(fav)
    
    # Batch fetch content for each type
    content_maps = {}
    coll_map = {'course': db.courses, 'audio': db.audios, 'article': db.articles}
    for ctype, type_favs in favs_by_type.items():
        coll = coll_map.get(ctype, db.audios)
        content_ids = [f['content_id'] for f in type_favs]
        contents = await coll.find({'id': {'$in': content_ids}}, {'_id': 0}).to_list(100)
        content_maps[ctype] = {c['id']: c for c in contents}
    
    # Build result
    result = []
    for fav in favs:
        ctype = fav.get('content_type', 'audio')
        content = content_maps.get(ctype, {}).get(fav['content_id'])
        if content:
            result.append({'favorite': fav, 'content': content})
    return result

@api_router.post("/user/favorites")
async def add_favorite(body: FavoriteRequest, request: Request):
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    existing = await db.user_favorites.find_one({'user_id': user['user_id'], 'content_id': body.content_id})
    if existing:
        return {'message': 'Déjà sauvegardé'}
    await db.user_favorites.insert_one({
        'user_id': user['user_id'],
        'content_id': body.content_id,
        'content_type': body.content_type,
        'saved_at': datetime.now(timezone.utc)
    })
    return {'message': 'Sauvegardé dans votre bibliothèque'}

@api_router.delete("/user/favorites/{content_type}/{content_id}")
async def remove_favorite(content_type: str, content_id: str, request: Request):
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    await db.user_favorites.delete_one({'user_id': user['user_id'], 'content_id': content_id})
    return {'message': 'Retiré de votre bibliothèque'}

# ─── User Referral Routes ─────────────────────────────────────────────────────

@api_router.get("/user/referral")
async def get_user_referral(request: Request):
    """Get user's referral code and stats."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    # Generate referral code if not exists
    if not user.get('referral_code'):
        referral_code = generate_referral_code(user['user_id'], user.get('name', 'User'))
        await db.users.update_one(
            {'user_id': user['user_id']},
            {'$set': {'referral_code': referral_code}}
        )
        user['referral_code'] = referral_code
    
    # Get referrals made by this user
    referrals = await db.referrals.find(
        {'referrer_id': user['user_id']},
        {'_id': 0}
    ).sort('created_at', -1).to_list(50)
    
    # Calculate stats
    total_referrals = len(referrals)
    converted_referrals = len([r for r in referrals if r.get('status') == 'converted'])
    pending_referrals = len([r for r in referrals if r.get('status') == 'pending'])
    
    return {
        'referral_code': user.get('referral_code'),
        'referral_count': user.get('referral_count', 0),
        'free_months_earned': user.get('free_months_earned', 0),
        'free_months_remaining': user.get('free_months_remaining', 0),
        'subscription_end_date': user.get('subscription_end_date'),
        'referred_by': user.get('referred_by'),
        'stats': {
            'total_referrals': total_referrals,
            'converted': converted_referrals,
            'pending': pending_referrals,
        },
        'referrals': referrals[:10],  # Last 10 referrals
    }

@api_router.get("/user/referrals")
async def get_user_referrals(request: Request):
    """Get detailed list of user's referrals (filleuls)."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    referrals = await db.referrals.find(
        {'referrer_id': user['user_id']},
        {'_id': 0}
    ).sort('created_at', -1).to_list(100)
    
    return {'referrals': referrals}

@api_router.post("/user/referral/validate")
async def validate_referral_code(body: ApplyReferralRequest):
    """Validate a referral code before registration."""
    code = body.referral_code.upper().strip()
    
    referrer = await db.users.find_one({'referral_code': code}, {'_id': 0, 'password_hash': 0})
    if not referrer:
        raise HTTPException(404, "Code de parrainage invalide")
    
    return {
        'valid': True,
        'referrer_name': referrer.get('name', 'Un membre Sijill'),
        'benefit': '1 mois gratuit pour vous et votre parrain'
    }

@api_router.post("/referrals/convert/{referee_id}")
async def convert_referral(referee_id: str, request: Request):
    """
    Called when a referred user (filleul) subscribes.
    Rewards the referrer (parrain) with 1 free month.
    """
    # This would typically be called by Stripe webhook or admin
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    # Find the referral
    referral = await db.referrals.find_one({'referee_id': referee_id, 'status': 'pending'})
    if not referral:
        raise HTTPException(404, "Parrainage non trouvé ou déjà converti")
    
    now = datetime.now(timezone.utc)
    
    # Update referral status
    await db.referrals.update_one(
        {'id': referral['id']},
        {'$set': {
            'status': 'converted',
            'converted_at': now,
            'referrer_rewarded': True,
        }}
    )
    
    # Reward the referrer with 1 free month
    referrer = await db.users.find_one({'user_id': referral['referrer_id']})
    if referrer:
        current_end = referrer.get('subscription_end_date') or now
        if current_end < now:
            current_end = now
        new_end = current_end + timedelta(days=30)
        
        await db.users.update_one(
            {'user_id': referral['referrer_id']},
            {
                '$inc': {
                    'referral_count': 1,
                    'free_months_earned': 1,
                    'free_months_remaining': 1,
                },
                '$set': {'subscription_end_date': new_end}
            }
        )
        logger.info(f"Referrer {referral['referrer_id']} rewarded with 1 free month")
        
        # Send conversion notification to referrer (if email configured)
        if is_email_configured():
            send_referral_conversion_notification(
                referrer_email=referrer.get('email', ''),
                referrer_name=referrer.get('name', 'Membre'),
                referee_name=referral.get('referee_name', 'Votre filleul'),
                free_months=1
            )
    
    return {'message': 'Parrainage converti, parrain récompensé'}

# ─── User Notification Preferences ─────────────────────────────────────────────

@api_router.get("/user/notifications/preferences")
async def get_notification_preferences(request: Request):
    """Get user notification preferences."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    prefs = await db.notification_preferences.find_one(
        {'user_id': user['user_id']}, 
        {'_id': 0}
    )
    
    # Return defaults if no preferences saved yet
    if not prefs:
        return {
            'new_courses': True,
            'new_episodes': True,
            'weekly_digest': True,
            'subscription_expiry': True,
            'subscription_reminder': True,
            'promotions': False,
        }
    
    return {
        'new_courses': prefs.get('new_courses', True),
        'new_episodes': prefs.get('new_episodes', True),
        'weekly_digest': prefs.get('weekly_digest', True),
        'subscription_expiry': prefs.get('subscription_expiry', True),
        'subscription_reminder': prefs.get('subscription_reminder', True),
        'promotions': prefs.get('promotions', False),
    }

@api_router.put("/user/notifications/preferences")
async def update_notification_preferences(request: Request):
    """Update user notification preferences."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    body = await request.json()
    
    await db.notification_preferences.update_one(
        {'user_id': user['user_id']},
        {'$set': {
            'user_id': user['user_id'],
            'new_courses': body.get('new_courses', True),
            'new_episodes': body.get('new_episodes', True),
            'weekly_digest': body.get('weekly_digest', True),
            'subscription_expiry': body.get('subscription_expiry', True),
            'subscription_reminder': body.get('subscription_reminder', True),
            'promotions': body.get('promotions', False),
            'updated_at': datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True
    )
    
    return {'message': 'Préférences sauvegardées'}

# ─── User Stats & Library ─────────────────────────────────────────────────────

@api_router.get("/user/stats")
async def get_user_stats(request: Request):
    """Get user statistics for profile page."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    user_id = user['user_id']
    
    # Count favorites
    favorites_count = await db.user_favorites.count_documents({'user_id': user_id})
    
    # Get progress data and calculate stats
    progress_items = await db.user_progress.find({'user_id': user_id}, {'_id': 0}).to_list(500)
    
    courses_followed = set()
    total_listening_seconds = 0
    completed_count = 0
    
    for p in progress_items:
        if p.get('content_type') == 'audio':
            # Get audio to find its duration and course
            audio = await db.audios.find_one({'id': p.get('content_id')}, {'_id': 0, 'duration': 1, 'course_id': 1})
            if audio:
                duration = audio.get('duration', 0)  # in seconds
                progress_pct = p.get('progress', 0)
                total_listening_seconds += int(duration * progress_pct)
                
                if audio.get('course_id'):
                    courses_followed.add(audio['course_id'])
                
                if p.get('completed'):
                    completed_count += 1
    
    # Convert seconds to hours
    total_listening_hours = round(total_listening_seconds / 3600, 1)
    
    return {
        'courses_followed': len(courses_followed),
        'listening_hours': total_listening_hours,
        'favorites_count': favorites_count,
        'completed_count': completed_count,
        'in_progress_count': len([p for p in progress_items if not p.get('completed') and p.get('progress', 0) > 0.01]),
    }

@api_router.get("/user/library")
async def get_user_library(request: Request):
    """Get user library data (in-progress, favorites, completed)."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    user_id = user['user_id']
    CURSUS_LETTERS = ['A', 'B', 'C', 'D', 'E', 'F', 'G']
    CURSUS_COLORS = ['#D97757', '#8B5CF6', '#EAD637', '#EC4899', '#04D182', '#06B6D4', '#F59E0B']
    
    # Pre-load cursus for enrichment
    all_cursus = await db.cursus.find({'is_active': True}, {'_id': 0}).sort('order', 1).to_list(20)
    cursus_map = {c['id']: c for c in all_cursus}
    
    def get_cursus_info(cursus_id: str):
        c = cursus_map.get(cursus_id, {})
        order = max(0, min(c.get('order', 1) - 1, len(CURSUS_LETTERS) - 1))
        return {
            'cursus_letter': CURSUS_LETTERS[order],
            'cursus_color': CURSUS_COLORS[order],
            'cursus_name': c.get('name', ''),
        }
    
    # Get progress items
    progress_items = await db.user_progress.find(
        {'user_id': user_id, 'content_type': 'audio'},
        {'_id': 0}
    ).sort('updated_at', -1).to_list(100)
    
    in_progress = []
    completed = []
    
    for p in progress_items:
        audio = await db.audios.find_one({'id': p.get('content_id')}, {'_id': 0})
        if not audio:
            continue
        
        # Get course info for cursus enrichment
        course = None
        if audio.get('course_id'):
            course = await db.courses.find_one({'id': audio['course_id']}, {'_id': 0, 'title': 1, 'cursus_id': 1})
        
        cursus_info = get_cursus_info(course.get('cursus_id', '') if course else '')
        duration_seconds = audio.get('duration', 0)
        duration_minutes = round(duration_seconds / 60)
        progress_pct = p.get('progress', 0)
        listened_minutes = round(duration_minutes * progress_pct)
        
        item = {
            'id': audio['id'],
            'title': clean_title(audio.get('title', '')),
            'cursus_letter': cursus_info['cursus_letter'],
            'cursus_color': cursus_info['cursus_color'],
            'cursus_name': cursus_info['cursus_name'],
            'course_title': clean_title(course.get('title', '') if course else ''),
            'episode_num': audio.get('episode_number', 1),
            'listened_minutes': listened_minutes,
            'total_minutes': duration_minutes,
            'progress': round(progress_pct * 100),
            'position': p.get('position', 0),
            'updated_at': p.get('updated_at', ''),
        }
        
        if p.get('completed'):
            completed.append(item)
        elif progress_pct > 0.01:
            in_progress.append(item)
    
    # Get favorites
    favs = await db.user_favorites.find({'user_id': user_id}, {'_id': 0}).sort('saved_at', -1).to_list(100)
    favorites = []
    
    for fav in favs:
        content_type = fav.get('content_type', 'audio')
        if content_type == 'audio':
            audio = await db.audios.find_one({'id': fav.get('content_id')}, {'_id': 0})
            if audio:
                course = None
                if audio.get('course_id'):
                    course = await db.courses.find_one({'id': audio['course_id']}, {'_id': 0, 'cursus_id': 1})
                
                cursus_info = get_cursus_info(course.get('cursus_id', '') if course else '')
                duration_minutes = round(audio.get('duration', 0) / 60)
                
                # Format saved_at as relative time
                saved_at = fav.get('saved_at')
                saved_date = 'Récemment'
                if saved_at:
                    if isinstance(saved_at, str):
                        saved_at = datetime.fromisoformat(saved_at.replace('Z', '+00:00'))
                    # Ensure saved_at is timezone-aware
                    if saved_at.tzinfo is None:
                        saved_at = saved_at.replace(tzinfo=timezone.utc)
                    days_ago = (datetime.now(timezone.utc) - saved_at).days
                    if days_ago == 0:
                        saved_date = "Aujourd'hui"
                    elif days_ago == 1:
                        saved_date = 'Hier'
                    elif days_ago < 7:
                        saved_date = f'Il y a {days_ago}j'
                    elif days_ago < 30:
                        weeks = days_ago // 7
                        saved_date = f'Il y a {weeks}s'
                    else:
                        saved_date = f'Il y a {days_ago // 30}m'
                
                favorites.append({
                    'id': audio['id'],
                    'title': clean_title(audio.get('title', '')),
                    'cursus_letter': cursus_info['cursus_letter'],
                    'cursus_color': cursus_info['cursus_color'],
                    'duration_minutes': duration_minutes,
                    'saved_date': saved_date,
                })
    
    # Calculate global progress
    total_progress = 0
    if in_progress or completed:
        all_items = in_progress + completed
        total_progress = round(sum(i['progress'] for i in all_items) / len(all_items)) if all_items else 0
    
    return {
        'in_progress': in_progress,
        'favorites': favorites,
        'completed': completed,
        'global_progress': total_progress,
    }

# ─── Recommendations ─────────────────────────────────────────────────────────

@api_router.get("/recommendations")
async def get_recommendations():
    courses = await db.courses.find({}, {'_id': 0}).limit(6).to_list(6)
    audios = await db.audios.find({}, {'_id': 0}).limit(6).to_list(6)
    return {'courses': courses, 'audios': audios}

# ─── Seed Data ────────────────────────────────────────────────────────────────

THUMBNAILS = [
    "https://images.unsplash.com/photo-1648443524209-d65b2930a743?w=600&q=80",
    "https://images.unsplash.com/photo-1739477274868-86a943b6cd5b?w=600&q=80",
    "https://images.unsplash.com/photo-1739478469935-65b0ac94355c?w=600&q=80",
    "https://images.unsplash.com/photo-1648410794337-b4394bda870f?w=600&q=80",
    "https://images.pexels.com/photos/19893407/pexels-photo-19893407.jpeg?auto=compress&cs=tinysrgb&w=600",
    "https://images.pexels.com/photos/28428586/pexels-photo-28428586.jpeg?auto=compress&cs=tinysrgb&w=600",
    "https://images.unsplash.com/photo-1507842217343-583bb7270b66?w=600&q=80",
    "https://images.unsplash.com/photo-1580094423654-d5ec2e4a2454?w=600&q=80",
]

SCHOLAR_PHOTOS = [
    "https://images.unsplash.com/photo-1628070435838-19eb835ad70d?w=400&q=80",
    "https://images.unsplash.com/photo-1573497019940-1c28c88b4f3e?w=400&q=80",
    "https://images.unsplash.com/photo-1507003211169-0a1dd7228f2d?w=400&q=80",
    "https://images.unsplash.com/photo-1438761681033-6461ffad8d80?w=400&q=80",
    "https://images.unsplash.com/photo-1500648767791-00dcc994a43e?w=400&q=80",
]

AUDIO_URLS = [
    "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-1.mp3",
    "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-2.mp3",
    "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-3.mp3",
    "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-4.mp3",
    "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-5.mp3",
    "https://www.soundhelix.com/examples/mp3/SoundHelix-Song-6.mp3",
]

async def seed_data():
    now = datetime.now(timezone.utc)

    # Scholars
    if await db.scholars.count_documents({}) == 0:
        scholars = [
            {"id": "sch-001", "name": "Prof. Mohammed Al-Fassi", "university": "Université Paris I Panthéon-Sorbonne", "bio": "Spécialiste du droit islamique classique et contemporain, le Prof. Al-Fassi est l'auteur de nombreux ouvrages de référence sur le fiqh et ses applications modernes. Titulaire d'une chaire à la Sorbonne depuis 2008.", "photo": SCHOLAR_PHOTOS[0], "specializations": ["Fiqh", "Jurisprudence islamique", "Droit comparé"], "content_count": 12},
            {"id": "sch-002", "name": "Dr. Leïla Bencherif", "university": "EHESS Paris", "bio": "Directrice de recherche à l'EHESS, le Dr. Bencherif explore la philosophie islamique médiévale et ses résonances contemporaines. Ses travaux sur Ibn Rushd et al-Ghazali font autorité dans le monde académique francophone.", "photo": SCHOLAR_PHOTOS[1], "specializations": ["Philosophie islamique", "Métaphysique", "Kalam"], "content_count": 18},
            {"id": "sch-003", "name": "Prof. Youssef El-Haddad", "university": "Université Lumière Lyon II", "bio": "Historien spécialisé dans la civilisation islamique médiévale et les échanges culturels entre l'Islam et l'Europe. Ses recherches sur al-Andalus ont transformé notre compréhension de l'histoire méditerranéenne.", "photo": SCHOLAR_PHOTOS[2], "specializations": ["Histoire de l'Islam", "Al-Andalus", "Civilisation médiévale"], "content_count": 9},
            {"id": "sch-004", "name": "Dr. Nadia Merah", "university": "Sciences Po Paris", "bio": "Sociologue et islamologue, le Dr. Merah étudie les dimensions spirituelles de l'Islam soufie ainsi que leurs expressions contemporaines en Europe. Chercheuse associée au CERI depuis 2015.", "photo": SCHOLAR_PHOTOS[3], "specializations": ["Soufisme", "Spiritualité islamique", "Islam en Europe"], "content_count": 14},
            {"id": "sch-005", "name": "Prof. Ali Benmakhlouf", "university": "Université Paris-Est Créteil", "bio": "Philosophe et logicien, le Prof. Benmakhlouf est l'un des grands spécialistes français de la logique arabe et de la pensée d'Ibn Khaldoun. Membre de l'Institut de France et auteur d'une vingtaine d'ouvrages.", "photo": SCHOLAR_PHOTOS[4], "specializations": ["Logique arabe", "Philosophie analytique", "Ibn Khaldoun"], "content_count": 11},
        ]
        await db.scholars.insert_many(scholars)
        logger.info("Scholars seeded")

    # Courses - SKIP if custom courses exist (cursus-falsafa)
    if await db.courses.count_documents({}) == 0 and await db.thematiques.find_one({'id': 'cursus-falsafa'}) is None:
        courses = [
            {"id": "crs-001", "title": "Introduction à la philosophie d'Averroès (Ibn Rushd)", "description": "Ce cours fondamental explore la pensée du plus grand philosophe andalou, sa lecture d'Aristote et son influence decisive sur la philosophie européenne médiévale. Une plongée dans l'une des œuvres les plus importantes de la pensée humaine.", "topic": "Philosophie islamique", "level": "Débutant", "language": "Français", "scholar_id": "sch-002", "scholar_name": "Dr. Leïla Bencherif", "duration": 480, "thumbnail": THUMBNAILS[0], "modules_count": 8, "tags": ["Averroès", "Philosophie", "Andalousie", "Aristote"], "type": "course", "published_at": now.isoformat()},
            {"id": "crs-002", "title": "Le Soufisme : des origines aux temps modernes", "description": "Une exploration académique du soufisme, de ses origines au VIIIe siècle jusqu'à ses expressions contemporaines. Le cours aborde les grands maîtres spirituels comme Rumi, Ibn Arabi et al-Hallaj, ainsi que les confréries et leur rôle social.", "topic": "Tasawwuf", "level": "Intermédiaire", "language": "Français", "scholar_id": "sch-004", "scholar_name": "Dr. Nadia Merah", "duration": 720, "thumbnail": THUMBNAILS[1], "modules_count": 12, "tags": ["Soufisme", "Rumi", "Ibn Arabi", "Spiritualité"], "type": "course", "published_at": now.isoformat()},
            {"id": "crs-003", "title": "Fiqh contemporain : droit islamique et modernité", "description": "Ce cours avancé examine les défis posés au droit islamique classique par la modernité. Bioéthique, finance islamique, droit de la famille : le Prof. Al-Fassi analyse les grandes questions juridiques contemporaines avec une rigueur scientifique exemplaire.", "topic": "Fiqh", "level": "Avancé", "language": "Français", "scholar_id": "sch-001", "scholar_name": "Prof. Mohammed Al-Fassi", "duration": 900, "thumbnail": THUMBNAILS[2], "modules_count": 15, "tags": ["Fiqh", "Droit", "Modernité", "Bioéthique"], "type": "course", "published_at": now.isoformat()},
            {"id": "crs-004", "title": "L'Islam en al-Andalus : 8 siècles de civilisation", "description": "De la conquête de 711 à la chute de Grenade en 1492, ce cours retrace l'histoire fascinante de la présence islamique en Espagne. Art, science, philosophie et convivencia : une civilisation d'une richesse exceptionnelle.", "topic": "Histoire de l'Islam", "level": "Débutant", "language": "Français", "scholar_id": "sch-003", "scholar_name": "Prof. Youssef El-Haddad", "duration": 600, "thumbnail": THUMBNAILS[3], "modules_count": 10, "tags": ["Al-Andalus", "Histoire", "Civilisation", "Espagne"], "type": "course", "published_at": now.isoformat()},
            {"id": "crs-005", "title": "Al-Ghazali et la renaissance spirituelle en Islam", "description": "Étude approfondie de l'œuvre d'al-Ghazali, le Hujjat al-Islam. Son Ihya Ulum al-Din reste l'une des synthèses les plus accomplies entre philosophie, théologie et spiritualité dans l'histoire de la pensée islamique.", "topic": "Tasawwuf", "level": "Intermédiaire", "language": "Français", "scholar_id": "sch-002", "scholar_name": "Dr. Leïla Bencherif", "duration": 540, "thumbnail": THUMBNAILS[4], "modules_count": 9, "tags": ["Al-Ghazali", "Philosophie", "Théologie", "Spiritualité"], "type": "course", "published_at": now.isoformat()},
            {"id": "crs-006", "title": "Ibn Khaldoun : pionnier des sciences sociales", "description": "La Muqaddima d'Ibn Khaldoun est un monument de la pensée humaine. Ce cours analyse en profondeur ses théories sur l'histoire, la sociologie, l'économie et la politique, des idées qui anticipent de plusieurs siècles les sciences sociales modernes.", "topic": "Philosophie islamique", "level": "Intermédiaire", "language": "Français", "scholar_id": "sch-005", "scholar_name": "Prof. Ali Benmakhlouf", "duration": 660, "thumbnail": THUMBNAILS[5], "modules_count": 11, "tags": ["Ibn Khaldoun", "Sociologie", "Histoire", "Sciences sociales"], "type": "course", "published_at": now.isoformat()},
            {"id": "crs-007", "title": "Sciences coraniques : introduction à l'exégèse (Tafsir)", "description": "Introduction rigoureuse aux sciences coraniques : histoire de la compilation du Coran, méthodes d'exégèse classiques et contemporaines, herméneutique coranique. Un cours indispensable pour comprendre la tradition exégétique islamique.", "topic": "Sciences coraniques", "level": "Débutant", "language": "Français", "scholar_id": "sch-001", "scholar_name": "Prof. Mohammed Al-Fassi", "duration": 420, "thumbnail": THUMBNAILS[6], "modules_count": 7, "tags": ["Coran", "Tafsir", "Exégèse", "Sciences islamiques"], "type": "course", "published_at": now.isoformat()},
            {"id": "crs-008", "title": "La théologie spéculative islamique (Kalam)", "description": "Le Kalam est la discipline théologique par excellence de l'Islam. Ce cours avancé explore les grandes écoles théologiques — Mu'tazila, Ash'ariyya, Maturidiyya — leurs méthodes et leurs débats fondateurs sur la nature de Dieu et de l'univers.", "topic": "Kalam", "level": "Avancé", "language": "Français", "scholar_id": "sch-002", "scholar_name": "Dr. Leïla Bencherif", "duration": 780, "thumbnail": THUMBNAILS[7], "modules_count": 13, "tags": ["Kalam", "Théologie", "Mu'tazila", "Ash'ariyya"], "type": "course", "published_at": now.isoformat()},
        ]
        await db.courses.insert_many(courses)
        logger.info("Courses seeded")

    # Audios
    if await db.audios.count_documents({}) == 0:
        audios = [
            {"id": "aud-001", "title": "La notion de Hikma dans la pensée islamique classique", "description": "Le concept de Hikma (sagesse) est au cœur de la tradition intellectuelle islamique. Cette conférence explore ses dimensions philosophiques, théologiques et spirituelles, de Platon à Ibn Rushd en passant par al-Farabi.", "scholar_id": "sch-002", "scholar_name": "Dr. Leïla Bencherif", "duration": 3420, "audio_url": AUDIO_URLS[0], "file_key": "podcasts/aud-001.mp3", "thumbnail": THUMBNAILS[0], "topic": "Philosophie islamique", "type": "podcast", "published_at": now.isoformat()},
            {"id": "aud-002", "title": "Ibn Khaldoun et la sociologie avant la lettre", "description": "Le Prof. Benmakhlouf analyse les grandes intuitions sociologiques d'Ibn Khaldoun : la théorie de l'asabiyya, les cycles historiques, l'analyse économique. Un génie méconnu qui mérite d'être redécouvert.", "scholar_id": "sch-005", "scholar_name": "Prof. Ali Benmakhlouf", "duration": 2880, "audio_url": AUDIO_URLS[1], "file_key": "podcasts/aud-002.mp3", "thumbnail": THUMBNAILS[1], "topic": "Philosophie islamique", "type": "podcast", "published_at": now.isoformat()},
            {"id": "aud-003", "title": "Al-Ghazali et le renouveau de la conscience spirituelle", "description": "Cette conférence magistrale explore comment al-Ghazali a réformé la spiritualité islamique en synthétisant philosophie grecque, théologie islamique et mystique soufie dans une vision cohérente et profonde.", "scholar_id": "sch-004", "scholar_name": "Dr. Nadia Merah", "duration": 3960, "audio_url": AUDIO_URLS[2], "file_key": "lectures/aud-003.mp3", "thumbnail": THUMBNAILS[2], "topic": "Tasawwuf", "type": "lecture", "published_at": now.isoformat()},
            {"id": "aud-004", "title": "Averroès et la transmission du savoir aristotélicien", "description": "Comment Averroès (Ibn Rushd) a-t-il sauvé et transmis la philosophie aristotélicienne à l'Europe médiévale ? Cette conférence retrace ce voyage intellectuel fascinant de Cordoue aux universités de Paris et d'Oxford.", "scholar_id": "sch-002", "scholar_name": "Dr. Leïla Bencherif", "duration": 2700, "audio_url": AUDIO_URLS[3], "file_key": "lectures/aud-004.mp3", "thumbnail": THUMBNAILS[3], "topic": "Philosophie islamique", "type": "lecture", "published_at": now.isoformat()},
            {"id": "aud-005", "title": "L'éthique dans le Coran : une lecture philosophique", "description": "Le Prof. Al-Fassi propose une lecture philosophique des dimensions éthiques du Coran : la justice, la dignité humaine, la responsabilité morale. Une approche rigoureuse qui dialogue avec les grandes traditions éthiques occidentales.", "scholar_id": "sch-001", "scholar_name": "Prof. Mohammed Al-Fassi", "duration": 3180, "audio_url": AUDIO_URLS[4], "file_key": "podcasts/aud-005.mp3", "thumbnail": THUMBNAILS[4], "topic": "Fiqh", "type": "podcast", "published_at": now.isoformat()},
            {"id": "aud-006", "title": "Récitation de Sourate Al-Baqara — Sheikh Abdul Rahman Al-Sudais", "description": "La Sourate Al-Baqara, deuxième et plus longue sourate du Coran, dans une récitation émouvante du Sheikh Abdul Rahman Al-Sudais, Imam de la Grande Mosquée de La Mecque. Translittération et traduction française disponibles.", "scholar_id": "sch-001", "scholar_name": "Sheikh Abdul Rahman Al-Sudais", "duration": 5400, "audio_url": AUDIO_URLS[5], "file_key": "quran/aud-006.mp3", "thumbnail": THUMBNAILS[5], "topic": "Sciences coraniques", "type": "quran", "published_at": now.isoformat()},
            {"id": "aud-007", "title": "Le dialogue interreligieux dans la tradition islamique", "description": "Contrairement aux idées reçues, la tradition islamique a une riche histoire de dialogue avec les autres traditions abrahamiques. Cette conférence explore cette histoire trop méconnue, des penseurs médiévaux aux expériences contemporaines.", "scholar_id": "sch-003", "scholar_name": "Prof. Youssef El-Haddad", "duration": 2520, "audio_url": AUDIO_URLS[0], "file_key": "podcasts/aud-007.mp3", "thumbnail": THUMBNAILS[6], "topic": "Histoire de l'Islam", "type": "podcast", "published_at": now.isoformat()},
            {"id": "aud-008", "title": "Rumi et la poésie mystique : entre Orient et Occident", "description": "Jalal ad-Din Rumi est aujourd'hui l'un des poètes les plus lus dans le monde. Le Dr. Merah analyse la profondeur mystique de ses œuvres, du Masnavi au Divan de Shams, et leur résonance universelle.", "scholar_id": "sch-004", "scholar_name": "Dr. Nadia Merah", "duration": 3300, "audio_url": AUDIO_URLS[1], "file_key": "lectures/aud-008.mp3", "thumbnail": THUMBNAILS[7], "topic": "Tasawwuf", "type": "lecture", "published_at": now.isoformat()},
            {"id": "aud-009", "title": "Récitation des 30 Juz — Sheikh Mishary Rashid Al-Afasy", "description": "Récitation complète du Coran par le Sheikh Mishary Rashid Al-Afasy, l'une des voix les plus appréciées dans le monde islamique. Une expérience sonore d'une beauté saisissante, accompagnée de notes de contexte académiques.", "scholar_id": "sch-001", "scholar_name": "Sheikh Mishary Rashid Al-Afasy", "duration": 7200, "audio_url": AUDIO_URLS[2], "file_key": "quran/aud-009.mp3", "thumbnail": THUMBNAILS[0], "topic": "Sciences coraniques", "type": "quran", "published_at": now.isoformat()},
            {"id": "aud-010", "title": "Islam et modernité : une tension créatrice", "description": "Comment les penseurs musulmans contemporains abordent-ils la modernité ? Cette conférence examine les différentes approches intellectuelles — réformistes, traditionalistes, progressistes — face aux défis du monde contemporain.", "scholar_id": "sch-001", "scholar_name": "Prof. Mohammed Al-Fassi", "duration": 2940, "audio_url": AUDIO_URLS[3], "file_key": "podcasts/aud-010.mp3", "thumbnail": THUMBNAILS[1], "topic": "Fiqh", "type": "podcast", "published_at": now.isoformat()},
            {"id": "aud-011", "title": "Ibn Arabi et l'ontologie soufie : wahdat al-wujud", "description": "La doctrine de l'unité de l'être (wahdat al-wujud) d'Ibn Arabi est l'une des contributions les plus originales et les plus controversées de la pensée islamique. Le Dr. Merah en décode les subtilités philosophiques avec une clarté remarquable.", "scholar_id": "sch-004", "scholar_name": "Dr. Nadia Merah", "duration": 4020, "audio_url": AUDIO_URLS[4], "file_key": "lectures/aud-011.mp3", "thumbnail": THUMBNAILS[2], "topic": "Tasawwuf", "type": "lecture", "published_at": now.isoformat()},
            {"id": "aud-012", "title": "Al-Andalus : chroniques d'une civilisation perdue", "description": "Documentaire audio retraçant l'histoire de la civilisation andalouse : ses bibliothèques, ses astronomes, ses médecins, ses poètes. Une plongée dans un monde où musulmans, chrétiens et juifs ont coexisté et créé ensemble.", "scholar_id": "sch-003", "scholar_name": "Prof. Youssef El-Haddad", "duration": 5040, "audio_url": AUDIO_URLS[5], "file_key": "documentaries/aud-012.mp3", "thumbnail": THUMBNAILS[3], "topic": "Histoire de l'Islam", "type": "documentary", "published_at": now.isoformat()},
        ]
        await db.audios.insert_many(audios)
        logger.info("Audios seeded")
    else:
        # Migration: add file_key to existing audio documents that don't have it
        r2_file_keys = {
            "aud-001": "podcasts/aud-001.mp3",
            "aud-002": "podcasts/aud-002.mp3",
            "aud-003": "lectures/aud-003.mp3",
            "aud-004": "lectures/aud-004.mp3",
            "aud-005": "podcasts/aud-005.mp3",
            "aud-006": "quran/aud-006.mp3",
            "aud-007": "podcasts/aud-007.mp3",
            "aud-008": "lectures/aud-008.mp3",
            "aud-009": "quran/aud-009.mp3",
            "aud-010": "podcasts/aud-010.mp3",
            "aud-011": "lectures/aud-011.mp3",
            "aud-012": "documentaries/aud-012.mp3",
        }
        for audio_id, file_key in r2_file_keys.items():
            await db.audios.update_one(
                {'id': audio_id, 'file_key': {'$exists': False}},
                {'$set': {'file_key': file_key}}
            )
        logger.info("Audio file_key migration complete")

    # Articles
    if await db.articles.count_documents({}) == 0:
        articles = [
            {"id": "art-001", "title": "L'humanisme islamique médiéval : de Bagdad à Oxford", "excerpt": "La Maison de la Sagesse de Bagdad a été le creuset d'une révolution intellectuelle sans précédent. Cet article retrace comment les savants musulmans ont préservé, enrichi et transmis le patrimoine grec à l'Europe.", "content": "La Maison de la Sagesse (Bayt al-Hikma) fondée sous le calife al-Ma'mun à Bagdad au IXe siècle représente l'un des moments les plus remarquables de l'histoire intellectuelle de l'humanité. Ce centre de traduction et de recherche a non seulement préservé les œuvres philosophiques et scientifiques de l'Antiquité, mais les a aussi enrichies d'apports originaux déterminants pour le développement des sciences modernes.\n\nLes traducteurs de Bagdad — Hunayn ibn Ishaq, Thabit ibn Qurra, al-Kindi et leurs contemporains — n'étaient pas de simples passeurs. Ils ont produit des synthèses originales, corrigé des erreurs de transmission, et enrichi considérablement les textes qu'ils traduisaient.\n\nAvicenne (Ibn Sina) a développé une philosophie naturelle qui a dominé la pensée européenne pendant quatre siècles. Averroès (Ibn Rushd) a produit des commentaires d'Aristote si influents qu'on l'appelait simplement 'Le Commentateur' dans les universités européennes. Al-Farabi a fondé la philosophie politique islamique tout en dialogue constant avec Platon et Aristote.\n\nCet héritage a voyagé de Bagdad à Cordoue, de Cordoue à Tolède, de Tolède à Paris et Oxford. Comprendre la philosophie médiévale européenne sans comprendre son substrat islamique, c'est se condamner à une incompréhension partielle de l'histoire des idées.", "scholar_id": "sch-003", "scholar_name": "Prof. Youssef El-Haddad", "reading_time": 8, "thumbnail": THUMBNAILS[6], "topic": "Histoire de l'Islam", "published_at": now.isoformat(), "type": "article"},
            {"id": "art-002", "title": "Averroès et la transmission du savoir aristotélicien en Europe", "excerpt": "Comment Ibn Rushd a-t-il sauvé Aristote ? Retour sur l'un des transferts culturels les plus décisifs de l'histoire intellectuelle occidentale et islamique.", "content": "Quand on lit saint Thomas d'Aquin, Albert le Grand ou Roger Bacon, on lit aussi — souvent sans le savoir — Ibn Rushd. Ce philosophe cordouan du XIIe siècle, connu en Europe sous le nom d'Averroès, a produit des commentaires si systématiques et si pénétrants des œuvres d'Aristote qu'ils sont devenus pendant des siècles la référence incontournable pour quiconque voulait comprendre le Stagirite.\n\nNé à Cordoue en 1126, Ibn Rushd a grandi dans une famille de juristes et s'est formé à toutes les disciplines du savoir islamique classique : droit malékite, théologie, médecine, mathématiques, astronomie et philosophie. C'est dans ce dernier domaine qu'il a laissé sa marque la plus durable.\n\nSes trois séries de commentaires d'Aristote — courts, moyens et longs — couvrent pratiquement toute l'œuvre du philosophe grec. Ils ont été traduits en latin et en hébreu dès le XIIe siècle, diffusant la philosophie aristotélicienne dans les universités naissantes d'Europe.\n\nParadoxalement, Averroès est parfois mieux connu en Europe qu'dans le monde arabe et islamique, où sa postérité a été plus complexe. Son conflit avec al-Ghazali sur la philosophie et la religion, son exil à Marrakech sur ordre du calife almohade, tout cela a contribué à une réception ambivalente.", "scholar_id": "sch-002", "scholar_name": "Dr. Leïla Bencherif", "reading_time": 12, "thumbnail": THUMBNAILS[7], "topic": "Philosophie islamique", "published_at": now.isoformat(), "type": "article"},
            {"id": "art-003", "title": "Le soufisme comme voie spirituelle : tradition et contemporanéité", "excerpt": "Le soufisme est souvent réduit à ses expressions les plus pittoresques. Cet article propose une exploration rigoureuse de cette tradition spirituelle et de ses défis contemporains.", "content": "Le soufisme — ou tasawwuf — désigne la dimension spirituelle et mystique de l'Islam. Loin des clichés exotiques, c'est une tradition intellectuelle et spirituelle d'une grande richesse, qui a produit certaines des plus belles œuvres poétiques et philosophiques de l'humanité.\n\nLes origines du soufisme sont disputées. Certains historiens le voient comme un développement interne à l'Islam, réaction à la mondanisation progressive de la communauté après les premières générations. D'autres soulignent les influences néoplatoniciennes, chrétiennes et bouddhistes dans sa formation. La réalité est sans doute plus complexe et plus intéressante : le soufisme est une synthèse originale.\n\nLes grands maîtres soufis — al-Hallaj, al-Muhasibi, Junayd de Bagdad, Rabi'a al-Adawiyya, Ibn Arabi, Rumi — ont développé des doctrines et des pratiques d'une sophistication remarquable. La notion de fana (extinction du moi dans le divin), de mahabbah (amour divin), de kashf (dévoilement spirituel) constituent un vocabulaire conceptuel d'une grande précision.\n\nAujourd'hui, le soufisme connaît à la fois un renouveau d'intérêt académique et des défis inédits. Les confréries soufies (turuq) — Qadiriyya, Naqshbandiyya, Shadhiliyya, Tijaniyya — restent des forces spirituelles et sociales importantes dans de nombreux pays musulmans, tout en s'adaptant au monde contemporain.", "scholar_id": "sch-004", "scholar_name": "Dr. Nadia Merah", "reading_time": 10, "thumbnail": THUMBNAILS[0], "topic": "Tasawwuf", "published_at": now.isoformat(), "type": "article"},
            {"id": "art-004", "title": "Islam et laïcité en France : enjeux académiques", "excerpt": "La question de l'Islam en France est souvent traitée dans la polémique. Cet article propose un regard académique sur les enjeux réels du vivre ensemble.", "content": "Le débat sur l'Islam en France souffre d'un déficit de rigueur analytique. Entre les discours alarmistes et les apologétiques naïfs, il manque souvent une approche rigoureusement empirique et conceptuellement précise.\n\nD'abord, quelques faits. L'Islam est la deuxième religion de France, avec une communauté estimée entre 5 et 6 millions de personnes. Cette communauté est extrêmement diverse : Marocains, Algériens, Tunisiens, Turcs, Africains subsahariens, convertis français — autant de pratiques, de courants théologiques et d'expériences sociales différents.\n\nLa laïcité française a une histoire complexe. Née dans le conflit avec l'Église catholique, elle s'est construite progressivement comme un principe de neutralité de l'État vis-à-vis des religions. Son application à l'Islam soulève des questions inédites — non parce que l'Islam serait incompatible avec la laïcité, mais parce que les conditions historiques et sociales sont différentes.\n\nLes chercheurs comme Olivier Roy, Gilles Kepel, Jocelyne Césari ont produit des travaux importants qui nuancent considérablement les discours politiques. Comprendre l'Islam en France nécessite de dépasser les catégories héritées du XIXe siècle et de se doter d'outils conceptuels adaptés à la réalité du XXIe siècle.", "scholar_id": "sch-001", "scholar_name": "Prof. Mohammed Al-Fassi", "reading_time": 9, "thumbnail": THUMBNAILS[1], "topic": "Fiqh", "published_at": now.isoformat(), "type": "article"},
            {"id": "art-005", "title": "Al-Andalus : mythe ou réalité de la convivencia ?", "excerpt": "La coexistence harmonieuse de musulmans, chrétiens et juifs en al-Andalus est-elle un mythe ou une réalité historique ? Retour sur un débat historiographique majeur.", "content": "La convivencia andalouse est devenue un véritable mythe fondateur pour les promoteurs du dialogue interreligieux — et une cible privilégiée pour les révisionnistes qui voient dans cette idéalisation une naïveté dangereuse. La réalité historique est, comme souvent, plus nuancée et plus intéressante que les deux camps ne le prétendent.\n\nIl est indéniable que la péninsule ibérique sous domination islamique (711-1492) a connu des périodes remarquables de collaboration intellectuelle entre musulmans, chrétiens et juifs. L'école de traduction de Tolède au XIIe siècle, où des savants de trois religions travaillaient ensemble à la traduction des œuvres grecques et arabes en latin, en est l'exemple le plus célèbre.\n\nMaïmonide, le grand philosophe juif, a écrit en arabe et s'est nourri de la philosophie islamique. Les médecins de Cordoue soignaient sans distinction de religion. Les poètes arabes, hébreux et romans se répondaient et s'influençaient mutuellement.\n\nMais la convivencia n'était pas une utopie de tolérance parfaite. Il y avait des moments de violence, des politiques de dhimma discriminatoires, des conversions forcées — notamment sous les Almohades. L'historienne María Rosa Menocal a peut-être idéalisé cette période, mais ses détracteurs risquent de commettre l'erreur inverse en niant ses réalisations réelles.", "scholar_id": "sch-003", "scholar_name": "Prof. Youssef El-Haddad", "reading_time": 11, "thumbnail": THUMBNAILS[2], "topic": "Histoire de l'Islam", "published_at": now.isoformat(), "type": "article"},
            {"id": "art-006", "title": "La logique dans la tradition islamique médiévale", "excerpt": "La logique aristotélicienne a connu un destin extraordinaire dans la civilisation islamique. Entre ceux qui la condamnaient et ceux qui la défendaient, un débat fondateur.", "content": "La réception de la logique aristotélicienne dans la pensée islamique médiévale est l'un des chapitres les plus fascinants de l'histoire des idées. Ce qui était initialement regardé avec méfiance — un outil de raisonnement développé par des penseurs grecs polythéistes — est devenu un instrument indispensable de la théologie, du droit et de la philosophie islamiques.\n\nAl-Farabi, le 'second maître' après Aristote, a été le premier à proposer une synthèse complète de la logique aristotélicienne adaptée au contexte islamique. Son Kitab al-Alfaz (Livre des termes) et ses commentaires des Analytiques ont posé les bases d'une tradition logique originale.\n\nAvicenne (Ibn Sina) a ensuite développé une logique proprement islamique, différant d'Aristote sur plusieurs points importants. Sa théorie des universaux, sa logique modale, son traitement du syllogisme ont enrichi considérablement la tradition aristotélicienne.\n\nLe Grand Débat a opposé les philosophes logiciens aux théologiens traditionnels, dont al-Ghazali. Dans son Incohérence des philosophes (Tahafut al-Falasifa), al-Ghazali attaque vigoureusement les prétentions des philosophes. Averroès lui répondra dans son Incohérence de l'incohérence, défendant la validité et la nécessité de la raison philosophique.", "scholar_id": "sch-005", "scholar_name": "Prof. Ali Benmakhlouf", "reading_time": 7, "thumbnail": THUMBNAILS[3], "topic": "Philosophie islamique", "published_at": now.isoformat(), "type": "article"},
        ]
        await db.articles.insert_many(articles)
        logger.info("Articles seeded")

    # Live Sessions
    if await db.live_sessions.count_documents({}) == 0:
        now_dt = datetime.now(timezone.utc)
        sessions = [
            {"id": "live-001", "title": "Masterclass : Introduction à la philosophie islamique", "description": "Le Dr. Bencherif vous invite à une masterclass exceptionnelle d'introduction à la philosophie islamique : ses grandes questions, ses grands penseurs, ses méthodes. Une session idéale pour les débutants et les curieux.", "scholar_id": "sch-002", "scholar_name": "Dr. Leïla Bencherif", "date": (now_dt + timedelta(days=7)).isoformat(), "duration": 90, "thumbnail": THUMBNAILS[0], "topic": "Philosophie islamique", "max_participants": 100, "registered_count": 47, "registered_users": []},
            {"id": "live-002", "title": "Séminaire : Islam et modernité en Europe — enjeux académiques", "description": "Une session interactive de débat académique autour des questions posées par la présence de l'Islam en Europe. Le Prof. Al-Fassi présentera ses recherches récentes et répondra aux questions du public.", "scholar_id": "sch-001", "scholar_name": "Prof. Mohammed Al-Fassi", "date": (now_dt + timedelta(days=14)).isoformat(), "duration": 120, "thumbnail": THUMBNAILS[1], "topic": "Fiqh", "max_participants": 80, "registered_count": 32, "registered_users": []},
            {"id": "live-003", "title": "Lecture commentée de la Muqaddima d'Ibn Khaldoun", "description": "Rejoignez le Prof. Benmakhlouf pour une lecture commentée des passages clés de la Muqaddima. Ce texte fondateur des sciences sociales sera analysé dans sa profondeur et dans sa modernité étonnante.", "scholar_id": "sch-005", "scholar_name": "Prof. Ali Benmakhlouf", "date": (now_dt + timedelta(days=21)).isoformat(), "duration": 90, "thumbnail": THUMBNAILS[2], "topic": "Philosophie islamique", "max_participants": 60, "registered_count": 28, "registered_users": []},
            {"id": "live-004", "title": "Atelier : Poésie soufie et traduction — Rumi en français", "description": "Un atelier pratique autour de la traduction et de l'interprétation de la poésie soufie. Le Dr. Merah présentera différentes traductions de Rumi et analysera les enjeux de la traduction de la poésie mystique.", "scholar_id": "sch-004", "scholar_name": "Dr. Nadia Merah", "date": (now_dt + timedelta(days=10)).isoformat(), "duration": 90, "thumbnail": THUMBNAILS[3], "topic": "Tasawwuf", "max_participants": 40, "registered_count": 38, "registered_users": []},
            {"id": "live-005", "title": "Table ronde : L'avenir des études islamiques en France", "description": "Nos cinq érudits se réunissent pour une table ronde exceptionnelle autour de l'avenir des études islamiques académiques en France. Perspectives, défis, opportunités : une conversation au sommet.", "scholar_id": "sch-003", "scholar_name": "Prof. Youssef El-Haddad", "date": (now_dt + timedelta(days=30)).isoformat(), "duration": 150, "thumbnail": THUMBNAILS[4], "topic": "Histoire de l'Islam", "max_participants": 200, "registered_count": 89, "registered_users": []},
        ]
        await db.live_sessions.insert_many(sessions)
        logger.info("Live sessions seeded")

    # ─── Migrations ──────────────────────────────────────────────────────────

    # 1. Add is_active to all content docs
    await db.audios.update_many({'is_active': {'$exists': False}}, {'$set': {'is_active': True}})
    await db.courses.update_many({'is_active': {'$exists': False}}, {'$set': {'is_active': True}})
    await db.scholars.update_many({'is_active': {'$exists': False}}, {'$set': {'is_active': True}})

    # Check if custom cursus exists - skip demo seeding if so
    custom_cursus = await db.cursus.find_one({'id': 'cursus-falsafa'})
    if not custom_cursus:
        # Also check old thematiques collection for backward compatibility
        custom_cursus = await db.thematiques.find_one({'id': 'cursus-falsafa'})
    logger.info(f"DEBUG: Checking for cursus-falsafa, found: {custom_cursus is not None}")
    
    # Migration: Set Meryem Sebti as default professor for all Cursus A courses
    sebti_update = await db.courses.update_many(
        {'$or': [{'cursus_id': 'cursus-falsafa'}, {'thematique_id': 'cursus-falsafa'}], 'scholar_id': {'$exists': False}},
        {'$set': {'scholar_id': 'sch-sebti', 'scholar_name': 'Meryem Sebti'}}
    )
    # Also update courses with "Divers auteurs"
    sebti_update2 = await db.courses.update_many(
        {'$or': [{'cursus_id': 'cursus-falsafa'}, {'thematique_id': 'cursus-falsafa'}], 'scholar_name': 'Divers auteurs'},
        {'$set': {'scholar_id': 'sch-sebti', 'scholar_name': 'Meryem Sebti'}}
    )
    if sebti_update.modified_count > 0 or sebti_update2.modified_count > 0:
        logger.info(f"Migration: Set Meryem Sebti as professor for Cursus A courses ({sebti_update.modified_count + sebti_update2.modified_count} updated)")
    
    # ─── Migration: Catalogue v3 (5→6 cursus) ──────────────────────────────
    # Update Cursus E: rename to "La Mystique islamique"
    v3_e = await db.cursus.update_one(
        {'id': 'cursus-spiritualites'},
        {'$set': {
            'name': 'La Mystique islamique',
            'description': 'Le premier taṣawwuf, Ibn ʿArabī, le soufisme iranien et les liens entre mystique et philosophie.',
            'hero_title': 'La Mystique islamique',
            'subtitle': 'Taṣawwuf',
        }}
    )
    if v3_e.modified_count > 0:
        logger.info("Migration v3: Cursus E renamed to 'La Mystique islamique'")

    # Create Cursus F if it doesn't exist
    existing_f = await db.cursus.find_one({'id': 'cursus-pensees-non-islamiques'})
    if not existing_f:
        await db.cursus.insert_one({
            'id': 'cursus-pensees-non-islamiques',
            'name': 'Pensées arabes non islamiques',
            'description': 'Le Kalām chrétien, les logiciens de Bagdad et la philosophie juive de langue arabe, de Saʿadya Gaon à Maïmonide.',
            'hero_title': 'Pensées arabes non islamiques',
            'subtitle': 'Philosophie juive et chrétienne',
            'order': 6,
            'is_active': True,
            'is_featured': False,
        })
        logger.info("Migration v3: Cursus F 'Pensées arabes non islamiques' created")

    # Update descriptions for A, B, C, D — ONLY when no admin edit has been
    # applied yet (cursus.seed_locked != True). Once an admin tweaks a cursus
    # subtitle/description in the panel, this seed must never overwrite it.
    await db.cursus.update_one(
        {'id': 'cursus-falsafa', 'seed_locked': {'$ne': True}},
        {'$set': {
            'description': "D'Al-Kindī à Mollā Sādrā, la falsafa, le post-avicennisme, la logique arabe, l'ismaélisme et les inclassables comme Ibn Khaldūn.",
            'subtitle': "Philosophie en terre d'islam",
        }}
    )
    await db.cursus.update_one(
        {'id': 'cursus-theologie', 'seed_locked': {'$ne': True}},
        {'$set': {
            'description': "Le Kalām dans ses trois périodes et l'histoire de la réflexion juridique (Uṣūl al-fiqh) : les quatre écoles et le droit musulman.",
            'subtitle': 'Kalām, Fiqh et fondements juridiques',
        }}
    )
    await db.cursus.update_one(
        {'id': 'cursus-sciences-islamiques', 'seed_locked': {'$ne': True}},
        {'$set': {
            'description': "Doxographie, transmission du Coran et du Hadith, historiographie (d'al-Ṭabarī à Ibn Khaldūn) et autobiographies dans le monde islamique.",
            'subtitle': 'Hadith, Coran, exégèse et historiographie',
        }}
    )
    await db.cursus.update_one(
        {'id': 'cursus-arts', 'seed_locked': {'$ne': True}},
        {'$set': {
            'description': 'Art islamique, poésie (arabe, persane, préislamique), pédagogie, sciences (biologie, astronomie, mathématiques), géographie et Adab.',
            'subtitle': 'Poésie, arts, pédagogie et sciences',
        }}
    )
    logger.info("Migration v3: Cursus descriptions updated (skipping admin-edited)")

    # ─── Migration v3b: Reassign courses to match new cursus structure ────
    # Move L'ismaélisme (cours-ismaelisme) from Cursus E → Cursus A
    v3b_1 = await db.courses.update_one(
        {'id': 'cours-ismaelisme', 'cursus_id': 'cursus-spiritualites'},
        {'$set': {
            'cursus_id': 'cursus-falsafa',
            'thematique_id': 'cursus-falsafa',
            'title': "Cours 8 : L'ismaélisme",
            'order': 8,
        }}
    )
    # Move Le Kalām chrétien (cours-kalam-chretien) from Cursus E → Cursus F
    v3b_2 = await db.courses.update_one(
        {'id': 'cours-kalam-chretien', 'cursus_id': {'$ne': 'cursus-pensees-non-islamiques'}},
        {'$set': {
            'cursus_id': 'cursus-pensees-non-islamiques',
            'thematique_id': 'cursus-pensees-non-islamiques',
            'title': 'Cours 23 : Le Kalām chrétien et les logiciens de Bagdad',
            'order': 23,
        }}
    )
    # Move La philosophie juive (cours-philo-juive) from Cursus E → Cursus F
    v3b_3 = await db.courses.update_one(
        {'id': 'cours-philo-juive', 'cursus_id': {'$ne': 'cursus-pensees-non-islamiques'}},
        {'$set': {
            'cursus_id': 'cursus-pensees-non-islamiques',
            'thematique_id': 'cursus-pensees-non-islamiques',
        }}
    )
    moved = v3b_1.modified_count + v3b_2.modified_count + v3b_3.modified_count
    if moved > 0:
        logger.info(f"Migration v3b: Reassigned {moved} courses to new cursus structure")
    # ─── End Migration v3b ────────────────────────────────────────────────
    # ─── End Migration v3 ─────────────────────────────────────────────────

    # ─── Migration v4: Catalogue de lancement Mai 2026 ────────────────────
    # New cursus structure (7 cursus): A=Histoire, B=Théologie, C=Sciences,
    # D=Arts, E=Falsafa (was A), F=Mystique (was E), G=Pensées non-islamiques
    
    # 1. Insert new Cursus A "Histoire du monde islamique"
    # Defensive: clean up any duplicate "Histoire" cursus that might exist with different ids
    duplicates = await db.cursus.delete_many({
        'id': {'$ne': 'cursus-histoire'},
        'name': {'$regex': '^histoire du monde islamique$', '$options': 'i'},
    })
    if duplicates.deleted_count > 0:
        logger.warning(f"Migration v4: removed {duplicates.deleted_count} duplicate Histoire cursus record(s) with different id")
    
    # Defensive: deduplicate cursus by `id` field (MongoDB allows multiple docs with same `id` since only `_id` is unique)
    # Find all duplicate ids and keep only the first occurrence
    pipeline = [
        {'$group': {'_id': '$id', 'docs': {'$push': '$_id'}, 'count': {'$sum': 1}}},
        {'$match': {'count': {'$gt': 1}}},
    ]
    dups_removed = 0
    async for group in db.cursus.aggregate(pipeline):
        # Keep the first _id, remove the rest
        keep_id = group['docs'][0]
        remove_ids = group['docs'][1:]
        if remove_ids:
            r = await db.cursus.delete_many({'_id': {'$in': remove_ids}})
            dups_removed += r.deleted_count
            logger.warning(f"Migration v4: deduplicated cursus id='{group['_id']}' — removed {r.deleted_count} extra record(s), kept _id={keep_id}")
    if dups_removed > 0:
        logger.warning(f"Migration v4: total {dups_removed} duplicate cursus removed by id-dedup")
    
    existing_histoire = await db.cursus.find_one({'id': 'cursus-histoire'})
    if not existing_histoire:
        await db.cursus.insert_one({
            'id': 'cursus-histoire',
            'name': 'Histoire du monde islamique',
            'description': "Des débuts de l'islam à l'époque ottomane, en passant par al-Andalus et les Mamelouks : les grandes époques de l'histoire islamique.",
            'hero_title': 'Histoire du monde islamique',
            'subtitle': 'Des origines à l\'époque ottomane',
            'icon': 'book-open',
            'order': 1,
            'is_active': True,
            'is_featured': False,
            'letter': 'A',
        })
        logger.info("Migration v4: Cursus A 'Histoire du monde islamique' created")
    
    # 2. Update orders + letters for all cursus
    cursus_order_map = [
        ('cursus-histoire', 1, 'A'),
        ('cursus-theologie', 2, 'B'),
        ('cursus-sciences-islamiques', 3, 'C'),
        ('cursus-arts', 4, 'D'),
        ('cursus-falsafa', 5, 'E'),
        ('cursus-spiritualites', 6, 'F'),
        ('cursus-pensees-non-islamiques', 7, 'G'),
    ]
    for cid, ordr, letter in cursus_order_map:
        await db.cursus.update_one({'id': cid}, {'$set': {'order': ordr, 'letter': letter}})
    logger.info("Migration v4: Cursus orders + letters updated (7 cursus A-G)")
    
    # 3. Create 4 placeholder courses for Cursus A (Histoire)
    histoire_courses = [
        {
            'id': 'cours-debuts-islam',
            'title': "Cours 1 : Les débuts de l'islam",
            'description': "Les origines de l'islam, du Prophète Muhammad aux premiers califats. Cours présenté par Hassan Bouali et Mehdi Ghouirgate.",
            'cursus_id': 'cursus-histoire',
            'thematique_id': 'cursus-histoire',
            'topic': 'Histoire',
            'level': 'Debutant',
            'language': 'Francais',
            'scholar_id': '',
            'scholar_name': 'Hassan Bouali · Mehdi Ghouirgate',
            'duration': 0,
            'thumbnail': '',
            'modules_count': 2,
            'tags': ['Histoire', 'Origines', 'Califat'],
            'is_active': True,
            'is_featured': False,
            'is_launch_catalog': True,
            'order': 1,
        },
        {
            'id': 'cours-andalus',
            'title': 'Cours 2 : Al-Andalus',
            'description': "Histoire d'al-Andalus, de la conquête omeyyade à la chute de Grenade. Par Mehdi Ghouirgate.",
            'cursus_id': 'cursus-histoire',
            'thematique_id': 'cursus-histoire',
            'topic': 'Histoire',
            'level': 'Debutant',
            'language': 'Francais',
            'scholar_id': '',
            'scholar_name': 'Mehdi Ghouirgate',
            'duration': 0,
            'thumbnail': '',
            'modules_count': 1,
            'tags': ['Histoire', 'Al-Andalus', 'Espagne musulmane'],
            'is_active': True,
            'is_featured': False,
            'is_launch_catalog': True,
            'youtube_url': 'https://youtu.be/cow2JfYaSC0',
            'order': 2,
        },
        {
            'id': 'cours-mamelouke',
            'title': "Cours 3 : L'époque mamelouke",
            'description': "Le sultanat mamelouk d'Égypte et de Syrie (1250-1517). Cours présenté par Sami Benkherfallah.",
            'cursus_id': 'cursus-histoire',
            'thematique_id': 'cursus-histoire',
            'topic': 'Histoire',
            'level': 'Intermediaire',
            'language': 'Francais',
            'scholar_id': '',
            'scholar_name': 'Sami Benkherfallah',
            'duration': 0,
            'thumbnail': '',
            'modules_count': 1,
            'tags': ['Histoire', 'Mamelouks', 'Égypte', 'Syrie'],
            'is_active': True,
            'is_featured': False,
            'is_launch_catalog': True,
            'coming_soon': True,
            'available_date': 'mai 2026',
            'order': 3,
        },
        {
            'id': 'cours-ottoman',
            'title': 'Cours 4 : Le monde ottoman',
            'description': "L'empire ottoman, de ses origines anatoliennes à son apogée méditerranéen. Cours présenté par Aysu Saban.",
            'cursus_id': 'cursus-histoire',
            'thematique_id': 'cursus-histoire',
            'topic': 'Histoire',
            'level': 'Intermediaire',
            'language': 'Francais',
            'scholar_id': '',
            'scholar_name': 'Aysu Saban',
            'duration': 0,
            'thumbnail': '',
            'modules_count': 1,
            'tags': ['Histoire', 'Ottomans', 'Empire'],
            'is_active': True,
            'is_featured': False,
            'is_launch_catalog': True,
            'coming_soon': True,
            'available_date': 'sept. 2026',
            'order': 4,
        },
    ]
    # Fields that the seed is allowed to keep refreshing on existing courses
    # (purely structural / catalogue placement — never content edited by admin).
    STRUCTURAL_FIELDS = {
        'cursus_id', 'thematique_id', 'order', 'is_launch_catalog',
        'coming_soon', 'available_date', 'is_active', 'is_featured',
    }
    for c in histoire_courses:
        existing = await db.courses.find_one(
            {'id': c['id']},
            {'_id': 0, 'seed_locked': 1, 'title': 1, 'description': 1, 'summary': 1, 'scholar_name': 1}
        )
        if not existing:
            # New course → insert with full seed values.
            await db.courses.insert_one(c)
            continue
        # Existing course → ONLY refresh structural fields. Never overwrite
        # title / description / summary / scholar_name (admins edit those in
        # the panel and the seed must not stomp their work on redeploys).
        structural = {k: v for k, v in c.items() if k in STRUCTURAL_FIELDS}
        if structural:
            await db.courses.update_one({'id': c['id']}, {'$set': structural})
    logger.info(f"Migration v4: {len(histoire_courses)} Histoire courses synced (content preserved)")
    
    # 4. Mark courses with is_launch_catalog flag based on Excel data
    launch_course_ids = {
        # Cursus E (Falsafa)
        'cours-traduction', 'cours-falsafa-grands', 'cours-post-avicennisme',
        'cours-falsafa-occident', 'cours-falsafa-persan', 'cours-inclassables',
        # Cursus B (Théologie)
        'cours-kalam', 'cours-fiqh',
        # Cursus C (Sciences)
        'cours-coran', 'cours-hadith', 'cours-historiographie',
        # Cursus D (Arts)
        'cours-art', 'cours-sciences',
        # Cursus F (Mystique)
        'cours-soufisme',
        # Cursus G (Pensées non-islamiques)
        'cours-philo-juive',
        # Cursus A (Histoire) - all 4 included above
        'cours-debuts-islam', 'cours-andalus', 'cours-mamelouke', 'cours-ottoman',
    }
    # Mark IN-launch courses
    in_launch = await db.courses.update_many(
        {'id': {'$in': list(launch_course_ids)}},
        {'$set': {'is_launch_catalog': True}}
    )
    # Mark NOT-in-launch courses  
    not_in_launch = await db.courses.update_many(
        {'id': {'$nin': list(launch_course_ids)}, 'is_launch_catalog': {'$ne': False}},
        {'$set': {'is_launch_catalog': False}}
    )
    logger.info(f"Migration v4: launch_catalog set on {in_launch.modified_count} (in) / {not_in_launch.modified_count} (out)")
    
    # 5. Set youtube_url on specific courses (where Excel has 1 episode/URL pair)
    youtube_course_map = {
        'cours-historiographie': 'https://youtu.be/RUc8p0K6Qg4',  # Ibn Khaldun by Ghouirgate
    }
    for cid, url in youtube_course_map.items():
        await db.courses.update_one({'id': cid}, {'$set': {'youtube_url': url}})
    
    # 6. Set youtube_url on specific audios from Excel
    # Format: { audio_id_substring: youtube_url }
    audio_youtube_map = [
        # Cursus A - Histoire (mapped to placeholder courses, audios may not exist yet)
        # cours-debuts-islam by Mehdi Ghouirgate
        ({'course_id': 'cours-debuts-islam'}, 'https://youtu.be/kB-fr8wwcAA'),
        # Cursus E - Falsafa
        ({'course_id': 'cours-traduction'}, 'https://www.youtube.com/watch?v=5EsSIUfeP-o'),
        # Al-Kindi episodes (cours-falsafa-grands)
        ({'id': {'$regex': 'al-kindi', '$options': 'i'}, 'episode_number': 1}, 'https://youtu.be/LDeseoNGAPQ'),
        ({'id': {'$regex': 'al-kindi', '$options': 'i'}, 'episode_number': 2}, 'https://youtu.be/YK7LJRJheDg'),
        # Avicenne ep1
        ({'id': {'$regex': '^aud_cours-falsafa-grands-avicenne-ep01$'}}, 'https://www.youtube.com/watch?v=PaqA6eCZSRY'),
        # Cursus G - Maïmonide episodes
        ({'id': {'$regex': 'maimonide', '$options': 'i'}, 'episode_number': 1}, 'https://youtu.be/kYWqboZxQP0'),
        ({'id': {'$regex': 'maimonide', '$options': 'i'}, 'episode_number': 2}, 'https://youtu.be/nkXImE6euX4'),
    ]
    yt_set = 0
    for query, url in audio_youtube_map:
        r = await db.audios.update_many(query, {'$set': {'youtube_url': url}})
        yt_set += r.modified_count
    logger.info(f"Migration v4: youtube_url set on {yt_set} audio(s)")

    # 7. Upsert Al-Kindī episodes 2 & 3 (video-only, audio not yet on R2)
    # ⚠️ Legacy: these audios are tied to `cours-falsafa-grands` (umbrella course
    # supprimé par Migration v15h). On skippe entièrement ce bloc si le cours
    # parent n'existe plus → évite la résurrection des audios à chaque boot.
    falsafa_grands_alive = await db.courses.find_one({'id': 'cours-falsafa-grands'}, {'_id': 1})
    if falsafa_grands_alive:
        al_kindi_episodes = [
            {'episode_number': 1, 'youtube_url': 'https://youtu.be/LDeseoNGAPQ', 'title': 'Al-Kindī — Épisode 1'},
            {'episode_number': 2, 'youtube_url': 'https://youtu.be/YK7LJRJheDg', 'title': 'Al-Kindī — Épisode 2'},
            {'episode_number': 3, 'youtube_url': 'https://www.youtube.com/watch?v=hBzgQV2XgrE', 'title': 'Al-Kindī — Épisode 3'},
        ]
        for ep in al_kindi_episodes:
            audio_id = f"aud_cours-falsafa-grands-al-kindi-ep{ep['episode_number']:02d}"
            # title/description/scholar_name are admin-editable: only seed at insert time.
            await db.audios.update_one(
                {'id': audio_id},
                {
                    '$set': {
                        'id': audio_id,
                        'course_id': 'cours-falsafa-grands',
                        'module_id': 'cours-falsafa-grands-mod-1',
                        'scholar_id': 'sch-006',
                        'episode_number': ep['episode_number'],
                        'duration': 0,
                        'youtube_url': ep['youtube_url'],
                        'type': 'episode',
                        'is_active': True,
                    },
                    '$setOnInsert': {
                        'title': ep['title'],
                        'description': '',
                        'scholar_name': 'Prof. Meryem Sebti',
                    },
                },
                upsert=True,
            )
        logger.info("Migration v4: Al-Kindī episodes 1-3 upserted with YouTube URLs")
    
    # 8. Set youtube_url on Histoire de l'art islamique épisode 1
    await db.audios.update_one(
        {'id': 'aud_cours-art-ep01'},
        {'$set': {'youtube_url': 'https://www.youtube.com/watch?v=5zqtVXI9S1Y'}}
    )
    logger.info("Migration v4: youtube_url set on cours-art ep01")
    
    # 9. Feature the new Histoire course on home page (replaces Falsafa featured)
    await db.courses.update_many({'is_featured': True}, {'$set': {'is_featured': False}})
    await db.courses.update_one(
        {'id': 'cours-debuts-islam'},
        {'$set': {'is_featured': True, 'hero_title': 'Histoire du monde islamique', 'hero_description': "Des débuts de l'islam à l'époque ottomane — parcourez l'histoire du monde islamique avec Hassan Bouali et Mehdi Ghouirgate."}}
    )
    logger.info("Migration v4: Histoire du monde islamique featured on home")
    # ─── End Migration v4 ──────────────────────────────────────────────────

    # ─── Migration v5: Deduplicate courses by id ──────────────────────────
    course_dup_pipeline = [
        {'$group': {'_id': '$id', 'docs': {'$push': '$_id'}, 'count': {'$sum': 1}}},
        {'$match': {'count': {'$gt': 1}}},
    ]
    course_dups_removed = 0
    async for group in db.courses.aggregate(course_dup_pipeline):
        # Get all duplicate documents to choose the "best" one to keep
        dup_docs = await db.courses.find({'_id': {'$in': group['docs']}}).to_list(20)
        # Score by content richness: featured > has_youtube_url > has_description > first
        def score(d):
            s = 0
            if d.get('is_featured'): s += 100
            if d.get('youtube_url'): s += 50
            if d.get('hero_title'): s += 20
            if d.get('description'): s += 10
            return s
        dup_docs.sort(key=score, reverse=True)
        keep = dup_docs[0]
        remove_ids = [d['_id'] for d in dup_docs[1:]]
        if remove_ids:
            r = await db.courses.delete_many({'_id': {'$in': remove_ids}})
            course_dups_removed += r.deleted_count
            logger.warning(f"Migration v5: deduplicated course id='{group['_id']}' — kept _id={keep['_id']}, removed {r.deleted_count}")
    if course_dups_removed > 0:
        logger.warning(f"Migration v5: total {course_dups_removed} duplicate courses removed")
    # ─── End Migration v5 ──────────────────────────────────────────────────

    # ─── Migration v6: Cleanup orphan audios (no module_id, no youtube_url) ────
    # User explicitly approved deleting old placeholder audios that don't appear in the launch catalog.
    # Curated audios (those with youtube_url manually set) are preserved.
    # Also preserve placeholder episodes created by v8/v9 (is_placeholder=True).
    orphan_cleanup = await db.audios.delete_many({
        '$and': [
            {'$or': [{'module_id': None}, {'module_id': ''}]},
            {'$or': [{'youtube_url': None}, {'youtube_url': ''}, {'youtube_url': {'$exists': False}}]},
            {'is_placeholder': {'$ne': True}},
        ]
    })
    if orphan_cleanup.deleted_count > 0:
        logger.warning(f"Migration v6: cleaned up {orphan_cleanup.deleted_count} orphan audio(s) without module_id and without youtube_url")
    # ─── End Migration v6 ──────────────────────────────────────────────────

    # ─── Migration v7: Mark launch-catalog modules per the May 2026 Excel ──
    # User-provided Excel `Sijill_Catalogue_Lancement_Mai2026_Emergent.xlsx` defines
    # which sub-modules of multi-module courses must appear in the public catalogue.
    # We mark them is_launch_catalog=True so they show up (as "Bientôt" if no episode yet)
    # and explicitly mark the others False to keep the catalogue clean.
    LAUNCH_MODULES_INCLUDE = {
        # cours-falsafa-grands → only Al-Kindī, Al-Fārābī, Avicenne (NOT "Avicenne dans le monde latin" etc.)
        'cours-falsafa-grands': ['cours-falsafa-grands-mod-1', 'cours-falsafa-grands-mod-2', 'cours-falsafa-grands-mod-3'],
    }
    for course_id, included_module_ids in LAUNCH_MODULES_INCLUDE.items():
        await db.modules.update_many(
            {'course_id': course_id, 'id': {'$in': included_module_ids}},
            {'$set': {'is_launch_catalog': True}}
        )
        await db.modules.update_many(
            {'course_id': course_id, 'id': {'$nin': included_module_ids}},
            {'$set': {'is_launch_catalog': False}}
        )
    # Ensure Al-Ghazālī module exists in cours-falsafa-grands (per Excel) — create if missing
    has_ghazali = await db.modules.find_one({'course_id': 'cours-falsafa-grands', 'id': 'cours-falsafa-grands-mod-ghazali'})
    if not has_ghazali:
        await db.modules.insert_one({
            'id': 'cours-falsafa-grands-mod-ghazali',
            'course_id': 'cours-falsafa-grands',
            'name': "Al-Ghazālī (m. 1111)",
            'title': "Al-Ghazālī (m. 1111)",
            'description': "",
            'order': 4,
            'is_active': True,
            'is_launch_catalog': True,
            'created_at': datetime.now(timezone.utc),
        })
        logger.info("Migration v7: created module 'Al-Ghazālī' under cours-falsafa-grands")
    else:
        await db.modules.update_one(
            {'id': 'cours-falsafa-grands-mod-ghazali'},
            {'$set': {'is_launch_catalog': True, 'is_active': True}}
        )
    logger.info("Migration v7: launch-catalog modules pinned per Excel (Falsafa: Kindī, Fārābī, Avicenne, Ghazālī)")
    # ─── End Migration v7 ──────────────────────────────────────────────────

    # ─── Migration v8: Placeholder episodes for Falsafa modules (URLs pending) ─
    # Per user Excel: Al-Fārābī (5 ép), Avicenne (5 ép), Al-Ghazālī (1 ép).
    # Episodes are created without youtube_url; URLs will be added later via Admin Panel.
    # The episodes will appear in the catalogue with their count; playback will display
    # a "Vidéo bientôt disponible" placeholder until URLs are filled in.
    #
    # ⚠️ Legacy: these placeholders are anchored on `cours-falsafa-grands` (the
    # umbrella course that Migration v15h deletes). Once that course is gone,
    # skip this whole block — otherwise every reboot re-creates 11 ghost audios
    # that v15h has to purge again. Causes log noise and pointless DB churn.
    falsafa_grands_alive_v8 = await db.courses.find_one({'id': 'cours-falsafa-grands'}, {'_id': 1})
    if falsafa_grands_alive_v8:
        PLACEHOLDER_EPISODES = [
            ('cours-falsafa-grands-mod-2', 'al-farabi', 'Al-Fārābī', 5),
            ('cours-falsafa-grands-mod-3', 'avicenne', 'Avicenne', 5),
            ('cours-falsafa-grands-mod-ghazali', 'al-ghazali', 'Al-Ghazālī', 1),
        ]
        placeholder_count = 0
        for module_id, slug, label, n_episodes in PLACEHOLDER_EPISODES:
            for ep in range(1, n_episodes + 1):
                audio_id = f"aud_cours-falsafa-grands-{slug}-ep{ep:02d}"
                existing = await db.audios.find_one({'id': audio_id})
                if existing:
                    # Ensure existing placeholders are active (idempotent fix for v8 first run)
                    if existing.get('is_active') is not True:
                        await db.audios.update_one({'id': audio_id}, {'$set': {'is_active': True}})
                    continue
                await db.audios.insert_one({
                    'id': audio_id,
                    'course_id': 'cours-falsafa-grands',
                    'module_id': module_id,
                    'title': f"{label} — Épisode {ep}",
                    'episode_number': ep,
                    'duration_seconds': 0,
                    'youtube_url': None,
                    'audio_url': None,
                    'file_key': None,
                    'is_placeholder': True,
                    'is_active': True,
                    'created_at': datetime.now(timezone.utc),
                })
                placeholder_count += 1
        if placeholder_count > 0:
            logger.info(f"Migration v8: {placeholder_count} placeholder episode(s) created (URLs pending)")
        # Re-link any pre-existing Avicenne ep1 that was created before module_id was set
        legacy_avicenne_ep1 = await db.audios.find_one({'id': 'aud_cours-falsafa-grands-avicenne-ep01'})
        if legacy_avicenne_ep1 and legacy_avicenne_ep1.get('module_id') != 'cours-falsafa-grands-mod-3':
            await db.audios.update_one(
                {'id': 'aud_cours-falsafa-grands-avicenne-ep01'},
                {'$set': {'module_id': 'cours-falsafa-grands-mod-3'}}
            )
            logger.info("Migration v8: re-linked aud_cours-falsafa-grands-avicenne-ep01 to mod-3")
    # ─── End Migration v8 ──────────────────────────────────────────────────

    # ─── Migration v9: Excel update Mai 2026 v2 ─────────────────────────────
    # Source: user-provided Sijill_Catalogue_Lancement_Mai2026_Emergent.xlsx (v2)
    # 1) Remove cours-post-avicennisme from launch catalogue (per user message)
    await db.courses.update_one(
        {'id': 'cours-post-avicennisme'},
        {'$set': {'is_launch_catalog': False}}
    )
    logger.info("Migration v9: cours-post-avicennisme removed from launch catalogue")

    # 2) Set scholar_name on each launch course (used by Catalogue card to show professor)
    SCHOLAR_BY_COURSE = {
        'cours-debuts-islam': 'Hassan Bouali · Mehdi Ghouirgate',
        'cours-andalus': 'Mehdi Ghouirgate',
        'cours-mamelouke': 'Sami Benkherfallah',
        'cours-ottoman': 'Aysu Saban',
        'cours-kalam': 'Ilyas Harifi',
        'cours-fiqh': 'Yannis Mahil',
        'cours-coran': 'Mehdi Azaiez',
        'cours-hadith': 'Hassan Chahdi',
        'cours-historiographie': 'Mehdi Ghouirgate',
        'cours-art': 'Camille Grandpierre',
        'cours-sciences': 'Marouane Ben Miled · Meyssa Ben Saad',
        'cours-traduction': 'Meryem Sebti',
        'cours-falsafa-grands': 'Meryem Sebti',
        'cours-falsafa-occident': 'Yassir Mechelloukh',
        'cours-falsafa-persan': 'Sajjad Rizvi',
        'cours-inclassables': 'Cédric Molino-Mochetto',
        'cours-soufisme': 'Gregory Vandamme',
        'cours-philo-juive': 'Géraldine Roux',
    }
    for cid, sname in SCHOLAR_BY_COURSE.items():
        # Skip if the course was edited via admin panel (seed_locked).
        await db.courses.update_one(
            {'id': cid, 'seed_locked': {'$ne': True}},
            {'$set': {'scholar_name': sname}},
        )
    logger.info(f"Migration v9: scholar_name set on {len(SCHOLAR_BY_COURSE)} courses")

    # 2bis) Set scholar_name on Falsafa modules per Excel (all by Meryem Sebti)
    SEBTI_MODULE_IDS = [
        'cours-falsafa-grands-mod-1',  # Al-Kindī
        'cours-falsafa-grands-mod-2',  # Al-Fārābī
        'cours-falsafa-grands-mod-3',  # Avicenne
        'cours-falsafa-grands-mod-ghazali',  # Al-Ghazālī
    ]
    await db.modules.update_many(
        {'id': {'$in': SEBTI_MODULE_IDS}},
        {'$set': {'scholar_name': 'Meryem Sebti'}}
    )
    logger.info(f"Migration v9: scholar_name='Meryem Sebti' set on {len(SEBTI_MODULE_IDS)} Falsafa modules")

    # 3) Hassan Bouali — 4 sub-episodes for "Les débuts de l'islam" + 1 Ghouirgate
    DEBUTS_ISLAM_EPISODES = [
        (1, "Muhammad et les débuts de l'islam", 'Hassan Bouali', None),
        (2, "Le califat au début de l'islam : de Médine à la sortie d'Arabie", 'Hassan Bouali', None),
        (3, "Les futūḥāt / conquêtes", 'Hassan Bouali', None),
        (4, "Le ḥajj et son intégration progressive à la piété musulmane", 'Hassan Bouali', None),
        (5, "Les débuts de l'islam (introduction générale)", 'Mehdi Ghouirgate', 'https://youtu.be/kB-fr8wwcAA'),
    ]
    for ep_num, ep_title, scholar, yt_url in DEBUTS_ISLAM_EPISODES:
        audio_id = f"aud_cours-debuts-islam-ep{ep_num:02d}"
        # title/scholar_name are admin-editable: only seed at insert time.
        await db.audios.update_one(
            {'id': audio_id},
            {
                '$set': {
                    'id': audio_id,
                    'course_id': 'cours-debuts-islam',
                    'module_id': None,
                    'episode_number': ep_num,
                    'duration_seconds': 0,
                    'youtube_url': yt_url,
                    'audio_url': None,
                    'file_key': None,
                    'is_placeholder': yt_url is None,
                    'is_active': True,
                },
                '$setOnInsert': {
                    'title': ep_title,
                    'scholar_name': scholar,
                    'created_at': datetime.now(timezone.utc),
                },
            },
            upsert=True
        )
    logger.info("Migration v9: cours-debuts-islam — 5 episodes upserted (4 Bouali + 1 Ghouirgate)")

    # 4) Update Al-Kindī — add ep4 with URL + ep5 placeholder per Excel
    # title/scholar_name are admin-editable: only seed at insert time.
    await db.audios.update_one(
        {'id': 'aud_cours-falsafa-grands-al-kindi-ep04'},
        {
            '$set': {
                'id': 'aud_cours-falsafa-grands-al-kindi-ep04',
                'course_id': 'cours-falsafa-grands',
                'module_id': 'cours-falsafa-grands-mod-1',
                'episode_number': 4,
                'youtube_url': 'https://www.youtube.com/watch?v=GeAHez38fzw',
                'is_active': True,
            },
            '$setOnInsert': {
                'title': 'Al-Kindī — Épisode 4',
                'scholar_name': 'Meryem Sebti',
                'created_at': datetime.now(timezone.utc),
            },
        },
        upsert=True
    )
    await db.audios.update_one(
        {'id': 'aud_cours-falsafa-grands-al-kindi-ep05'},
        {
            '$set': {
                'id': 'aud_cours-falsafa-grands-al-kindi-ep05',
                'course_id': 'cours-falsafa-grands',
                'module_id': 'cours-falsafa-grands-mod-1',
                'episode_number': 5,
                'youtube_url': None,
                'is_placeholder': True,
                'is_active': True,
            },
            '$setOnInsert': {
                'title': 'Al-Kindī — Épisode 5',
                'scholar_name': 'Meryem Sebti',
                'created_at': datetime.now(timezone.utc),
            },
        },
        upsert=True
    )
    logger.info("Migration v9: Al-Kindī ep4 (URL) + ep5 (placeholder) upserted")

    # 5) Histoire de l'art islamique — 4 episodes per Excel (replace ep1 URL, add ep2 URL, ep3-4 placeholders)
    ART_EPISODES = [
        (1, "Histoire de l'art islamique — Épisode 1", 'https://www.youtube.com/watch?v=tNzuUgeMGb4'),
        (2, "Histoire de l'art islamique — Épisode 2", 'https://www.youtube.com/watch?v=QDK4iVJ7b2k'),
        (3, "Histoire de l'art islamique — Épisode 3", None),
        (4, "Histoire de l'art islamique — Épisode 4", None),
    ]
    for ep_num, ep_title, yt in ART_EPISODES:
        audio_id = f"aud_cours-art-ep{ep_num:02d}"
        # title/scholar_name are admin-editable: only seed at insert time.
        await db.audios.update_one(
            {'id': audio_id},
            {
                '$set': {
                    'id': audio_id,
                    'course_id': 'cours-art',
                    'module_id': None,
                    'episode_number': ep_num,
                    'youtube_url': yt,
                    'is_placeholder': yt is None,
                    'is_active': True,
                },
                '$setOnInsert': {
                    'title': ep_title,
                    'scholar_name': 'Camille Grandpierre',
                    'created_at': datetime.now(timezone.utc),
                },
            },
            upsert=True
        )
    logger.info("Migration v9: cours-art — 4 episodes upserted (ep1+ep2 URL, ep3+ep4 placeholders)")

    # 6) Maïmonide ep1 + ep2 (cours-philo-juive) — upsert both (some older DBs miss ep1)
    # title/scholar_name are admin-editable: only seed at insert time.
    await db.audios.update_one(
        {'id': 'aud_cours-philo-juive-maimonide-ep01'},
        {
            '$set': {
                'id': 'aud_cours-philo-juive-maimonide-ep01',
                'course_id': 'cours-philo-juive',
                'module_id': 'cours-philo-juive-mod-9',
                'episode_number': 1,
                'youtube_url': 'https://youtu.be/kYWqboZxQP0',
                'is_active': True,
            },
            '$setOnInsert': {
                'title': 'Moïse Maïmonide — Épisode 1',
                'scholar_name': 'Géraldine Roux',
                'created_at': datetime.now(timezone.utc),
            },
        },
        upsert=True
    )
    await db.audios.update_one(
        {'id': 'aud_cours-philo-juive-maimonide-ep02'},
        {
            '$set': {
                'id': 'aud_cours-philo-juive-maimonide-ep02',
                'course_id': 'cours-philo-juive',
                'module_id': 'cours-philo-juive-mod-9',
                'episode_number': 2,
                'youtube_url': 'https://youtu.be/nkXImE6euX4',
                'is_active': True,
            },
            '$setOnInsert': {
                'title': 'Moïse Maïmonide — Épisode 2',
                'scholar_name': 'Géraldine Roux',
                'created_at': datetime.now(timezone.utc),
            },
        },
        upsert=True
    )
    logger.info("Migration v9: cours-philo-juive ep1 + ep2 upserted")

    # 6bis) Pin Maïmonide module (mod-9) as launch + relink existing Maïmonide episodes to it
    await db.modules.update_one(
        {'id': 'cours-philo-juive-mod-9'},
        {'$set': {'is_launch_catalog': True, 'scholar_name': 'Géraldine Roux'}}
    )
    # Mark all OTHER cours-philo-juive modules explicitly NOT in launch (clean catalogue)
    await db.modules.update_many(
        {'course_id': 'cours-philo-juive', 'id': {'$ne': 'cours-philo-juive-mod-9'}},
        {'$set': {'is_launch_catalog': False}}
    )
    # Relink the 2 Maïmonide audios to the proper module
    await db.audios.update_many(
        {'id': {'$in': ['aud_cours-philo-juive-maimonide-ep01', 'aud_cours-philo-juive-maimonide-ep02']}},
        {'$set': {'module_id': 'cours-philo-juive-mod-9'}}
    )
    logger.info("Migration v9: Maïmonide module pinned as launch + 2 episodes re-linked to mod-9")

    # 6ter) PILOT — Maïmonide: video R2 + script per episode + biblio + glossaire (course-level)
    # Files in R2 under cursus-f-nonarabe/24-philosophie-juive/maimonide/
    MAIMONIDE_R2_PREFIX = "cursus-f-nonarabe/24-philosophie-juive/maimonide"

    def _probe_r2_key(key: str) -> bool:
        """HEAD-check a candidate key in R2; returns True if the file exists."""
        if not r2_client or not key:
            return False
        try:
            r2_client.head_object(Bucket=R2_BUCKET, Key=key)
            return True
        except Exception:
            return False

    for _ep_num in (1, 2):
        _audio_id = f"aud_cours-philo-juive-maimonide-ep0{_ep_num}"
        # Probe several known naming conventions, .m4a first (current upload), .mp3 fallback
        _candidates = [
            f"{MAIMONIDE_R2_PREFIX}/Penseenonarabe-maimounide-{_ep_num}.m4a",
            f"{MAIMONIDE_R2_PREFIX}/episode{_ep_num}_maimounide.m4a",
            f"{MAIMONIDE_R2_PREFIX}/episode{_ep_num}_maimounide.mp3",
        ]
        _audio_key = next((k for k in _candidates if _probe_r2_key(k)), None)
        _has_audio = bool(_audio_key)
        # If neither exists yet, store the expected default so admins can see the path
        if not _audio_key:
            _audio_key = _candidates[0]
        await db.audios.update_one(
            {'id': _audio_id},
            {'$set': {
                'r2_video_key': f"{MAIMONIDE_R2_PREFIX}/episode{_ep_num}_maimounide.mp4",
                'r2_audio_key': _audio_key,
                'has_r2_audio': _has_audio,
                'episode_resources': [
                    {'type': 'script', 'label': "Script de l'épisode", 'r2_key': f"{MAIMONIDE_R2_PREFIX}/episode{_ep_num}_maimounide.pdf", 'mime': 'application/pdf'},
                ],
            }}
        )
    # Course-level resources for Maïmonide (biblio + glossaire shared across all episodes)
    await db.courses.update_one(
        {'id': 'cours-philo-juive'},
        {'$set': {
            'course_resources': [
                {'type': 'biblio', 'label': 'Bibliographie sélective', 'r2_key': f"{MAIMONIDE_R2_PREFIX}/bibliographie_maimounide.pdf", 'mime': 'application/pdf'},
                {'type': 'glossaire', 'label': 'Glossaire des termes', 'r2_key': f"{MAIMONIDE_R2_PREFIX}/glossaire-maimounide.pdf", 'mime': 'application/pdf'},
            ],
        }}
    )
    logger.info("Migration v9 PILOT: Maïmonide R2 video + scripts + biblio + glossaire wired")
    # Persist r2_prefix so Migration v11 will pick it up
    await db.courses.update_one(
        {'id': 'cours-philo-juive'},
        {'$set': {'r2_prefix': MAIMONIDE_R2_PREFIX + '/'}}
    )
    # ─── End Migration v9 ──────────────────────────────────────────────────

    # ─── Migration v12: Persist r2_prefix on launch-catalog courses (May 2026) ──
    # Mapping derived from /app/Sijill_Catalogue_Lancement_Mai2026 + R2 listing.
    # Note: Al-Kindī, Al-Fārābī, Avicenne are MODULES under cours-falsafa-grands.
    #       Droit musulman is a MODULE under cours-fiqh.
    #       Ibn Khaldūn historiographie is a MODULE under cours-historiographie.
    # The R2 prefix on a course covers all sub-modules under it.
    V12_PREFIXES = {
        'cours-traduction': 'cursus-a-falsafa/01-mouvement-traduction/',
        'cours-falsafa-grands': 'cursus-a-falsafa/02-falsafa/',         # covers al-kindi/, al-farabi/, avicenne/
        'cours-fiqh': 'cursus-b-theologie-droit/10-usul-al-fiqh/',       # covers droit-musulman/
        'cours-historiographie': 'cursus-c-sciences-islamiques/14-historiographie/',  # covers ibn-khaldun-histoire/
        'cours-art': 'cursus-d-arts-litterature/16-histoire-art/',
        'cours-andalus': 'cursus-histoire/andalous/',
        'cours-debuts-islam': 'cursus-histoire/debuts-islam/',
        'cours-philo-juive': 'cursus-f-nonarabe/24-philosophie-juive/',  # covers maimonide/
    }
    v12_set = 0
    for cid, prefix in V12_PREFIXES.items():
        r = await db.courses.update_one({'id': cid}, {'$set': {'r2_prefix': prefix}})
        if r.matched_count:
            v12_set += 1
    logger.info(f"Migration v12: r2_prefix set on {v12_set}/{len(V12_PREFIXES)} launch courses")
    # ─── End Migration v12 ─────────────────────────────────────────────────

    # ─── Migration v11: Auto-sync R2 media for ALL courses with r2_prefix ──
    try:
        v11_courses = await db.courses.find(
            {'r2_prefix': {'$exists': True, '$ne': ''}},
            {'_id': 0, 'id': 1, 'title': 1, 'r2_prefix': 1},
        ).to_list(500)
        v11_synced = 0
        for c in v11_courses:
            prefix = (c.get('r2_prefix') or '').strip().strip('/')
            if not prefix:
                continue
            if not prefix.endswith('/'):
                prefix = prefix + '/'
            try:
                detections = _build_r2_detections(prefix)
                summary = await _apply_r2_detections(c['id'], detections)
                v11_synced += 1
                logger.info(
                    f"Migration v11: synced {c['id']} ({c.get('title') or '?'}) — "
                    f"course_docs={summary['course_resources_count']}, "
                    f"videos={summary['audios_with_video']}, audios={summary['audios_with_audio']}, "
                    f"scripts={summary['audios_with_script']}, unclassified={summary['unclassified_count']}"
                )
            except Exception as e:
                logger.warning(f"Migration v11: failed to sync {c['id']}: {e}")
        logger.info(f"Migration v11: R2 auto-sync complete ({v11_synced}/{len(v11_courses)} courses)")
    except Exception as e:
        logger.warning(f"Migration v11 skipped: {e}")
    # ─── End Migration v11 ─────────────────────────────────────────────────

    # ─── Migration v13: published_at default for launch-catalog episodes ────
    # For all audios attached to a launch-catalog course that don't have an explicit
    # `published_at`, default it to "2026-05" so the frontend can show a date badge
    # ("À venir · Mai 2026") on episodes that aren't yet available. Episodes with
    # actual media (r2_audio_key, youtube_url, audio_url) keep status='available'.
    try:
        launch_course_ids_v13 = [c['id'] async for c in db.courses.find(
            {'is_launch_catalog': True}, {'_id': 0, 'id': 1}
        )]
        if launch_course_ids_v13:
            res_v13 = await db.audios.update_many(
                {
                    'course_id': {'$in': launch_course_ids_v13},
                    '$or': [
                        {'published_at': {'$exists': False}},
                        {'published_at': None},
                        {'published_at': ''},
                    ],
                },
                {'$set': {'published_at': '2026-05'}},
            )
            logger.info(f"Migration v13: published_at='2026-05' set on {res_v13.modified_count} launch-catalog episodes")
    except Exception as e:
        logger.warning(f"Migration v13 skipped: {e}")
    # ─── End Migration v13 ────────────────────────────────────────────────

    # ─── Migration v14: Split falsafa-grands + apply Excel catalog YT links ───
    try:
        # 1) NEW courses to create (or upsert) in cursus-falsafa
        v14_new_courses = [
            ('cours-al-kindi',           'Al-Kindī',                              'Le « philosophe des Arabes » : vie, pensée et héritage.',                'cursus-a-falsafa/02-falsafa/al-kindi/',   2),
            ('cours-al-farabi',          'Al-Fārābī',                             'Le « Second Maître » : logique, métaphysique et cité vertueuse.',        'cursus-a-falsafa/02-falsafa/al-farabi/',  3),
            ('cours-avicenne',           'Avicenne (Ibn Sīnā)',                    "Le penseur encyclopédique qui marqua l'Orient et l'Occident.",           'cursus-a-falsafa/02-falsafa/avicenne/',   4),
            ('cours-al-ghazali',         'Al-Ghazālī',                            "L'imam-philosophe : critique de la falsafa et renouveau spirituel.",      'cursus-a-falsafa/02-falsafa/al-ghazali/', 5),
            ('cours-falsafa-occident',   "La falsafa en Occident musulman",       "Ibn Bajja, Ibn Ṭufayl et Averroès — la philosophie en al-Andalus.",      'cursus-a-falsafa/03-occident-musulman/',  6),
            ('cours-falsafa-inclassables', 'Les inclassables (Ibn Khaldūn)',      "La pensée d'Ibn Khaldūn entre histoire, philosophie et sociologie.",     'cursus-a-falsafa/07-inclassables/',       7),
            ('cours-falsafa-persane',    "Renouveau de la philosophie persane",   "Mullā Ṣadrā et la philosophie islamique tardive.",                       'cursus-a-falsafa/06-renouveau-persan/',   8),
        ]
        for cid, title, summary, r2_prefix, order in v14_new_courses:
            existing_v14 = await db.courses.find_one(
                {'id': cid}, {'_id': 0, 'title': 1, 'summary': 1, 'description': 1}
            )
            structural_fields = {
                'cursus_id': 'cursus-falsafa',
                'r2_prefix': r2_prefix,
                'is_launch_catalog': True,
                'is_active': True,
                'order': order,
            }
            if not existing_v14:
                # New course → insert with full seed content.
                await db.courses.insert_one({
                    'id': cid, 'title': title, 'summary': summary, 'description': summary,
                    'modules': [{'id': f'{cid}-mod-1', 'order': 1, 'title': title}],
                    'created_at': datetime.now(timezone.utc).isoformat(),
                    **structural_fields,
                })
            else:
                # Existing → ONLY refresh structural fields. Never overwrite
                # admin-edited title / summary / description.
                await db.courses.update_one({'id': cid}, {'$set': structural_fields})

        # 2) Deactivate the old umbrella course
        await db.courses.update_one(
            {'id': 'cours-falsafa-grands'},
            {'$set': {'is_launch_catalog': False, 'is_active': False}},
        )

        # 3) Move existing audios from cours-falsafa-grands → their dedicated course
        FALSAFA_ROUTING = [
            ('al-kindi',                'cours-al-kindi'),
            ('al-farabi',               'cours-al-farabi'),
            ('avicenne',                'cours-avicenne'),
            ('avicenne-monde-latin',    'cours-avicenne'),
            ('disciples-avicenne',      'cours-avicenne'),
            ('al-ghazali',              'cours-al-ghazali'),
        ]
        for sf_pat, new_course_id in FALSAFA_ROUTING:
            res_move = await db.audios.update_many(
                {
                    'course_id': 'cours-falsafa-grands',
                    '$or': [
                        {'r2_subprefix': {'$regex': f'^{sf_pat}/?$'}},
                        {'r2_audio_key': {'$regex': f'/{sf_pat}/'}},
                    ],
                },
                {'$set': {
                    'course_id': new_course_id,
                    'module_id': f'{new_course_id}-mod-1',
                    'r2_subprefix': f'{sf_pat}/',
                }},
            )
            if res_move.modified_count:
                logger.info(f"Migration v14: moved {res_move.modified_count} audio(s) -> {new_course_id} (sf={sf_pat})")
        # Deactivate remaining orphan audios still tagged on the old umbrella course
        await db.audios.update_many(
            {'course_id': 'cours-falsafa-grands'},
            {'$set': {'is_active': False}},
        )

        # 4) Apply YouTube URLs from the Excel catalog
        EXCEL_YT_MAP = {
            ('cours-al-kindi', 1):  'https://youtu.be/LDeseoNGAPQ',
            ('cours-al-kindi', 2):  'https://youtu.be/YK7LJRJheDg',
            ('cours-al-kindi', 3):  'https://youtu.be/hBzgQV2XgrE',
            ('cours-al-kindi', 4):  'https://youtu.be/GeAHez38fzw',
            ('cours-al-kindi', 5):  'https://youtu.be/dQ14KnQonMU',
            ('cours-al-farabi', 1): 'https://youtu.be/c5CKp355l0U',
            ('cours-al-farabi', 2): 'https://youtu.be/L7uD6n-D4-g',
            ('cours-al-farabi', 3): 'https://youtu.be/ehwtqMRApY0',
            ('cours-avicenne', 1):  'https://youtu.be/PaqA6eCZSRY',
            ('cours-traduction', 1):'https://youtu.be/5EsSIUfeP-o',
            ('cours-fiqh', 1):      'https://youtu.be/42OhGxr4THU',
            ('cours-fiqh', 2):      'https://youtu.be/H7Q2kepCGQY',
            ('cours-fiqh', 3):      'https://youtu.be/PSuT_A7IMMA',
            ('cours-fiqh', 4):      'https://youtu.be/j7dWjZQXo08',
            ('cours-fiqh', 5):      'https://youtu.be/vITKWxnhIi8',
            ('cours-fiqh', 6):      'https://youtu.be/rvjZ4RkDDnw',
            ('cours-art', 1):       'https://youtu.be/5zqtVXI9S1Y',
            ('cours-art', 2):       'https://youtu.be/QDK4iVJ7b2k',
            ('cours-art', 3):       'https://youtu.be/fFDsyYhCnFo',
            ('cours-art', 4):       'https://youtu.be/A9J_CkjT_hU',
            ('cours-historiographie', 1): 'https://youtu.be/RUc8p0K6Qg4',
            ('cours-philo-juive', 1): 'https://youtu.be/kYWqboZxQP0',
            ('cours-philo-juive', 2): 'https://youtu.be/nkXImE6euX4',
            ('cours-debuts-islam', 1): 'https://youtu.be/kB-fr8wwcAA',
            ('cours-andalus', 1):   'https://youtu.be/cow2JfYaSC0',
        }
        yt_applied = 0
        for (cid, ep_num), yt in EXCEL_YT_MAP.items():
            res_yt = await db.audios.update_one(
                {'course_id': cid, 'episode_number': ep_num},
                {'$set': {'youtube_url': yt, 'is_active': True}},
            )
            if res_yt.modified_count:
                yt_applied += 1
                continue
            # No audio doc yet — create a YT-only placeholder
            new_id = f'aud_{cid}-yt-ep{ep_num:02d}'
            cdoc = await db.courses.find_one({'id': cid}, {'_id': 0, 'modules': 1, 'title': 1})
            mid = ((cdoc or {}).get('modules') or [{}])[0].get('id') if cdoc else None
            ctitle = (cdoc or {}).get('title', 'Épisode')
            await db.audios.update_one(
                {'id': new_id},
                {
                    '$set': {
                        'id': new_id, 'course_id': cid, 'module_id': mid,
                        'episode_number': ep_num,
                        'youtube_url': yt, 'is_active': True, 'published_at': '2026-05',
                    },
                    '$setOnInsert': {
                        'title': f"{ctitle} — Épisode {ep_num}",
                        'audio_url': '',
                        'created_at': datetime.now(timezone.utc).isoformat(),
                    },
                },
                upsert=True,
            )
            yt_applied += 1
        logger.info(f"Migration v14: applied {yt_applied} YouTube URLs from Excel")

        # 5) cours-debuts-islam — Now fully managed by Migration v15g (which
        # applies the canonical Excel titles ep1-5 and locks them). The legacy
        # logic that used ep_num 3-6 with shifted numbering has been removed
        # to prevent it from re-creating stale ep6 duplicates on every reboot.

        # 6) cours-andalus — Ghouirgate intro
        await db.audios.update_one(
            {'course_id': 'cours-andalus', 'episode_number': 1},
            {
                '$set': {
                    'youtube_url': 'https://youtu.be/cow2JfYaSC0', 'is_active': True,
                    'published_at': '2026-05',
                },
                '$setOnInsert': {
                    'title': 'Al-Andalus — Introduction',
                    'intervenant': 'Mehdi Ghouirgate',
                },
            },
        )
        logger.info("Migration v14: falsafa split + YT applied + debuts-islam cleaned")

        # 7) v14b CLEANUP — deduplicate per (course_id, episode_number)
        # Keep the best audio per (course, ep): prefer one with both YT + R2_audio,
        # then R2 audio only, then YT only. Delete the rest.
        v14_target_courses = [c[0] for c in v14_new_courses] + [
            'cours-traduction', 'cours-fiqh', 'cours-debuts-islam',
            'cours-andalus', 'cours-historiographie', 'cours-art', 'cours-philo-juive',
        ]
        for cid in v14_target_courses:
            ep_groups: dict = {}
            async for a in db.audios.find({'course_id': cid}, {'_id': 0}):
                ep = a.get('episode_number') or 0
                ep_groups.setdefault(ep, []).append(a)
            for ep, group in ep_groups.items():
                if len(group) <= 1:
                    continue
                # Score each audio: 4=R2+YT, 3=R2 only, 2=YT only, 1=neither
                def _score(a):
                    has_r2 = bool(a.get('r2_audio_key'))
                    has_yt = bool(a.get('youtube_url'))
                    return (2 if has_r2 else 0) + (1 if has_yt else 0)
                group.sort(key=lambda a: (-_score(a), -(len(a.get('title') or '')), a.get('id') or ''))
                # Merge: keep top, copy missing YT/R2 from others, delete the rest
                keep = group[0]
                merged_set = {}
                for other in group[1:]:
                    if not keep.get('youtube_url') and other.get('youtube_url'):
                        merged_set['youtube_url'] = other['youtube_url']
                        keep['youtube_url'] = other['youtube_url']
                    if not keep.get('r2_audio_key') and other.get('r2_audio_key'):
                        merged_set['r2_audio_key'] = other['r2_audio_key']
                        merged_set['has_r2_audio'] = True
                        keep['r2_audio_key'] = other['r2_audio_key']
                if merged_set:
                    await db.audios.update_one({'id': keep['id']}, {'$set': merged_set})
                victim_ids = [o['id'] for o in group[1:] if o.get('id')]
                if victim_ids:
                    res_del = await db.audios.delete_many({'id': {'$in': victim_ids}})
                    logger.info(f"Migration v14b: cleaned {res_del.deleted_count} dup audio(s) on {cid} ep{ep} (kept {keep.get('id')})")
        # Special: drop ghost audios on cours-avicenne that aren't real avicenne content
        await db.audios.delete_many({
            'course_id': 'cours-avicenne',
            'title': {'$regex': '^(Avicenne Monde Latin|Disciples Avicenne|Épisode \\d+$)'},
        })
        # Special: cours-traduction Excel has only ep1, drop fake ep2 placeholders
        await db.audios.delete_many({
            'course_id': 'cours-traduction', 'episode_number': {'$gt': 1},
        })
        logger.info("Migration v14b: dedup + ghost cleanup complete")

        # v14c — Final fixes:
        # - cours-debuts-islam: episodes are now fully managed by Migration v15g
        #   from the user's Excel (eps 1-5 with canonical titles + seed_locked).
        #   The legacy v14c logic that deleted ep2 and constrained eps to
        #   {1, 3, 4, 5, 6} has been disabled.
        # - cours-historiographie: ensure ep1 has intervenant + Excel title at INSERT only.
        # We leave title untouched after creation (admin-editable).
        await db.audios.update_one(
            {'course_id': 'cours-historiographie', 'episode_number': 1},
            {'$set': {'intervenant': 'Mehdi Ghouirgate'}},
        )
        # - cours-al-ghazali: create an "À venir" placeholder ep1 (Commandé per Excel)
        ag_audio = await db.audios.find_one({'course_id': 'cours-al-ghazali', 'episode_number': 1})
        if not ag_audio:
            ag_course = await db.courses.find_one({'id': 'cours-al-ghazali'}, {'_id': 0, 'modules': 1})
            ag_mid = ((ag_course or {}).get('modules') or [{}])[0].get('id')
            await db.audios.update_one(
                {'id': 'aud_cours-al-ghazali-ep01'},
                {
                    '$set': {
                        'id': 'aud_cours-al-ghazali-ep01', 'course_id': 'cours-al-ghazali',
                        'module_id': ag_mid, 'episode_number': 1,
                        'is_active': True, 'published_at': '2026-05',
                    },
                    '$setOnInsert': {
                        'title': 'Al-Ghazālī — Épisode 1',
                        'intervenant': 'Meryem Sebti',
                        'audio_url': '',
                        'created_at': datetime.now(timezone.utc).isoformat(),
                    },
                },
                upsert=True,
            )
        logger.info("Migration v14c: final polish applied")

        # v14d — ensure each new course has an entry in the `modules` collection
        # (the catalogue endpoint joins audios → modules → courses to count episodes)
        for cid, title, _summary, _r2, _order in v14_new_courses:
            mid = f'{cid}-mod-1'
            await db.modules.update_one(
                {'id': mid},
                {'$set': {
                    'id': mid, 'course_id': cid, 'order': 1,
                    'name': title, 'title': title,
                    'is_active': True, 'is_launch_catalog': True,
                }},
                upsert=True,
            )
        logger.info("Migration v14d: ensured modules collection for new falsafa courses")

        # v14e — attach orphan audios (module_id=None) to the matching module by title slug.
        # For cours-fiqh: Droit Musulman audios → module 'Droit musulman' (cours-fiqh-mod-2).
        # For cours-historiographie ep1 'Ibn Khaldūn' → module 'Ibn Khaldūn (...)' (mod-2).
        # Otherwise fallback to the first active module of the course.
        async def _attach_audio_to_best_module(course_id: str):
            mods = await db.modules.find(
                {'course_id': course_id, 'is_active': {'$ne': False}},
                {'_id': 0, 'id': 1, 'name': 1, 'title': 1, 'order': 1}
            ).sort('order', 1).to_list(50)
            if not mods:
                return
            audios = await db.audios.find(
                {'course_id': course_id, '$or': [{'module_id': None}, {'module_id': {'$exists': False}}]},
                {'_id': 0, 'id': 1, 'title': 1, 'episode_number': 1}
            ).to_list(200)
            def _norm(s): return (s or '').lower().replace('-', ' ').replace('_', ' ')
            for a in audios:
                atit = _norm(a.get('title'))
                target = None
                for m in mods:
                    mtit = _norm(m.get('name') or m.get('title'))
                    if not mtit:
                        continue
                    key = mtit.split('(')[0].strip()
                    if key and key in atit:
                        target = m; break
                if not target:
                    target = mods[0]
                await db.audios.update_one(
                    {'id': a['id']},
                    {'$set': {'module_id': target['id']}},
                )
        for cid in ['cours-fiqh', 'cours-historiographie', 'cours-philo-juive',
                    'cours-art', 'cours-debuts-islam', 'cours-andalus', 'cours-traduction']:
            await _attach_audio_to_best_module(cid)
        logger.info("Migration v14e: attached orphan audios to best matching module")
    except Exception as e:
        logger.warning(f"Migration v14 failed: {e}", exc_info=True)

    # ─── Migration v15 — orphan audio purge + historiographie split ────────
    # Purge audios whose r2_audio_key no longer exists in R2 (user deleted in Cloudflare).
    # Also rename cours-historiographie to focus on Ibn Khaldūn (1 character = 1 course).
    try:
        if r2_client:
            # Build a set of existing R2 keys (cache per prefix to keep it fast)
            from botocore.exceptions import ClientError as _CE
            existing_cache: dict = {}
            def _r2_key_exists(key: str) -> bool:
                if not key:
                    return False
                if key in existing_cache:
                    return existing_cache[key]
                try:
                    r2_client.head_object(Bucket=R2_BUCKET, Key=key)
                    existing_cache[key] = True
                    return True
                except _CE:
                    existing_cache[key] = False
                    return False
                except Exception:
                    existing_cache[key] = True  # be conservative on transient errors
                    return True
            # Scan all audios that claim an r2_audio_key
            purged_total = 0
            async for a in db.audios.find(
                {'r2_audio_key': {'$exists': True, '$ne': None}},
                {'_id': 0, 'id': 1, 'course_id': 1, 'r2_audio_key': 1, 'youtube_url': 1, 'title': 1}
            ):
                key = a.get('r2_audio_key')
                if not _r2_key_exists(key):
                    # If the audio also has no YT URL → DELETE it (truly orphan)
                    if not a.get('youtube_url'):
                        await db.audios.delete_one({'id': a['id']})
                        purged_total += 1
                        logger.info(f"Migration v15: purged orphan audio {a['id']} ({a.get('title','')[:40]}) — R2 key missing: {key}")
                    else:
                        # Has YT → just clear the broken r2_audio_key, keep the audio
                        await db.audios.update_one(
                            {'id': a['id']},
                            {'$unset': {'r2_audio_key': '', 'has_r2_audio': ''}},
                        )
                        logger.info(f"Migration v15: cleared dangling r2_audio_key on {a['id']} (kept for YT)")
            logger.info(f"Migration v15: purged {purged_total} orphan audio(s)")

        # Rename cours-historiographie to focus on Ibn Khaldūn (only character with R2 content).
        # IMPORTANT: respect `seed_locked` — once an admin has edited the title /
        # summary / description in the admin panel, this migration MUST NOT
        # overwrite their changes on subsequent reboots. We only push the
        # `r2_prefix` (structural, not editorial) unconditionally.
        await db.courses.update_one(
            {'id': 'cours-historiographie', 'seed_locked': {'$ne': True}},
            {'$set': {
                'title': 'Ibn Khaldūn — Historiographie',
                'summary': "La pensée historiographique d'Ibn Khaldūn et la Muqaddima.",
                'description': "La pensée historiographique d'Ibn Khaldūn et la Muqaddima : philosophie de l'histoire, théorie de la 'aṣabiyya, cycles dynastiques.",
            }},
        )
        # Structural field (R2 path) — always pushed.
        await db.courses.update_one(
            {'id': 'cours-historiographie'},
            {'$set': {'r2_prefix': 'cursus-c-sciences-islamiques/14-historiographie/ibn-khaldun-histoire/'}},
        )
        # Update the audio title to match — also respects per-audio seed_locked.
        await db.audios.update_one(
            {'course_id': 'cours-historiographie', 'episode_number': 1, 'seed_locked': {'$ne': True}},
            {'$set': {
                'title': "Ibn Khaldūn philosophe et historien",
                'intervenant': 'Mehdi Ghouirgate',
            }},
        )
        # R2 sub-prefix is structural — always pushed.
        await db.audios.update_one(
            {'course_id': 'cours-historiographie', 'episode_number': 1},
            {'$set': {'r2_subprefix': 'ibn-khaldun-histoire/'}},
        )
        logger.info("Migration v15: historiographie restructured around Ibn Khaldūn")
    except Exception as e:
        logger.warning(f"Migration v15 failed: {e}", exc_info=True)
    # ─── End Migration v15 ─────────────────────────────────────────────────

    # ─── Migration v15b — Historiographie: 1 cours = 1 personnage ──────────
    # Keep only the Ibn Khaldūn module (mod-2). Remove the 4 placeholder
    # modules (Ibn Baṭṭūṭa, Al-Maqrīzī, Al-Ṭabarī, Ibn Kathīr) since we don't
    # have delivery dates for those personalities yet. The course now exposes
    # "1 module · 1 épisode" — Ibn Khaldūn only.
    try:
        kept_module_id = 'cours-historiographie-mod-2'
        modules_to_drop = [
            'cours-historiographie-mod-1',
            'cours-historiographie-mod-3',
            'cours-historiographie-mod-4',
            'cours-historiographie-mod-5',
        ]
        # Detach any audio still attached to a dropped module (defensive).
        await db.audios.update_many(
            {'course_id': 'cours-historiographie', 'module_id': {'$in': modules_to_drop}},
            {'$set': {'module_id': kept_module_id}},
        )
        drop_res = await db.modules.delete_many({'id': {'$in': modules_to_drop}})
        # Ensure the surviving Ibn Khaldūn module is order=1 and properly named.
        await db.modules.update_one(
            {'id': kept_module_id},
            {'$set': {
                'name': 'Ibn Khaldūn (1332–1406)',
                'scholar_name': 'Mehdi Ghouirgate',
                'order': 1,
                'episode_count': 1,
            }},
        )
        # Sync the parent course's modules_count.
        await db.courses.update_one(
            {'id': 'cours-historiographie'},
            {'$set': {'modules_count': 1}},
        )
        logger.info(f"Migration v15b: dropped {drop_res.deleted_count} placeholder module(s) from cours-historiographie")
    except Exception as e:
        logger.warning(f"Migration v15b failed: {e}", exc_info=True)
    # ─── End Migration v15b ────────────────────────────────────────────────

    # ─── Migration v15c — Ensure debuts-islam has a module ─────────────────
    # The course `cours-debuts-islam` historically had `modules_count: 2` on
    # the course doc but ZERO entries in db.modules → public page reported
    # "0 module · 5 épisodes". Seed a single canonical module and re-attach
    # the orphan audios so the catalogue is consistent.
    try:
        debuts_module_id = 'cours-debuts-islam-mod-1'
        await db.modules.update_one(
            {'id': debuts_module_id},
            {'$set': {
                'id': debuts_module_id,
                'course_id': 'cours-debuts-islam',
                'name': "Les débuts de l'islam",
                'scholar_name': 'Hassan Bouali · Mehdi Ghouirgate',
                'order': 1,
                'is_active': True,
                'episode_count': 5,
                'notes': '',
                'created_at': datetime.now(timezone.utc),
            }},
            upsert=True,
        )
        # Attach orphan audios (module_id missing or null) to this module.
        attach_res = await db.audios.update_many(
            {'course_id': 'cours-debuts-islam', '$or': [
                {'module_id': None}, {'module_id': {'$exists': False}}, {'module_id': ''},
            ]},
            {'$set': {'module_id': debuts_module_id}},
        )
        logger.info(f"Migration v15c: ensured cours-debuts-islam module + attached {attach_res.modified_count} audio(s)")
    except Exception as e:
        logger.warning(f"Migration v15c failed: {e}", exc_info=True)
    # ─── End Migration v15c ────────────────────────────────────────────────

    # ─── Migration v15d — Apply Sijill_Catalogue_Editorial.docx (Feb 2026) ─
    # One-shot import of the editorial catalogue provided by the team. Sets
    # the official title / summary / description / scholar_name on each
    # course AND flips `seed_locked: True` so no future seed can overwrite.
    # Idempotent: gated by the dedicated `migration_v15d_applied` flag in
    # the `runtime_flags` collection so it runs exactly once per database.
    try:
        flag_doc = await db.runtime_flags.find_one({'_id': 'migration_v15d_applied'})
        if not flag_doc:
            EDITORIAL_CATALOG = [
                {
                    'id': 'cours-debuts-islam',
                    'title': "Cours 1 : Les débuts de l'islam",
                    'scholar_name': 'Hassan Bouali · Mehdi Ghouirgate',
                    'summary': "En quatre épisodes thématiques, Hassan Bouali approfondit les grandes questions des débuts de l'islam : la vie du Prophète, la formation du califat, les conquêtes et l'intégration progressive du ḥajj à la piété musulmane.",
                    'description': "Hassan Bouali, historien médiéviste et chercheur associé au CEFREPA (CNRS) et à l'Université Paris Nanterre, aborde les débuts de l'islam par quatre entrées thématiques complémentaires.\nIl commence par la figure de Muhammad et les conditions d'émergence de l'islam en Arabie au VIIe siècle, puis explore la formation du califat de Médine à la sortie d'Arabie.\nIl retrace ensuite les futūḥāt — les grandes conquêtes — avant de conclure par le ḥajj et son intégration progressive à la piété musulmane.\n\nCe cours couvre :\n• Muhammad et les débuts de l'islam (juillet 2026)\n• Le califat : de Médine à la sortie d'Arabie (sept. 2026)\n• Les futūḥāt (oct. 2026)\n• Le ḥajj (nov. 2026)",
                },
                {
                    'id': 'cours-andalus',
                    'title': "Cours 2 : Al-Andalus",
                    'scholar_name': 'Mehdi Ghouirgate',
                    'summary': "L'unique civilisation arabo-islamique durablement implantée sur le sol européen. Du VIIIe au XVe siècle, al-Andalus fut un espace d'échanges entre musulmans, chrétiens et juifs, un foyer de transmission du savoir et un carrefour intellectuel qui façonna l'histoire européenne.",
                    'description': "Al-Andalus représente l'une des expériences historiques les plus singulières du Moyen Âge.\nNée d'une conquête rapide en 711, consolidée par la dynastie omeyyade exilée de Damas, rayonnante sous le califat de Cordoue, puis fragmentée en royaumes de taifas avant de s'éteindre avec la chute de Grenade en 1492.\nMehdi Ghouirgate propose une lecture panoramique et rigoureuse de cette civilisation : la genèse, l'apogée avec Averroès et Maïmonide, le rôle de médiateur entre Orient et Occident, les empires berbères, l'Alhambra nasride et les débats contemporains sur l'héritage andalou.\n\nCe cours couvre :\n• La genèse d'al-Andalus\n• L'apogée : le califat de Cordoue\n• Al-Andalus comme médiateur Orient-Occident\n• Les empires berbères et l'Alhambra\n• L'héritage contesté",
                },
                {
                    'id': 'cours-historiographie',
                    'title': "Ibn Khaldūn — Biographie",
                    'scholar_name': 'Mehdi Ghouirgate',
                    'summary': "Né en 1332 dans une famille de la noblesse andalouse exilée, Ibn Khaldūn occupe dans l'histoire de la pensée une place absolument unique. Ni théologien, ni philosophe classique, ni simple chroniqueur — il invente une méthode entièrement nouvelle pour comprendre pourquoi les civilisations naissent, prospèrent et s'effondrent.",
                    'description': "Ibn Khaldūn est l'une des figures les plus originales de toute l'histoire intellectuelle mondiale.\nSa vie est d'abord celle d'un homme de pouvoir : secrétaire, diplomate, émissaire, prisonnier, puis conseiller de sultans au Maghreb, en al-Andalus et en Égypte.\nCes déplacements constants font de lui un observateur direct des mécanismes tribaux et des cycles de grandeur et de déclin des dynasties.\nC'est cette expérience de terrain qui nourrit sa notion fondamentale d'ʿasabiyya — la force de cohésion qui permet à un groupe de conquérir le pouvoir.\nMehdi Ghouirgate, lauréat du Prix de la biographie littéraire 2025 de l'Académie française pour son ouvrage Ibn Khaldūn (CNRS Éditions, 2024), retrace dans cet épisode son itinéraire exceptionnel.\n\nCe cours couvre :\n• Formation et noblesse andalouse\n• Carrière politique tumultueuse\n• L'ʿasabiyya\n• Grenade et l'Égypte\n• La Muqaddima",
                },
                {
                    'id': 'cours-fiqh',
                    'title': "Histoire du droit musulman",
                    'scholar_name': 'Yannis Mahil',
                    'summary': "Le droit musulman — fiqh et sharīʿa — est à la fois un système juridique, un fait social et un fait religieux vivant. En six épisodes, Yannis Mahil en retrace l'histoire, les fondements théoriques, les mécanismes d'élaboration et les dynamiques contemporaines.",
                    'description': "Qu'est-ce que le droit musulman ? Loin d'être un code figé, le fiqh constitue l'un des édifices intellectuels les plus complexes et les plus vivants de la civilisation islamique — un système de normes élaboré sur quatorze siècles.\nYannis Mahil, maître de conférences en études islamiques à GISTU University, propose une introduction rigoureuse et accessible.\nIl commence par les fondements conceptuels — fiqh, sharīʿa, ijtihad — avant d'en retracer l'histoire et d'explorer les uṣūl al-fiqh, cette théorie générale du droit.\nLe cours s'achève sur l'ijtihad contemporain et les maqāṣid al-sharīʿa.\n\nCe cours couvre :\n• Introduction au droit musulman\n• L'ijtihad\n• Histoire du droit : périodisation\n• Uṣūl al-fiqh I : sources primaires\n• Uṣūl al-fiqh II : qiyās et sources secondaires\n• Ijtihad contemporain et maqāṣid",
                },
                {
                    'id': 'cours-coran',
                    'title': "Histoire de la constitution du Coran",
                    'scholar_name': 'Mehdi Azaiez',
                    'summary': "Comment le texte coranique s'est-il constitué ? Mehdi Azaiez, professeur d'islamologie à l'UCLouvain et spécialiste des études coraniques, retrace l'histoire de la formation du Coran.",
                    'description': "Mehdi Azaiez, professeur d'islamologie à l'Université catholique de Louvain et auteur du Contre-discours coranique (De Gruyter, 2015), propose une introduction aux études coraniques et à l'histoire de la constitution du texte coranique.\nCe cours aborde la question de la formation du corpus coranique, les débats historiographiques et les méthodes critiques contemporaines.",
                },
                {
                    'id': 'cours-hadith',
                    'title': "Le corpus du hadith",
                    'scholar_name': 'Hassan Chahdi',
                    'summary': "Comment les paroles et actes du Prophète ont-ils été collectés, transmis et codifiés ? Hassan Chahdi retrace l'histoire du corpus du hadith en trois épisodes.",
                    'description': "Hassan Chahdi est maître de conférences en islamologie à l'Université de Lorraine, docteur de l'EPHE et lauréat du prix de la meilleure thèse francophone (GIS-Moyen Orient).\nEn trois épisodes, il retrace l'histoire de la transmission du hadith — les paroles et actes du Prophète Muhammad — depuis les premières transmissions orales jusqu'à la codification des grands recueils canoniques.\nIl aborde les méthodes critiques d'authentification et leur rôle dans l'élaboration du droit islamique.",
                },
                {
                    'id': 'cours-sciences',
                    'title': "Histoire des sciences arabes",
                    'scholar_name': 'Marouane Ben Miled · Meyssa Ben Saâd',
                    'summary': "Comment les savants arabes médiévaux ont-ils pensé et transformé les sciences héritées de l'Antiquité ? Marouane Ben Miled retrace l'histoire des mathématiques arabes et Meyssa Ben Saâd celle des sciences naturelles à travers l'œuvre d'Al-Jāḥiẓ.",
                    'description': "Marouane Ben Miled, enseignant-chercheur à l'École Nationale d'Ingénieurs de Tunis (ENIT), chercheur au LAMSIN et associé au CGGG (CNRS), retrace en trois épisodes l'histoire des mathématiques arabes — algèbre, géométrie, arithmétique, trigonométrie — en montrant comment les savants arabes ont non seulement transmis le savoir grec mais l'ont profondément transformé, créant de nouvelles disciplines dont hérita l'Europe moderne.\n\nMeyssa Ben Saâd, docteure en histoire des sciences spécialisée en histoire de la zoologie arabe médiévale, chercheuse associée à l'UMR SPHERE (CNRS, Université Paris Cité) et membre de la Faculté de l'Institut Supérieur d'Éducation Spécialisée de l'Université de la Manouba (Tunis), consacre deux épisodes à l'histoire des sciences naturelles dans le monde islamique médiéval et à l'œuvre d'Al-Jāḥiẓ — penseur du IXe siècle auteur du Livre des animaux, l'une des premières tentatives systématiques de classification du monde animal dans la tradition arabe.",
                },
                {
                    'id': 'cours-art',
                    'title': "Les arts de l'Islam",
                    'scholar_name': 'Camille Grandpierre',
                    'summary': "Des mosquées de Cordoue aux manuscrits enluminés de Hérat — les arts de l'Islam forment l'un des ensembles artistiques les plus vastes et les plus divers de l'histoire humaine. Camille Grandpierre en propose une introduction rigoureuse et illustrée.",
                    'description': "Les arts de l'Islam ne forment pas un ensemble homogène.\nIls s'étendent sur quatorze siècles et un territoire immense, portés par des dynasties, des langues et des traditions artistiques radicalement différentes.\nCamille Grandpierre, doctorante à Sorbonne Université et lauréate d'un contrat doctoral de l'INHA, propose en quatre épisodes une introduction rigoureuse et illustrée.\nElle commence par clarifier les contours de la discipline — qu'est-ce que les arts de l'Islam ? Sont-ils vraiment aniconiques ? — avant d'aborder l'architecture, les arts des objets (céramique, verre, métal, bois, ivoire, textiles), et enfin les arts du livre et de la calligraphie.\n\nCe cours couvre :\n• Introduction aux arts de l'Islam\n• L'architecture : mosquées et espaces du sacré\n• Les objets de forme : céramique, verre, métal, textiles\n• Les arts du livre : calligraphie, enluminure, manuscrits",
                },
                {
                    'id': 'cours-traduction',
                    'title': "Le grand mouvement de traduction",
                    'scholar_name': 'Meryem Sebti',
                    'summary': "Bagdad, IXe siècle. Des textes grecs vieux de mille ans traversent les frontières, changent de langue, et renaissent dans un univers culturel radicalement différent. Meryem Sebti retrace le moment fondateur de la philosophie en terre d'islam.",
                    'description': "Comment la philosophie grecque est-elle parvenue au monde islamique ? La réponse est bien plus complexe qu'un simple transfert.\nLa philosophie qui arrive en arabe n'est pas celle de Platon et d'Aristote tels qu'ils enseignaient au IVe siècle avant notre ère — c'est une pensée déjà transformée par six siècles de commentaires dans les grandes écoles d'Alexandrie et d'Athènes.\nMeryem Sebti retrace les étapes de ce transfert exceptionnel : le rôle fondateur d'Alexandrie, le maillon décisif du syriaque, le mécénat abbasside sous al-Maʾmūn et les grandes figures de traducteurs comme Ḥunayn ibn Isḥāq.\nCe mouvement n'est pas une réception passive — il est une transformation active qui donne naissance à la falsafa.\n\nCe cours couvre :\n• La philosophie de l'Antiquité tardive\n• Le maillon syriaque\n• Bagdad abbasside et la Maison de la Sagesse\n• Naissance de la falsafa",
                },
                {
                    'id': 'cours-al-ghazali',
                    'title': "Introduction à Al-Ghazālī",
                    'scholar_name': 'Meryem Sebti',
                    'summary': "Al-Ghazālī occupe une place singulière dans l'histoire de la pensée islamique : à la fois philosophe, théologien et mystique, il est celui qui formula la critique la plus radicale de la falsafa tout en en étant profondément nourri. Meryem Sebti présente cette figure essentielle.",
                    'description': "Al-Ghazālī (1058-1111) est l'une des figures les plus complexes et les plus influentes de la pensée islamique.\nAuteur de L'Incohérence des philosophes, il se distingue par une critique radicale de la falsafa aristotélicienne, tout en ayant lui-même une formation philosophique poussée.\nMeryem Sebti, directrice de recherche au CNRS, propose dans cet épisode une introduction à sa pensée et à sa place dans l'histoire de la philosophie islamique.",
                },
                {
                    'id': 'cours-al-kindi',
                    'title': "Al-Kindī",
                    'scholar_name': 'Meryem Sebti',
                    'summary': "Premier philosophe de la tradition arabe, al-Kindī incarne le moment fondateur de la falsafa. Meryem Sebti lui consacre cinq épisodes explorant sa vie, sa défense de la philosophie, sa cosmologie, sa théorie de l'intellect et son éthique.",
                    'description': "Al-Kindī occupe dans la tradition philosophique arabe une position rigoureusement unique : il en fut le premier penseur.\nNé vers 800 à Bagdad, il traverse l'un des siècles les plus fertiles de l'histoire intellectuelle.\nEn cinq épisodes, Meryem Sebti propose une lecture complète de sa pensée : vie et contexte, défense de la falsafa, cosmologie et sciences, intellect et connaissance, éthique et héritage kindien.",
                },
                {
                    'id': 'cours-al-farabi',
                    'title': "Al-Fārābī",
                    'scholar_name': 'Meryem Sebti',
                    'summary': "Al-Fārābī — le « Deuxième Maître » après Aristote — est l'architecte de la philosophie politique islamique et le grand systématisateur de la falsafa. Meryem Sebti lui consacre cinq épisodes.",
                    'description': "Al-Fārābī (872-950) est l'un des penseurs les plus originaux de la tradition philosophique islamique.\nLogicien, métaphysicien et philosophe politique, il est le premier à avoir systématiquement intégré la philosophie d'Aristote dans un cadre islamique cohérent.\nEn cinq épisodes, Meryem Sebti explore l'ensemble de sa pensée : logique, métaphysique, philosophie politique (La Cité vertueuse) et philosophie de la musique.",
                },
                {
                    'id': 'cours-avicenne',
                    'title': "Avicenne (Ibn Sīnā)",
                    'scholar_name': 'Meryem Sebti',
                    'summary': "« Chef et prince des philosophes » — c'est ainsi que Roger Bacon désignait Avicenne. Le philosophe le plus important du monde islamique, dont l'œuvre marque un tournant majeur : il y a un avant et un après Avicenne dans l'histoire de la pensée.",
                    'description': "Avicenne (980-1037) aspirait à créer une synthèse philosophique capable de se substituer à celle d'Aristote. Il y parvint.\nSon Livre de la guérison est l'une des œuvres les plus monumentales de toute la philosophie mondiale.\nMeryem Sebti lui consacre cinq épisodes : vie et contexte, métaphysique et théologie, psychologie et théorie de l'intellect, philosophie de la nature, héritage avicennien en Orient et en Occident.",
                },
                {
                    'id': 'cours-falsafa-occident',
                    'title': "La falsafa en Occident musulman",
                    'scholar_name': 'Yassir Mechelloukh',
                    'summary': "De Saragosse à Cordoue, la falsafa s'épanouit dans l'Occident musulman avec trois figures majeures. Yassir Mechelloukh explore en sept épisodes cette tradition philosophique andalouse qui influencera durablement la scolastique latine.",
                    'description': "Yassir Mechelloukh, doctorant en histoire de la philosophie arabe à l'INALCO et à l'Université Paris 1, explore en sept épisodes la tradition philosophique de l'Occident islamique.\nIl commence par Ibn Bajja (Avenpace) et sa réflexion sur l'intellect et le sage solitaire, poursuit avec Ibn Ṭufayl et son roman philosophique Hayy ibn Yaqẓān, avant de consacrer trois épisodes à Averroès (Ibn Rushd) — le Grand Commentateur d'Aristote dont les traductions latines transformèrent la philosophie médiévale européenne.",
                },
                {
                    'id': 'cours-inclassables',
                    'title': "La pensée d'Ibn Khaldūn",
                    'scholar_name': 'Cédric Molino-Machetto',
                    'summary': "Comment Ibn Khaldūn pense-t-il l'histoire des civilisations ? Cédric Molino-Machetto explore en deux épisodes la Muqaddima — l'œuvre philosophique la plus originale du monde islamique médiéval.",
                    'description': "Cédric Molino-Machetto, agrégé et docteur en philosophie, spécialiste de philosophie arabe prémoderne, consacre deux épisodes à la pensée philosophique d'Ibn Khaldūn.\nIl explore la Muqaddima comme œuvre philosophique : la théorie de l'ʿasabiyya, la science de la civilisation humaine (ʿilm al-ʿumrān), la philosophie de l'histoire et la place d'Ibn Khaldūn dans la tradition de la falsafa.",
                },
                {
                    'id': 'cours-falsafa-persan',
                    'title': "Le renouveau de la philosophie persane : Mullā Ṣadrā",
                    'scholar_name': 'Sajjad H. Rizvi',
                    'summary': "Mullā Ṣadrā (1571-1640) est la figure centrale du renouveau de la philosophie islamique en Iran safavide. Sajjad Rizvi, professeur à l'Université d'Exeter, lui consacre trois épisodes.",
                    'description': "Sajjad H. Rizvi, Professor of Islamic Intellectual History à l'Université d'Exeter et directeur du Global & Area Studies, est l'un des spécialistes mondiaux de Mullā Ṣadrā.\nEn trois épisodes, il explore la « sagesse transcendante » (al-ḥikma al-mutaʿāliya) de Mullā Ṣadrā — une synthèse extraordinaire entre falsafa avicennienne, soufisme d'Ibn ʿArabī et théologie shiite — et son influence durable sur la philosophie islamique jusqu'à nos jours.",
                },
                {
                    'id': 'cours-philo-juive',
                    'title': "La philosophie de Maïmonide",
                    'scholar_name': 'Géraldine Roux',
                    'summary': "Philosophe juif du XIIe siècle écrivant en arabe, nourri de la pensée grecque transmise par la tradition islamique — Maïmonide se tient à la croisée de trois traditions intellectuelles. Géraldine Roux retrace son itinéraire et explore le doute comme méthode philosophique dans le Guide des Égarés.",
                    'description': "Maïmonide (1138-1204) occupe dans l'histoire de la pensée une place singulière : né à Cordoue, formé dans un monde arabo-islamique, il écrit son œuvre majeure en arabe tout en pensant dans un cadre juif.\nGéraldine Roux, agrégée et docteure en philosophie, autrice de Maïmonide ou la nostalgie de la sagesse (Seuil, 2017), retrace sa biographie et montre comment la philosophie gréco-arabe irrigue l'élaboration conceptuelle du Guide des Égarés.\nElle explore ensuite la notion de « perplexité » comme méthode philosophique à part entière, puis aborde les deux piliers de la pensée maïmonidienne : l'interprétation allégorique des Écritures et la définition de la croyance.\nMaïmonide distingue les croyances vraies, fondées sur la démonstration rationnelle, des croyances fausses ou simplement utiles à l'ordre social. Cette théorie influencera Spinoza, les philosophes des Lumières et continue d'alimenter les débats contemporains sur les rapports entre raison et révélation.\n\nCe cours couvre :\n• Biographie\n• Influence de la philosophie gréco-arabe\n• Le chemin des manuscrits grecs\n• Le doute et la perplexité comme méthode\n• Interprétation allégorique et philosophique\n• Croyances vraies et croyances fausses\n• La postérité de Maïmonide",
                },
            ]
            applied = 0
            for entry in EDITORIAL_CATALOG:
                cid = entry.pop('id')
                # seed_locked=True freezes the course against any future seed sweep.
                entry['seed_locked'] = True
                r = await db.courses.update_one({'id': cid}, {'$set': entry})
                if r.matched_count:
                    applied += 1
                else:
                    logger.warning(f"Migration v15d: course '{cid}' not found, skipped")
            await db.runtime_flags.insert_one({
                '_id': 'migration_v15d_applied',
                'applied_at': datetime.now(timezone.utc),
                'courses_updated': applied,
            })
            logger.info(f"Migration v15d: editorial catalogue applied to {applied} course(s)")
    except Exception as e:
        logger.warning(f"Migration v15d failed: {e}", exc_info=True)
    # ─── End Migration v15d ────────────────────────────────────────────────

    # ─── Migration v15e — Statut Excel (Lancement + Vision cible) ──────────
    # Aligne la base sur le fichier Excel Sijill_Catalogue_Lancement_Mai2026.
    # Réutilise le mécanisme existant `coming_soon` + `available_date` et
    # ajoute `recruiting=True` pour les cours dont l'intervenant est encore
    # à recruter (affichage grisé côté front).
    # NON-gated: les champs structurels (statut + cursus + ordre +
    # is_launch_catalog) sont réappliqués à chaque démarrage pour
    # contrecarrer les seeds antérieurs (v4) qui les écrasent. Les titres /
    # summary / description / scholar_name ne sont posés QU'À L'INSERTION
    # initiale (puis verrouillés par seed_locked dès qu'un admin édite).
    try:
        # Tuple: (course_id, cursus_id, order, is_launch_catalog, coming_soon,
        #        available_date, recruiting, default_title, default_scholar,
        #        default_summary)
        V15E_STATUS = [
            # ═══ Catalogue de lancement Mai 2026 (sheet 1 Excel) — is_launch_catalog=True ═══
            # A. Histoire
            ('cours-debuts-islam',       'cursus-histoire', 1, True, False, None,         False, None, None, None),
            ('cours-mamelouke',          'cursus-histoire', 2, True, True,  'mai 2026',   False, None, None, None),
            ('cours-andalus',            'cursus-histoire', 3, True, False, None,         False, None, None, None),
            ('cours-ottoman',            'cursus-histoire', 4, True, True,  'sept. 2026', False, None, None, None),
            # B. Théologie et Droit
            ('cours-kalam',              'cursus-theologie', 1, True, True, 'mai 2026',   False, None, None, None),
            ('cours-fiqh',               'cursus-theologie', 2, True, False, None,        False, None, None, None),
            # C. Sciences islamiques (launch)
            ('cours-coran',              'cursus-sciences-islamiques', 1, True, True, 'mai 2026', False, None, None, None),
            ('cours-hadith',             'cursus-sciences-islamiques', 2, True, True, 'mai 2026', False, None, None, None),
            ('cours-historiographie',    'cursus-sciences-islamiques', 3, True, False, None,      False, None, None, None),
            # D. Arts (launch)
            ('cours-art',                'cursus-arts', 1, True, False, None,        False, None, None, None),
            ('cours-maths-arabes',       'cursus-arts', 2, True, True,  'mai 2026',  False, 'Histoire des mathématiques arabes', 'Marouane Ben Miled', "Trois épisodes consacrés à l'histoire des mathématiques arabes."),
            ('cours-sciences-naturelles','cursus-arts', 3, True, True,  'mai 2026',  False, 'Histoire des sciences naturelles', 'Meyssa Ben Saad', "Introduction à l'histoire des sciences naturelles arabes — al-Jāḥiẓ et la classification des animaux."),
            # E. Falsafa (launch)
            ('cours-traduction',         'cursus-falsafa', 1, True, False, None,        False, None, None, None),
            ('cours-al-kindi',           'cursus-falsafa', 2, True, False, None,        False, None, None, None),
            ('cours-al-farabi',          'cursus-falsafa', 3, True, False, None,        False, None, None, None),
            ('cours-avicenne',           'cursus-falsafa', 4, True, False, None,        False, None, None, None),
            ('cours-al-ghazali',         'cursus-falsafa', 5, True, True,  'mai 2026', False, None, None, None),
            ('cours-falsafa-occident',   'cursus-falsafa', 6, True, True,  'mai 2026', False, None, None, None),
            ('cours-falsafa-inclassables','cursus-falsafa', 7, True, True, 'mai 2026', False, None, None, None),
            ('cours-falsafa-persane',    'cursus-falsafa', 8, True, True,  'mai 2026', False, None, None, None),
            # F. Mystique (launch)
            ('cours-soufisme',           'cursus-spiritualites', 1, True, True, 'mai 2026', False, None, None, None),
            # G. Pensées arabes non islamiques (launch)
            ('cours-philo-juive',        'cursus-pensees-non-islamiques', 1, True, False, None, False, None, None, None),

            # ═══ Vision cible (sheet 2 Excel) — is_launch_catalog=False (visible /cursus uniquement) ═══
            # C. Sciences islamiques (vision cible)
            ('cours-doxographie',        'cursus-sciences-islamiques', 4, False, True, 'prochainement', True, 'Histoire de la doxographie', 'Intervenant à recruter', "Histoire de la doxographie islamique : Ibn al-Nadīm, Ḥājjī Khalīfa, Ṭāshköprīzāde."),
            ('cours-autobiographies',    'cursus-sciences-islamiques', 5, False, True, 'prochainement', True, 'Les autobiographies dans le monde islamique', 'Intervenant à recruter', "Tradition des autobiographies dans la littérature arabo-islamique."),
            # D. Arts (vision cible)
            ('cours-poesie',             'cursus-arts', 4, False, True, 'date à fixer',  False, None, None, None),
            ('cours-urjuza',             'cursus-arts', 5, False, True, 'prochainement', True,  None, None, None),
            ('cours-geographie',         'cursus-arts', 6, False, True, 'prochainement', True,  None, None, None),
            ('cours-adab',               'cursus-arts', 7, False, True, 'date à fixer',  False, None, None, None),
            # E. Falsafa (vision cible)
            ('cours-post-avicennisme',   'cursus-falsafa', 9,  False, True, 'date à fixer', False, None, None, None),
            ('cours-logique',            'cursus-falsafa', 10, False, True, 'date à fixer', False, None, None, None),
            ('cours-ismaelisme',         'cursus-falsafa', 11, False, True, 'date à fixer', False, None, None, None),
            # G. Pensées arabes non islamiques (vision cible)
            ('cours-kalam-chretien',     'cursus-pensees-non-islamiques', 2, False, True, 'date à fixer', False, None, None, None),
        ]
        applied_v15e = 0
        inserted_v15e = 0
        for cid, cursus_id, order, is_launch, coming_soon, avail_date, recruiting, def_title, def_scholar, def_summary in V15E_STATUS:
            structural = {
                'cursus_id': cursus_id,
                'order': order,
                'is_launch_catalog': is_launch,
                'is_active': True,
                'coming_soon': coming_soon,
                'available_date': avail_date,
                'recruiting': recruiting,
            }
            existing = await db.courses.find_one({'id': cid}, {'_id': 0, 'id': 1})
            if existing:
                # Update ONLY structural status. Never touch admin-owned content
                # (title / summary / description / scholar_name).
                await db.courses.update_one({'id': cid}, {'$set': structural})
                applied_v15e += 1
            else:
                # Insert with default content (admin will refine later via panel).
                insert_doc = {
                    'id': cid,
                    'title': def_title or cid.replace('cours-', '').replace('-', ' ').title(),
                    'scholar_name': def_scholar or 'Intervenant à confirmer',
                    'summary': def_summary or '',
                    'description': def_summary or '',
                    'created_at': datetime.now(timezone.utc).isoformat(),
                    **structural,
                }
                await db.courses.insert_one(insert_doc)
                inserted_v15e += 1

        # Désactiver les cours obsolètes (doublons / umbrella) qui ne sont
        # plus dans le catalogue cible mais qui pourraient encore traîner.
        for obsolete_id in ('cours-sciences', 'cours-inclassables', 'cours-falsafa-persan', 'cours-falsafa-grands'):
            await db.courses.update_one(
                {'id': obsolete_id},
                {'$set': {'is_launch_catalog': False, 'is_active': False, 'coming_soon': False}}
            )

        # Cleanup résidu de test sur cours-andalus.available_date='TEST_date'.
        await db.courses.update_one(
            {'id': 'cours-andalus', 'available_date': 'TEST_date'},
            {'$set': {'available_date': None}}
        )

        logger.info(f"Migration v15e: status applied — {applied_v15e} updated, {inserted_v15e} inserted")
    except Exception as e:
        logger.warning(f"Migration v15e failed: {e}", exc_info=True)
    # ─── End Migration v15e ────────────────────────────────────────────────

    # ─── Migration v15f — Nettoyage titres « Cours N : » + seed_locked ─────
    # Beaucoup de cours ont conservé les titres legacy « Cours 8 : L'ismaélisme »
    # depuis un vieux seed. v15d n'avait nettoyé que 7 cours. v15f termine le
    # travail sur les ~15 cours restants ET pose seed_locked=True pour qu'aucun
    # seed futur ne puisse les écraser à nouveau au reboot.
    #
    # Sécurité : on n'écrase un titre QUE si :
    #   - le cours n'est pas déjà seed_locked (admin owner)
    #   - ET son titre actuel matche le pattern legacy `^Cours \d+\s*:`
    #     OU son titre est vide.
    # Sinon on respecte l'édition admin et on pose juste seed_locked=True.
    #
    # Non-gated (rejoue à chaque démarrage) pour rattraper progressivement
    # d'éventuels résidus, mais l'effet est nul si tout est déjà propre +
    # locké (idempotent par construction).
    try:
        import re as _re_v15f
        LEGACY_RE = _re_v15f.compile(r"^\s*Cours\s+\d+\s*:\s*", _re_v15f.IGNORECASE)
        V15F_TITLES = [
            # (course_id, clean_title)
            # ─── Cursus Histoire ───
            ('cours-debuts-islam',       "Les débuts de l'islam"),
            ('cours-mamelouke',          "L'époque mamelouke"),
            ('cours-andalus',            "Al-Andalus"),
            ('cours-ottoman',            "Le monde ottoman"),
            # ─── Cursus Théologie ───
            ('cours-kalam',              "Le Kalām — Trois périodes"),
            # ─── Cursus Sciences islamiques ───
            ('cours-doxographie',        "Histoire de la doxographie"),
            ('cours-autobiographies',    "Les autobiographies dans le monde islamique"),
            # ─── Cursus Arts ───
            ('cours-maths-arabes',       "Histoire des mathématiques arabes"),
            ('cours-sciences-naturelles',"Histoire des sciences naturelles"),
            ('cours-poesie',             "La poésie arabe et persane"),
            ('cours-urjuza',             "Les Urjūzā — Enseigner par la poésie"),
            ('cours-geographie',         "La géographie islamique"),
            ('cours-adab',               "Adab et sciences médicales"),
            # ─── Cursus Falsafa ───
            ('cours-post-avicennisme',   "Le post-avicennisme"),
            ('cours-logique',            "La logique arabe"),
            ('cours-ismaelisme',         "L'ismaélisme"),
            ('cours-falsafa-inclassables', "Les inclassables — Ibn Khaldūn"),
            ('cours-falsafa-persane',    "Le renouveau de la philosophie persane — Mullā Ṣadrā"),
            # ─── Cursus Mystique ───
            ('cours-soufisme',           "La mystique islamique — Soufisme"),
            # ─── Cursus Pensées non islamiques ───
            ('cours-kalam-chretien',     "Le Kalām chrétien et les logiciens de Bagdad"),
        ]
        cleaned_v15f = 0
        locked_only_v15f = 0
        for cid, clean_title in V15F_TITLES:
            doc = await db.courses.find_one({'id': cid}, {'_id': 0, 'title': 1, 'seed_locked': 1})
            if not doc:
                continue
            already_locked = bool(doc.get('seed_locked'))
            current_title = (doc.get('title') or '').strip()
            is_legacy = bool(LEGACY_RE.match(current_title)) or not current_title

            # IMPORTANT: even if seed_locked is True, if the title still matches
            # the legacy pattern « Cours N : ... » we force-override (this state
            # only happens because Migration v15d locked the wrong legacy text).
            if is_legacy:
                await db.courses.update_one(
                    {'id': cid},
                    {'$set': {'title': clean_title, 'seed_locked': True}},
                )
                cleaned_v15f += 1
            elif not already_locked:
                # Title looks custom (admin-edited or already clean) — just lock it.
                await db.courses.update_one(
                    {'id': cid},
                    {'$set': {'seed_locked': True}},
                )
                locked_only_v15f += 1
            # else: clean & already locked → nothing to do
        if cleaned_v15f or locked_only_v15f:
            logger.info(f"Migration v15f: title cleanup — {cleaned_v15f} legacy titles rewritten, {locked_only_v15f} clean titles locked")
    except Exception as e:
        logger.warning(f"Migration v15f failed: {e}", exc_info=True)
    # ─── End Migration v15f ────────────────────────────────────────────────

    # ─── Migration v15g — Titres d'épisodes depuis Excel + lock complet ────
    # Apporte les VRAIS titres d'épisodes du fichier Excel utilisateur
    # (Sijill_Catalogue_Lancement_Mai2026) sur la collection `audios`.
    # Force aussi seed_locked=True au niveau audio pour bloquer définitivement
    # toute future réécriture par un seed/migration.
    # Non-gated, idempotent : ne fait rien si tout est déjà clean + locked.
    # Force-override sur les titres génériques placeholder (regex), sinon
    # respect total d'une édition admin.
    try:
        import re as _re_v15g
        # Patterns considérés comme "génériques placeholder" → on peut écraser
        # même si seed_locked=True (legacy lock posé avec un mauvais titre).
        GENERIC_PATTERNS = [
            _re_v15g.compile(r"^Histoire de l'art islamique\s*—\s*Épisode\s+\d+\s*$"),
            _re_v15g.compile(r"^\s*Épisode\s+\d+\s*$"),
            _re_v15g.compile(r"^Cours\s+\d+\s*:.*Épisode\s+\d+\s*$"),
        ]
        # (course_id, episode_number, clean_title, scholar_name)
        # Tous les titres viennent du fichier Excel sheet 1 « Catalogue de
        # lancement Mai 2026 » fourni par l'utilisateur.
        EPISODE_TITLES = [
            # ─── A. Histoire ───
            ('cours-debuts-islam', 1, "Les débuts de l'islam",                                                     'Mehdi Ghouirgate'),
            ('cours-debuts-islam', 2, "Muhammad et les débuts de l'islam",                                         'Hassan Bouali'),
            ('cours-debuts-islam', 3, "Le califat au début de l'islam : de Médine à la sortie d'Arabie",            'Hassan Bouali'),
            ('cours-debuts-islam', 4, "Les futūḥāt / conquêtes",                                                    'Hassan Bouali'),
            ('cours-debuts-islam', 5, "Le ḥajj et son intégration progressive à la piété musulmane",                'Hassan Bouali'),
            ('cours-mamelouke',    1, "L'époque mamelouke — Épisode 1",                                             'Sami Benkherfallah'),
            ('cours-mamelouke',    2, "L'époque mamelouke — Épisode 2",                                             'Sami Benkherfallah'),
            ('cours-mamelouke',    3, "L'époque mamelouke — Épisode 3",                                             'Sami Benkherfallah'),
            ('cours-andalus',      1, "Al-Andalus : histoire d'une civilisation islamique sur continent européen",  'Mehdi Ghouirgate'),
            ('cours-ottoman',      1, "Le monde ottoman — Épisode 1",                                               'Aysu Saban'),
            ('cours-ottoman',      2, "Le monde ottoman — Épisode 2",                                               'Aysu Saban'),
            ('cours-ottoman',      3, "Le monde ottoman — Épisode 3",                                               'Aysu Saban'),
            # ─── B. Théologie et Droit ───
            ('cours-kalam',        1, "Histoire du Kalām — Épisode 1",                                              'Ilyas Harifi'),
            ('cours-kalam',        2, "Histoire du Kalām — Épisode 2",                                              'Ilyas Harifi'),
            ('cours-kalam',        3, "Histoire du Kalām — Épisode 3",                                              'Ilyas Harifi'),
            ('cours-kalam',        4, "Histoire du Kalām — Épisode 4",                                              'Ilyas Harifi'),
            ('cours-kalam',        5, "Histoire du Kalām — Épisode 5",                                              'Ilyas Harifi'),
            ('cours-fiqh',         1, "Introduction au droit musulman",                                             'Dr. Yannis Mahil'),
            ('cours-fiqh',         2, "Introduction au droit musulman — L'Ijtihad",                                 'Dr. Yannis Mahil'),
            ('cours-fiqh',         3, "Histoire du droit musulman : périodisation et historiographie",              'Dr. Yannis Mahil'),
            ('cours-fiqh',         4, "Uṣūl al-fiqh I : définition, histoire et sources primaires",                 'Dr. Yannis Mahil'),
            ('cours-fiqh',         5, "Uṣūl al-fiqh II : le qiyās et les sources méthodologiques secondaires",      'Dr. Yannis Mahil'),
            ('cours-fiqh',         6, "Ijtihad contemporain et Maqāṣid al-sharīʿa",                                 'Dr. Yannis Mahil'),
            # ─── C. Sciences islamiques ───
            ('cours-coran',        1, "Histoire de la constitution du Coran",                                       'Mehdi Azaiez'),
            ('cours-hadith',       1, "Le corpus du hadith — Épisode 1",                                            'Hassan Chahdi'),
            ('cours-hadith',       2, "Le corpus du hadith — Épisode 2",                                            'Hassan Chahdi'),
            ('cours-hadith',       3, "Le corpus du hadith — Épisode 3",                                            'Hassan Chahdi'),
            ('cours-historiographie', 1, "Ibn Khaldūn : éléments biographiques",                                    'Mehdi Ghouirgate'),
            # ─── D. Arts ───
            ('cours-maths-arabes', 1, "Histoire des mathématiques arabes — Épisode 1",                              'Marouane Ben Miled'),
            ('cours-maths-arabes', 2, "Histoire des mathématiques arabes — Épisode 2",                              'Marouane Ben Miled'),
            ('cours-maths-arabes', 3, "Histoire des mathématiques arabes — Épisode 3",                              'Marouane Ben Miled'),
            ('cours-sciences-naturelles', 1, "Introduction — Histoire des sciences naturelles",                     'Meyssa Ben Saad'),
            ('cours-sciences-naturelles', 2, "Al-Jāḥiẓ et la classification des animaux",                           'Meyssa Ben Saad'),
            ('cours-art',          1, "Introduction aux arts de l'Islam",                                           'Camille Grandpierre'),
            ('cours-art',          2, "L'architecture",                                                              'Camille Grandpierre'),
            ('cours-art',          3, "Les objets de forme",                                                         'Camille Grandpierre'),
            ('cours-art',          4, "Les arts du livre et la calligraphie",                                        'Camille Grandpierre'),
            # ─── E. Falsafa ───
            ('cours-traduction',   1, "Le grand mouvement de traduction : quand le savoir grec devint arabe",        'Meryem Sebti'),
            ('cours-al-kindi',     1, "Al-Kindī : vie, contexte et héritage",                                        'Meryem Sebti'),
            ('cours-al-kindi',     2, "Pourquoi philosopher ? Al-Kindī et la défense de la falsafa",                 'Meryem Sebti'),
            ('cours-al-kindi',     3, "Cosmologie et sciences chez al-Kindī",                                        'Meryem Sebti'),
            ('cours-al-kindi',     4, "L'intellect et la connaissance",                                              'Meryem Sebti'),
            ('cours-al-kindi',     5, "Éthique et thérapie de l'âme — Héritage et tradition kindienne",              'Meryem Sebti'),
            ('cours-al-farabi',    1, "Al-Fārābī — Vie, contexte et héritage",                                       'Meryem Sebti'),
            ('cours-al-farabi',    2, "La place de la logique dans la pensée d'al-Fārābī",                           'Meryem Sebti'),
            ('cours-al-farabi',    3, "Métaphysique de Fārābī",                                                      'Meryem Sebti'),
            ('cours-al-farabi',    4, "Al-Fārābī — Épisode 4",                                                       'Meryem Sebti'),
            ('cours-al-farabi',    5, "Al-Fārābī — Épisode 5",                                                       'Meryem Sebti'),
            ('cours-avicenne',     1, "Avicenne : une vie entre savoir, pouvoir et errance",                         'Meryem Sebti'),
            ('cours-avicenne',     2, "Avicenne — Épisode 2",                                                        'Meryem Sebti'),
            ('cours-avicenne',     3, "Avicenne — Épisode 3",                                                        'Meryem Sebti'),
            ('cours-avicenne',     4, "Avicenne — Épisode 4",                                                        'Meryem Sebti'),
            ('cours-avicenne',     5, "Avicenne — Épisode 5",                                                        'Meryem Sebti'),
            ('cours-al-ghazali',   1, "Al-Ghazālī — Épisode 1",                                                      'Meryem Sebti'),
            ('cours-falsafa-occident',    1, "La pensée d'Ibn Bajja — Épisode 1",                                    'Yassir Mechelloukh'),
            ('cours-falsafa-occident',    2, "La pensée d'Ibn Bajja — Épisode 2",                                    'Yassir Mechelloukh'),
            ('cours-falsafa-occident',    3, "La pensée d'Ibn Ṭufayl — Épisode 1",                                   'Yassir Mechelloukh'),
            ('cours-falsafa-occident',    4, "La pensée d'Ibn Ṭufayl — Épisode 2",                                   'Yassir Mechelloukh'),
            ('cours-falsafa-occident',    5, "La pensée d'Averroès — Épisode 1",                                     'Yassir Mechelloukh'),
            ('cours-falsafa-occident',    6, "La pensée d'Averroès — Épisode 2",                                     'Yassir Mechelloukh'),
            ('cours-falsafa-occident',    7, "La pensée d'Averroès — Épisode 3",                                     'Yassir Mechelloukh'),
            ('cours-falsafa-inclassables', 1, "La pensée d'Ibn Khaldūn — Épisode 1",                                 'Cédric Molino Mochetto'),
            ('cours-falsafa-inclassables', 2, "La pensée d'Ibn Khaldūn — Épisode 2",                                 'Cédric Molino Mochetto'),
            ('cours-falsafa-persane',     1, "La pensée de Mullā Ṣadrā — Épisode 1",                                 'Sajjad Rizvi'),
            ('cours-falsafa-persane',     2, "La pensée de Mullā Ṣadrā — Épisode 2",                                 'Sajjad Rizvi'),
            ('cours-falsafa-persane',     3, "La pensée de Mullā Ṣadrā — Épisode 3",                                 'Sajjad Rizvi'),
            # ─── F. Mystique ───
            ('cours-soufisme',     1, "Les débuts du soufisme",                                                      'Gregory Vandamme'),
            ('cours-soufisme',     2, "La pensée d'Ibn ʿArabī — Épisode 1",                                          'Gregory Vandamme'),
            ('cours-soufisme',     3, "La pensée d'Ibn ʿArabī — Épisode 2",                                          'Gregory Vandamme'),
            ('cours-soufisme',     4, "La pensée d'Ibn ʿArabī — Épisode 3",                                          'Gregory Vandamme'),
            # ─── G. Pensées non islamiques ───
            ('cours-philo-juive',  1, "Introduction à la philosophie de Maïmonide : le doute comme méthode",         'Géraldine Roux'),
            ('cours-philo-juive',  2, "Allégorie, croyance et postérité de Maïmonide",                               'Géraldine Roux'),
        ]
        rewritten_v15g = 0
        locked_only_v15g = 0
        inserted_v15g = 0
        for cid, ep_num, clean_title, scholar in EPISODE_TITLES:
            doc = await db.audios.find_one(
                {'course_id': cid, 'episode_number': ep_num},
                {'_id': 0, 'id': 1, 'title': 1, 'seed_locked': 1},
            )
            if not doc:
                # Episode missing in DB → create it from Excel (placeholder audio).
                new_id = f"aud_{cid}-ep{ep_num:02d}"
                await db.audios.insert_one({
                    'id': new_id,
                    'course_id': cid,
                    'episode_number': ep_num,
                    'title': clean_title,
                    'scholar_name': scholar,
                    'intervenant': scholar,
                    'is_active': True,
                    'is_placeholder': True,
                    'published_at': '2026-05',
                    'seed_locked': True,
                    'created_at': datetime.now(timezone.utc),
                })
                inserted_v15g += 1
                continue
            current = (doc.get('title') or '').strip()
            is_generic = any(p.match(current) for p in GENERIC_PATTERNS) or not current
            already_locked = bool(doc.get('seed_locked'))

            if current == clean_title and already_locked:
                continue  # nothing to do
            if is_generic or current != clean_title and not already_locked:
                # Set the clean title + scholar + lock
                await db.audios.update_one(
                    {'id': doc['id']},
                    {'$set': {'title': clean_title, 'scholar_name': scholar, 'seed_locked': True}},
                )
                rewritten_v15g += 1
            elif not already_locked:
                # Title is custom (admin edit) but not locked yet → just lock
                await db.audios.update_one(
                    {'id': doc['id']},
                    {'$set': {'seed_locked': True}},
                )
                locked_only_v15g += 1

        # Cleanup: remove stray cours-debuts-islam ep6 (legacy duplicate of ep5)
        stray = await db.audios.find_one({
            'course_id': 'cours-debuts-islam',
            'episode_number': 6,
            'seed_locked': {'$ne': True},
        }, {'_id': 0, 'id': 1, 'title': 1})
        if stray:
            await db.audios.delete_one({'id': stray['id']})
            inserted_v15g += 0  # cosmetic
            logger.info(f"Migration v15g: removed stray duplicate {stray['id']}")

        if rewritten_v15g or locked_only_v15g or inserted_v15g:
            logger.info(f"Migration v15g: episode titles — {rewritten_v15g} rewritten, {locked_only_v15g} locked, {inserted_v15g} created (stray_removed: {1 if stray else 0})")
    except Exception as e:
        logger.warning(f"Migration v15g failed: {e}", exc_info=True)
    # ─── End Migration v15g ────────────────────────────────────────────────


    # ─── Migration v15h — Cleanup doublons Ibn Khaldūn + falsafa-grands ─────
    # Applique, de façon idempotente, les corrections manuelles validées par
    # l'utilisatrice sur la preview (fév. 2026) :
    #
    #   1. Supprime `cours-inclassables` (doublon obsolète + 5 modules orphelins)
    #   2. Supprime `cours-falsafa-grands` (« Cours 2 : Falsafa — Les grands
    #      philosophes ») + ses 7 audios sans fichier R2 + ses 11 modules
    #   3. Supprime `cours-falsafa-persan` (variante orthographique singulier
    #      sans « e » — le bon ID est `cours-falsafa-persane`)
    #   4. Renomme `cours-falsafa-inclassables` → « Les inclassables » (sans
    #      le suffixe parasite « (Ibn Khaldūn) ») + scholar Cédric Molino-
    #      Mochetto + seed_locked=True. Ne touche pas si admin a déjà reverrouillé.
    #   5. Met les 2 épisodes « La pensée d'Ibn Khaldūn — Épisode 1/2 » en
    #      `published_at='2026-07'` + seed_locked=True (juillet 2026).
    #
    # Idempotent : chaque action peut être ré-exécutée sans effet de bord.
    try:
        v15h_summary = {'deleted_courses': 0, 'deleted_audios': 0, 'deleted_modules': 0,
                        'renamed': 0, 'episodes_dated': 0}
        # 1) cours-inclassables — doublon obsolète
        a = await db.audios.delete_many({'course_id': 'cours-inclassables'})
        m = await db.modules.delete_many({'course_id': 'cours-inclassables'})
        c = await db.courses.delete_one({'id': 'cours-inclassables'})
        v15h_summary['deleted_courses'] += c.deleted_count
        v15h_summary['deleted_audios'] += a.deleted_count
        v15h_summary['deleted_modules'] += m.deleted_count

        # 2) cours-falsafa-grands — umbrella course remplacé par cours dédiés
        a = await db.audios.delete_many({'course_id': 'cours-falsafa-grands'})
        m = await db.modules.delete_many({'course_id': 'cours-falsafa-grands'})
        c = await db.courses.delete_one({'id': 'cours-falsafa-grands'})
        v15h_summary['deleted_courses'] += c.deleted_count
        v15h_summary['deleted_audios'] += a.deleted_count
        v15h_summary['deleted_modules'] += m.deleted_count

        # 3) cours-falsafa-persan — variante orthographique (le bon est ...persane)
        a = await db.audios.delete_many({'course_id': 'cours-falsafa-persan'})
        m = await db.modules.delete_many({'course_id': 'cours-falsafa-persan'})
        c = await db.courses.delete_one({'id': 'cours-falsafa-persan'})
        v15h_summary['deleted_courses'] += c.deleted_count
        v15h_summary['deleted_audios'] += a.deleted_count
        v15h_summary['deleted_modules'] += m.deleted_count

        # 4) Renommer cours-falsafa-inclassables. On NE force que si pas déjà
        # renommé (le titre legacy contient « (Ibn Khaldūn) ») → reste
        # idempotent et respectueux d'éventuelles éditions admin futures.
        existing = await db.courses.find_one(
            {'id': 'cours-falsafa-inclassables'},
            {'_id': 0, 'title': 1, 'seed_locked': 1}
        )
        if existing and '(Ibn Khaldūn)' in (existing.get('title') or ''):
            rr = await db.courses.update_one(
                {'id': 'cours-falsafa-inclassables'},
                {'$set': {
                    'title': 'Les inclassables',
                    'summary': "La pensée d'Ibn Khaldūn entre histoire, philosophie et sociologie.",
                    'description': "Cédric Molino-Mochetto, agrégé et docteur en philosophie, spécialiste de philosophie arabe prémoderne, consacre deux épisodes à Ibn Khaldūn — figure inclassable entre histoire, philosophie et sociologie.",
                    'scholar_id': 'sch-molino',  # Cédric Molino Mochetto (overrides legacy sch-sebti)
                    'scholar_name': 'Cédric Molino-Mochetto',
                    'co_scholar_ids': [],
                    'seed_locked': True,
                }},
            )
            v15h_summary['renamed'] += rr.modified_count
        # Always make sure the scholar pointer is correct (defensive — even if title
        # was already cleaned). The catalogue's _scholar_label reads scholar_id first.
        await db.courses.update_one(
            {'id': 'cours-falsafa-inclassables', 'scholar_id': {'$ne': 'sch-molino'}},
            {'$set': {'scholar_id': 'sch-molino', 'scholar_name': 'Cédric Molino-Mochetto', 'co_scholar_ids': []}},
        )

        # 5) Épisodes de cours-falsafa-inclassables → juillet 2026 + verrou
        # On évite de re-écraser un `published_at` déjà différent de 2026-05
        # (le default seedé). Si l'admin a déjà changé en autre chose, on
        # respecte cette valeur.
        ep_res = await db.audios.update_many(
            {'course_id': 'cours-falsafa-inclassables',
             '$or': [{'published_at': '2026-05'}, {'published_at': None}, {'published_at': {'$exists': False}}]},
            {'$set': {
                'published_at': '2026-07',
                'intervenant': 'Cédric Molino-Mochetto',
                'scholar_name': 'Cédric Molino-Mochetto',
                'seed_locked': True,
            }},
        )
        v15h_summary['episodes_dated'] += ep_res.modified_count

        # 6) cours-falsafa-persane — intervenant principal = Sajjad H. Rizvi
        #    (legacy seed l'attribuait à tort à Meryem Sebti). On ne force que
        #    si le scholar_id n'a pas déjà été corrigé.
        persane_fix = await db.courses.update_one(
            {'id': 'cours-falsafa-persane', 'scholar_id': {'$ne': 'sch-rizvi'}},
            {'$set': {
                'scholar_id': 'sch-rizvi',
                'scholar_name': 'Sajjad H. Rizvi',
                'co_scholar_ids': [],
                'seed_locked': True,
            }},
        )
        if persane_fix.modified_count:
            v15h_summary['persane_scholar_fixed'] = persane_fix.modified_count

        # 7) Nettoyage des registrations d'episode_resources mal taggées :
        # les PDF/JPG/PNG enregistrés avec `type='script'` (legacy seed Al-Kindi
        # `kindi_manuscript_1.pdf`) doivent disparaître pour que l'auto-detect
        # de `/courses/{id}/resources` les ré-émette avec le bon `type='manuscript'`
        # ou `type='image'` et un label lisible. Idempotent : MongoDB `$pull` ne
        # fait rien si aucun élément ne matche.
        mistagged = await db.audios.update_many(
            {'episode_resources': {'$elemMatch': {
                'type': 'script',
                'r2_key': {'$regex': r'\.(pdf|jpe?g|png)$', '$options': 'i'},
            }}},
            {'$pull': {'episode_resources': {
                'type': 'script',
                'r2_key': {'$regex': r'\.(pdf|jpe?g|png)$', '$options': 'i'},
            }}},
        )
        if mistagged.modified_count:
            v15h_summary['mistagged_resources_pulled'] = mistagged.modified_count

        # 8) r2_prefix manquants sur cours-mamelouke / cours-ottoman → empêchait
        # leurs frises (map_mamelouks.html, sijill_map_ottoman.html) d'apparaître
        # dans l'onglet Frise. Idempotent : on ne touche que si null/vide.
        r2_prefix_fixes = [
            ('cours-mamelouke', 'cursus-histoire/mamelouke/'),
            ('cours-ottoman', 'cursus-histoire/ottomane/'),
        ]
        rp_fixed = 0
        for cid, prefix in r2_prefix_fixes:
            r = await db.courses.update_one(
                {'id': cid, '$or': [{'r2_prefix': None}, {'r2_prefix': ''}, {'r2_prefix': {'$exists': False}}]},
                {'$set': {'r2_prefix': prefix}},
            )
            rp_fixed += r.modified_count
        if rp_fixed:
            v15h_summary['r2_prefix_set'] = rp_fixed

        logger.info(f"Migration v15h: cleanup Ibn Khaldūn & falsafa-grands — {v15h_summary}")
    except Exception as e:
        logger.warning(f"Migration v15h failed: {e}", exc_info=True)
    # ─── End Migration v15h ────────────────────────────────────────────────



    # ─── Migration v10: Seed all scholars from Excel + link courses ────────
    SCHOLARS_SEED = [
        ('sch-bouali', 'Hassan Bouali', 'Docteur en histoire médiévale',
         "Hassan Bouali est historien médiéviste islamisant, docteur en histoire médiévale de l'Université Paris Nanterre. Chercheur associé au Centre français de Recherche de la Péninsule Arabique (CEFREPA, CNRS) et au Centre d'Histoire des Sociétés Médiévales et Modernes (Mémo), ses travaux portent sur les débuts de l'islam et l'histoire politique et religieuse des premiers siècles de l'ère islamique."),
        ('sch-ghouirgate', 'Mehdi Ghouirgate', "Professeur d'histoire médiévale",
         "Mehdi Ghouirgate est professeur à l'Université Bordeaux-Montaigne, titulaire d'une Habilitation à Diriger les Recherches obtenue à l'EHESS en 2019. Spécialiste de l'histoire médiévale du Maghreb et d'al-Andalus, il est l'auteur de nombreux ouvrages et articles publiés dans des revues internationales. Son ouvrage « Ibn Khaldūn : itinéraire d'un penseur maghrébin » (CNRS Éditions, 2024) a reçu le Prix de la biographie littéraire 2025 de l'Académie française."),
        ('sch-benkherfallah', 'Sami Benkherfallah', None,
         "Spécialiste de l'époque mamelouke. Biographie complète en cours de confirmation."),
        ('sch-saban', 'Aysu Saban', 'Doctorante contractuelle en histoire et philologie',
         "Aysu Saban est doctorante contractuelle à l'EPHE/PSL sous la direction d'Alexandre Papas (CETOBaC) et de Nicolas Michel (IREMAM). Sa thèse, intitulée provisoirement « L'ascension de la maison d'ʿOsmān de c. 1300 à 1453 », propose une étude prosopographique et cartographique de la dynastie ottomane à ses débuts. Formée en histoire byzantine et ottomane à l'Université Paris I Panthéon-Sorbonne et en paléographie ottomane à l'EPHE, elle est également chargée de l'inventorisation du fonds Beldiceanu au Centre d'études ottomanes du Collège de France."),
        ('sch-harifi', 'Ilyas Harifi', None,
         "Spécialiste de l'histoire du Kalām. Biographie complète en cours de confirmation."),
        ('sch-mahil', 'Dr. Yannis Mahil', 'Maître de conférences en études islamiques',
         "Maître de conférences en études islamiques à la GIBTU University (Faculty of Theology, Department of Islamic Legal Studies), il est spécialiste du droit islamique et de la théorie des sources (Uṣūl al-fiqh)."),
        ('sch-azaiez', 'Mehdi Azaiez', "Professeur d'islamologie",
         "Mehdi Azaiez est professeur d'islamologie à l'Université catholique de Louvain (UCLouvain), spécialiste des études coraniques et de l'histoire de l'islam ancien. Il est l'auteur de plusieurs ouvrages de référence, dont « Le contre-discours coranique » (De Gruyter, 2015) et a dirigé le projet international Qur'ān Seminar. Il est chercheur associé à l'IREMAM et membre du comité des publications de l'International Quranic Studies Association (IQSA)."),
        ('sch-chahdi', 'Hassan Chahdi', 'Maître de conférences en islamologie',
         "Hassan Chahdi est maître de conférences en islamologie au département de théologie de l'Université de Lorraine (EA 3943 Écritures). Docteur de l'EPHE en histoire et philologie (2016), il a soutenu une thèse sur le muṣḥaf dans les débuts de l'islam. Lauréat du prix de la meilleure thèse francophone (GIS-Moyen Orient et IISMM) et du prix Schneider/Aguirre-Basualdo (Chancellerie des universités de Paris). Ancien chercheur au Collège de France (Chaire du Coran) et chargé de conférence à l'EPHE. Membre du CISA et du comité scientifique du BCAI. Ses travaux en cours portent sur la reconstitution du codex coranique perdu d'Ibn Masʿūd."),
        ('sch-grandpierre', 'Camille Grandpierre', "Doctorante contractuelle — INHA",
         "Camille Grandpierre est doctorante à Sorbonne Université, lauréate d'un contrat doctoral de l'Institut national d'histoire de l'art (INHA) depuis 2024. Elle prépare une thèse sous la direction d'Éloïse Brac de la Perrière (Sorbonne Université) et d'Alexandre Papas (EHESS), consacrée aux Dīvāns, des anthologies poétiques produites aux XIVe et XVe siècles dans le monde turco-persan. Formée en arts de l'Islam, elle est également titulaire d'un master en langue et littérature persanes (INALCO). Après avoir enseigné à l'École du Louvre entre 2017 et 2024, elle participe actuellement à l'INHA aux programmes CallFront (calligraphies arabes aux frontières du monde islamique) et RePaZ (patrimoine matériel en zones de conflit)."),
        ('sch-benmiled', 'Marouane Ben Miled', 'Enseignant-chercheur',
         "Enseignant-chercheur à l'École Nationale d'Ingénieurs de Tunis (ENIT) — Université de Tunis El Manar, il est spécialiste de l'histoire des mathématiques arabes. Chercheur LAMSIN (Histoire des Mathématiques) et associé CGGG (CNRS)."),
        ('sch-bensaad', 'Meyssa Ben Saad', 'Docteure en histoire des sciences',
         "Meyssa Ben Saâd est docteure en histoire des sciences, spécialisée en histoire de la zoologie arabe médiévale. Chercheure associée à l'UMR SPHERE (CNRS, Université Paris Cité), elle est également membre de la Faculté de l'Institut Supérieur d'Éducation Spécialisée de l'Université de la Manouba (Tunis). Ses recherches portent sur l'histoire naturelle dans le monde islamique médiéval, en particulier l'œuvre d'Al-Jāḥiẓ."),
        ('sch-sebti', 'Meryem Sebti', 'Directrice de recherche au CNRS',
         "Meryem Sebti est directrice de recherche au CNRS, spécialiste de philosophie médiévale arabe. Ses travaux portent principalement sur la pensée d'al-Kindī, d'al-Fārābī et d'Avicenne, ainsi que sur la transmission de la philosophie grecque dans le monde islamique. Elle est co-fondatrice de Sijill Project."),
        ('sch-mechelloukh', 'Yassir Mechelloukh', 'Doctorant et professeur certifié de philosophie',
         "Yassir Mechelloukh est doctorant en histoire de la philosophie arabe à l'INALCO et à l'Université Paris 1 Panthéon-Sorbonne, et professeur certifié de philosophie. Ses recherches portent sur la première métaphysique d'Averroès. Il est membre du CERMOM (EA 4091), du groupe GRAMATA et de l'UMR SPHERE 7219."),
        ('sch-rizvi', 'Sajjad H. Rizvi', 'Professor of Islamic Intellectual History',
         "Sajjad H. Rizvi est professeur d'histoire intellectuelle islamique à l'Université d'Exeter, où il dirige le Global and Area Studies de l'Institute of Arab and Islamic Studies. Formé à Oxford (BA, MA, MPhil) et Cambridge (PhD), il est l'un des spécialistes mondiaux de la tradition philosophique islamique orientale et de l'œuvre de Mullā Ṣadrā. Fellow de la Royal Society of Arts (FRSA), de la Royal Historical Society (FRHistS) et de la Royal Asiatic Society (FRAS)."),
        ('sch-molino', 'Cédric Molino Mochetto', 'Agrégé et docteur en philosophie',
         "Cédric Molino-Machetto est agrégé et docteur en philosophie. Il enseigne en lycée et à l'Université Toulouse Jean-Jaurès. Spécialiste de philosophie arabe et d'histoire de la pensée politique en contexte islamique prémoderne, il a publié plusieurs articles dans ce domaine."),
        ('sch-vandamme', 'Gregory Vandamme', 'Professeur et spécialiste du soufisme',
         "Grégory Vandamme est professeur et spécialiste du soufisme et de la pensée islamique médiévale. Ses recherches portent notamment sur la mystique musulmane et les grandes figures du soufisme, en particulier Ibn ʿArabī. Il s'intéresse aux dimensions spirituelles, symboliques et philosophiques des textes islamiques, ainsi qu'à leur dialogue avec d'autres traditions religieuses."),
        ('sch-roux', 'Géraldine Roux', 'Agrégée et docteure en philosophie',
         "Agrégée et docteure en philosophie, spécialiste de philosophie juive médiévale et autrice de \"Maïmonide ou la nostalgie de la sagesse\". Professeure agrégée de philosophie et chercheure associée au CRIMEL (EA 3311, Université de Reims Champagne-Ardenne)."),
    ]
    for sid, sname, stitle, sbio in SCHOLARS_SEED:
        await db.scholars.update_one(
            {'id': sid},
            {'$setOnInsert': {
                'id': sid,
                'name': sname,
                'title': stitle,
                'bio': sbio,
                'photo_url': None,
                'is_active': True,
                'created_at': datetime.now(timezone.utc),
            }},
            upsert=True
        )
        # Only keep is_active in sync. Content fields (name/title/bio) are
        # admin-owned: once a scholar doc exists, the seed must NEVER overwrite
        # them — admin edits in the panel are the source of truth.
        await db.scholars.update_one(
            {'id': sid},
            {'$set': {'is_active': True}}
        )
    # Link scholar_id on each course (the catalogue card now shows scholar_name from this)
    SCHOLAR_LINK = {
        'cours-debuts-islam': 'sch-bouali',  # primary; Ghouirgate is co-intervenant
        'cours-andalus': 'sch-ghouirgate',
        'cours-mamelouke': 'sch-benkherfallah',
        'cours-ottoman': 'sch-saban',
        'cours-kalam': 'sch-harifi',
        'cours-fiqh': 'sch-mahil',
        'cours-coran': 'sch-azaiez',
        'cours-hadith': 'sch-chahdi',
        'cours-historiographie': 'sch-ghouirgate',
        'cours-art': 'sch-grandpierre',
        'cours-sciences': 'sch-benmiled',
        'cours-traduction': 'sch-sebti',
        'cours-falsafa-grands': 'sch-sebti',
        'cours-falsafa-occident': 'sch-mechelloukh',
        'cours-falsafa-persan': 'sch-rizvi',
        'cours-inclassables': 'sch-molino',
        'cours-soufisme': 'sch-vandamme',
        'cours-philo-juive': 'sch-roux',
    }
    for cid, sid in SCHOLAR_LINK.items():
        await db.courses.update_one({'id': cid}, {'$set': {'scholar_id': sid}})
    # Co-intervenants per Excel (multi-scholar courses)
    CO_SCHOLARS = {
        'cours-debuts-islam': ['sch-ghouirgate'],  # Bouali (primary) + Ghouirgate (co)
        'cours-sciences': ['sch-bensaad'],          # Ben Miled (primary) + Ben Saad (co)
    }
    for cid, co_ids in CO_SCHOLARS.items():
        await db.courses.update_one({'id': cid}, {'$set': {'co_scholar_ids': co_ids}})
    # Ensure absent field is empty list on others to keep payloads predictable
    await db.courses.update_many(
        {'co_scholar_ids': {'$exists': False}},
        {'$set': {'co_scholar_ids': []}}
    )
    logger.info(f"Migration v10: seeded {len(SCHOLARS_SEED)} scholars + linked to {len(SCHOLAR_LINK)} courses + {len(CO_SCHOLARS)} co-intervenant(s)")
    # Clean up duplicate scholars (legacy random IDs created before the named-id seed)
    legacy_dup = await db.scholars.delete_many({
        'id': {'$nin': [s[0] for s in SCHOLARS_SEED]},
        'name': {'$in': [s[1] for s in SCHOLARS_SEED]}
    })
    if legacy_dup.deleted_count > 0:
        logger.info(f"Migration v10: removed {legacy_dup.deleted_count} duplicate legacy scholar(s)")
    # ─── End Migration v10 ─────────────────────────────────────────────────

    if custom_cursus:
        logger.info("Custom cursus 'cursus-falsafa' found - skipping all demo course/audio seeding")
        # Migrate plans to fondateur pricing (Standard Annuel also active for /tarification)
        for plan_id, plan_data in DEFAULT_PLANS.items():
            await db.plans.update_one(
                {'plan_id': plan_id},
                {'$set': {**plan_data, 'plan_id': plan_id}},
                upsert=True
            )
        await db.plans.update_many(
            {'plan_id': {'$nin': list(DEFAULT_PLANS.keys())}},
            {'$set': {'is_active': False}}
        )
        logger.info("Plans migration complete")
        logger.info("Database seeding complete")
        return

    # 2. Set admin role for admin email
    result = await db.users.update_one(
        {'email': 'loubna.serrar@gmail.com'},
        {'$set': {'role': 'admin'}}
    )
    if result.modified_count > 0:
        logger.info("Admin role set for loubna.serrar@gmail.com")

    # 3. Add/Update Meryem Sebti
    meryem_data = {
        "id": "sch-006",
        "name": "Prof. Meryem Sebti",
        "university": "CNRS / EPHE (École Pratique des Hautes Études)",
        "bio": (
            "Directrice de recherche au CNRS et chargée de conférences invitée à l'EPHE, "
            "Meryem Sebti a principalement travaillé sur Avicenne, notamment sur sa doctrine "
            "de l'âme et sa prophétologie. Spécialiste de la philosophie arabe médiévale, "
            "elle est l'auteure de nombreux ouvrages de référence dont "
            "Avicenne – Prophétie et gouvernement du monde (Cerf, 2021), "
            "Noétique et théorie de la connaissance dans la philosophie arabe (Vrin, 2020), "
            "et Avicenne. L'âme humaine (PUF, 2000)."
        ),
        "photo": "https://customer-assets.emergentagent.com/job_057211bf-8567-4749-b2d1-1f73d9b86661/artifacts/6hhq4gdt_Cercle_Meriemv2.jpg",
        "specializations": ["Avicenne", "Noétique", "Prophétologie", "Philosophie arabe"],
        "content_count": 1,
        "is_active": True,
    }
    await db.scholars.update_one(
        {'id': 'sch-006'},
        {'$set': meryem_data},
        upsert=True
    )
    logger.info("Scholar Meryem Sebti added/updated")

    # 4. Add Henry Corbin
    corbin_data = {
        "id": "sch-007",
        "name": "Henry Corbin",
        "university": "EPHE / Université de Téhéran",
        "bio": (
            "Henry Corbin (1903-1978), philosophe, germaniste, iranologue, arabisant, étudiera avec acharnement "
            "les textes des écoles philosophiques de l'Iran du 12e au 19e siècle. Tout était à faire en ce domaine : "
            "établir les textes, les éditer, les traduire, les présenter. Il fera connaître les richesses de l'école "
            "d'Ispahan à l'Occident et à l'Orient. Ses ouvrages restent à nos jours l'une des plus grandes expositions "
            "du soufisme persan classique disponible pour un public occidental.\n\n"
            "Grâce à lui, l'Occident découvrira trois grandes figures : Sohravardî, le grand platonicien de Perse "
            "(L'Archange empourpré, 1976), Ibn 'Arabi, le grand maître (L'Imagination créatrice dans le soufisme "
            "d'Ibn 'Arabī 1958 ; 2e éd. 1977), et Mollā Sādrā Shīrāzī (Livre des pénétrations métaphysiques, 1964). "
            "Il laissera également deux grandes références majeures en philosophie islamique : Histoire de la "
            "philosophie islamique, et En Islam iranien (t. I et II, 1971 ; t. III et IV, 1973).\n\n"
            "La bibliographie complète, ainsi que de nombreuses ressources sont disponibles sur le site de "
            "l'association des amis de Corbin : https://www.amiscorbin.com/"
        ),
        "photo": "https://customer-assets.emergentagent.com/job_057211bf-8567-4749-b2d1-1f73d9b86661/artifacts/hxtabq8s_1973_HCorbine04d31.jpeg",
        "specializations": ["Philosophie iranienne", "Soufisme persan", "Sohravardî", "Ibn Arabi", "Mollā Sādrā"],
        "content_count": 14,
        "is_active": True,
    }
    await db.scholars.update_one(
        {'id': 'sch-007'},
        {'$set': corbin_data},
        upsert=True
    )
    logger.info("Scholar Henry Corbin added/updated")

    # Skip creating demo courses if our custom cursus exists
    if await db.thematiques.find_one({'id': 'cursus-falsafa'}):
        logger.info("Custom cursus found - skipping demo course seeding")
        logger.info("Database seeding complete")
        return

    # 5. Create Course for "Philosophie" folder (Meryem Sebti)
    philosophie_course = {
        "id": "crs-philo-sebti",
        "title": "Philosophie",
        "description": (
            "Cours de philosophie arabe et islamique par le Prof. Meryem Sebti. "
            "Ce cycle explore les grandes traditions philosophiques du monde arabo-musulman, "
            "avec un accent particulier sur Avicenne, la noétique et la prophétologie. "
            "Une plongée académique rigoureuse dans la pensée islamique médiévale."
        ),
        "topic": "Philosophie islamique",
        "level": "Intermédiaire",
        "language": "Français",
        "scholar_id": "sch-006",
        "scholar_name": "Prof. Meryem Sebti",
        "duration": 0,
        "thumbnail": "https://images.unsplash.com/photo-1507842217343-583bb7270b66?w=600&q=80",
        "modules_count": 0,
        "tags": ["Philosophie", "Avicenne", "Noétique", "Philosophie arabe"],
        "type": "course",
        "r2_folder": "Philosophie",
        "published_at": now.isoformat(),
        "is_active": True,
    }
    await db.courses.update_one(
        {'id': 'crs-philo-sebti'},
        {'$set': philosophie_course},
        upsert=True
    )
    logger.info("Course 'Philosophie' (Meryem Sebti) created/updated")

    # 6. Create Course for "Henry Corbin" folder
    corbin_course = {
        "id": "crs-henry-corbin",
        "title": "Cycle Henry Corbin",
        "description": (
            "Cycle complet sur l'œuvre et la pensée de Henry Corbin, le grand iranologue et philosophe. "
            "Ce cours en 14 épisodes explore les thèmes majeurs de son travail : la philosophie iranienne, "
            "le soufisme persan, Sohravardî, Ibn 'Arabi et Mollā Sādrā. Un voyage intellectuel extraordinaire "
            "à travers l'une des plus grandes expositions du soufisme persan classique disponible en français."
        ),
        "topic": "Philosophie islamique",
        "level": "Avancé",
        "language": "Français",
        "scholar_id": "sch-007",
        "scholar_name": "Henry Corbin",
        "duration": 0,  # Will be calculated from episodes
        "thumbnail": "https://customer-assets.emergentagent.com/job_057211bf-8567-4749-b2d1-1f73d9b86661/artifacts/hxtabq8s_1973_HCorbine04d31.jpeg",
        "modules_count": 14,
        "tags": ["Henry Corbin", "Philosophie iranienne", "Soufisme", "Sohravardî", "Ibn Arabi"],
        "type": "course",
        "r2_folder": "Henry Corbin",
        "published_at": now.isoformat(),
        "is_active": True,
    }
    await db.courses.update_one(
        {'id': 'crs-henry-corbin'},
        {'$set': corbin_course},
        upsert=True
    )
    logger.info("Course 'Cycle Henry Corbin' created/updated")

    # 7. Create Audio entries for Henry Corbin episodes
    corbin_episodes = [
        {"ep": 1, "title": "Introduction à la pensée de Henry Corbin", "file_key": "Henry Corbin/Cycle_HCorbin_episode1.m4a", "duration": 420},
        {"ep": 2, "title": "Sohravardî et la philosophie de l'illumination", "file_key": "Henry Corbin/Cycle_HCorbin_episode2.m4a", "duration": 470},
        {"ep": 3, "title": "Le monde imaginal ('alam al-mithal)", "file_key": "Henry Corbin/Cycle_HCorbin_episode3.m4a", "duration": 470},
        {"ep": 4, "title": "Ibn 'Arabi : L'imagination créatrice", "file_key": "Henry Corbin/Cycle_HCorbin_episode4.m4a", "duration": 640},
        {"ep": 5, "title": "Mollā Sādrā Shīrāzī et les pénétrations métaphysiques", "file_key": "Henry Corbin/Cycle_HCorbin_episode5.m4a", "duration": 750},
        {"ep": 6, "title": "L'école d'Ispahan : contexte historique", "file_key": "Henry Corbin/Cycle_HCorbin_episode6.m4a", "duration": 1080},
        {"ep": 7, "title": "L'Archange empourpré : mystique de la lumière", "file_key": "Henry Corbin/Cycle_HCorbin_episode7.m4a", "duration": 2100},
        {"ep": 8, "title": "Soufisme persan et gnose islamique", "file_key": "Henry Corbin/Cycle_HCorbin_episode8.m4a", "duration": 2400},
        {"ep": 9, "title": "En Islam iranien : le shi'isme duodécimain", "file_key": "Henry Corbin/Cycle_HCorbin_episode9.m4a", "duration": 2280},
        {"ep": 10, "title": "La prophétologie et l'eschatologie", "file_key": "Henry Corbin/Cycle_HCorbin_episode10.m4a", "duration": 4800},
        {"ep": 11, "title": "L'herméneutique spirituelle du Coran", "file_key": "Henry Corbin/Cycle_HCorbin_episode11.m4a", "duration": 2940},
        {"ep": 12, "title": "Le temps cyclique et l'hiérohistoire", "file_key": "Henry Corbin/Cycle_HCorbin_episode12.m4a", "duration": 3990},
        {"ep": 13, "title": "Corbin et la phénoménologie religieuse", "file_key": "Henry Corbin/Cycle_HCorbin_episode13.m4a", "duration": 360},
        {"ep": 14, "title": "Conclusion : L'héritage de Henry Corbin", "file_key": "Henry Corbin/Cycle_HCorbin_episode14.m4a", "duration": 2094},
    ]

    for ep in corbin_episodes:
        audio_doc = {
            "id": f"aud-corbin-{ep['ep']:02d}",
            "title": f"Épisode {ep['ep']} — {ep['title']}",
            "description": f"Cycle Henry Corbin, épisode {ep['ep']}. {ep['title']}.",
            "scholar_id": "sch-007",
            "scholar_name": "Henry Corbin",
            "duration": ep['duration'],
            "audio_url": "",
            "file_key": ep['file_key'],
            "thumbnail": "https://customer-assets.emergentagent.com/job_057211bf-8567-4749-b2d1-1f73d9b86661/artifacts/hxtabq8s_1973_HCorbine04d31.jpeg",
            "topic": "Philosophie islamique",
            "type": "lecture",
            "course_id": "crs-henry-corbin",
            "episode_number": ep['ep'],
            "published_at": now.isoformat(),
            "is_active": True,
        }
        await db.audios.update_one(
            {'id': f"aud-corbin-{ep['ep']:02d}"},
            {'$set': audio_doc},
            upsert=True
        )
    logger.info("Henry Corbin audio episodes created/updated")

    # 8. Reassign existing Philosophie islamique courses to Meryem Sebti (except Corbin's)
    phil_update = await db.courses.update_many(
        {'topic': 'Philosophie islamique', 'scholar_id': {'$nin': ['sch-006', 'sch-007']}},
        {'$set': {'scholar_id': 'sch-006', 'scholar_name': 'Prof. Meryem Sebti'}}
    )
    if phil_update.modified_count > 0:
        logger.info(f"Reassigned {phil_update.modified_count} philosophy courses to Meryem Sebti")

    # 9. Migrate plans to fondateur pricing
    for plan_id, plan_data in DEFAULT_PLANS.items():
        await db.plans.update_one(
            {'plan_id': plan_id},
            {'$set': {**plan_data, 'plan_id': plan_id, 'is_active': plan_data.get('is_fondateur', False)}},
            upsert=True
        )
    # Deactivate old plans
    await db.plans.update_many(
        {'plan_id': {'$nin': list(DEFAULT_PLANS.keys())}},
        {'$set': {'is_active': False}}
    )
    logger.info("Plans migration complete")

    logger.info("Database seeding complete")

# ─── Admin Routes ─────────────────────────────────────────────────────────────

# Admin: Stats
@api_router.get("/admin/stats")
async def admin_stats(request: Request):
    await require_admin(request)
    audios_total = await db.audios.count_documents({})
    audios_active = await db.audios.count_documents({'is_active': True})
    scholars_total = await db.scholars.count_documents({})
    courses_total = await db.courses.count_documents({})
    courses_active = await db.courses.count_documents({'is_active': True})
    users_total = await db.users.count_documents({})
    return {
        'audios': {'total': audios_total, 'active': audios_active},
        'scholars': {'total': scholars_total},
        'courses': {'total': courses_total, 'active': courses_active},
        'users': {'total': users_total},
    }

# Admin: Listening Statistics Dashboard
@api_router.get("/admin/listening-stats")
async def admin_listening_stats(request: Request, period: str = "all"):
    """Get detailed listening statistics by cursus, course, module, and professor."""
    await require_admin(request)
    
    # Calculate date filter based on period
    now = datetime.now(timezone.utc)
    date_filter = {}
    if period == "7days":
        date_filter = {"updated_at": {"$gte": now - timedelta(days=7)}}
    elif period == "month":
        date_filter = {"updated_at": {"$gte": now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)}}
    elif period == "year":
        date_filter = {"updated_at": {"$gte": now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)}}
    # "all" = no filter
    
    # Get all progress data
    progress_query = {"content_type": "audio"}
    if date_filter:
        progress_query.update(date_filter)
    
    progress_items = await db.user_progress.find(progress_query, {"_id": 0}).to_list(10000)
    
    # Get all audios, courses, cursus, and scholars for reference
    audios = {a['id']: a for a in await db.audios.find({}, {"_id": 0}).to_list(1000)}
    courses = {c['id']: c for c in await db.courses.find({}, {"_id": 0}).to_list(100)}
    cursus_list = {c['id']: c for c in await db.thematiques.find({}, {"_id": 0}).to_list(50)}
    scholars = {s['id']: s for s in await db.scholars.find({}, {"_id": 0}).to_list(100)}
    
    # Aggregate stats
    stats_by_cursus = {}
    stats_by_course = {}
    stats_by_audio = {}
    stats_by_professor = {}
    daily_stats = {}
    
    total_listening_seconds = 0
    total_plays = 0
    
    for p in progress_items:
        audio_id = p.get('content_id')
        audio = audios.get(audio_id, {})
        if not audio:
            continue
        
        duration = audio.get('duration', 0)
        progress_pct = p.get('progress', 0)
        listening_seconds = int(duration * progress_pct)
        total_listening_seconds += listening_seconds
        total_plays += 1
        
        # Get related entities
        course_id = audio.get('course_id', '')
        course = courses.get(course_id, {})
        cursus_id = course.get('cursus_id', '')
        cursus = cursus_list.get(cursus_id, {})
        professor_id = audio.get('scholar_id', '') or course.get('scholar_id', '')
        professor = scholars.get(professor_id, {})
        
        # Aggregate by cursus
        if cursus_id:
            if cursus_id not in stats_by_cursus:
                stats_by_cursus[cursus_id] = {
                    'id': cursus_id,
                    'name': cursus.get('name', cursus_id),
                    'listening_seconds': 0,
                    'plays': 0
                }
            stats_by_cursus[cursus_id]['listening_seconds'] += listening_seconds
            stats_by_cursus[cursus_id]['plays'] += 1
        
        # Aggregate by course
        if course_id:
            if course_id not in stats_by_course:
                stats_by_course[course_id] = {
                    'id': course_id,
                    'title': course.get('title', course_id),
                    'cursus_name': cursus.get('name', ''),
                    'listening_seconds': 0,
                    'plays': 0
                }
            stats_by_course[course_id]['listening_seconds'] += listening_seconds
            stats_by_course[course_id]['plays'] += 1
        
        # Aggregate by audio (module/episode)
        if audio_id not in stats_by_audio:
            stats_by_audio[audio_id] = {
                'id': audio_id,
                'title': audio.get('title', audio_id),
                'course_title': course.get('title', ''),
                'listening_seconds': 0,
                'plays': 0
            }
        stats_by_audio[audio_id]['listening_seconds'] += listening_seconds
        stats_by_audio[audio_id]['plays'] += 1
        
        # Aggregate by professor
        if professor_id:
            if professor_id not in stats_by_professor:
                stats_by_professor[professor_id] = {
                    'id': professor_id,
                    'name': professor.get('name', professor_id),
                    'photo': professor.get('photo_url', ''),
                    'listening_seconds': 0,
                    'plays': 0
                }
            stats_by_professor[professor_id]['listening_seconds'] += listening_seconds
            stats_by_professor[professor_id]['plays'] += 1
        
        # Daily stats for chart
        updated_at = p.get('updated_at')
        if updated_at:
            if isinstance(updated_at, str):
                updated_at = datetime.fromisoformat(updated_at.replace('Z', '+00:00'))
            day_key = updated_at.strftime('%Y-%m-%d')
            if day_key not in daily_stats:
                daily_stats[day_key] = {'date': day_key, 'listening_seconds': 0, 'plays': 0}
            daily_stats[day_key]['listening_seconds'] += listening_seconds
            daily_stats[day_key]['plays'] += 1
    
    # Sort and limit results
    cursus_sorted = sorted(stats_by_cursus.values(), key=lambda x: x['listening_seconds'], reverse=True)
    courses_sorted = sorted(stats_by_course.values(), key=lambda x: x['listening_seconds'], reverse=True)[:20]
    audios_sorted = sorted(stats_by_audio.values(), key=lambda x: x['listening_seconds'], reverse=True)[:20]
    professors_sorted = sorted(stats_by_professor.values(), key=lambda x: x['listening_seconds'], reverse=True)
    
    # Sort daily stats by date
    daily_sorted = sorted(daily_stats.values(), key=lambda x: x['date'])
    
    # Convert seconds to hours for display
    def to_hours(seconds):
        return round(seconds / 3600, 1)
    
    for item in cursus_sorted:
        item['listening_hours'] = to_hours(item['listening_seconds'])
    for item in courses_sorted:
        item['listening_hours'] = to_hours(item['listening_seconds'])
    for item in audios_sorted:
        item['listening_hours'] = to_hours(item['listening_seconds'])
    for item in professors_sorted:
        item['listening_hours'] = to_hours(item['listening_seconds'])
    for item in daily_sorted:
        item['listening_hours'] = to_hours(item['listening_seconds'])
    
    return {
        'period': period,
        'total': {
            'listening_hours': to_hours(total_listening_seconds),
            'listening_seconds': total_listening_seconds,
            'plays': total_plays,
            'unique_users': len(set(p.get('user_id') for p in progress_items))
        },
        'by_cursus': cursus_sorted,
        'by_course': courses_sorted,
        'by_audio': audios_sorted,
        'by_professor': professors_sorted,
        'daily': daily_sorted
    }

# Admin: Audio CRUD
@api_router.get("/admin/audios")
async def admin_list_audios(request: Request):
    await require_admin(request)
    audios = await db.audios.find({}, {'_id': 0}).to_list(500)
    for a in audios:
        a['stream_url'] = resolve_audio_url(a)
    return audios

@api_router.post("/admin/audios")
async def admin_create_audio(body: AudioCreate, request: Request):
    await require_admin(request)
    audio_id = f"aud-{uuid.uuid4().hex[:8]}"
    now = datetime.now(timezone.utc)
    doc = {**body.model_dump(), 'id': audio_id, 'published_at': now.isoformat()}
    await db.audios.insert_one(doc)
    doc.pop('_id', None)
    doc['stream_url'] = resolve_audio_url(doc)
    return doc

@api_router.put("/admin/audios/{audio_id}")
async def admin_update_audio(audio_id: str, body: AudioUpdate, request: Request):
    await require_admin(request)
    update = {k: v for k, v in body.model_dump().items() if v is not None}
    if not update:
        raise HTTPException(400, "Aucune mise à jour fournie")
    result = await db.audios.update_one({'id': audio_id}, {'$set': update})
    if result.matched_count == 0:
        raise HTTPException(404, "Audio non trouvé")
    doc = await db.audios.find_one({'id': audio_id}, {'_id': 0})
    doc['stream_url'] = resolve_audio_url(doc)
    return doc

@api_router.delete("/admin/audios/{audio_id}")
async def admin_delete_audio(audio_id: str, request: Request):
    await require_admin(request)
    result = await db.audios.delete_one({'id': audio_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Audio non trouvé")
    return {'message': 'Audio supprimé'}

@api_router.patch("/admin/audios/{audio_id}/toggle")
async def admin_toggle_audio(audio_id: str, request: Request):
    await require_admin(request)
    doc = await db.audios.find_one({'id': audio_id})
    if not doc:
        raise HTTPException(404, "Audio non trouvé")
    new_status = not doc.get('is_active', True)
    await db.audios.update_one({'id': audio_id}, {'$set': {'is_active': new_status}})
    return {'id': audio_id, 'is_active': new_status}

@api_router.patch("/admin/audios/{audio_id}/seed-lock")
async def admin_toggle_audio_seed_lock(audio_id: str, request: Request):
    """Toggle the `seed_locked` flag on an audio episode.

    When TRUE the title is protected: no startup migration / seed can rewrite
    it on reboot. Admin uses the lock icon in /api/admin-panel/tree to free
    a title temporarily for editing, then re-lock once done.
    """
    await require_admin(request)
    doc = await db.audios.find_one({'id': audio_id}, {'_id': 0, 'seed_locked': 1})
    if not doc:
        raise HTTPException(404, "Audio non trouvé")
    new_state = not bool(doc.get('seed_locked'))
    await db.audios.update_one({'id': audio_id}, {'$set': {'seed_locked': new_state}})
    return {'id': audio_id, 'seed_locked': new_state}

@api_router.patch("/admin/courses/{course_id}/seed-lock")
async def admin_toggle_course_seed_lock(course_id: str, request: Request):
    """Toggle the `seed_locked` flag on a course (protects title/description/scholar)."""
    await require_admin(request)
    doc = await db.courses.find_one({'id': course_id}, {'_id': 0, 'seed_locked': 1})
    if not doc:
        raise HTTPException(404, "Cours non trouvé")
    new_state = not bool(doc.get('seed_locked'))
    await db.courses.update_one({'id': course_id}, {'$set': {'seed_locked': new_state}})
    return {'id': course_id, 'seed_locked': new_state}

@api_router.post("/admin/audios/bulk-toggle")
async def admin_bulk_toggle_audios(request: Request):
    """Activate or deactivate all audios for a given course or cursus."""
    await require_admin(request)
    body = await request.json()
    course_id = body.get('course_id')
    cursus_id = body.get('cursus_id')
    is_active = body.get('is_active', True)

    query = {}
    if course_id:
        query['course_id'] = course_id
    elif cursus_id:
        query['cursus_id'] = cursus_id
    else:
        raise HTTPException(400, "course_id ou cursus_id requis")

    result = await db.audios.update_many(query, {'$set': {'is_active': is_active}})
    return {'modified': result.modified_count, 'is_active': is_active}

# Admin: Scholar CRUD
@api_router.get("/admin/scholars")
async def admin_list_scholars(request: Request):
    await require_admin(request)
    scholars = await db.scholars.find({}, {'_id': 0}).to_list(100)
    return scholars

@api_router.post("/admin/scholars")
async def admin_create_scholar(body: ScholarCreate, request: Request):
    await require_admin(request)
    scholar_id = f"sch-{uuid.uuid4().hex[:6]}"
    doc = {**body.model_dump(), 'id': scholar_id, 'content_count': 0, 'is_active': True}
    await db.scholars.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put("/admin/scholars/{scholar_id}")
async def admin_update_scholar(scholar_id: str, body: ScholarUpdate, request: Request):
    await require_admin(request)
    update = {k: v for k, v in body.model_dump().items() if v is not None}
    # Defensive: an empty string for `photo` means "no change" (the form left it empty),
    # NOT "clear the photo". Same for `name` to avoid accidental wipes.
    for protected_field in ('photo', 'name'):
        if update.get(protected_field) == '':
            update.pop(protected_field, None)
    if not update:
        raise HTTPException(400, "Aucune mise à jour fournie")
    # Lock this scholar against the on-startup seed so admin edits persist
    # across redeploys (see Migration v10 / SCHOLARS_SEED).
    update['seed_locked'] = True
    result = await db.scholars.update_one({'id': scholar_id}, {'$set': update})
    if result.matched_count == 0:
        raise HTTPException(404, "Érudit non trouvé")
    doc = await db.scholars.find_one({'id': scholar_id}, {'_id': 0})
    return doc

@api_router.delete("/admin/scholars/{scholar_id}")
async def admin_delete_scholar(scholar_id: str, request: Request):
    await require_admin(request)
    result = await db.scholars.delete_one({'id': scholar_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Érudit non trouvé")
    return {'message': 'Érudit supprimé'}


# ─── R2 Auto-Detection Admin endpoints (Phase 3) ────────────────────────────

@api_router.post("/admin/courses/{course_id}/r2-prefix")
async def admin_set_course_r2_prefix(course_id: str, body: dict, request: Request):
    """Set the R2 prefix for a course (e.g. 'cursus-X/NN-foo/' or 'cursus-X/NN-foo/bar/').
    The prefix is the path under which to scan for media files.
    """
    await require_admin(request)
    prefix = (body or {}).get('r2_prefix', '').strip().strip('/')
    if not prefix:
        # Allow clearing the prefix
        await db.courses.update_one({'id': course_id}, {'$unset': {'r2_prefix': ''}})
        return {'course_id': course_id, 'r2_prefix': None}
    course = await db.courses.find_one({'id': course_id}, {'_id': 0, 'id': 1})
    if not course:
        raise HTTPException(404, "Cours non trouvé")
    # Append trailing slash for prefix scans
    if not prefix.endswith('/'):
        prefix = prefix + '/'
    await db.courses.update_one({'id': course_id}, {'$set': {'r2_prefix': prefix}})
    return {'course_id': course_id, 'r2_prefix': prefix}


@api_router.get("/admin/courses/{course_id}/r2-detection")
async def admin_get_r2_detection(course_id: str, request: Request, prefix: Optional[str] = None):
    """Preview R2 auto-detection for a course (read-only). Pass ?prefix=… to override
    the stored r2_prefix without persisting it."""
    await require_admin(request)
    course = await db.courses.find_one({'id': course_id}, {'_id': 0})
    if not course:
        raise HTTPException(404, "Cours non trouvé")
    effective_prefix = (prefix or course.get('r2_prefix') or '').strip().strip('/')
    if not effective_prefix:
        raise HTTPException(400, "Aucun r2_prefix défini pour ce cours")
    if not effective_prefix.endswith('/'):
        effective_prefix = effective_prefix + '/'
    detections = _build_r2_detections(effective_prefix)
    audios = await db.audios.find(
        {'course_id': course_id},
        {'_id': 0, 'id': 1, 'episode_number': 1, 'title': 1}
    ).to_list(200)
    audios.sort(key=lambda a: a.get('episode_number') or 0)
    return {
        'course_id': course_id,
        'course_title': course.get('title') or course.get('name'),
        'prefix': effective_prefix,
        'episodes_in_db': audios,
        'detected': {
            'videos': {str(k): v for k, v in detections['videos'].items()},
            'audios': {str(k): v for k, v in detections['audios'].items()},
            'episode_docs': {str(k): v for k, v in detections['episode_docs'].items()},
            'course_docs': detections['course_docs'],
            'unclassified': detections['unclassified'],
        },
        'totals': {
            'videos': len(detections['videos']),
            'audios': len(detections['audios']),
            'episode_doc_groups': len(detections['episode_docs']),
            'course_docs': len(detections['course_docs']),
            'unclassified': len(detections['unclassified']),
        },
    }


@api_router.post("/admin/courses/{course_id}/sync-r2")
async def admin_sync_r2_for_course(course_id: str, request: Request, body: Optional[dict] = None):
    """Apply R2 auto-detection for a single course. Body may set { r2_prefix } to
    persist a new prefix BEFORE running the detection. Returns a summary."""
    await require_admin(request)
    course = await db.courses.find_one({'id': course_id}, {'_id': 0})
    if not course:
        raise HTTPException(404, "Cours non trouvé")
    new_prefix = (body or {}).get('r2_prefix') if body else None
    effective_prefix = (new_prefix or course.get('r2_prefix') or '').strip().strip('/')
    if not effective_prefix:
        raise HTTPException(400, "Aucun r2_prefix défini pour ce cours")
    if not effective_prefix.endswith('/'):
        effective_prefix = effective_prefix + '/'
    detections = _build_r2_detections(effective_prefix)
    summary = await _apply_r2_detections(course_id, detections)
    return {
        'course_id': course_id,
        'prefix': effective_prefix,
        'summary': summary,
        'unclassified': detections['unclassified'],
    }


@api_router.post("/admin/sync-r2-all")
async def admin_sync_r2_all(request: Request):
    """Apply R2 auto-detection to all courses that have a stored r2_prefix. Returns per-course summaries."""
    await require_admin(request)
    courses = await db.courses.find({'r2_prefix': {'$exists': True, '$ne': ''}}, {'_id': 0}).to_list(500)
    results = []
    for c in courses:
        prefix = (c.get('r2_prefix') or '').strip().strip('/')
        if not prefix:
            continue
        if not prefix.endswith('/'):
            prefix = prefix + '/'
        detections = _build_r2_detections(prefix)
        summary = await _apply_r2_detections(c['id'], detections)
        results.append({
            'course_id': c['id'],
            'title': c.get('title') or c.get('name'),
            'prefix': prefix,
            'summary': summary,
        })
    return {'count': len(results), 'results': results}


@api_router.delete("/admin/courses/{course_id}/episode-resources/{audio_id}/{file_basename}")
async def admin_unlink_episode_resource(course_id: str, audio_id: str, file_basename: str, request: Request):
    """Manual override: remove a single episode_resources entry by its filename."""
    await require_admin(request)
    a = await db.audios.find_one({'id': audio_id, 'course_id': course_id}, {'_id': 0, 'episode_resources': 1})
    if not a:
        raise HTTPException(404, "Épisode non trouvé")
    new_list = [r for r in (a.get('episode_resources') or []) if not r.get('r2_key', '').endswith('/' + file_basename) and not r.get('r2_key', '').endswith(file_basename)]
    await db.audios.update_one({'id': audio_id}, {'$set': {'episode_resources': new_list}})
    return {'audio_id': audio_id, 'episode_resources': new_list}


@api_router.delete("/admin/courses/{course_id}/course-resources/{file_basename}")
async def admin_unlink_course_resource(course_id: str, file_basename: str, request: Request):
    """Manual override: remove a single course_resources entry by its filename."""
    await require_admin(request)
    c = await db.courses.find_one({'id': course_id}, {'_id': 0, 'course_resources': 1})
    if not c:
        raise HTTPException(404, "Cours non trouvé")
    new_list = [r for r in (c.get('course_resources') or []) if not r.get('r2_key', '').endswith('/' + file_basename) and not r.get('r2_key', '').endswith(file_basename)]
    await db.courses.update_one({'id': course_id}, {'$set': {'course_resources': new_list}})
    return {'course_id': course_id, 'course_resources': new_list}


# Admin: Course CRUD
@api_router.get("/admin/courses")
async def admin_list_courses(request: Request):
    await require_admin(request)
    courses = await db.courses.find({}, {'_id': 0}).to_list(200)
    # Count modules for each course
    for c in courses:
        c['module_count'] = await db.modules.count_documents({'course_id': c['id']})
    return courses

@api_router.post("/admin/courses")
async def admin_create_course(body: CourseCreate, request: Request):
    await require_admin(request)
    course_id = f"crs-{uuid.uuid4().hex[:8]}"
    now = datetime.now(timezone.utc)
    doc = {**body.model_dump(), 'id': course_id, 'type': 'course', 'published_at': now.isoformat(), 'is_active': True}
    await db.courses.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put("/admin/courses/{course_id}")
async def admin_update_course(course_id: str, body: CourseUpdate, request: Request):
    await require_admin(request)
    update = {k: v for k, v in body.model_dump().items() if v is not None}
    if not update:
        raise HTTPException(400, "Aucune mise à jour fournie")
    # Lock against on-startup seed overwrites (description / summary / scholar_name etc).
    update['seed_locked'] = True
    result = await db.courses.update_one({'id': course_id}, {'$set': update})
    if result.matched_count == 0:
        raise HTTPException(404, "Cours non trouvé")
    doc = await db.courses.find_one({'id': course_id}, {'_id': 0})
    return doc

@api_router.delete("/admin/courses/{course_id}")
async def admin_delete_course(course_id: str, request: Request):
    await require_admin(request)
    result = await db.courses.delete_one({'id': course_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Cours non trouvé")
    return {'message': 'Cours supprimé'}

@api_router.patch("/admin/courses/{course_id}/toggle")
async def admin_toggle_course(course_id: str, request: Request):
    await require_admin(request)
    doc = await db.courses.find_one({'id': course_id})
    if not doc:
        raise HTTPException(404, "Cours non trouvé")
    new_status = not doc.get('is_active', True)
    await db.courses.update_one({'id': course_id}, {'$set': {'is_active': new_status}})
    return {'id': course_id, 'is_active': new_status}

# ─── R2 Folder Sync for Courses ───────────────────────────────────────────────

class SyncR2FolderRequest(BaseModel):
    r2_folder: str  # e.g. "Philosophie" or "Henry Corbin"

@api_router.get("/admin/r2/folders")
async def list_r2_folders(request: Request):
    """List all folders (prefixes) in the R2 bucket."""
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    try:
        # List objects and extract unique folder prefixes
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, MaxKeys=1000)
        folders = set()
        for obj in response.get('Contents', []):
            key = obj['Key']
            if '/' in key:
                folder = key.split('/')[0]
                folders.add(folder)
        return {'folders': sorted(list(folders)), 'bucket': R2_BUCKET}
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

@api_router.get("/admin/r2/folder/{folder_name}/files")
async def list_r2_folder_files(folder_name: str, request: Request):
    """List all audio files in a specific R2 folder."""
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    try:
        prefix = f"{folder_name}/"
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix=prefix, MaxKeys=500)
        files = []
        for obj in response.get('Contents', []):
            key = obj['Key']
            size = obj['Size']
            if size > 0:  # Skip folder markers
                filename = key.replace(prefix, '')
                files.append({
                    'key': key,
                    'filename': filename,
                    'size': size,
                    'size_mb': round(size / (1024*1024), 1),
                })
        # Sort by filename to get proper episode order
        files.sort(key=lambda x: x['filename'])
        return {'folder': folder_name, 'files': files, 'count': len(files)}
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

# ─── Timeline Routes ───────────────────────────────────────────────────────────

@api_router.get("/timeline/{cursus_letter}/access-url")
async def get_timeline_access_url(cursus_letter: str, request: Request):
    """Issue a short-lived signed URL to view the timeline HTML in a new tab.
    Requires authentication + active subscription/trial/admin/free_access."""
    user = await require_subscriber(request)
    letter = cursus_letter.upper().strip()
    if letter not in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        raise HTTPException(400, "Cursus invalide")
    token = create_jwt({
        'sub': user['user_id'],
        'cursus_letter': letter,
        'scope': 'content_access',
        'exp': int((datetime.now(timezone.utc) + timedelta(hours=1)).timestamp()),
    })
    scheme = request.headers.get('x-forwarded-proto', 'https')
    host = request.headers.get('x-forwarded-host') or request.headers.get('host', '')
    return {
        'url': f"{scheme}://{host}/api/timeline/{letter}?t={token}",
        'expires_in': 3600,
    }

@api_router.get("/timeline/file/{filename}/access-url")
async def get_timeline_file_access_url(filename: str, request: Request, key: Optional[str] = None):
    """Issue a short-lived signed URL for a specific timeline file.

    `key` is the optional full R2 key for files outside `Timeline/`
    (new course/cursus-scoped frises). Forwarded verbatim in the signed URL.
    """
    user = await require_subscriber(request)
    if not filename.endswith('.html'):
        filename = f"{filename}.html"
    token = create_jwt({
        'sub': user['user_id'],
        'filename': filename,
        'scope': 'content_access',
        'exp': int((datetime.now(timezone.utc) + timedelta(hours=1)).timestamp()),
    })
    scheme = request.headers.get('x-forwarded-proto', 'https')
    host = request.headers.get('x-forwarded-host') or request.headers.get('host', '')
    url = f"{scheme}://{host}/api/timeline/file/{filename}?t={token}"
    if key:
        from urllib.parse import quote
        url += f"&key={quote(key, safe='/')}"
    return {'url': url, 'expires_in': 3600}

@api_router.get("/timeline/{cursus_letter}")
async def get_timeline_html(cursus_letter: str, request: Request, t: Optional[str] = None):
    """Get timeline HTML content for a cursus (A, B, C, D, E).
    Requires authentication + active subscription, either via Authorization header or signed token (?t=)."""
    await verify_content_access(request, t)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    # Normalize the cursus letter
    letter = cursus_letter.upper().strip()
    if letter not in ['A', 'B', 'C', 'D', 'E']:
        raise HTTPException(400, "Cursus invalide. Utilisez A, B, C, D ou E.")
    
    # Check if there's a manual assignment in the database
    # Get all entries for this cursus and sort by updated_at (most recent first)
    # Also prioritize files with 'map' in the name
    db_entries = await db.timeline_resources.find(
        {'cursus_letter': letter, 'type': 'timeline'}
    ).sort('updated_at', -1).to_list(10)
    
    r2_keys_to_try = []
    if db_entries:
        # Prioritize map files, otherwise use the most recently updated
        map_entry = next((e for e in db_entries if 'map' in e.get('filename', '').lower()), None)
        if map_entry and map_entry.get('filename'):
            r2_keys_to_try.append(f"Timeline/{map_entry['filename']}")
        # Add all other files as fallbacks
        for entry in db_entries:
            key = f"Timeline/{entry.get('filename', '')}"
            if key not in r2_keys_to_try:
                r2_keys_to_try.append(key)
    
    # Always add the default as last fallback
    default_key = f"Timeline/sijill_timeline_cursus_{letter.lower()}.html"
    if default_key not in r2_keys_to_try:
        r2_keys_to_try.append(default_key)
    
    # Try each key until we find one that works
    for r2_key in r2_keys_to_try:
        try:
            response = r2_client.get_object(Bucket=R2_BUCKET, Key=r2_key)
            html_content = response['Body'].read().decode('utf-8')
            return HTMLResponse(content=html_content, media_type="text/html")
        except ClientError as e:
            error_code = e.response.get('Error', {}).get('Code', 'Unknown')
            if error_code != 'NoSuchKey':
                raise HTTPException(500, f"Erreur R2: {str(e)}")
            # File not found, try next one
            continue
    
    # None of the files were found
    raise HTTPException(404, f"Timeline non trouvée pour le cursus {letter}")

@api_router.get("/timelines")
async def list_available_timelines():
    """List all available timeline files."""
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    try:
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix='Timeline/', MaxKeys=100)
        timelines = []
        for obj in response.get('Contents', []):
            key = obj['Key']
            if key.endswith('.html') and 'timeline' in key.lower():
                # Extract cursus letter from filename
                filename = key.split('/')[-1]
                # sijill_timeline_cursus_a.html -> A
                import re
                match = re.search(r'cursus_([a-f])\.html', filename.lower())
                if match:
                    letter = match.group(1).upper()
                    timelines.append({
                        'cursus_letter': letter,
                        'cursus_name': {
                            'A': 'La Falsafa — Philosophie islamique',
                            'B': 'Théologie et Droit',
                            'C': 'Sciences islamiques et traditions',
                            'D': 'Arts, Littérature et Sciences',
                            'E': 'Philosophies et spiritualités'
                        }.get(letter, f'Cursus {letter}'),
                        'r2_key': key,
                        'url': f'/api/timeline/{letter}'
                    })
        return {'timelines': sorted(timelines, key=lambda x: x['cursus_letter'])}
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

@api_router.get("/timelines/cursus/{cursus_id}")
async def get_cursus_timelines(cursus_id: str):
    """List frises (HTML maps) for all courses of a cursus.

    Sources frises exclusively from the new course/cursus folder layout
    (`<cursus_root>/(sijill_)?map_*.html` for cursus-level,
    `<cursus_root>/<course>/(sijill_)?map_*.html` for course-specific).
    Legacy `Timeline/` folder + `db.timeline_resources` are no longer
    consulted (Fév. 2026 cleanup — those files are now obsolete).
    """
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    cursus_courses = await db.courses.find(
        {'cursus_id': cursus_id, 'r2_prefix': {'$exists': True, '$nin': [None, '']}},
        {'_id': 0, 'id': 1, 'r2_prefix': 1, 'title': 1},
    ).to_list(100)
    if not cursus_courses:
        return {'timelines': [], 'count': 0}
    cursus_root = _derive_cursus_r2_root(cursus_courses)
    items = await _collect_frise_files_for_cursus(cursus_root, cursus_courses)
    return {'timelines': items, 'count': len(items)}


@api_router.get("/courses/{course_id}/frises")
async def get_course_frises(course_id: str):
    """List frises (chronological maps, HTML) for a specific course.

    Convention (Fév. 2026):
    - `<cursus_root>/(sijill_)?map_*.html` → cursus-level frises shared by all
      courses of the cursus (e.g. `cursus-histoire/map_cursus_a_histoire.html`).
    - `<course_r2_prefix>/(sijill_)?map_*.html` → course-specific frises
      (e.g. `cursus-histoire/andalous/sijill_map_andalous.html`).

    The cursus root is derived from the first path segment of the course's
    `r2_prefix` (so no extra DB field needed on the cursus). Each entry exposes
    a signed URL via `/timeline/file/...?t=...` so the front-end iframe works
    with the existing access-control plumbing.
    """
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    course = await db.courses.find_one({'id': course_id}, {'_id': 0, 'id': 1, 'r2_prefix': 1, 'cursus_id': 1, 'title': 1})
    if not course or not (course.get('r2_prefix') or '').strip():
        return {'frises': [], 'count': 0}
    cursus_courses = await db.courses.find(
        {'cursus_id': course['cursus_id'], 'r2_prefix': {'$exists': True, '$nin': [None, '']}},
        {'_id': 0, 'id': 1, 'r2_prefix': 1, 'title': 1},
    ).to_list(100)
    cursus_root = _derive_cursus_r2_root(cursus_courses or [course])
    items = await _collect_frise_files_for_course(cursus_root, course)
    return {'frises': items, 'count': len(items)}


def _derive_cursus_r2_root(cursus_courses: List[dict]) -> str:
    """Return the cursus folder prefix on R2 (always ends with '/').

    Heuristic: first path segment of any non-empty course r2_prefix. The
    convention is one cursus = one top-level folder (`cursus-histoire/`,
    `cursus-a-falsafa/`, etc.).
    """
    for c in cursus_courses or []:
        p = (c.get('r2_prefix') or '').strip().strip('/')
        if p:
            return p.split('/', 1)[0] + '/'
    return ''


async def _collect_frise_files_for_cursus(cursus_root: str, cursus_courses: List[dict]) -> List[dict]:
    """Aggregate cursus-level frise + every course-specific frise of the cursus."""
    out: List[dict] = []
    seen = set()
    # 1. Cursus-level frises (immediate children of cursus_root only)
    out.extend(await _list_frises_at_root(cursus_root, scope='cursus', seen=seen))
    # 2. Per-course frises
    for c in cursus_courses:
        prefix = (c.get('r2_prefix') or '').strip()
        if not prefix:
            continue
        if not prefix.endswith('/'):
            prefix += '/'
        out.extend(await _list_frises_at_root(prefix, scope='course', seen=seen, course=c))
    return out


async def _collect_frise_files_for_course(cursus_root: str, course: dict) -> List[dict]:
    """Frises shown on a single course detail page: the cursus-shared frise(s)
    first, then the course-specific frises."""
    out: List[dict] = []
    seen = set()
    out.extend(await _list_frises_at_root(cursus_root, scope='cursus', seen=seen))
    prefix = (course.get('r2_prefix') or '').strip()
    if prefix and not prefix.endswith('/'):
        prefix += '/'
    if prefix:
        out.extend(await _list_frises_at_root(prefix, scope='course', seen=seen, course=course))
    return out


async def _list_frises_at_root(prefix: str, *, scope: str, seen: set,
                                course: Optional[dict] = None) -> List[dict]:
    """List `(sijill_)?map_*.html` at the IMMEDIATE level of `prefix`
    (no recursion). Returns normalised frise dicts. The display title is
    extracted from each HTML's `<title>` tag (falls back to a cleaned-up
    filename if the file has no title or the fetch fails).
    """
    import re
    if not prefix:
        return []
    if not prefix.endswith('/'):
        prefix += '/'
    try:
        resp = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix=prefix, Delimiter='/', MaxKeys=200)
    except Exception:
        return []
    pattern = re.compile(r'^(?:sijill_)?map_.*\.html$', re.IGNORECASE)
    out = []
    for obj in resp.get('Contents', []):
        key = obj['Key']
        if key in seen:
            continue
        rel = key[len(prefix):]
        if not rel or '/' in rel:
            continue
        if not pattern.match(rel):
            continue
        seen.add(key)
        stem = rel.rsplit('.', 1)[0]
        # Try to pull the human title from the HTML <title> tag (cached)
        html_title = await _extract_html_title(key, obj.get('LastModified'))
        if html_title:
            title = html_title
        else:
            m = re.match(r'^(?:sijill_)?map[_-](.+)$', stem, re.IGNORECASE)
            slug = m.group(1) if m else stem
            label = slug.replace('_', ' ').replace('-', ' ').strip()
            label = ' '.join(w.capitalize() for w in label.split()) or rel
            title = f"Frise — {label}"
        entry = {
            'id': stem.lower().replace('_', '-'),
            'filename': rel,
            'r2_key': key,
            'title': title,
            'scope': scope,
            'url': f'/api/timeline/file/{rel}?key={key}',
        }
        if course:
            entry.update({'course_id': course.get('id'), 'course_title': course.get('title')})
        out.append(entry)
    return out


# In-memory cache keyed by (r2_key, last_modified.isoformat()).
# Avoids re-fetching the HTML on every list request — invalidated automatically
# when the file is re-uploaded to R2 (LastModified changes).
_FRISE_TITLE_CACHE: dict = {}


async def _extract_html_title(r2_key: str, last_modified=None) -> Optional[str]:
    """Fetch the first 4 KB of an R2 HTML file and extract its `<title>` tag.

    Returns the trimmed title or `None` if not found / on error.
    """
    cache_key = (r2_key, last_modified.isoformat() if last_modified else '')
    if cache_key in _FRISE_TITLE_CACHE:
        return _FRISE_TITLE_CACHE[cache_key]
    title: Optional[str] = None
    try:
        # Range request — pull just enough to contain <title>...</title>
        resp = r2_client.get_object(Bucket=R2_BUCKET, Key=r2_key, Range='bytes=0-4096')
        head = resp['Body'].read().decode('utf-8', errors='ignore')
        import re
        m = re.search(r'<title[^>]*>(.*?)</title>', head, re.IGNORECASE | re.DOTALL)
        if m:
            title = m.group(1).strip()
            # Collapse any internal whitespace / newlines
            title = re.sub(r'\s+', ' ', title)
            # If empty after strip, treat as missing
            if not title:
                title = None
    except Exception:
        title = None
    _FRISE_TITLE_CACHE[cache_key] = title
    return title


@api_router.get("/timeline/file/{filename}")
async def get_timeline_by_filename(filename: str, request: Request, t: Optional[str] = None, key: Optional[str] = None):
    """Get timeline HTML by filename. Requires authentication + active subscription
    (header or signed token).

    When `key` is provided, it is the FULL r2_key (e.g. `cursus-histoire/andalous/sijill_map_andalous.html`),
    enabling the new course/cursus-scoped frises. Falls back to `Timeline/<filename>`
    for the legacy layout.
    """
    await verify_content_access(request, t)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")

    # Sanitize filename
    if not filename.endswith('.html'):
        filename = f"{filename}.html"

    # Decide which R2 key to fetch. If `key` is given AND it resolves to the
    # provided filename, use it; otherwise fall back to legacy Timeline/<filename>.
    if key and key.endswith(filename) and (key.startswith('cursus-') or key.startswith('Timeline/')):
        r2_key = key
    else:
        r2_key = f"Timeline/{filename}"

    try:
        response = r2_client.get_object(Bucket=R2_BUCKET, Key=r2_key)
        html_content = response['Body'].read().decode('utf-8')
        return HTMLResponse(content=html_content, media_type="text/html")
    except ClientError as e:
        error_code = e.response.get('Error', {}).get('Code', 'Unknown')
        if error_code == 'NoSuchKey':
            raise HTTPException(404, f"Timeline non trouvée: {filename}")
        raise HTTPException(500, f"Erreur R2: {str(e)}")

@api_router.put("/admin/resources/timeline/{resource_id:path}")
async def update_timeline_resource(resource_id: str, request: Request):
    """Update timeline resource metadata (title, display_order, cursus_letter)."""
    await require_admin(request)
    
    body = await request.json()
    
    # Support both formats: filename.html or resource-id without extension
    filename = resource_id if resource_id.endswith('.html') else resource_id.replace('-', '_') + '.html'
    
    update_data = {
        'filename': filename,
        'type': 'timeline',
        'updated_at': datetime.now(timezone.utc).isoformat()
    }
    
    if 'title' in body:
        update_data['title'] = body['title']
    if 'display_order' in body:
        update_data['display_order'] = int(body['display_order']) if body['display_order'] else 0
    if 'cursus_letter' in body:
        cursus = body['cursus_letter'].upper().strip() if body['cursus_letter'] else None
        if cursus and cursus not in ['A', 'B', 'C', 'D', 'E']:
            raise HTTPException(400, "Cursus invalide. Utilisez A, B, C, D ou E.")
        update_data['cursus_letter'] = cursus
    
    result = await db.timeline_resources.update_one(
        {'filename': filename, 'type': 'timeline'},
        {'$set': update_data},
        upsert=True
    )
    
    return {
        'message': 'Timeline mise à jour',
        'resource_id': resource_id,
        'filename': filename,
        'modified': result.modified_count > 0 or result.upserted_id is not None
    }

# ─── Context Resources (Word Documents) ────────────────────────────────────────

@api_router.get("/resources/context")
async def list_context_resources():
    """List all course-context (Word) resources across the catalog.

    NEW (fév. 2026): contextes désormais stockés DANS LE DOSSIER R2 DE CHAQUE COURS,
    avec un préfixe de nom `Contexte_*.docx`. Plus de scan global du dossier `Timeline/`.
    On boucle sur tous les cours qui ont un `r2_prefix` et on agrège leurs fichiers
    `Contexte_*.docx`. resource_id = `{course_id}__{slug}` (stable, unique, partageable).
    """
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    try:
        resources = await _collect_contexte_files()
        # Sort by course title then filename for stable display
        resources.sort(key=lambda x: (x.get('course_id', ''), x.get('filename', '')))
        return {'resources': resources, 'count': len(resources)}
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")


async def _collect_contexte_files(course_filter: Optional[str] = None,
                                  cursus_filter: Optional[str] = None) -> List[dict]:
    """Scan each course's r2_prefix for `Contexte_*.docx` files.

    Returns a normalised list of dicts (id, title, course_id, r2_key, etc.).
    Optionally filtered by a single `course_filter` (course_id) or `cursus_filter`.
    """
    import re
    query = {'r2_prefix': {'$exists': True, '$nin': [None, '']}}
    if course_filter:
        query['id'] = course_filter
    if cursus_filter:
        query['cursus_id'] = cursus_filter
    courses = await db.courses.find(query, {
        '_id': 0, 'id': 1, 'title': 1, 'cursus_id': 1, 'r2_prefix': 1
    }).to_list(200)

    out: List[dict] = []
    for c in courses:
        prefix = c.get('r2_prefix') or ''
        if not prefix.endswith('/'):
            prefix += '/'
        try:
            resp = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix=prefix, MaxKeys=300)
        except Exception:
            continue
        for obj in resp.get('Contents', []):
            key = obj['Key']
            filename = key.split('/')[-1]
            if filename.startswith('~$') or not filename.lower().endswith('.docx'):
                continue
            # Only files starting with "Contexte_" (case-insensitive)
            m = re.match(r'^Contexte_(.+)\.docx$', filename, re.IGNORECASE)
            if not m:
                continue
            slug = m.group(1)
            # Pretty subject: al-kindi → Al Kindi
            subject = slug.replace('-', ' ').replace('_', ' ').strip()
            subject = ' '.join(w.capitalize() for w in subject.split())
            resource_id = f"{c['id']}__{slug.lower()}"
            db_entry = await db.context_resources.find_one({'resource_id': resource_id}, {'_id': 0})
            out.append({
                'id': resource_id,
                'filename': filename,
                'r2_key': key,
                'course_id': c['id'],
                'course_title': c.get('title'),
                'cursus_id': c.get('cursus_id'),
                'subject': (db_entry or {}).get('subject') or subject,
                'title': (db_entry or {}).get('title') or subject,
                'description': (db_entry or {}).get('description', ''),
                'credits': (db_entry or {}).get('credits', ''),
                # module_number is kept for back-compat with the front-end's display
                'module_number': (db_entry or {}).get('module_number', 1),
                'size': obj['Size'],
                'url': f'/api/resources/context/{resource_id}',
            })
    return out


@api_router.get("/resources/context/cursus/{cursus_id}")
async def list_context_resources_by_cursus(cursus_id: str):
    """List context (Word) resources for all courses of a given cursus.

    Now derived live from each course's R2 `r2_prefix` (looking for files
    matching `Contexte_*.docx`). See `_collect_contexte_files`.
    """
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    try:
        resources = await _collect_contexte_files(cursus_filter=cursus_id)
        resources.sort(key=lambda x: (x.get('course_id', ''), x.get('filename', '')))
        return {'resources': resources, 'count': len(resources)}
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")


@api_router.get("/courses/{course_id}/context")
async def list_context_resources_by_course(course_id: str):
    """Course-scoped context resources — only the `Contexte_*.docx` files that
    live inside this course's R2 folder. Used by the course detail page's
    « Contexte » tab so it shows only the files relevant to this course."""
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    try:
        resources = await _collect_contexte_files(course_filter=course_id)
        resources.sort(key=lambda x: x.get('filename', ''))
        return {'resources': resources, 'count': len(resources)}
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")


@api_router.get("/resources/context/{resource_id}")
async def get_context_resource(resource_id: str, request: Request):
    """Return the parsed content of a context resource.

    Resource IDs follow the new convention `{course_id}__{slug}` (see
    `_collect_contexte_files`). For back-compat we still try a relaxed lookup
    so links generated before the rename keep working.
    """
    await require_subscriber(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")

    target_key = None
    target_entry = None
    try:
        # Preferred path: parse the new {course_id}__{slug} id directly.
        if '__' in resource_id:
            course_id, slug = resource_id.split('__', 1)
            for r in await _collect_contexte_files(course_filter=course_id):
                if r['id'] == resource_id:
                    target_key = r['r2_key']
                    target_entry = r
                    break
        # Fallback: scan everything (covers old links/manual /resources/context/<file>)
        if not target_key:
            for r in await _collect_contexte_files():
                if r['id'] == resource_id or r['filename'].replace('.docx', '').lower() == resource_id.lower():
                    target_key = r['r2_key']
                    target_entry = r
                    break

        if not target_key:
            raise HTTPException(404, f"Ressource non trouvée: {resource_id}")

        response = r2_client.get_object(Bucket=R2_BUCKET, Key=target_key)
        docx_content = response['Body'].read()

        from docx import Document as DocxDocument
        from io import BytesIO as BytesIOClass
        doc = DocxDocument(BytesIOClass(docx_content))

        # Parse document content
        content_blocks = []
        current_section = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else 'Normal'
            if 'Heading' in style or text.startswith('🏛') or text.startswith('🧠') or text.startswith('📅'):
                current_section = text
                content_blocks.append({'type': 'heading', 'text': text,
                                       'level': 1 if 'Heading 1' in style else 2})
            elif text.startswith('•') or text.startswith('-'):
                content_blocks.append({'type': 'list_item', 'text': text.lstrip('•- '),
                                       'section': current_section})
            else:
                content_blocks.append({'type': 'paragraph', 'text': text,
                                       'section': current_section})

        return {
            'id': resource_id,
            'title': target_entry.get('title') if target_entry else resource_id,
            'subject': target_entry.get('subject') if target_entry else resource_id,
            'description': target_entry.get('description', '') if target_entry else '',
            'credits': target_entry.get('credits', '') if target_entry else '',
            'course_id': target_entry.get('course_id') if target_entry else None,
            'cursus_id': target_entry.get('cursus_id') if target_entry else None,
            'module_number': target_entry.get('module_number', 1) if target_entry else 1,
            # Field name kept as `content` for frontend back-compat
            # (ResourceViewer.jsx reads `data.content`).
            'content': content_blocks,
        }
    except HTTPException:
        raise
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")
    except Exception as e:
        logger.exception(f"get_context_resource error: {e}")
        raise HTTPException(500, f"Erreur de lecture du document: {str(e)}")


# ─── Audio Resources (Conferences) ─────────────────────────────────────────────

@api_router.get("/resources/audio")
async def list_audio_resources(request: Request):
    """List all audio resources (conferences) from the audio folder.
    Strips stream_url for non-subscribers so they cannot bypass the paywall."""
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    import re
    
    try:
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix='audio/', MaxKeys=200)
        resources = []
        
        for obj in response.get('Contents', []):
            key = obj['Key']
            filename = key.split('/')[-1]
            
            # Skip non-audio files
            if not any(filename.lower().endswith(ext) for ext in ['.mp3', '.m4a', '.wav', '.ogg', '.aac']):
                continue
            
            # Parse filename: Conf_Averroes_Brenet_module4.m4a
            # Format: Conf_{Subject}_{Speaker}_module{N}.ext
            match = re.match(r'Conf_([^_]+)_([^_]+)_module(\d+)\.(mp3|m4a|wav|ogg|aac)', filename, re.IGNORECASE)
            
            if match:
                subject = match.group(1).replace('-', ' ')
                speaker = match.group(2).replace('-', ' ')
                module_num = int(match.group(3))
                ext = match.group(4).lower()
            else:
                # Fallback: try to extract any useful info
                subject = filename.replace('_', ' ').rsplit('.', 1)[0]
                speaker = ''
                module_num = 0
                ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
            
            resource_id = filename.rsplit('.', 1)[0].lower().replace(' ', '-').replace('_', '-')
            
            # Check for custom data in DB
            db_entry = await db.audio_resources.find_one({'resource_id': resource_id}, {'_id': 0})
            
            resource_data = {
                'id': resource_id,
                'filename': filename,
                'r2_key': key,
                'subject': db_entry.get('subject', subject) if db_entry else subject,
                'speaker': db_entry.get('speaker', speaker) if db_entry else speaker,
                'module_number': db_entry.get('module_number', module_num) if db_entry else module_num,
                'cursus_letter': db_entry.get('cursus_letter') if db_entry else None,
                'title': db_entry.get('title', f"Conférence : {subject}" + (f" par {speaker}" if speaker else "")) if db_entry else f"Conférence : {subject}" + (f" par {speaker}" if speaker else ""),
                'description': db_entry.get('description', '') if db_entry else '',
                'credits': db_entry.get('credits', '') if db_entry else '',
                'size': obj['Size'],
                'size_mb': round(obj['Size'] / (1024*1024), 1),
                'format': ext,
                'stream_url': f'/api/resources/audio/stream/{filename}'
            }
            resources.append(resource_data)
        
        # Sort by module number then subject
        resources.sort(key=lambda x: (x['module_number'], x['subject']))
        # Strip stream_url for non-subscribers (paywall)
        user = await get_current_user(request)
        is_subscriber = False
        if user:
            if user.get('role') == 'admin' or user.get('free_access'):
                is_subscriber = True
            else:
                now = datetime.now(timezone.utc)
                for fld in ('trial', 'subscription'):
                    v = user.get(fld)
                    if v and v.get('expires_at'):
                        exp = v['expires_at']
                        if isinstance(exp, str):
                            exp = datetime.fromisoformat(exp.replace('Z', '+00:00'))
                        if exp > now:
                            is_subscriber = True
                            break
        if not is_subscriber:
            for r in resources:
                r.pop('stream_url', None)
                r['locked'] = True
        return {'resources': resources, 'count': len(resources)}
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

@api_router.get("/resources/audio/{filename}/access-url")
async def get_audio_resource_access_url(filename: str, request: Request):
    """Issue a short-lived signed URL to stream an audio resource (conférence).
    Requires authentication + active subscription/trial/admin/free_access."""
    user = await require_subscriber(request)
    token = create_jwt({
        'sub': user['user_id'],
        'filename': filename,
        'scope': 'audio_resource_stream',
        'exp': int((datetime.now(timezone.utc) + timedelta(hours=1)).timestamp()),
    })
    scheme = request.headers.get('x-forwarded-proto', 'https')
    host = request.headers.get('x-forwarded-host') or request.headers.get('host', '')
    return {
        'url': f"{scheme}://{host}/api/resources/audio/stream/{filename}?t={token}",
        'expires_in': 3600,
    }

@api_router.get("/resources/audio/stream/{filename}")
async def stream_audio_resource(filename: str, request: Request, t: Optional[str] = None):
    """Stream an audio resource file from R2.
    Requires either a valid signed token (?t=, scope=audio_resource_stream) or Authorization header with active subscription."""
    # Authentication gate
    token = t or request.query_params.get('t')
    if token:
        payload = verify_jwt(token)
        if not payload or payload.get('scope') != 'audio_resource_stream' or payload.get('filename') != filename:
            raise HTTPException(403, "Jeton invalide ou expiré")
    else:
        await require_subscriber(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    r2_key = f"audio/{filename}"
    
    try:
        # Get file metadata first
        head = r2_client.head_object(Bucket=R2_BUCKET, Key=r2_key)
        file_size = head['ContentLength']
        content_type = head.get('ContentType', 'audio/mpeg')
        
        # Determine content type from extension
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        content_types = {
            'mp3': 'audio/mpeg',
            'm4a': 'audio/mp4',
            'wav': 'audio/wav',
            'ogg': 'audio/ogg',
            'aac': 'audio/aac'
        }
        content_type = content_types.get(ext, 'audio/mpeg')
        
        # Handle range requests for seeking
        range_header = request.headers.get('range')
        
        if range_header:
            # Parse range header
            import re
            range_match = re.match(r'bytes=(\d+)-(\d*)', range_header)
            if range_match:
                start = int(range_match.group(1))
                end = int(range_match.group(2)) if range_match.group(2) else file_size - 1
                
                # Fetch partial content
                response = r2_client.get_object(
                    Bucket=R2_BUCKET, 
                    Key=r2_key,
                    Range=f'bytes={start}-{end}'
                )
                
                return Response(
                    content=response['Body'].read(),
                    status_code=206,
                    headers={
                        'Content-Type': content_type,
                        'Content-Range': f'bytes {start}-{end}/{file_size}',
                        'Accept-Ranges': 'bytes',
                        'Content-Length': str(end - start + 1)
                    }
                )
        
        # Full file request
        response = r2_client.get_object(Bucket=R2_BUCKET, Key=r2_key)
        
        return Response(
            content=response['Body'].read(),
            headers={
                'Content-Type': content_type,
                'Accept-Ranges': 'bytes',
                'Content-Length': str(file_size)
            }
        )
        
    except ClientError as e:
        error_code = e.response.get('Error', {}).get('Code', 'Unknown')
        if error_code == 'NoSuchKey':
            raise HTTPException(404, f"Fichier audio non trouvé: {filename}")
        raise HTTPException(500, f"Erreur R2: {str(e)}")


# ─── Course Resources (Pilot Maïmonide) ────────────────────────────────────

@api_router.get("/courses/{course_id}/resources")
async def list_course_resources(course_id: str, request: Request):
    """List all resources attached to a course (course-level + per-episode).

    Combines two sources:
    1. **Explicit DB registrations** (`course.course_resources` + `audio.episode_resources`)
    2. **Auto-detection** from the course's R2 folder for *manuscripts* (PDF) and
       *images* (JPG/PNG). Lets the user simply drop a `*_manuscript*.pdf` or
       `*.jpg` next to the audio files and have it appear automatically, without
       having to register it in the admin. Files already registered in (1) are
       deduped (registration wins → keeps any custom label).

    Files explicitly excluded from auto-detection:
    - Audio (handled by /audios)
    - DOCX scripts (registered per episode in `episode_resources`)
    - `Contexte_*.docx` (served by the dedicated /context endpoint)

    Requires authentication + active subscription/trial/admin/free_access.
    """
    await require_subscriber(request)
    course = await db.courses.find_one({'id': course_id}, {'_id': 0})
    if not course:
        raise HTTPException(404, "Cours non trouvé")
    items = []
    seen_keys = set()

    for r in (course.get('course_resources') or []):
        if not r.get('r2_key') or r['r2_key'] in seen_keys:
            continue
        # Hide the legacy Contexte_*.docx registration — it now lives in the /context tab.
        fn = r['r2_key'].split('/')[-1]
        if fn.lower().startswith('contexte_'):
            seen_keys.add(r['r2_key'])
            continue
        items.append({
            'r2_key': r['r2_key'],
            'type': r.get('type'),
            'label': r.get('label') or fn,
            'mime': r.get('mime') or 'application/octet-stream',
            'scope': 'course',
        })
        seen_keys.add(r['r2_key'])
    audios_cur = db.audios.find(
        {'course_id': course_id, 'is_active': True, 'episode_resources': {'$exists': True, '$ne': []}},
        {'_id': 0},
    )
    audios = await audios_cur.to_list(200)
    audios.sort(key=lambda a: (a.get('episode_number') or 0, a.get('order') or 0))
    for a in audios:
        for r in (a.get('episode_resources') or []):
            if not r.get('r2_key') or r['r2_key'] in seen_keys:
                continue
            fn = r['r2_key'].split('/')[-1]
            # Skip mis-tagged manuscripts/images that historically landed in
            # episode_resources as type=script — they'll be re-emitted below with
            # the correct type by the auto-detect block.
            if fn.lower().endswith(('.pdf', '.jpg', '.jpeg', '.png')) and r.get('type') == 'script':
                continue
            items.append({
                'r2_key': r['r2_key'],
                'type': r.get('type'),
                'label': r.get('label') or fn,
                'mime': r.get('mime') or 'application/octet-stream',
                'scope': 'episode',
                'audio_id': a['id'],
                'audio_title': a.get('title'),
                'episode_number': a.get('episode_number'),
            })
            seen_keys.add(r['r2_key'])

    # ─── Auto-detection from R2 ────────────────────────────────────────────
    # Pulls every supplementary asset that lives in the course's R2 folder
    # (PDFs = manuscripts, JPG/PNG = manuscripts/illustrations).
    auto_items = await _autodetect_course_assets(course, exclude_keys=seen_keys)
    items.extend(auto_items)

    return {'resources': items, 'count': len(items)}


async def _autodetect_course_assets(course: dict, exclude_keys: set) -> List[dict]:
    """Scan the course's R2 folder for manuscripts / images that should appear
    in the Resources tab without requiring DB registration.

    Heuristics for the label:
    - `kindi_manuscript_1.pdf` → "Manuscrit — kindi 1"
    - `Al-Kindi_-_Cambridge,_Trinity_College_Library,_..._MS_R.15.17_..._BEIC_11521297.jpg`
      → "Manuscrit — Cambridge, Trinity College Library, MS R.15.17 (BEIC 11521297)"
    - Anything else → filename without extension, underscores/dashes replaced by spaces.

    Heuristics for the episode binding (so the asset displays under the right
    « Épisode N » group): we look for `ep<N>` or `episode<N>` in the filename.
    Otherwise the asset becomes `scope='course'`.
    """
    if not r2_client:
        return []
    prefix = (course.get('r2_prefix') or '').strip()
    if not prefix:
        return []
    if not prefix.endswith('/'):
        prefix += '/'
    try:
        resp = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix=prefix, MaxKeys=500)
    except Exception:
        return []

    import re
    audios_by_ep = {}
    audios_cur = db.audios.find({'course_id': course['id']}, {'_id': 0, 'id': 1, 'title': 1, 'episode_number': 1})
    for a in await audios_cur.to_list(200):
        if a.get('episode_number'):
            audios_by_ep[int(a['episode_number'])] = a

    out: List[dict] = []
    for obj in resp.get('Contents', []):
        key = obj['Key']
        if key in exclude_keys:
            continue
        filename = key.split('/')[-1]
        lower = filename.lower()
        if filename.startswith('~$') or not filename:
            continue
        # Reject: audio + DOCX scripts + Contexte_*
        if lower.endswith(('.m4a', '.mp3', '.wav', '.ogg', '.aac')):
            continue
        if lower.startswith('contexte_'):
            continue
        if lower.startswith('script-') or lower.startswith('script_') or '/script-' in key.lower():
            continue
        # Only emit visual/textual assets
        if lower.endswith('.pdf'):
            mime, asset_type = 'application/pdf', 'manuscript'
        elif lower.endswith(('.jpg', '.jpeg')):
            mime, asset_type = 'image/jpeg', 'image'
        elif lower.endswith('.png'):
            mime, asset_type = 'image/png', 'image'
        elif lower.startswith('bibliographie_') and lower.endswith('.docx'):
            mime, asset_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'bibliographie'
        elif lower.startswith('bibliographie_') and lower.endswith('.pdf'):
            mime, asset_type = 'application/pdf', 'bibliographie'
        else:
            continue

        # Try to bind to an episode if the filename mentions one
        ep_num = None
        m = re.search(r'(?:ep|episode)[\s_-]*(\d+)', lower)
        if m:
            try:
                ep_num = int(m.group(1))
            except ValueError:
                ep_num = None

        # Build a human label.
        stem = filename.rsplit('.', 1)[0]
        label = stem.replace('_', ' ').replace('-', ' ').strip()
        # Drop boilerplate ("BEIC 11521297" id, weird dashes) and prepend prefix
        if asset_type == 'manuscript':
            label = f"Manuscrit — {label}"
        elif asset_type == 'image':
            label = label.capitalize() if label else filename
        elif asset_type == 'bibliographie':
            # `bibliographie_falsafa` → "Bibliographie — Falsafa"
            label = label[len('bibliographie '):].strip() if label.lower().startswith('bibliographie ') else label
            label = f"Bibliographie — {' '.join(w.capitalize() for w in label.split())}" if label else "Bibliographie"

        entry = {
            'r2_key': key,
            'type': asset_type,
            'label': label,
            'mime': mime,
            'auto_detected': True,
        }
        if ep_num and ep_num in audios_by_ep:
            entry.update({
                'scope': 'episode',
                'audio_id': audios_by_ep[ep_num]['id'],
                'audio_title': audios_by_ep[ep_num].get('title'),
                'episode_number': ep_num,
            })
        else:
            entry['scope'] = 'course'
        out.append(entry)

    # ─── Module-level bibliographies (shared across all courses of the module) ──
    # Scan the IMMEDIATE PARENT of the course's r2_prefix for `bibliographie_*`
    # files. Example: `cursus-a-falsafa/02-falsafa/al-kindi/` → also look at
    # `cursus-a-falsafa/02-falsafa/` for `bibliographie_falsafa.docx` etc.
    # Without this layer the user would have to copy/paste the same biblio file
    # into every course folder of a module.
    parent_prefix = prefix.rstrip('/').rsplit('/', 1)[0]
    if parent_prefix:
        parent_prefix = parent_prefix + '/'
        try:
            presp = r2_client.list_objects_v2(
                Bucket=R2_BUCKET, Prefix=parent_prefix, Delimiter='/', MaxKeys=200,
            )
        except Exception:
            presp = {'Contents': []}
        for obj in presp.get('Contents', []):
            key = obj['Key']
            if key in exclude_keys:
                continue
            rel = key[len(parent_prefix):]
            if not rel or '/' in rel:  # only immediate children
                continue
            lower = rel.lower()
            if not lower.startswith('bibliographie_'):
                continue
            if lower.endswith('.docx'):
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif lower.endswith('.pdf'):
                mime = 'application/pdf'
            else:
                continue
            stem = rel.rsplit('.', 1)[0]
            tail = stem[len('bibliographie_'):].replace('_', ' ').replace('-', ' ').strip()
            tail = ' '.join(w.capitalize() for w in tail.split()) if tail else 'Module'
            out.append({
                'r2_key': key,
                'type': 'bibliographie',
                'label': f"Bibliographie — {tail}",
                'mime': mime,
                'auto_detected': True,
                'scope': 'module',
            })

    return out


@api_router.post("/courses/{course_id}/resource-access-url")
async def get_course_resource_access_url(course_id: str, body: dict, request: Request):
    """Issue a signed URL to fetch a course-scoped R2 resource (PDF/DOCX).
    The r2_key MUST be whitelisted in this course's resources."""
    user = await require_subscriber(request)
    r2_key = (body or {}).get('r2_key')
    if not r2_key or not isinstance(r2_key, str):
        raise HTTPException(400, "r2_key requis")
    course = await db.courses.find_one({'id': course_id}, {'_id': 0, 'course_resources': 1, 'r2_prefix': 1})
    found = None
    for r in (course.get('course_resources') or []) if course else []:
        if r.get('r2_key') == r2_key:
            found = r
            break
    if not found:
        audio = await db.audios.find_one(
            {'course_id': course_id, 'episode_resources.r2_key': r2_key},
            {'_id': 0, 'episode_resources': 1},
        )
        for r in (audio.get('episode_resources') or []) if audio else []:
            if r.get('r2_key') == r2_key:
                found = r
                break
    if not found:
        # Auto-detected file: not registered in DB but lives inside the course's
        # R2 folder (manuscript PDF / image). Allow as long as the key is
        # genuinely under that prefix → no privilege escalation.
        prefix = (course.get('r2_prefix') or '').strip() if course else ''
        if prefix and not prefix.endswith('/'):
            prefix += '/'
        # Also allow files at the IMMEDIATE PARENT level (= module-shared
        # bibliographies), but only if they match the `bibliographie_*` naming
        # convention. Anything else at the parent level is rejected to avoid
        # leaking sibling-course content.
        parent_prefix = prefix.rstrip('/').rsplit('/', 1)[0] + '/' if prefix else ''
        fn = r2_key.split('/')[-1].lower()
        in_course = bool(prefix and r2_key.startswith(prefix))
        in_module = bool(
            parent_prefix
            and r2_key.startswith(parent_prefix)
            and '/' not in r2_key[len(parent_prefix):]
            and fn.startswith('bibliographie_')
            and fn.endswith(('.docx', '.pdf'))
        )
        if in_course or in_module:
            mime_guess = (
                'application/pdf' if fn.endswith('.pdf') else
                'image/jpeg' if fn.endswith(('.jpg', '.jpeg')) else
                'image/png' if fn.endswith('.png') else
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if fn.endswith('.docx') else
                None
            )
            if mime_guess:
                found = {'r2_key': r2_key, 'mime': mime_guess, 'label': r2_key.split('/')[-1]}
    if not found:
        raise HTTPException(404, "Ressource non rattachée à ce cours")
    mime = found.get('mime') or 'application/octet-stream'
    token = create_jwt({
        'sub': user['user_id'],
        'r2_key': r2_key,
        'mime': mime,
        'scope': 'course_resource',
        'exp': int((datetime.now(timezone.utc) + timedelta(hours=1)).timestamp()),
    })
    scheme = request.headers.get('x-forwarded-proto', 'https')
    host = request.headers.get('x-forwarded-host') or request.headers.get('host', '')
    is_docx = 'wordprocessingml' in mime or mime == 'application/msword'
    return {
        'url': f"{scheme}://{host}/api/files/r2-stream?t={token}",
        'html_url': (f"{scheme}://{host}/api/files/r2-html?t={token}" if is_docx else None),
        'mime': mime,
        'label': found.get('label'),
        'expires_in': 3600,
    }


@api_router.api_route("/files/r2-stream", methods=["GET", "HEAD"])
async def stream_r2_resource(request: Request, t: Optional[str] = None):
    """Stream a course-scoped R2 file (PDF/DOCX/MP3/MP4) using a signed token.
    The token must have scope=course_resource (issued by /resource-access-url
    or by /audios/{id}/audio-access-url for episode audio)."""
    token = t or request.query_params.get('t')
    if not token:
        raise HTTPException(401, "Jeton requis")
    payload = verify_jwt(token)
    if not payload or payload.get('scope') not in ('course_resource', 'episode_audio'):
        raise HTTPException(403, "Jeton invalide ou expiré")
    r2_key = payload.get('r2_key')
    if not r2_key:
        raise HTTPException(403, "Jeton invalide")
    mime = payload.get('mime') or 'application/octet-stream'
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")

    range_header = request.headers.get('Range')
    try:
        kwargs: dict = {'Bucket': R2_BUCKET, 'Key': r2_key}
        if range_header:
            kwargs['Range'] = range_header
        resp = r2_client.get_object(**kwargs)
        body_bytes = resp['Body'].read()
        headers = {
            'Content-Type': mime,
            'Accept-Ranges': 'bytes',
            'Cache-Control': 'private, max-age=3600',
            'Access-Control-Allow-Origin': '*',
        }
        if mime == 'application/pdf':
            headers['Content-Disposition'] = 'inline'
        if 'ContentLength' in resp:
            headers['Content-Length'] = str(resp['ContentLength'])
        if 'ContentRange' in resp:
            headers['Content-Range'] = resp['ContentRange']
        from fastapi.responses import Response
        status_code = 206 if range_header else 200
        return Response(content=body_bytes, status_code=status_code, headers=headers, media_type=mime)
    except ClientError as e:
        code = e.response.get('Error', {}).get('Code', '')
        if code in ('NoSuchKey', '404'):
            raise HTTPException(404, "Fichier non disponible")
        logging.getLogger(__name__).error(f"R2 stream error for {r2_key}: {e}")
        raise HTTPException(500, "Erreur de lecture du fichier")
    except Exception as e:
        logging.getLogger(__name__).error(f"R2 stream error for {r2_key}: {e}")
        raise HTTPException(500, "Erreur de lecture du fichier")


# In-memory cache for PDF → article rendering (course-scoped, lasts a process lifetime)
_pdf_article_cache: dict = {}

def _docx_to_text(docx_bytes: bytes) -> str:
    """Convert DOCX → structured plain text.

    Uses python-docx to inspect each paragraph's style. Headings get explicit
    markers (`[H1]`, `[H2]`, `[H3]`) so the downstream parser can render them
    as proper titles instead of mangling them into prose. Line breaks inside
    a paragraph (Shift+Enter in Word) are kept as `\u2028` so the front-end
    can render them as `<br/>` while paragraph breaks remain as `\n\n`.

    Falls back to mammoth (raw text only) if python-docx fails for any reason.
    """
    try:
        from docx import Document  # python-docx
        doc = Document(io.BytesIO(docx_bytes))
        out_lines: list[str] = []
        for para in doc.paragraphs:
            text = (para.text or '').strip()
            style_name = (para.style.name if para.style else '') or ''
            if not text:
                # Blank paragraph → keep one blank line (will become \n\n).
                out_lines.append('')
                continue
            # Detect heading level from the style name.
            lvl = None
            sn = style_name.lower()
            if sn.startswith('heading 1') or sn == 'title':
                lvl = 1
            elif sn.startswith('heading 2'):
                lvl = 2
            elif sn.startswith('heading 3') or sn.startswith('heading 4'):
                lvl = 3
            if lvl:
                out_lines.append(f"[H{lvl}]{text}")
            else:
                out_lines.append(text)
        return '\n\n'.join(out_lines)
    except Exception as e:
        logging.getLogger(__name__).warning(f"python-docx extract failed, falling back to mammoth: {e}")
    try:
        import mammoth
        result = mammoth.extract_raw_text(io.BytesIO(docx_bytes))
        return result.value or ""
    except Exception as e:
        logging.getLogger(__name__).warning(f"DOCX extract error: {e}")
        return ""


def _pdf_to_article(pdf_bytes: bytes, label: str, *, mime: str = 'application/pdf') -> dict:
    """Convert PDF or DOCX to a blog-style article structure {title, lead, sections}.
    Robust to glossary-style content (Term : definition spanning multiple lines)
    and to running prose. Adds bold on glossary term entries automatically.
    Each paragraph has shape "TERM\u00a0: definition" or plain prose.
    """
    is_docx = ('wordprocessingml' in mime) or (mime == 'application/msword')
    if is_docx:
        text = _docx_to_text(pdf_bytes)
    else:
        try:
            from pdfminer.high_level import extract_text
        except Exception as e:
            raise HTTPException(500, f"PDF parser indisponible: {e}")
        text = extract_text(io.BytesIO(pdf_bytes)) or ""
    text = (text
            .replace('\ufb01', 'fi').replace('\ufb02', 'fl')
            .replace('\u00a0', ' ')
            .replace('\r', '\n'))

    # ── New path (DOCX with explicit heading markers) ──────────────────────
    # When the input was a DOCX processed by _docx_to_text, each Word
    # paragraph is on its own line, separated by a blank line, and headings
    # are prefixed with [H1]/[H2]/[H3]. Parse that structure directly — it
    # is far more accurate than the heuristic block-merging used for PDF.
    if is_docx and ('[H1]' in text or '[H2]' in text or '[H3]' in text or '\n\n' in text):
        sections_docx: list = []
        lead_docx: str | None = None
        current_sec = {'title': None, 'paragraphs': [], 'level': 2}
        # Split on the blank lines we inserted between paragraphs.
        raw_paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
        for raw in raw_paragraphs:
            m = re.match(r'^\[H([123])\](.+)$', raw, flags=re.DOTALL)
            if m:
                level = int(m.group(1))
                heading_text = m.group(2).strip()
                if level == 1:
                    # H1 is treated as the article main title — only used to
                    # close any open section and start a new one (the article
                    # title is set by the caller from the resource metadata).
                    if current_sec['title'] or current_sec['paragraphs']:
                        sections_docx.append(current_sec)
                    current_sec = {'title': heading_text, 'paragraphs': [], 'level': 1}
                elif level == 2:
                    if current_sec['title'] or current_sec['paragraphs']:
                        sections_docx.append(current_sec)
                    current_sec = {'title': heading_text, 'paragraphs': [], 'level': 2}
                else:  # H3 → inline sub-heading kept as a paragraph with marker
                    current_sec['paragraphs'].append(f'[H3]{heading_text}')
            else:
                current_sec['paragraphs'].append(raw)
        if current_sec['title'] or current_sec['paragraphs']:
            sections_docx.append(current_sec)
        # Promote first paragraph to lead if it reads like an introduction.
        if sections_docx and sections_docx[0]['paragraphs']:
            first = sections_docx[0]['paragraphs'][0]
            if not first.startswith('[H') and 60 <= len(first) <= 800 and first[:1].isupper():
                lead_docx = first
                sections_docx[0]['paragraphs'] = sections_docx[0]['paragraphs'][1:]
        # Strip leading empty sections.
        while sections_docx and not sections_docx[0]['title'] and not sections_docx[0]['paragraphs']:
            sections_docx.pop(0)
        return {
            'title': label,
            'lead': lead_docx,
            'sections': sections_docx,
            'word_count': sum(len(p.split()) for s in sections_docx for p in s['paragraphs']) + (len((lead_docx or '').split())),
        }
    # ── End new DOCX path ──────────────────────────────────────────────────

    # Split into raw lines, normalise inner whitespace
    raw_lines = [re.sub(r'\s+', ' ', ln).strip() for ln in text.split('\n')]

    # ── Strip repeated headers/footers (PDF page boilerplate) ──────────────
    # PDF exports often carry an author byline like "By Yannis Mahil" or a
    # running title at the top/bottom of every page. pdfminer extracts these
    # verbatim, polluting the article body. Heuristic: any short line
    # (≤80 chars, ≥4 chars) that appears 3+ times across the document AND
    # is not glossary-shaped is treated as boilerplate and silently dropped.
    from collections import Counter as _Counter
    line_counts = _Counter(ln for ln in raw_lines if 4 <= len(ln) <= 80)
    boilerplate = {
        ln for ln, n in line_counts.items()
        if n >= 3
        and ':' not in ln[:60]  # spare glossary "Term : def" entries
        and not re.match(r'^\d+\s*[\.\)]', ln)  # spare numbered headings
    }
    if boilerplate:
        logging.getLogger(__name__).info(
            f"PDF/DOCX parser: stripped {len(boilerplate)} repeated header/footer line(s)"
        )
        raw_lines = [ln for ln in raw_lines if ln not in boilerplate]
    # Always strip standalone page numbers like "Page 1", "1 / 24", "- 12 -"
    PAGE_NUM_RE = re.compile(r'^(?:Page\s+)?\d+(?:\s*[/–\-]\s*\d+)?$|^[-–—]\s*\d+\s*[-–—]$', re.IGNORECASE)
    raw_lines = [ln for ln in raw_lines if not PAGE_NUM_RE.match(ln)]
    # Re-derive the source `text` so the downstream block-merger stays
    # consistent with the cleaned lines (preserves blank-line counts).
    text = '\n'.join(raw_lines)
    # ── End boilerplate strip ──────────────────────────────────────────────

    # Build "blocks" by merging consecutive non-empty lines, BUT start a new block
    # whenever a line looks like a glossary term entry ("Terme : …") or like a heading
    # or whenever the previous line ended with a sentence terminator AND the new line
    # starts with a capitalised word (likely a new paragraph).
    GLOSS_RE = re.compile(r'^[A-ZÀ-ÝŒÇ][^.:;\n]{1,60}\s*:\s+\S')
    HEADING_RE = re.compile(r'^(I{1,3}V?|IV|VI{0,3}|IX|XI{0,3}|XV?I{0,3})\s*[\.\)]\s+\S|^\d+\s*[\.\)]\s+\S')
    TERMINAL = ('.', '!', '?', '»', '…', ';')

    def is_heading_block(s: str) -> bool:
        if len(s) < 4 or len(s) >= 90 or s.endswith('.') or s.endswith('…'):
            return False
        letters = [c for c in s if c.isalpha()]
        if not letters:
            return False
        upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
        return upper_ratio > 0.55 or bool(HEADING_RE.match(s))

    blocks: list = []
    current_lines: list = []

    def flush_block():
        nonlocal current_lines
        if current_lines:
            blocks.append(' '.join(current_lines).strip())
            current_lines = []

    # Count consecutive blank lines so we can distinguish a soft visual line break
    # (single \n\n from pdfminer) from an actual paragraph break (3+ newlines or
    # 2 blank lines).
    text_lines = text.split('\n')
    i = 0
    blank_run = 0
    while i < len(text_lines):
        ln_raw = text_lines[i]
        ln = re.sub(r'\s+', ' ', ln_raw).strip()
        if not ln:
            blank_run += 1
            i += 1
            continue
        # Decide based on what we just saw
        if not current_lines:
            current_lines.append(ln)
        else:
            prev = current_lines[-1]
            hard_break = blank_run >= 2  # 2+ blank lines = real paragraph break
            starts_new = False
            if hard_break:
                starts_new = True
            elif GLOSS_RE.match(ln):
                starts_new = True
            elif is_heading_block(ln):
                starts_new = True
            elif prev.endswith(TERMINAL) and ln[:1].isalpha() and ln[:1].isupper():
                # Previous sentence finished AND new line starts a new sentence
                starts_new = True
            if starts_new:
                flush_block()
                current_lines.append(ln)
            else:
                current_lines.append(ln)
        blank_run = 0
        i += 1
    flush_block()

    blocks = [b for b in blocks if b]

    # Group blocks into sections by detecting headings between them
    sections: list = []
    current = {'title': None, 'paragraphs': []}
    def flush_sec():
        if current['title'] or current['paragraphs']:
            sections.append({'title': current['title'], 'paragraphs': list(current['paragraphs'])})
    for block in blocks:
        if is_heading_block(block):
            flush_sec()
            current = {'title': block, 'paragraphs': []}
        else:
            current['paragraphs'].append(block)
    flush_sec()

    # Lead = first paragraph if it looks like an introduction (60–600 chars, sentence-shaped)
    lead = None
    if sections and sections[0]['paragraphs']:
        first = sections[0]['paragraphs'][0]
        if 60 < len(first) < 600 and ('. ' in first or first.endswith('.')):
            lead = first
            sections[0]['paragraphs'] = sections[0]['paragraphs'][1:]
    while sections and not sections[0]['title'] and not sections[0]['paragraphs']:
        sections.pop(0)

    return {
        'title': label,
        'lead': lead,
        'sections': sections,
        'word_count': sum(len(p.split()) for s in sections for p in s['paragraphs']) + (len((lead or '').split())),
    }


@api_router.get("/courses/{course_id}/resource-article")
async def get_course_resource_article(course_id: str, r2_key: str, request: Request):
    """Return the PDF resource rendered as a blog-style article (no direct PDF access).
    Whitelisted to course/episode resources only. Requires active subscription."""
    user = await require_subscriber(request)
    # Validate r2_key membership
    course = await db.courses.find_one({'id': course_id}, {'_id': 0, 'course_resources': 1})
    found = None
    scope = None
    audio_title = None
    episode_number = None
    for r in (course.get('course_resources') or []) if course else []:
        if r.get('r2_key') == r2_key:
            found = r; scope = 'course'; break
    if not found:
        audio = await db.audios.find_one(
            {'course_id': course_id, 'episode_resources.r2_key': r2_key},
            {'_id': 0, 'episode_resources': 1, 'title': 1, 'episode_number': 1},
        )
        for r in (audio.get('episode_resources') or []) if audio else []:
            if r.get('r2_key') == r2_key:
                found = r; scope = 'episode'
                audio_title = audio.get('title')
                episode_number = audio.get('episode_number')
                break
    if not found:
        # Auto-detected docs (bibliographies module-level, Contexte_, manuscripts).
        # Accept if the key lies within the course's r2_prefix OR at the
        # immediate parent (module-shared bibliographies only).
        course = await db.courses.find_one({'id': course_id}, {'_id': 0, 'r2_prefix': 1})
        prefix = (course.get('r2_prefix') or '').strip() if course else ''
        if prefix and not prefix.endswith('/'):
            prefix += '/'
        parent_prefix = prefix.rstrip('/').rsplit('/', 1)[0] + '/' if prefix else ''
        fn = r2_key.split('/')[-1].lower()
        in_course = bool(prefix and r2_key.startswith(prefix))
        in_module = bool(
            parent_prefix
            and r2_key.startswith(parent_prefix)
            and '/' not in r2_key[len(parent_prefix):]
            and fn.startswith('bibliographie_')
            and fn.endswith(('.docx', '.pdf'))
        )
        if in_course or in_module:
            if fn.endswith('.pdf'):
                mg = 'application/pdf'
            elif fn.endswith('.docx'):
                mg = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                mg = None
            if mg:
                found = {'r2_key': r2_key, 'mime': mg, 'label': r2_key.split('/')[-1]}
                scope = 'module' if in_module else 'course'
    if not found:
        raise HTTPException(404, "Ressource non rattachée à ce cours")

    mime = found.get('mime') or ''
    res_type = found.get('type') or ''
    # Slides → not converted; client must call /resource-access-url for inline PDF viewer.
    if res_type == 'slides':
        raise HTTPException(400, "Les slides sont rendus en PDF inline (utilisez /resource-access-url)")
    is_pdf = 'pdf' in mime
    is_docx = ('wordprocessingml' in mime) or (mime == 'application/msword')
    if not (is_pdf or is_docx):
        raise HTTPException(400, "Conversion article disponible pour PDF et DOCX uniquement")
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    cache_key = f"{course_id}::{r2_key}::{mime}"
    if cache_key not in _pdf_article_cache:
        try:
            obj = r2_client.get_object(Bucket=R2_BUCKET, Key=r2_key)
            data = obj['Body'].read()
            _pdf_article_cache[cache_key] = _pdf_to_article(data, found.get('label') or r2_key.split('/')[-1], mime=mime)
        except ClientError as e:
            code = e.response.get('Error', {}).get('Code', '')
            if code in ('NoSuchKey', '404'):
                raise HTTPException(404, "Document non disponible")
            logging.getLogger(__name__).error(f"R2 fetch error for article {r2_key}: {e}")
            raise HTTPException(500, "Erreur de lecture du document")
        except HTTPException:
            raise
        except Exception as e:
            logging.getLogger(__name__).error(f"PDF→article error for {r2_key}: {e}")
            raise HTTPException(500, "Erreur de conversion du document")
    article = dict(_pdf_article_cache[cache_key])
    article.update({
        'r2_key': r2_key,
        'course_id': course_id,
        'scope': scope,
        'type': found.get('type'),
        'audio_title': audio_title,
        'episode_number': episode_number,
        'course_title': course.get('title') if course else None,
    })
    return article


# --- PROTECTED PDF DOWNLOAD (with user watermark) -----------------------------
def _build_protected_pdf(article: dict, user_name: str, user_email: str) -> bytes:
    """Render the article (script/glossaire/biblio) as a Sijill Prestige-styled PDF
    with diagonal watermark + footer carrying the subscriber's name/email on every page."""
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor, Color
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageTemplate, Paragraph,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.fonts import addMapping

    # Register Unicode-aware fonts once (EB Garamond for body, DejaVu Sans for UI/UI text).
    fonts_dir = Path(__file__).parent / 'fonts'
    if 'EBGaramond' not in pdfmetrics.getRegisteredFontNames():
        try:
            pdfmetrics.registerFont(TTFont('EBGaramond', str(fonts_dir / 'EBGaramond-Regular.ttf')))
            pdfmetrics.registerFont(TTFont('EBGaramond-Bold', str(fonts_dir / 'EBGaramond-Bold.ttf')))
            pdfmetrics.registerFont(TTFont('EBGaramond-Italic', str(fonts_dir / 'EBGaramond-Italic.ttf')))
            addMapping('EBGaramond', 0, 0, 'EBGaramond')
            addMapping('EBGaramond', 1, 0, 'EBGaramond-Bold')
            addMapping('EBGaramond', 0, 1, 'EBGaramond-Italic')
            addMapping('EBGaramond', 1, 1, 'EBGaramond-Bold')
            pdfmetrics.registerFont(TTFont('DejaVuSans', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
            addMapping('DejaVuSans', 0, 0, 'DejaVuSans')
            addMapping('DejaVuSans', 1, 0, 'DejaVuSans-Bold')
            addMapping('DejaVuSans', 0, 1, 'DejaVuSans')
        except Exception as e:
            logging.getLogger(__name__).warning(f"Font registration failed, falling back: {e}")

    CREAM = HexColor('#F4EDE0')
    INK = HexColor('#1A1611')
    GREEN = HexColor('#1FAE6B')
    MUTED = HexColor('#6B5E4A')

    type_labels = {
        'script': "Script de l'épisode",
        'glossaire': "Glossaire",
        'biblio': "Bibliographie",
        'bibliographie': "Bibliographie",
        'document': "Document",
    }
    type_label = type_labels.get(article.get('type') or '', 'Document pédagogique')
    title = article.get('title') or 'Document'
    course_title = article.get('course_title') or ''
    audio_title = article.get('audio_title') or ''
    sections = article.get('sections') or []
    is_glossary = (article.get('type') == 'glossaire')

    buffer = BytesIO()

    def _draw_page_chrome(canv, doc):
        page_w, page_h = A4
        # Cream background
        canv.setFillColor(CREAM)
        canv.rect(0, 0, page_w, page_h, stroke=0, fill=1)
        # Diagonal watermark with subscriber name
        canv.saveState()
        canv.setFillColor(Color(0.35, 0.30, 0.20, alpha=0.10))
        canv.setFont('DejaVuSans-Bold', 36)
        canv.translate(page_w / 2, page_h / 2)
        canv.rotate(-30)
        wm = (user_name or user_email or 'Sijill').upper()
        for dy in (-200, -60, 80, 220):
            canv.drawCentredString(0, dy, wm)
        canv.restoreState()
        # Header rule + meta
        canv.setStrokeColor(MUTED)
        canv.setLineWidth(0.4)
        canv.line(20 * mm, page_h - 18 * mm, page_w - 20 * mm, page_h - 18 * mm)
        canv.setFillColor(MUTED)
        canv.setFont('DejaVuSans', 8)
        canv.drawString(20 * mm, page_h - 14 * mm, 'SIJILL PROJECT')
        canv.drawRightString(page_w - 20 * mm, page_h - 14 * mm, type_label.upper())
        # Footer
        canv.line(20 * mm, 18 * mm, page_w - 20 * mm, 18 * mm)
        canv.setFont('DejaVuSans', 7.5)
        canv.drawString(20 * mm, 12 * mm, f"Document réservé · Lecture par {user_name} <{user_email}>")
        canv.drawRightString(page_w - 20 * mm, 12 * mm, f"Page {doc.page}")
        canv.setFont('EBGaramond-Italic', 7.5)
        canv.drawCentredString(page_w / 2, 7 * mm, "Reproduction interdite — usage strictement personnel")

    frame = Frame(
        22 * mm, 22 * mm,
        A4[0] - 44 * mm, A4[1] - 44 * mm,
        leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
        showBoundary=0,
    )
    doc = BaseDocTemplate(
        buffer, pagesize=A4,
        leftMargin=22 * mm, rightMargin=22 * mm,
        topMargin=22 * mm, bottomMargin=22 * mm,
        title=title, author='Sijill Project',
    )
    doc.addPageTemplates([PageTemplate(id='prestige', frames=[frame], onPage=_draw_page_chrome)])

    base = getSampleStyleSheet()
    s_pill = ParagraphStyle('pill', parent=base['BodyText'],
        textColor=GREEN, fontName='DejaVuSans-Bold', fontSize=8.5, leading=10,
        spaceAfter=10, alignment=TA_LEFT)
    s_title = ParagraphStyle('title', parent=base['Title'],
        textColor=INK, fontName='EBGaramond-Bold', fontSize=26, leading=30,
        spaceAfter=4, alignment=TA_LEFT)
    s_meta = ParagraphStyle('meta', parent=base['BodyText'],
        textColor=MUTED, fontName='EBGaramond-Italic', fontSize=11, leading=14,
        spaceAfter=14, alignment=TA_LEFT)
    s_lead = ParagraphStyle('lead', parent=base['BodyText'],
        textColor=INK, fontName='EBGaramond-Italic', fontSize=12, leading=18,
        spaceAfter=14, leftIndent=8, alignment=TA_LEFT)
    s_h2 = ParagraphStyle('h2', parent=base['Heading2'],
        textColor=GREEN, fontName='EBGaramond-Bold', fontSize=14, leading=18,
        spaceBefore=16, spaceAfter=8, alignment=TA_LEFT)
    s_p = ParagraphStyle('p', parent=base['BodyText'],
        textColor=INK, fontName='EBGaramond', fontSize=11, leading=17,
        spaceAfter=8, alignment=TA_LEFT, firstLineIndent=14)
    s_gloss = ParagraphStyle('gloss', parent=s_p,
        firstLineIndent=0, spaceAfter=6)

    GLOSS_RE = re.compile(r'^([A-ZÀ-ÝŒÇ][^.:;\n]{1,60}?)\s*:\s+(.*)$')

    def _esc(s: str) -> str:
        return (s or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    story = []
    pill_text = type_label.upper()
    if audio_title:
        pill_text += f"  ·  {_esc(audio_title)}"
    if article.get('word_count'):
        wc = article['word_count']
        rt = max(1, round(wc / 220))
        pill_text += f"  ·  {wc} mots · {rt} min"
    story.append(Paragraph(pill_text, s_pill))
    story.append(Paragraph(_esc(title), s_title))
    if course_title:
        story.append(Paragraph(_esc(course_title), s_meta))
    if article.get('lead'):
        story.append(Paragraph(_esc(article['lead']), s_lead))

    for sec in sections:
        sec_title = sec.get('title')
        paras = sec.get('paragraphs') or []
        if sec_title:
            story.append(Paragraph(_esc(sec_title), s_h2))
        for p in paras:
            text = (p or '').strip()
            if not text:
                continue
            if is_glossary:
                m = GLOSS_RE.match(text)
                if m:
                    term = _esc(m.group(1).strip())
                    rest = _esc(m.group(2))
                    story.append(Paragraph(f"<b>{term}</b> &mdash; {rest}", s_gloss))
                    continue
            story.append(Paragraph(_esc(text), s_p))

    if not story:
        story.append(Paragraph("Document vide.", s_p))

    doc.build(story)
    return buffer.getvalue()


@api_router.get("/courses/{course_id}/resource-pdf")
async def get_course_resource_pdf(course_id: str, r2_key: str, request: Request):
    """Return a watermarked PDF rendition of a script/glossaire/biblio.
    The PDF embeds the subscriber's name + email as watermark/footer."""
    user = await require_subscriber(request)
    course = await db.courses.find_one({'id': course_id}, {'_id': 0, 'course_resources': 1, 'title': 1})
    found = None
    scope = None
    audio_title = None
    episode_number = None
    for r in (course.get('course_resources') or []) if course else []:
        if r.get('r2_key') == r2_key:
            found = r; scope = 'course'; break
    if not found:
        audio = await db.audios.find_one(
            {'course_id': course_id, 'episode_resources.r2_key': r2_key},
            {'_id': 0, 'episode_resources': 1, 'title': 1, 'episode_number': 1},
        )
        for r in (audio.get('episode_resources') or []) if audio else []:
            if r.get('r2_key') == r2_key:
                found = r; scope = 'episode'
                audio_title = audio.get('title')
                episode_number = audio.get('episode_number')
                break
    if not found:
        raise HTTPException(404, "Ressource non rattachée à ce cours")

    mime = found.get('mime') or ''
    res_type = found.get('type') or ''
    if res_type == 'slides':
        raise HTTPException(400, "Les slides ne sont pas téléchargeables")
    is_pdf = 'pdf' in mime
    is_docx = ('wordprocessingml' in mime) or (mime == 'application/msword')
    if not (is_pdf or is_docx):
        raise HTTPException(400, "Téléchargement disponible pour PDF et DOCX uniquement")
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")

    cache_key = f"{course_id}::{r2_key}::{mime}"
    if cache_key not in _pdf_article_cache:
        try:
            obj = r2_client.get_object(Bucket=R2_BUCKET, Key=r2_key)
            data = obj['Body'].read()
            _pdf_article_cache[cache_key] = _pdf_to_article(data, found.get('label') or r2_key.split('/')[-1], mime=mime)
        except ClientError as e:
            code = e.response.get('Error', {}).get('Code', '')
            if code in ('NoSuchKey', '404'):
                raise HTTPException(404, "Document non disponible")
            raise HTTPException(500, "Erreur de lecture du document")
        except Exception as e:
            logging.getLogger(__name__).error(f"PDF→article error for {r2_key}: {e}")
            raise HTTPException(500, "Erreur de conversion du document")

    article = dict(_pdf_article_cache[cache_key])
    article.update({
        'type': found.get('type'),
        'audio_title': audio_title,
        'episode_number': episode_number,
        'course_title': course.get('title') if course else None,
    })

    user_name = user.get('name') or (user.get('email') or '').split('@')[0]
    user_email = user.get('email') or ''
    try:
        pdf_bytes = _build_protected_pdf(article, user_name, user_email)
    except Exception as e:
        logging.getLogger(__name__).error(f"PDF render error for {r2_key}: {e}")
        raise HTTPException(500, "Erreur de génération du PDF")

    safe_label = re.sub(r'[^A-Za-z0-9_.-]+', '_', (found.get('label') or article.get('title') or 'document'))
    filename = f"sijill-{safe_label}.pdf"
    return Response(
        content=pdf_bytes,
        media_type='application/pdf',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Cache-Control': 'no-store, no-cache, must-revalidate',
            'X-Robots-Tag': 'noindex, nofollow',
        },
    )






@api_router.get("/files/r2-html")
async def r2_resource_as_html(request: Request, t: Optional[str] = None):
    """Convert a Word (.docx) resource to clean HTML using mammoth.
    Requires the same signed token (scope=course_resource)."""
    token = t or request.query_params.get('t')
    if not token:
        raise HTTPException(401, "Jeton requis")
    payload = verify_jwt(token)
    if not payload or payload.get('scope') != 'course_resource':
        raise HTTPException(403, "Jeton invalide ou expiré")
    r2_key = payload.get('r2_key')
    mime = payload.get('mime') or ''
    if 'wordprocessingml' not in mime and mime != 'application/msword':
        raise HTTPException(400, "Conversion HTML disponible uniquement pour les documents Word")
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    try:
        import mammoth
        resp = r2_client.get_object(Bucket=R2_BUCKET, Key=r2_key)
        body_bytes = resp['Body'].read()
        result = mammoth.convert_to_html(io.BytesIO(body_bytes))
        return {
            'html': result.value,
            'messages': [str(m) for m in (result.messages or [])][:10],
        }
    except ClientError as e:
        code = e.response.get('Error', {}).get('Code', '')
        if code in ('NoSuchKey', '404'):
            raise HTTPException(404, "Fichier non disponible")
        raise HTTPException(500, "Erreur R2")
    except Exception as e:
        logging.getLogger(__name__).error(f"DOCX→HTML error: {e}")
        raise HTTPException(500, "Erreur de conversion du document Word")


@api_router.get("/audios/{audio_id}/audio-access-url")
async def get_episode_audio_access_url(audio_id: str, request: Request):
    """Issue a short-lived signed URL to stream the episode-level R2 audio (.mp3/.m4a).
    Distinct from /audios/{id}/stream-url (which uses file_key); this one uses r2_audio_key
    set on Maïmonide pilot. Requires authentication + active subscription."""
    user = await require_subscriber(request)
    a = await db.audios.find_one({'id': audio_id}, {'_id': 0})
    if not a:
        raise HTTPException(404, "Audio non trouvé")
    if not a.get('has_r2_audio') or not a.get('r2_audio_key'):
        raise HTTPException(404, "Audio podcast non disponible")
    # Per-audio access check
    access = await check_user_access(user['user_id'], content_type='audio', content_id=audio_id)
    if not access.get('has_access'):
        raise HTTPException(403, {"error": "subscription_required", "reason": access.get('reason', 'no_access')})
    r2_key = a['r2_audio_key']
    ext = (r2_key.rsplit('.', 1)[-1] or '').lower()
    mime_map = {
        'mp3': 'audio/mpeg', 'm4a': 'audio/mp4', 'wav': 'audio/wav',
        'ogg': 'audio/ogg', 'aac': 'audio/aac',
    }
    mime = mime_map.get(ext, 'audio/mpeg')
    token = create_jwt({
        'sub': user['user_id'],
        'r2_key': r2_key,
        'mime': mime,
        'audio_id': audio_id,
        'scope': 'episode_audio',
        'exp': int((datetime.now(timezone.utc) + timedelta(hours=1)).timestamp()),
    })
    scheme = request.headers.get('x-forwarded-proto', 'https')
    host = request.headers.get('x-forwarded-host') or request.headers.get('host', '')
    return {
        'audio_id': audio_id,
        'stream_url': f"{scheme}://{host}/api/files/r2-stream?t={token}",
        'mime': mime,
        'expires_in': 3600,
    }


@api_router.get("/admin/resources/timeline")
async def admin_list_timeline_resources(request: Request):
    """Admin: List all timeline resources (HTML + DOCX + Audio) for management."""
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    import re
    
    try:
        # Get Timeline folder files
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix='Timeline/', MaxKeys=200)
        
        html_files = []
        docx_files = []
        
        for obj in response.get('Contents', []):
            key = obj['Key']
            filename = key.split('/')[-1]
            
            # Skip temp files
            if filename.startswith('~$'):
                continue
            
            file_info = {
                'key': key,
                'filename': filename,
                'size': obj['Size'],
                'size_kb': round(obj['Size'] / 1024, 1),
                'last_modified': obj['LastModified'].isoformat() if obj.get('LastModified') else None
            }
            
            if filename.endswith('.html'):
                # Parse cursus letter - support formats like:
                # sijill_timeline_cursus_a.html
                # sijill_timeline_cursus_a_map.html
                # cursus_a.html, etc.
                match = re.search(r'cursus_([a-f])(?:_[a-z]+)?\.html', filename.lower())
                if match:
                    file_info['cursus_letter'] = match.group(1).upper()
                file_info['type'] = 'timeline_html'
                
                # Check if we have a manual assignment in DB
                db_entry = await db.timeline_resources.find_one({'filename': filename, 'type': 'timeline'}, {'_id': 0})
                if db_entry:
                    if db_entry.get('cursus_letter'):
                        file_info['cursus_letter'] = db_entry['cursus_letter']
                        file_info['manual_assignment'] = True
                    if db_entry.get('title'):
                        file_info['title'] = db_entry['title']
                    file_info['display_order'] = db_entry.get('display_order', 99)
                else:
                    file_info['display_order'] = 99
                
                html_files.append(file_info)
            elif filename.endswith('.docx'):
                # Skip temp files
                if filename.startswith('~$'):
                    continue
                
                file_info['type'] = 'context_docx'
                
                # Parse new format: sijill_{cursus}_m{NN}_{penseur}.docx
                new_match = re.match(r'sijill_([a-f])_m(\d+)_(.+)\.docx', filename, re.IGNORECASE)
                # Parse old format: Timeline_Module{N}_{Penseur}.docx
                old_match = re.match(r'Timeline_Module(\d+)_(.+)\.docx', filename)
                
                if new_match:
                    file_info['cursus_letter'] = new_match.group(1).upper()
                    file_info['module_number'] = int(new_match.group(2))
                    subject_raw = new_match.group(3)
                    file_info['subject'] = subject_raw.replace('-', ' ').replace('_', ' ').title()
                elif old_match:
                    file_info['module_number'] = int(old_match.group(1))
                    file_info['subject'] = old_match.group(2).replace('_', ' ').replace('-', ' ')
                else:
                    file_info['module_number'] = 0
                    file_info['subject'] = filename.replace('.docx', '').replace('_', ' ')
                
                # Check for custom data in DB
                resource_id = filename.replace('.docx', '').lower().replace(' ', '-').replace('_', '-')
                db_entry = await db.context_resources.find_one({'resource_id': resource_id}, {'_id': 0})
                
                if db_entry:
                    file_info['title'] = db_entry.get('title', file_info.get('subject', ''))
                    file_info['description'] = db_entry.get('description', '')
                    file_info['credits'] = db_entry.get('credits', '')
                    if db_entry.get('module_number'):
                        file_info['module_number'] = db_entry['module_number']
                    if db_entry.get('subject'):
                        file_info['subject'] = db_entry['subject']
                    if db_entry.get('cursus_letter'):
                        file_info['cursus_letter'] = db_entry['cursus_letter']
                else:
                    file_info['title'] = file_info.get('subject', filename.replace('.docx', ''))
                    file_info['description'] = ''
                    file_info['credits'] = ''
                
                docx_files.append(file_info)
        
        # Get audio files from audio/ folder
        audio_response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix='audio/', MaxKeys=200)
        audio_files = []
        
        for obj in audio_response.get('Contents', []):
            key = obj['Key']
            filename = key.split('/')[-1]
            
            # Skip non-audio files
            if not any(filename.lower().endswith(ext) for ext in ['.mp3', '.m4a', '.wav', '.ogg', '.aac']):
                continue
            
            resource_id = filename.rsplit('.', 1)[0].lower().replace(' ', '-').replace('_', '-')
            ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
            
            # Parse filename: Conf_Averroes_Brenet_module4.m4a
            match = re.match(r'Conf_([^_]+)_([^_]+)_module(\d+)\.(mp3|m4a|wav|ogg|aac)', filename, re.IGNORECASE)
            
            if match:
                subject = match.group(1).replace('-', ' ')
                speaker = match.group(2).replace('-', ' ')
                module_num = int(match.group(3))
            else:
                subject = filename.replace('_', ' ').rsplit('.', 1)[0]
                speaker = ''
                module_num = 0
            
            # Check for custom data in DB
            db_entry = await db.audio_resources.find_one({'resource_id': resource_id}, {'_id': 0})
            
            audio_info = {
                'id': resource_id,
                'filename': filename,
                'r2_key': key,
                'subject': db_entry.get('subject', subject) if db_entry else subject,
                'speaker': db_entry.get('speaker', speaker) if db_entry else speaker,
                'module_number': db_entry.get('module_number', module_num) if db_entry else module_num,
                'title': db_entry.get('title', f"Conférence : {subject}") if db_entry else f"Conférence : {subject}",
                'description': db_entry.get('description', '') if db_entry else '',
                'credits': db_entry.get('credits', '') if db_entry else '',
                'size': obj['Size'],
                'size_mb': round(obj['Size'] / (1024*1024), 1),
                'format': ext,
                'stream_url': f'/api/resources/audio/stream/{filename}'
            }
            audio_files.append(audio_info)
        
        # Sort
        html_files.sort(key=lambda x: x.get('cursus_letter', 'Z'))
        docx_files.sort(key=lambda x: (x.get('module_number', 99), x.get('subject', '')))
        audio_files.sort(key=lambda x: (x['module_number'], x['subject']))
        
        return {
            'timelines': html_files,
            'context_docs': docx_files,
            'audio_resources': audio_files,
            'total_timelines': len(html_files),
            'total_context_docs': len(docx_files),
            'total_audios': len(audio_files)
        }
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

@api_router.post("/admin/resources/sync-timeline")
async def admin_sync_timeline_resources(request: Request):
    """Admin: Sync timeline resources to database for linking with courses/modules."""
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    import re
    
    try:
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix='Timeline/', MaxKeys=200)
        
        synced_timelines = 0
        synced_contexts = 0
        
        for obj in response.get('Contents', []):
            key = obj['Key']
            filename = key.split('/')[-1]
            
            if filename.startswith('~$'):
                continue
            
            if filename.endswith('.html') and 'timeline' in filename.lower():
                # Sync timeline HTML - support various naming formats
                match = re.search(r'cursus_([a-f])(?:_[a-z]+)?\.html', filename.lower())
                letter = match.group(1).upper() if match else None
                
                # Check if manual assignment exists
                existing = await db.timeline_resources.find_one({'filename': filename, 'type': 'timeline'})
                if existing and existing.get('cursus_letter'):
                    letter = existing['cursus_letter']
                
                await db.timeline_resources.update_one(
                    {'filename': filename, 'type': 'timeline'},
                    {'$set': {
                        'type': 'timeline',
                        'cursus_letter': letter,
                        'r2_key': key,
                        'filename': filename,
                        'updated_at': datetime.now(timezone.utc).isoformat()
                    }},
                    upsert=True
                )
                synced_timelines += 1
                    
            elif filename.endswith('.docx'):
                # Sync context document
                match = re.match(r'Timeline_Module(\d+)_(.+)\.docx', filename)
                if match:
                    module_num = int(match.group(1))
                    subject = match.group(2).replace('_', ' ').replace('-', ' ')
                    resource_id = filename.replace('.docx', '').lower().replace(' ', '-').replace('_', '-')
                    
                    await db.timeline_resources.update_one(
                        {'type': 'context', 'resource_id': resource_id},
                        {'$set': {
                            'type': 'context',
                            'resource_id': resource_id,
                            'module_number': module_num,
                            'subject': subject,
                            'title': f"Contexte historique : {subject}",
                            'r2_key': key,
                            'filename': filename,
                            'updated_at': datetime.now(timezone.utc).isoformat()
                        }},
                        upsert=True
                    )
                    synced_contexts += 1
        
        return {
            'message': 'Synchronisation terminée',
            'synced_timelines': synced_timelines,
            'synced_contexts': synced_contexts
        }
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

@api_router.patch("/admin/resources/timeline/{filename}/assign-cursus")
async def assign_timeline_cursus(filename: str, request: Request):
    """Manually assign a cursus letter to a timeline file."""
    await require_admin(request)
    
    body = await request.json()
    cursus_letter = body.get('cursus_letter', '').upper().strip()
    
    if cursus_letter and cursus_letter not in ['A', 'B', 'C', 'D', 'E']:
        raise HTTPException(400, "Cursus invalide. Utilisez A, B, C, D ou E.")
    
    # Update or create the timeline resource entry
    result = await db.timeline_resources.update_one(
        {'filename': filename, 'type': 'timeline'},
        {'$set': {
            'cursus_letter': cursus_letter if cursus_letter else None,
            'updated_at': datetime.now(timezone.utc).isoformat()
        }},
        upsert=True
    )
    
    return {
        'message': f'Cursus {"assigné: " + cursus_letter if cursus_letter else "retiré"} pour {filename}',
        'filename': filename,
        'cursus_letter': cursus_letter if cursus_letter else None
    }

@api_router.put("/admin/resources/audio/{resource_id}")
async def update_audio_resource(resource_id: str, request: Request):
    """Update audio resource metadata (title, description, credits, etc.)."""
    await require_admin(request)
    
    body = await request.json()
    
    update_data = {
        'resource_id': resource_id,
        'updated_at': datetime.now(timezone.utc).isoformat()
    }
    
    # Allow updating these fields
    allowed_fields = ['title', 'description', 'credits', 'speaker', 'subject', 'module_number', 'cursus_id']
    for field in allowed_fields:
        if field in body:
            update_data[field] = body[field]
    
    result = await db.audio_resources.update_one(
        {'resource_id': resource_id},
        {'$set': update_data},
        upsert=True
    )
    
    return {
        'message': 'Audio mis à jour',
        'resource_id': resource_id,
        'updated_fields': [f for f in allowed_fields if f in body]
    }

@api_router.put("/admin/resources/context/{resource_id}")
async def update_context_resource(resource_id: str, request: Request):
    """Update context document metadata (title, description, credits, etc.)."""
    await require_admin(request)
    
    body = await request.json()
    
    # Normalize resource_id to use dashes (matching the public API format)
    normalized_id = resource_id.replace('_', '-').lower()
    
    update_data = {
        'resource_id': normalized_id,
        'updated_at': datetime.now(timezone.utc).isoformat()
    }
    
    # Allow updating these fields
    allowed_fields = ['title', 'description', 'credits', 'subject', 'module_number', 'cursus_id', 'cursus_letter']
    for field in allowed_fields:
        if field in body:
            update_data[field] = body[field]
    
    result = await db.context_resources.update_one(
        {'resource_id': normalized_id},
        {'$set': update_data},
        upsert=True
    )
    
    return {
        'message': 'Fiche mise à jour',
        'resource_id': normalized_id,
        'updated_fields': [f for f in allowed_fields if f in body]
    }

@api_router.get("/admin/resources/audio/{resource_id}")
async def get_audio_resource_admin(resource_id: str, request: Request):
    """Get audio resource details for editing."""
    await require_admin(request)
    
    db_entry = await db.audio_resources.find_one({'resource_id': resource_id}, {'_id': 0})
    
    if not db_entry:
        return {
            'resource_id': resource_id,
            'title': '',
            'description': '',
            'credits': '',
            'speaker': '',
            'subject': '',
            'module_number': 0,
            'cursus_id': ''
        }
    
    return db_entry

@api_router.post("/admin/courses/{course_id}/sync-r2")
async def sync_course_with_r2(course_id: str, body: SyncR2FolderRequest, request: Request):
    """
    Sync a course with an R2 folder: scan the folder and create audio episodes.
    Accepts multiple file naming patterns: episode-01, episode01, ep01, 01.m4a, etc.
    """
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    # Get the course
    course = await db.courses.find_one({'id': course_id})
    if not course:
        raise HTTPException(404, "Cours non trouvé")
    
    r2_folder = body.r2_folder.strip().rstrip('/')
    if not r2_folder:
        raise HTTPException(400, "Dossier R2 requis")
    
    # Extract folder name for episode titles
    folder_name = r2_folder.split('/')[-1].replace('-', ' ').replace('_', ' ').title()
    
    try:
        # List files in the R2 folder
        prefix = f"{r2_folder}/"
        logger.info(f"Sync R2: Listing files in bucket={R2_BUCKET}, prefix={prefix}")
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix=prefix, MaxKeys=500)
        
        files = []
        contents = response.get('Contents', [])
        logger.info(f"Sync R2: Found {len(contents)} objects")
        
        for obj in contents:
            key = obj['Key']
            size = obj['Size']
            logger.info(f"Sync R2: Checking file: {key} (size={size})")
            if size > 0 and (key.endswith('.m4a') or key.endswith('.mp3') or key.endswith('.wav')):
                filename = key.replace(prefix, '')
                # Extract episode number from filename with multiple patterns
                import re
                # Try various patterns: episode-01, episode01, ep-01, ep01, 01.m4a, piste-01, etc.
                match = re.search(r'(?:episode|ep|piste)?[-_]?(\d+)', filename, re.IGNORECASE)
                logger.info(f"Sync R2: Filename={filename}, match={match.group(1) if match else 'None'}")
                if match:
                    ep_num = int(match.group(1))
                    files.append({
                        'key': key,
                        'filename': filename,
                        'size': size,
                        'episode_number': ep_num,
                    })
        
        logger.info(f"Sync R2: Final files count: {len(files)}")
        if not files:
            raise HTTPException(400, f"Aucun fichier audio trouvé dans '{r2_folder}/' (formats: .m4a, .mp3, .wav)")
        
        # Sort by episode number
        files.sort(key=lambda x: x['episode_number'])
        
        # Create or update audio entries for each episode
        created = 0
        updated = 0
        now = datetime.now(timezone.utc)
        
        for f in files:
            audio_id = f"aud-{course_id.replace('crs-', '')}-{f['episode_number']:02d}"
            
            # Check if audio already exists to preserve description and other custom fields
            existing_audio = await db.audios.find_one({'id': audio_id})
            
            # Use folder name for title: "Nom Dossier ep1" instead of "Épisode 1 — Cours Title"
            default_title = f"{folder_name} ep{f['episode_number']}"
            
            audio_doc = {
                'id': audio_id,
                'title': existing_audio.get('title') if existing_audio and existing_audio.get('title') and not existing_audio.get('title').startswith('Épisode') and not existing_audio.get('title').startswith('1.') else default_title,
                'description': existing_audio.get('description') if existing_audio and existing_audio.get('description') else '',
                'scholar_id': course.get('scholar_id', ''),
                'scholar_name': course.get('scholar_name', ''),
                'duration': existing_audio.get('duration', 0) if existing_audio else 0,
                'audio_url': '',
                'file_key': f['key'],
                'thumbnail': existing_audio.get('thumbnail') if existing_audio and existing_audio.get('thumbnail') else course.get('thumbnail', ''),
                'topic': existing_audio.get('topic') if existing_audio and existing_audio.get('topic') else course.get('topic', ''),
                'type': 'lecture',
                'course_id': course_id,
                'episode_number': f['episode_number'],
                'published_at': existing_audio.get('published_at') if existing_audio else now.isoformat(),
                'is_active': existing_audio.get('is_active', True) if existing_audio else True,
            }
            
            result = await db.audios.update_one(
                {'id': audio_id},
                {'$set': audio_doc},
                upsert=True
            )
            
            if result.upserted_id:
                created += 1
            elif result.modified_count > 0:
                updated += 1
        
        # Update the course with r2_folder and modules_count
        await db.courses.update_one(
            {'id': course_id},
            {'$set': {
                'r2_folder': r2_folder,
                'modules_count': len(files),
            }}
        )
        
        return {
            'message': f"Synchronisation terminée pour '{r2_folder}'",
            'course_id': course_id,
            'r2_folder': r2_folder,
            'episodes_created': created,
            'episodes_updated': updated,
            'total_episodes': len(files),
        }
        
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

# Mapping R2 folders to course IDs
R2_TO_COURSE_MAPPING = {
    '01-mouvement-traduction': 'cours-traduction',
    '02-falsafa': 'cours-falsafa-grands',
    '03-post-avicennisme': 'cours-post-avicennisme',
    '04-falsafa-occident-musulman': 'cours-falsafa-occident',
    '05-renouveau-falsafa-persan': 'cours-falsafa-persan',
    '06-logique-arabe': 'cours-logique',
    '07-inclassables': 'cours-inclassables',
    '08-kalam': 'cours-kalam',
    '09-usul-al-fiqh': 'cours-fiqh',
    '10-doxographie': 'cours-doxographie',
    '11-transmission-coran': 'cours-coran',
    '12-transmission-hadith': 'cours-hadith',
    '13-historiographie': 'cours-historiographie',
    '14-autobiographies': 'cours-autobiographies',
    '15-histoire-art': 'cours-art',
    '16-poesie': 'cours-poesie',
    '17-histoire-pedagogie': 'cours-urjuza',
    '18-sciences': 'cours-sciences',
    '19-geographie-islamique': 'cours-geographie',
    '20-adab-sciencesmedicales': 'cours-adab',
    '21-kalam-chretien': 'cours-kalam-chretien',
    '22-mystique-islamique': 'cours-soufisme',
    '23-ismaelisme': 'cours-ismaelisme',
    '24-philosophie-juive': 'cours-philo-juive',
}

# R2 cursus folder → cursus_id mapping (7 cursus)
R2_CURSUS_MAPPING = {
    'cursus-a-histoire': 'cursus-histoire',
    'cursus-b-theologie-droit': 'cursus-theologie',
    'cursus-c-sciences-islamiques': 'cursus-sciences-islamiques',
    'cursus-d-arts-litterature': 'cursus-arts',
    'cursus-e-falsafa': 'cursus-falsafa',
    'cursus-f-mystique': 'cursus-spiritualites',
    'cursus-g-pensees-non-islamiques': 'cursus-pensees-non-islamiques',
    # Legacy folder names (backward compatibility)
    'cursus-a-falsafa': 'cursus-falsafa',
    'cursus-e-spiritualites': 'cursus-spiritualites',
    'cursus-f-pensees-non-islamiques': 'cursus-pensees-non-islamiques',
}

# ========== PROFESSOR PHOTO SYNC ==========
@api_router.post("/admin/scholars/{scholar_id}/upload-photo")
async def admin_upload_scholar_photo(scholar_id: str, request: Request, file: UploadFile = File(...)):
    """Upload a photo file directly to R2 (Prof/{slug}.{ext}) and set it on the scholar.
    The uploaded URL is stored in the scholar's `photo` field using the public R2 CDN URL.
    """
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    scholar = await db.scholars.find_one({'id': scholar_id}, {'_id': 0})
    if not scholar:
        raise HTTPException(404, "Érudit non trouvé")
    ctype = (file.content_type or '').lower()
    allowed = {'image/jpeg', 'image/jpg', 'image/png', 'image/webp'}
    if ctype not in allowed:
        raise HTTPException(400, "Format accepté : JPG, PNG ou WebP")
    ext_map = {'image/jpeg': 'jpg', 'image/jpg': 'jpg', 'image/png': 'png', 'image/webp': 'webp'}
    ext = ext_map[ctype]
    import re, unicodedata
    def slugify(s: str) -> str:
        s = unicodedata.normalize('NFD', s or '')
        s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
        return re.sub(r'[^a-z0-9]+', '-', s.lower()).strip('-')
    slug = slugify(scholar.get('name') or scholar_id)
    key = f"Prof/{slug}.{ext}"
    data = await file.read()
    if len(data) > 5 * 1024 * 1024:
        raise HTTPException(413, "Photo trop volumineuse (max 5 Mo)")
    try:
        r2_client.put_object(Bucket=R2_BUCKET, Key=key, Body=data, ContentType=ctype,
                             CacheControl='public, max-age=31536000')
    except Exception as e:
        raise HTTPException(500, f"Échec upload R2: {e}")
    public_base = os.environ.get('R2_PUBLIC_BASE_URL', '').rstrip('/')
    photo_url = f"{public_base}/{key}" if public_base else key
    await db.scholars.update_one({'id': scholar_id}, {'$set': {'photo': photo_url}})
    return {'photo': photo_url, 'r2_key': key, 'size': len(data)}


@api_router.post("/admin/sync-professor-photos")
async def sync_professor_photos(request: Request):
    """Sync professor photos from R2 Prof/ folder.
    Matching is tolerant to: accents, common typos (Granpierre↔Grandpierre, Ghouirate↔Ghouirgate),
    and uses Levenshtein-like distance via shared prefix matching."""
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")

    import re, unicodedata

    def normalize(s: str) -> str:
        """Strip accents + lowercase + remove non-alphanum."""
        s = unicodedata.normalize('NFD', s)
        s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
        return re.sub(r'[^a-z0-9]', '', s.lower())

    try:
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix='Prof/', MaxKeys=100)
        # Pre-load all scholars once
        all_scholars = await db.scholars.find({}, {'_id': 0}).to_list(200)
        scholar_norms = [(s, normalize(s.get('name', ''))) for s in all_scholars]

        updated = 0
        skipped = []
        for obj in response.get('Contents', []):
            key = obj['Key']
            filename = key.replace('Prof/', '')
            if not filename or filename.endswith('/'):
                continue
            # Extract candidate name from filename
            stem = filename.rsplit('.', 1)[0]
            stem = re.sub(r'^(Prof_|\d+_)', '', stem)
            # Convert CamelCase to spaces, then normalize
            spaced = re.sub(r'([a-z])([A-Z])', r'\1 \2', stem)
            target = normalize(spaced)
            if not target:
                skipped.append({'file': filename, 'reason': 'empty target'})
                continue

            # Find best scholar by largest shared prefix length (handles typos like Ghouirate ↔ Ghouirgate)
            best, best_score = None, 0
            for s, sn in scholar_norms:
                if not sn:
                    continue
                # Skip very short normalized names to avoid false positives
                if len(sn) < 4:
                    continue
                # Score = length of common prefix (works well for misspellings missing a letter)
                score = 0
                for a, b in zip(target, sn):
                    if a == b:
                        score += 1
                    else:
                        break
                # Bonus if either fully contains the other (Granpierre ⊂ Grandpierre)
                if target in sn or sn in target:
                    score = max(score, min(len(target), len(sn)))
                # Require ≥ 70% of the shorter normalized name to be common
                threshold = max(5, int(min(len(target), len(sn)) * 0.7))
                if score >= threshold and score > best_score:
                    best, best_score = s, score

            if best:
                photo_url = f"/api/images/{filename}"
                await db.scholars.update_one(
                    {'id': best['id']},
                    {'$set': {'photo_url': photo_url, 'photo': photo_url, 'photo_key': key}}
                )
                logger.info(f"sync_professor_photos: {filename} → {best['name']} (score={best_score})")
                updated += 1
            else:
                skipped.append({'file': filename, 'reason': 'no scholar match'})
                logger.warning(f"sync_professor_photos: no match for {filename}")

        return {
            'message': 'Synchronisation des photos terminée',
            'photos_updated': updated,
            'photos_skipped': skipped,
            'total_files_in_r2': len(response.get('Contents', [])),
        }
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

# ========== BIBLIOGRAPHY SYNC ==========
@api_router.post("/admin/sync-bibliographies")
async def sync_bibliographies(request: Request):
    """
    Sync bibliography files (.docx) from R2 Biblio/ folder.
    Uses positional matching: files are sorted per cursus and matched
    to courses in order, handling both global and local module numbering.
    """
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    try:
        from docx import Document
        from io import BytesIO
        import re
        from collections import defaultdict
        
        cursus_mapping = {
            'A': 'cursus-histoire',
            'B': 'cursus-theologie',
            'C': 'cursus-sciences-islamiques',
            'D': 'cursus-arts',
            'E': 'cursus-falsafa',
            'F': 'cursus-spiritualites',
            'G': 'cursus-pensees-non-islamiques',
        }
        
        # Pre-load all courses grouped by cursus, in their natural DB order
        cursus_courses = {}
        for letter, cursus_id in cursus_mapping.items():
            courses = await db.courses.find(
                {'cursus_id': cursus_id}, {'_id': 0, 'id': 1, 'title': 1}
            ).to_list(50)
            cursus_courses[letter] = courses
            logger.info(f"Cursus {letter} ({cursus_id}): {len(courses)} courses")
        
        # List and group R2 files by cursus letter
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix='Biblio/', MaxKeys=200)
        
        files_by_cursus = defaultdict(list)
        for obj in response.get('Contents', []):
            key = obj['Key']
            if not key.endswith('.docx'):
                continue
            filename = key.replace('Biblio/', '')
            match = re.match(r'Biblio_Module(\d+)\s*_Cursus([A-E])\.docx', filename)
            if not match:
                logger.warning(f"Skipping unrecognized biblio file: {filename}")
                continue
            module_num = int(match.group(1))
            cursus_letter = match.group(2)
            files_by_cursus[cursus_letter].append((module_num, key))
        
        created = 0
        updated = 0
        skipped = 0
        
        for letter in sorted(files_by_cursus.keys()):
            files = sorted(files_by_cursus[letter], key=lambda x: x[0])
            courses = cursus_courses.get(letter, [])
            cursus_id = cursus_mapping.get(letter)
            
            if not cursus_id:
                continue
            
            # Match files to courses positionally (N-th file → N-th course)
            for position, (module_num, key) in enumerate(files):
                if position >= len(courses):
                    logger.warning(f"Extra biblio file {key} (pos {position}) — cursus {letter} only has {len(courses)} courses. Skipping.")
                    skipped += 1
                    continue
                
                course = courses[position]
                course_id = course['id']
                course_title = course['title']
                
                # Compute the display module number (global = cursus start + position)
                # This is for display only, the real link is via course_id
                global_module_num = position + 1  # 1-based within cursus
                
                try:
                    file_response = r2_client.get_object(Bucket=R2_BUCKET, Key=key)
                    docx_content = file_response['Body'].read()
                    doc = Document(BytesIO(docx_content))
                    
                    content_parts = []
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            is_heading = any(run.bold for run in para.runs if run.text.strip())
                            if is_heading:
                                content_parts.append(f"## {text}")
                            else:
                                content_parts.append(text)
                    
                    content = "\n\n".join(content_parts)
                    
                    biblio_id = f"biblio-{letter.lower()}-mod{global_module_num:02d}"
                    
                    biblio_doc = {
                        'id': biblio_id,
                        'title': f"Bibliographie — {course_title}",
                        'content': content,
                        'content_html': content.replace('\n\n## ', '\n\n<h3>').replace('## ', '<h3>').replace('\n\n', '</h3>\n\n<p>') + '</p>' if '## ' in content else f"<p>{content.replace(chr(10)+chr(10), '</p><p>')}</p>",
                        'cursus_id': cursus_id,
                        'cursus_letter': letter,
                        'course_id': course_id,
                        'module_number': global_module_num,
                        'file_key': key,
                        'updated_at': datetime.now(timezone.utc).isoformat(),
                    }
                    
                    result = await db.bibliographies.update_one(
                        {'id': biblio_id},
                        {'$set': biblio_doc},
                        upsert=True
                    )
                    
                    if result.upserted_id:
                        created += 1
                    else:
                        updated += 1
                        
                    logger.info(f"Synced bibliography: {biblio_id} → {course_id} ({course_title[:40]})")
                    
                except Exception as e:
                    logger.error(f"Error processing {key}: {e}")
                    continue
        
        # Clean up old orphaned biblios with wrong IDs from previous sync
        all_valid_ids = []
        for letter in sorted(files_by_cursus.keys()):
            files = sorted(files_by_cursus[letter], key=lambda x: x[0])
            courses = cursus_courses.get(letter, [])
            for position in range(min(len(files), len(courses))):
                all_valid_ids.append(f"biblio-{letter.lower()}-mod{position+1:02d}")
        
        orphans = await db.bibliographies.find(
            {'id': {'$nin': all_valid_ids}}, {'_id': 0, 'id': 1}
        ).to_list(100)
        
        deleted = 0
        if orphans:
            orphan_ids = [o['id'] for o in orphans]
            result = await db.bibliographies.delete_many({'id': {'$in': orphan_ids}})
            deleted = result.deleted_count
            logger.info(f"Cleaned up {deleted} orphaned bibliographies: {orphan_ids}")
        
        return {
            'message': 'Synchronisation des bibliographies terminée',
            'bibliographies_created': created,
            'bibliographies_updated': updated,
            'skipped': skipped,
            'deleted_orphans': deleted,
            'total': created + updated
        }
        
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

@api_router.get("/bibliographies")
async def list_bibliographies(request: Request, cursus_id: str = None, course_id: str = None):
    """List bibliographies, optionally filtered by cursus or course.
    Strips full 'content' for non-subscribers to avoid leaking premium material."""
    query = {
        'content': {'$exists': True, '$ne': ''},  # Only new format with content
        'module_number': {'$exists': True}
    }
    if cursus_id:
        query['cursus_id'] = cursus_id
    if course_id:
        query['course_id'] = course_id
    
    biblios = await db.bibliographies.find(query, {'_id': 0}).sort('module_number', 1).to_list(100)
    
    # Determine subscriber status (no exception if guest)
    is_subscriber = False
    user = await get_current_user(request)
    if user:
        if user.get('role') == 'admin' or user.get('free_access'):
            is_subscriber = True
        else:
            now = datetime.now(timezone.utc)
            for fld in ('trial', 'subscription'):
                v = user.get(fld)
                if v and v.get('expires_at'):
                    exp = v['expires_at']
                    if isinstance(exp, str):
                        exp = datetime.fromisoformat(exp.replace('Z', '+00:00'))
                    if exp > now:
                        is_subscriber = True
                        break
    
    if not is_subscriber:
        for b in biblios:
            b.pop('content', None)
            b['locked'] = True
    return biblios

@api_router.get("/bibliographies/{biblio_id}")
async def get_bibliography(biblio_id: str, request: Request):
    """Get a specific bibliography by ID. Requires authentication + active subscription."""
    await require_subscriber(request)
    biblio = await db.bibliographies.find_one({'id': biblio_id}, {'_id': 0})
    if not biblio:
        raise HTTPException(404, "Bibliographie non trouvée")
    return biblio

@api_router.post("/admin/sync-preview")
async def sync_preview_r2(request: Request):
    """
    Preview what would change if a full R2 sync was executed.
    Returns lists of files to create, update, and orphans to delete.
    """
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")

    try:
        to_create = []
        to_update = []
        to_delete = []
        all_r2_keys = set()
        synced_audio_ids = set()

        for cursus_folder, cursus_id in R2_CURSUS_MAPPING.items():
            response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix=f"{cursus_folder}/", MaxKeys=500)
            for obj in response.get('Contents', []):
                key = obj['Key']
                size = obj['Size']
                if size == 0 or not (key.endswith('.m4a') or key.endswith('.mp3')):
                    continue
                all_r2_keys.add(key)
                rel_path = key[len(cursus_folder) + 1:]
                parts = rel_path.split('/')
                if len(parts) < 2:
                    continue
                module_folder = parts[0].strip().lower()
                import unicodedata
                module_folder_normalized = unicodedata.normalize('NFD', module_folder)
                module_folder_normalized = ''.join(c for c in module_folder_normalized if unicodedata.category(c) != 'Mn').strip()
                course_id = R2_TO_COURSE_MAPPING.get(module_folder) or R2_TO_COURSE_MAPPING.get(module_folder_normalized)
                if not course_id:
                    match = re.match(r'^(\d+)-', module_folder_normalized)
                    if match:
                        num_prefix = match.group(1)
                        for mk, mv in R2_TO_COURSE_MAPPING.items():
                            if mk.startswith(f"{num_prefix}-"):
                                course_id = mv
                                break
                if not course_id:
                    continue
                episode_number = 1
                if len(parts) == 3:
                    subfolder_slug = parts[1].replace(' ', '-').lower()
                    match = re.search(r'episode-?(\d+)', parts[2], re.IGNORECASE)
                    if match:
                        episode_number = int(match.group(1))
                    audio_id = f"aud_{course_id}-{subfolder_slug}-ep{episode_number:02d}"
                else:
                    match = re.search(r'episode-?(\d+)', parts[1], re.IGNORECASE)
                    if match:
                        episode_number = int(match.group(1))
                    audio_id = f"aud_{course_id}-ep{episode_number:02d}"
                audio_id = re.sub(r'[^a-z0-9_-]', '', audio_id)
                synced_audio_ids.add(audio_id)
                existing = await db.audios.find_one({'id': audio_id}, {'_id': 0, 'id': 1})
                filename = key.split('/')[-1]
                if existing:
                    to_update.append(filename)
                else:
                    to_create.append(filename)

        # Find orphans
        all_audios = await db.audios.find({'file_key': {'$exists': True, '$ne': ''}}, {'_id': 0, 'id': 1, 'title': 1, 'file_key': 1}).to_list(2000)
        for audio in all_audios:
            fk = audio.get('file_key', '')
            aid = audio.get('id', '')
            if not fk:
                continue
            if fk.startswith('cursus-') and aid not in synced_audio_ids:
                to_delete.append(audio.get('title', aid))
            elif fk.startswith('Audio/'):
                to_delete.append(audio.get('title', aid))

        return {'to_create': to_create, 'to_update': to_update, 'to_delete': to_delete}

    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

# ── Manifest Parser & API ──────────────────────────────────────────────────

CURSUS_LETTER_MAP = {
    'A': 'cursus-histoire',
    'B': 'cursus-theologie',
    'C': 'cursus-sciences-islamiques',
    'D': 'cursus-arts',
    'E': 'cursus-falsafa',
    'F': 'cursus-spiritualites',
    'G': 'cursus-pensees-non-islamiques',
}

def parse_manifest_docx(file_bytes: bytes) -> dict:
    """Parse the Sijill manifest .docx into a structured hierarchy."""
    from docx import Document
    import io
    doc = Document(io.BytesIO(file_bytes))

    # Parse paragraphs to find cursus and course headings
    cursus_headings = []  # [(para_index, letter, name)]
    course_headings = []  # [(para_index, number, title, cursus_letter)]

    current_cursus_letter = None
    for pi, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        # Cursus heading: starts with ◆ or contains "A. ", "B. " etc at start
        m = re.match(r'[◆◇●■□▪▸►]\s*([A-E])\.\s*(.+)', text)
        if m:
            current_cursus_letter = m.group(1)
            cursus_headings.append((pi, m.group(1), m.group(2).strip()))
            continue
        # Course heading: starts with a number followed by "."
        m2 = re.match(r'^(\d{1,2})\.\s+(.+)', text)
        if m2 and current_cursus_letter:
            course_headings.append((pi, int(m2.group(1)), m2.group(2).strip(), current_cursus_letter))

    # Build course-to-table mapping based on document order
    # Tables appear after their course heading paragraph
    # Table 0 is the summary table, skip it
    content_tables = doc.tables[1:]  # Skip summary table

    # Map each course heading to its table
    result = {
        'cursus': [],
        'total_modules': 0,
        'total_episodes': 0,
    }

    # Group courses by cursus
    cursus_courses = {}
    for _, course_num, course_title, cursus_letter in course_headings:
        if cursus_letter not in cursus_courses:
            cursus_courses[cursus_letter] = []
        cursus_courses[cursus_letter].append({
            'number': course_num,
            'title': course_title,
            'modules': [],
        })

    # Assign tables to courses (one table per course, in order)
    table_idx = 0
    for cursus_letter in ['A', 'B', 'C', 'D', 'E']:
        courses = cursus_courses.get(cursus_letter, [])
        for course in courses:
            if table_idx < len(content_tables):
                table = content_tables[table_idx]
                table_idx += 1
                for row in table.rows[1:]:  # Skip header row
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) >= 4:
                        num_str = cells[0]
                        module_name = cells[1]
                        professor = cells[2]
                        ep_str = cells[3]
                        notes = cells[4] if len(cells) > 4 else ''
                        try:
                            ep_count = int(re.search(r'\d+', ep_str).group()) if re.search(r'\d+', ep_str) else 1
                        except (ValueError, AttributeError):
                            ep_count = 1
                        module = {
                            'number': num_str,
                            'name': module_name,
                            'professor': professor,
                            'episodes': ep_count,
                            'notes': notes if notes != '—' else '',
                        }
                        course['modules'].append(module)
                        result['total_modules'] += 1
                        result['total_episodes'] += ep_count

    # Build final structure
    for _, letter, name in cursus_headings:
        courses = cursus_courses.get(letter, [])
        total_eps = sum(sum(m['episodes'] for m in c['modules']) for c in courses)
        total_mods = sum(len(c['modules']) for c in courses)
        result['cursus'].append({
            'letter': letter,
            'cursus_id': CURSUS_LETTER_MAP.get(letter, ''),
            'name': name,
            'courses_count': len(courses),
            'modules_count': total_mods,
            'episodes_count': total_eps,
            'courses': courses,
        })

    return result


@api_router.post("/admin/manifest/upload")
async def upload_manifest(request: Request, file: UploadFile = File(...)):
    """Upload and parse a manifest .docx file."""
    await require_admin(request)

    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "Le fichier doit être un .docx")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(400, "Fichier trop volumineux (max 10MB)")

    try:
        parsed = parse_manifest_docx(content)
    except Exception as e:
        logger.error(f"Manifest parse error: {e}")
        raise HTTPException(400, f"Erreur de parsing: {str(e)}")

    # Store in DB (replace previous manifest)
    manifest_doc = {
        'filename': file.filename,
        'uploaded_at': datetime.now(timezone.utc).isoformat(),
        'data': parsed,
    }
    await db.manifest.delete_many({})
    await db.manifest.insert_one(manifest_doc)

    return {
        'filename': file.filename,
        'cursus_count': len(parsed['cursus']),
        'total_modules': parsed['total_modules'],
        'total_episodes': parsed['total_episodes'],
        'cursus': [{
            'letter': c['letter'],
            'name': c['name'],
            'courses': c['courses_count'],
            'modules': c['modules_count'],
            'episodes': c['episodes_count'],
        } for c in parsed['cursus']],
    }


@api_router.get("/admin/manifest")
async def get_manifest(request: Request):
    """Get the current manifest data."""
    await require_admin(request)
    doc = await db.manifest.find_one({}, {'_id': 0})
    if not doc:
        return {'manifest': None}
    return {'manifest': doc}

# ── Blog Parser & API ──────────────────────────────────────────────────────

def parse_blog_docx(file_bytes: bytes, file_key: str) -> dict:
    """Parse a Sijill Times blog article .docx into structured data."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))

    # ── Paragraphs: extract metadata and body ──
    paras = [(p.text.strip(), any(r.bold for r in p.runs if r.bold)) for p in doc.paragraphs if p.text.strip()]

    number = 0
    date_ah = ''
    date_ce = ''
    epoch = ''
    title = ''
    tags = []
    body_sections = []
    current_section = None

    SECTION_MAP = {
        'I': "TERRES D'ISLAM",
        'II': 'VIE INTELLECTUELLE',
        'III': 'LES ÉCHANGES',
        'IV': 'LE RESTE DU MONDE',
        'V': "CE QUE ÇA CHANGE",
        'VI': 'POUR ALLER PLUS LOIN',
    }

    for text, is_bold in paras:
        # P0: "SIJILL TIMES · Chroniques..."
        if 'SIJILL TIMES' in text:
            continue

        # P1: "Le monde en… · Numéro N"
        nm = re.search(r'Numéro\s+(\d+)', text)
        if nm and not number:
            number = int(nm.group(1))
            continue

        # P3: "Le monde en 370 — Époque buyide"
        m_head = re.match(r'Le monde en\s+(\d+)\s*[—–-]\s*(.+)', text)
        if m_head:
            year = int(m_head.group(1))
            date_ah = f"{year} AH"
            epoch = m_head.group(2).strip()
            continue

        # P4: Tags line (multiple · separators, no bold)
        if not is_bold and '·' in text and len(text.split('·')) >= 3 and not tags:
            tags = [t.strip() for t in text.split('·') if t.strip()]
            continue

        # Bold title (first bold paragraph that's not a section header)
        if is_bold and not title:
            sec_match = re.match(r'^(I{1,3}V?|IV|VI?)\.\s+', text)
            if not sec_match:
                title = text
                continue

        # Section header: "I. TERRES D'ISLAM", "II. VIE INTELLECTUELLE"
        sec_match = re.match(r'^(I{1,3}V?|IV|VI?)\.\s+(.+)', text)
        if sec_match and is_bold:
            if current_section:
                body_sections.append(current_section)
            roman = sec_match.group(1)
            sec_name = SECTION_MAP.get(roman, sec_match.group(2).strip())
            current_section = {'type': 'section', 'roman': roman, 'title': sec_name, 'content': ''}
            continue

        # Content paragraphs
        if current_section:
            if current_section['content']:
                current_section['content'] += '\n\n'
            current_section['content'] += text
        elif title:
            # Pre-section content = intro
            if body_sections and body_sections[-1]['type'] == 'intro':
                body_sections[-1]['content'] += '\n\n' + text
            else:
                body_sections.append({'type': 'intro', 'title': '', 'content': text})

    if current_section:
        body_sections.append(current_section)

    # ── Tables ──
    # T0: Introduction/hook (1x1)
    hook = doc.tables[0].cell(0, 0).text.strip() if len(doc.tables) > 0 else ''

    # T1: Carte politique (6x1)
    carte = []
    if len(doc.tables) > 1:
        for row in doc.tables[1].rows:
            txt = row.cells[0].text.strip()
            if txt and 'CARTE POLITIQUE' not in txt:
                carte.append(txt)

    # T2: Portrait (2x1)
    portrait_header = ''
    portrait_text = ''
    if len(doc.tables) > 2:
        portrait_header = doc.tables[2].cell(0, 0).text.strip()
        portrait_text = doc.tables[2].cell(1, 0).text.strip() if len(doc.tables[2].rows) > 1 else ''

    # T3: References (3x2)
    references = []
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        for row in t3.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[1]:
                references.append({'type': cells[0], 'text': cells[1]})

    # ── Compute CE date from AH ──
    ah_num = re.search(r'(\d+)', date_ah)
    year_ah = int(ah_num.group(1)) if ah_num else 0
    if year_ah > 0:
        year_ce = round(year_ah * 0.970229 + 621.5)
        date_ce = f"{year_ce} CE"

    # ── Generate ID (use number to avoid collisions when multiple articles share the same year) ──
    article_id = f"mondeen-{number}" if number else f"mondeen-{year_ah}"

    # ── Image key (check if jpeg exists) ──
    base = file_key.replace('.docx', '')
    image_key = base + '.jpeg'

    # ── SEO description ──
    seo_desc = hook[:200].rsplit(' ', 1)[0] + '...' if len(hook) > 50 else (title or '')

    return {
        'id': article_id,
        'series': 'Le monde en…',
        'series_subtitle': 'Chroniques de la civilisation islamique',
        'number': number,
        'year_ah': year_ah,
        'date_ah': date_ah,
        'date_ce': date_ce,
        'epoch': epoch,
        'title': title,
        'tags': tags,
        'hook': hook,
        'carte_politique': carte,
        'portrait_header': portrait_header,
        'portrait': portrait_text,
        'references': references,
        'body_sections': body_sections,
        'seo_description': seo_desc,
        'image_key': image_key,
        'file_key': file_key,
        'is_active': True,
    }


@api_router.get("/blog/image/{article_id}")
async def blog_image(article_id: str):
    """Serve blog article illustration image from R2."""
    article = await db.blog_articles.find_one({'id': article_id}, {'_id': 0, 'image_key': 1})
    if not article or not article.get('image_key'):
        raise HTTPException(404, "Image introuvable")
    try:
        resp = r2_client.get_object(Bucket=R2_BUCKET, Key=article['image_key'])
        content = resp['Body'].read()
        return Response(content=content, media_type='image/jpeg',
                       headers={'Cache-Control': 'public, max-age=86400'})
    except Exception:
        raise HTTPException(404, "Image introuvable dans R2")


@api_router.post("/admin/blog/sync-r2")
async def sync_blog_r2(request: Request):
    """Sync blog articles from R2 Blog/ folder."""
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")

    try:
        response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix='Blog/', MaxKeys=200)
        files = [obj for obj in response.get('Contents', [])
                 if obj['Key'].endswith('.docx') and obj['Size'] > 500
                 and '/~$' not in obj['Key']]

        created = 0
        updated = 0
        errors = []
        synced_ids = set()

        for obj in files:
            key = obj['Key']
            try:
                file_resp = r2_client.get_object(Bucket=R2_BUCKET, Key=key)
                content = file_resp['Body'].read()
                parsed = parse_blog_docx(content, key)
                synced_ids.add(parsed['id'])

                existing = await db.blog_articles.find_one({'id': parsed['id']})
                if existing:
                    # Preserve is_active status
                    parsed['is_active'] = existing.get('is_active', True)
                    parsed['published_at'] = existing.get('published_at', datetime.now(timezone.utc).isoformat())
                    await db.blog_articles.update_one({'id': parsed['id']}, {'$set': parsed})
                    updated += 1
                else:
                    parsed['published_at'] = datetime.now(timezone.utc).isoformat()
                    await db.blog_articles.insert_one(parsed)
                    created += 1
            except Exception as e:
                errors.append(f"{key}: {str(e)}")
                logger.error(f"Blog sync error for {key}: {e}")

        # Clean orphans
        deleted = 0
        all_articles = await db.blog_articles.find({}, {'_id': 0, 'id': 1}).to_list(500)
        for art in all_articles:
            if art['id'] not in synced_ids:
                await db.blog_articles.delete_one({'id': art['id']})
                deleted += 1

        return {'created': created, 'updated': updated, 'deleted': deleted, 'errors': errors, 'total': len(files)}

    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")


@api_router.get("/admin/blog")
async def admin_list_blog(request: Request):
    """Admin: list all blog articles including drafts."""
    await require_admin(request)
    articles = await db.blog_articles.find({}, {'_id': 0}).sort('number', 1).to_list(200)
    return articles


@api_router.patch("/admin/blog/{article_id}/toggle")
async def admin_toggle_blog(article_id: str, request: Request):
    """Toggle blog article active status."""
    await require_admin(request)
    article = await db.blog_articles.find_one({'id': article_id})
    if not article:
        raise HTTPException(404, "Article introuvable")
    new_status = not article.get('is_active', True)
    await db.blog_articles.update_one({'id': article_id}, {'$set': {'is_active': new_status}})
    return {'id': article_id, 'is_active': new_status}


@api_router.get("/blog")
async def public_list_blog():
    """Public: list active blog articles (free access)."""
    articles = await db.blog_articles.find(
        {'is_active': True},
        {'_id': 0, 'id': 1, 'series': 1, 'number': 1, 'date_ah': 1, 'date_ce': 1, 'epoch': 1,
         'subtitle': 1, 'title': 1, 'tags': 1, 'author': 1, 'seo_description': 1, 'published_at': 1}
    ).sort('number', 1).to_list(200)
    return articles


@api_router.get("/blog/{article_id}")
async def public_get_blog(article_id: str):
    """Public: get a single blog article (free access, full content)."""
    article = await db.blog_articles.find_one(
        {'id': article_id, 'is_active': True},
        {'_id': 0}
    )
    if not article:
        raise HTTPException(404, "Article introuvable")
    return article

@api_router.post("/admin/sync-all-r2")
async def sync_all_r2_audio(request: Request):
    """
    Sync all audio files from R2 with database.
    Scans cursus-*/ folders. Path pattern:
      cursus-{x}-{name}/{NN}-{module-slug}/[subfolder/]episode-{N}.m4a
    Also cleans up orphaned DB entries (files deleted from R2).
    """
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    try:
        stats = {'total_files': 0, 'created': 0, 'updated': 0, 'skipped': 0, 'deleted_orphans': 0, 'errors': []}
        now = datetime.now(timezone.utc)
        
        # Collect ALL R2 audio file keys
        all_r2_keys = set()
        synced_audio_ids = set()
        
        # Scan each cursus folder
        for cursus_folder, cursus_id in R2_CURSUS_MAPPING.items():
            response = r2_client.list_objects_v2(Bucket=R2_BUCKET, Prefix=f"{cursus_folder}/", MaxKeys=500)
            
            for obj in response.get('Contents', []):
                key = obj['Key']
                size = obj['Size']
                
                if size == 0 or not (key.endswith('.m4a') or key.endswith('.mp3')):
                    continue
                
                all_r2_keys.add(key)
                stats['total_files'] += 1
                
                # Parse: cursus-a-falsafa/01-mouvement-traduction/episode-01.m4a
                # or:    cursus-a-falsafa/02-falsafa/al-kindi/episode-01.m4a
                rel_path = key[len(cursus_folder) + 1:]  # remove cursus folder + /
                parts = rel_path.split('/')
                
                if len(parts) < 2:
                    stats['skipped'] += 1
                    continue
                
                # Normalize module folder: strip accents, spaces, trailing chars
                module_folder_raw = parts[0]
                module_folder = module_folder_raw.strip().lower()
                # Remove accents for matching
                import unicodedata
                module_folder_normalized = unicodedata.normalize('NFD', module_folder)
                module_folder_normalized = ''.join(c for c in module_folder_normalized if unicodedata.category(c) != 'Mn')
                module_folder_normalized = module_folder_normalized.strip()
                
                # Try to find course_id from mapping
                course_id = None
                module_num_str = None
                
                # First try exact match
                course_id = R2_TO_COURSE_MAPPING.get(module_folder)
                
                # Try normalized match
                if not course_id:
                    course_id = R2_TO_COURSE_MAPPING.get(module_folder_normalized)
                
                # Try fuzzy match: extract number prefix and check all mappings
                if not course_id:
                    match = re.match(r'^(\d+)-', module_folder_normalized)
                    if match:
                        num_prefix = match.group(1)
                        for mk, mv in R2_TO_COURSE_MAPPING.items():
                            if mk.startswith(f"{num_prefix}-"):
                                course_id = mv
                                break
                
                if not course_id:
                    stats['skipped'] += 1
                    stats['errors'].append(f"No mapping for folder: {module_folder_raw}")
                    continue
                
                # Get course from DB
                course = await db.courses.find_one({'id': course_id}, {'_id': 0})
                if not course:
                    stats['skipped'] += 1
                    stats['errors'].append(f"Course not found: {course_id}")
                    continue
                
                # Parse subfolder and episode
                episode_title = ""
                episode_number = 1
                
                if len(parts) == 3:
                    # Has subfolder: module-folder/subfolder/episode-N.m4a
                    subfolder = parts[1]
                    filename = parts[2]
                    episode_title = subfolder.replace('-', ' ').title()
                    match = re.search(r'episode-?(\d+)', filename, re.IGNORECASE)
                    if match:
                        episode_number = int(match.group(1))
                elif len(parts) == 2:
                    # No subfolder: module-folder/episode-N.m4a
                    filename = parts[1]
                    match = re.search(r'episode-?(\d+)', filename, re.IGNORECASE)
                    if match:
                        episode_number = int(match.group(1))
                
                # Generate audio ID
                subfolder_slug = parts[1].replace(' ', '-').lower() if len(parts) == 3 else ''
                if subfolder_slug:
                    audio_id = f"aud_{course_id}-{subfolder_slug}-ep{episode_number:02d}"
                else:
                    audio_id = f"aud_{course_id}-ep{episode_number:02d}"
                audio_id = re.sub(r'[^a-z0-9_-]', '', audio_id)
                
                # Build title
                if episode_title:
                    title = episode_title
                else:
                    title = clean_title(course.get('title', ''))
                
                audio_doc = {
                    'id': audio_id,
                    'title': title,
                    'description': f"{title} — {clean_title(course.get('title', ''))}",
                    'scholar_id': course.get('scholar_id', ''),
                    'scholar_name': course.get('scholar_name', ''),
                    'duration': 0,
                    'audio_url': '',
                    'file_key': key,
                    'thumbnail': course.get('thumbnail', ''),
                    'topic': course.get('topic', ''),
                    'type': 'lecture',
                    'course_id': course_id,
                    'cursus_id': cursus_id,
                    'episode_number': episode_number,
                    'published_at': now.isoformat(),
                    'is_active': True,
                }
                
                # Preserve existing duration if already set
                existing = await db.audios.find_one({'id': audio_id}, {'_id': 0, 'duration': 1})
                if existing and existing.get('duration', 0) > 0:
                    audio_doc['duration'] = existing['duration']
                
                result = await db.audios.update_one(
                    {'id': audio_id},
                    {'$set': audio_doc},
                    upsert=True
                )
                
                synced_audio_ids.add(audio_id)
                
                if result.upserted_id:
                    stats['created'] += 1
                else:
                    stats['updated'] += 1
        
        # Clean up orphaned DB entries:
        # 1. file_key no longer exists in R2
        # 2. audio ID was not part of this sync (legacy entries with wrong course mapping)
        all_audios = await db.audios.find({'file_key': {'$exists': True, '$ne': ''}}, {'_id': 0, 'id': 1, 'file_key': 1}).to_list(2000)
        orphan_ids = []
        for audio in all_audios:
            fk = audio.get('file_key', '')
            aid = audio.get('id', '')
            if not fk:
                continue
            # If file_key points to a cursus-* path, it must be in synced_audio_ids
            if fk.startswith('cursus-'):
                if aid not in synced_audio_ids:
                    orphan_ids.append(aid)
            # Legacy Audio/ paths: always orphaned since we don't use Audio/ anymore
            elif fk.startswith('Audio/'):
                orphan_ids.append(aid)
        
        if orphan_ids:
            result = await db.audios.delete_many({'id': {'$in': orphan_ids}})
            stats['deleted_orphans'] = result.deleted_count
            logger.info(f"Deleted {result.deleted_count} orphaned audio entries: {orphan_ids[:10]}")
        
        return {
            'message': 'Synchronisation R2 terminée',
            'stats': stats
        }
        
    except ClientError as e:
        raise HTTPException(500, f"Erreur R2: {str(e)}")

@api_router.post("/admin/update-professor-photos")
async def update_professor_photos(request: Request):
    """Update professor photos from R2 images folder."""
    await require_admin(request)
    if not r2_client:
        raise HTTPException(503, "R2 non configuré")
    
    # Mapping of R2 filenames to scholar IDs
    photo_mapping = {
        'Prof_MeryemSebti.jpg': 'sch-sebti',
        'Prof_ChristianJambet.png': 'sch-jambet',
        'Prof_EricGeoffroy.jpg': 'sch-geoffroy',
        'Prof_1973_HCorbin.jpeg': 'sch-corbin',
    }
    
    updated = []
    for filename, scholar_id in photo_mapping.items():
        r2_key = f"images/{filename}"
        photo_url = f"https://{R2_PUBLIC_URL}/{r2_key}" if R2_PUBLIC_URL else f"/api/audios/stream/{r2_key}"
        
        result = await db.scholars.update_one(
            {'id': scholar_id},
            {'$set': {'photo': photo_url}},
            upsert=False
        )
        
        if result.modified_count > 0:
            updated.append(scholar_id)
    
    return {
        'message': 'Photos des professeurs mises à jour',
        'updated': updated
    }

@api_router.get("/admin/courses/{course_id}/episodes")
async def get_course_episodes(course_id: str, request: Request):
    """Get all audio episodes linked to a course."""
    await require_admin(request)
    episodes = await db.audios.find(
        {'course_id': course_id},
        {'_id': 0}
    ).sort('episode_number', 1).to_list(100)
    
    for ep in episodes:
        ep['stream_url'] = resolve_audio_url(ep)
    
    return {'course_id': course_id, 'episodes': episodes, 'count': len(episodes)}


# ─── Admin: Top 10 Courses Management ─────────────────────────────────────────

@api_router.get("/admin/top10")
async def get_admin_top10(request: Request):
    """Get Top 10 config and auto-ranked courses by play_count."""
    await require_admin(request)
    config = await db.config.find_one({'key': 'top10_courses'}, {'_id': 0})
    manual_ids = config.get('course_ids', []) if config else []
    all_courses = await db.courses.find({'is_active': True}, {'_id': 0}).sort('play_count', -1).limit(20).to_list(20)
    return {
        'manual_ids': manual_ids,
        'all_courses': all_courses,
    }


@api_router.put("/admin/top10")
async def update_admin_top10(body: Top10UpdateRequest, request: Request):
    """Set the Top 10 courses manually."""
    await require_admin(request)
    await db.config.update_one(
        {'key': 'top10_courses'},
        {'$set': {
            'key': 'top10_courses',
            'course_ids': body.course_ids,
            'updated_at': datetime.now(timezone.utc).isoformat()
        }},
        upsert=True
    )
    return {'message': 'Top 10 mis à jour', 'course_ids': body.course_ids}

# ─── Admin: Highlight/Featured Management ──────────────────────────────────────

@api_router.get("/admin/highlight")
async def get_highlight_config(request: Request):
    """Get the current highlight configuration."""
    await require_admin(request)
    
    # Get highlight config
    config = await db.config.find_one({'key': 'highlight_config'}, {'_id': 0})
    if not config:
        config = {
            'key': 'highlight_config',
            'mode': 'manual',  # 'manual' or 'random'
        }
    
    # Get current featured course
    featured_course = await db.courses.find_one({'is_featured': True, 'is_active': True}, {'_id': 0, 'id': 1, 'title': 1})
    
    # Get current featured cursus
    featured_cursus = await db.cursus.find_one({'is_featured': True, 'is_active': True}, {'_id': 0, 'id': 1, 'name': 1})
    
    # Get all courses and cursus for selection
    all_courses = await db.courses.find({'is_active': True}, {'_id': 0, 'id': 1, 'title': 1, 'is_featured': 1}).to_list(100)
    all_cursus = await db.cursus.find({'is_active': True}, {'_id': 0, 'id': 1, 'name': 1, 'is_featured': 1}).to_list(50)
    
    return {
        'mode': config.get('mode', 'manual'),
        'featured_course': featured_course,
        'featured_cursus': featured_cursus,
        'all_courses': all_courses,
        'all_cursus': all_cursus
    }

@api_router.put("/admin/highlight/mode")
async def set_highlight_mode(request: Request):
    """Set highlight mode (manual or random)."""
    await require_admin(request)
    body = await request.json()
    mode = body.get('mode', 'manual')
    
    if mode not in ['manual', 'random']:
        raise HTTPException(400, "Mode invalide. Utilisez 'manual' ou 'random'.")
    
    await db.config.update_one(
        {'key': 'highlight_config'},
        {'$set': {
            'key': 'highlight_config',
            'mode': mode,
            'updated_at': datetime.now(timezone.utc).isoformat()
        }},
        upsert=True
    )
    
    return {'message': f'Mode highlight changé en: {mode}', 'mode': mode}

@api_router.patch("/admin/courses/{course_id}/set-featured")
async def set_course_featured(course_id: str, request: Request):
    """Set a course as featured (unfeaturing any previous)."""
    await require_admin(request)
    
    # First, unfeatured all courses and cursus
    await db.courses.update_many({}, {'$set': {'is_featured': False}})
    await db.cursus.update_many({}, {'$set': {'is_featured': False}})
    
    # Set this course as featured
    result = await db.courses.update_one(
        {'id': course_id},
        {'$set': {'is_featured': True}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(404, "Cours non trouvé")
    
    return {'message': f'Cours {course_id} mis en avant', 'course_id': course_id}

@api_router.patch("/admin/cursus/{cursus_id}/set-featured")
async def set_cursus_featured(cursus_id: str, request: Request):
    """Set a cursus as featured (unfeaturing any previous)."""
    await require_admin(request)
    
    # First, unfeatured all courses and cursus
    await db.courses.update_many({}, {'$set': {'is_featured': False}})
    await db.cursus.update_many({}, {'$set': {'is_featured': False}})
    
    # Set this cursus as featured
    result = await db.cursus.update_one(
        {'id': cursus_id},
        {'$set': {'is_featured': True}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(404, "Cursus non trouvé")
    
    return {'message': f'Cursus {cursus_id} mis en avant', 'cursus_id': cursus_id}

@api_router.delete("/admin/highlight/clear")
async def clear_featured(request: Request):
    """Clear all featured items."""
    await require_admin(request)
    
    await db.courses.update_many({}, {'$set': {'is_featured': False}})
    await db.cursus.update_many({}, {'$set': {'is_featured': False}})
    
    return {'message': 'Highlight effacé'}

# ─── Admin: Scholar Toggle ─────────────────────────────────────────────────────

@api_router.patch("/admin/scholars/{scholar_id}/toggle")
async def admin_toggle_scholar(scholar_id: str, request: Request):
    await require_admin(request)
    doc = await db.scholars.find_one({'id': scholar_id})
    if not doc:
        raise HTTPException(404, "Érudit non trouvé")
    new_status = not doc.get('is_active', True)
    await db.scholars.update_one({'id': scholar_id}, {'$set': {'is_active': new_status}})
    return {'id': scholar_id, 'is_active': new_status}

# ─── Admin: User Management ────────────────────────────────────────────────────

@api_router.get("/admin/users")
async def admin_list_users(request: Request):
    await require_admin(request)
    users = await db.users.find({}, {'_id': 0, 'password_hash': 0}).to_list(500)
    return users

@api_router.get("/admin/kpis")
async def admin_kpis(request: Request):
    """Aggregate commercial KPIs for the admin dashboard.

    Returns estimated MRR (EUR), active subscriber count (by plan interval),
    free-trial users, total preregistrations, gift cards purchased and
    redeemed. All numbers are derived live from MongoDB collections.
    """
    await require_admin(request)
    now = datetime.now(timezone.utc)

    # Load active plans price/interval map.
    plans = await db.plans.find({}, {'_id': 0, 'plan_id': 1, 'price': 1, 'interval': 1}).to_list(100)
    plan_map = {p['plan_id']: p for p in plans}

    monthly_active = 0
    yearly_active = 0
    gift_active = 0
    free_access = 0
    trial_active = 0
    mrr = 0.0

    async for u in db.users.find(
        {},
        {'_id': 0, 'subscription_status': 1, 'subscription_plan': 1, 'subscription_source': 1,
         'subscription_end_date': 1, 'free_access': 1, 'has_free_access': 1,
         'trial': 1, 'trial_end_date': 1}
    ):
        if u.get('free_access') or u.get('has_free_access'):
            free_access += 1

        status = u.get('subscription_status')
        end = u.get('subscription_end_date')
        # Normalise end date.
        end_dt = None
        if isinstance(end, datetime):
            end_dt = end if end.tzinfo else end.replace(tzinfo=timezone.utc)
        elif isinstance(end, str):
            try:
                end_dt = datetime.fromisoformat(end.replace('Z', '+00:00'))
            except Exception:
                end_dt = None
        is_active = status == 'active' and (end_dt is None or end_dt > now)

        if is_active:
            plan_id = u.get('subscription_plan')
            plan = plan_map.get(plan_id) if plan_id else None
            source = u.get('subscription_source')
            price = float((plan or {}).get('price') or 0)
            interval = (plan or {}).get('interval') or 'month'
            if source == 'gift_card':
                gift_active += 1
            elif interval == 'year':
                yearly_active += 1
                mrr += price / 12.0
            else:
                monthly_active += 1
                mrr += price

        # Trial detection (legacy & new schemas)
        trial = u.get('trial') or {}
        t_end = u.get('trial_end_date') or trial.get('end_date')
        if isinstance(t_end, str):
            try:
                t_end = datetime.fromisoformat(t_end.replace('Z', '+00:00'))
            except Exception:
                t_end = None
        if isinstance(t_end, datetime):
            if (t_end if t_end.tzinfo else t_end.replace(tzinfo=timezone.utc)) > now:
                trial_active += 1

    preregistrations = await db.preregistrations.count_documents({})
    gift_purchased = await db.gift_cards.count_documents({'status': {'$in': ['paid', 'redeemed']}})
    gift_redeemed = await db.gift_cards.count_documents({'status': 'redeemed'})
    gift_revenue_doc = await db.gift_cards.aggregate([
        {'$match': {'status': {'$in': ['paid', 'redeemed']}}},
        {'$group': {'_id': None, 'total': {'$sum': '$amount_paid'}}},
    ]).to_list(1)
    gift_revenue = float((gift_revenue_doc[0]['total'] if gift_revenue_doc else 0) or 0)

    return {
        'mrr_eur': round(mrr, 2),
        'arr_eur': round(mrr * 12, 2),
        'subscribers_active_total': monthly_active + yearly_active + gift_active,
        'subscribers_monthly': monthly_active,
        'subscribers_yearly': yearly_active,
        'subscribers_gift': gift_active,
        'free_access_users': free_access,
        'trial_active': trial_active,
        'preregistrations_total': preregistrations,
        'gift_cards_purchased': gift_purchased,
        'gift_cards_redeemed': gift_redeemed,
        'gift_cards_revenue_eur': round(gift_revenue, 2),
        'generated_at': now.isoformat(),
    }

@api_router.get("/admin/users/{user_id}")
async def admin_get_user(user_id: str, request: Request):
    await require_admin(request)
    user = await db.users.find_one({'user_id': user_id}, {'_id': 0, 'password_hash': 0})
    if not user:
        raise HTTPException(404, "Utilisateur non trouvé")
    return user

@api_router.post("/admin/users/{user_id}/grant-access")
async def admin_grant_free_access(user_id: str, request: Request):
    """Grant free access to a user."""
    await require_admin(request)
    now = datetime.now(timezone.utc)
    result = await db.users.update_one(
        {'user_id': user_id},
        {'$set': {
            'free_access': True,
            'has_free_access': True,  # For backwards compatibility
            'free_access_granted_at': now.isoformat()
        }}
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Utilisateur non trouvé")
    logger.info(f"Free access granted to user {user_id}")
    return {'message': 'Accès gratuit accordé', 'user_id': user_id}

@api_router.post("/admin/users/{user_id}/revoke-access")
async def admin_revoke_access(user_id: str, request: Request):
    """Revoke all access from a user (free access and subscription)."""
    await require_admin(request)
    result = await db.users.update_one(
        {'user_id': user_id},
        {
            '$set': {'free_access': False, 'has_free_access': False},
            '$unset': {'free_access_granted_at': '', 'subscription': '', 'trial': ''}
        }
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Utilisateur non trouvé")
    return {'message': 'Accès révoqué', 'user_id': user_id}

class ExtendSubscriptionRequest(BaseModel):
    days: int

class GrantSubscriptionRequest(BaseModel):
    plan_id: str  # monthly or annual

@api_router.post("/admin/users/{user_id}/extend-subscription")
async def admin_extend_subscription(user_id: str, body: ExtendSubscriptionRequest, request: Request):
    """Extend a user's subscription by a number of days."""
    await require_admin(request)
    user = await db.users.find_one({'user_id': user_id})
    if not user:
        raise HTTPException(404, "Utilisateur non trouvé")
    
    now = datetime.now(timezone.utc)
    current_expires = None
    
    def parse_date(date_val):
        """Parse a date value to a timezone-aware datetime."""
        if date_val is None:
            return None
        if isinstance(date_val, datetime):
            if date_val.tzinfo is None:
                return date_val.replace(tzinfo=timezone.utc)
            return date_val
        if isinstance(date_val, str):
            try:
                dt = datetime.fromisoformat(date_val.replace('Z', '+00:00'))
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=timezone.utc)
                return dt
            except:
                return None
        return None
    
    # Check if user has an existing subscription or trial
    if user.get('subscription') and user['subscription'].get('expires_at'):
        current_expires = parse_date(user['subscription']['expires_at'])
    elif user.get('trial') and user['trial'].get('expires_at'):
        current_expires = parse_date(user['trial']['expires_at'])
    
    # Calculate new expiration date
    if current_expires and current_expires > now:
        # Extend from current expiration
        new_expires = current_expires + timedelta(days=body.days)
    else:
        # Start fresh from now
        new_expires = now + timedelta(days=body.days)
    
    # Update subscription
    await db.users.update_one(
        {'user_id': user_id},
        {'$set': {
            'subscription': {
                'plan_id': user.get('subscription', {}).get('plan_id', 'manual'),
                'expires_at': new_expires.isoformat(),
                'status': 'active',
                'extended_by_admin': True,
                'extended_at': now.isoformat()
            }
        }}
    )
    
    logger.info(f"Subscription extended for user {user_id} by {body.days} days until {new_expires.isoformat()}")
    return {
        'message': f'Abonnement prolongé de {body.days} jours',
        'user_id': user_id,
        'new_expires_at': new_expires.isoformat()
    }

@api_router.post("/admin/users/{user_id}/grant-subscription")
async def admin_grant_subscription(user_id: str, body: GrantSubscriptionRequest, request: Request):
    """Grant a subscription to a user (monthly or annual)."""
    await require_admin(request)
    user = await db.users.find_one({'user_id': user_id})
    if not user:
        raise HTTPException(404, "Utilisateur non trouvé")
    
    now = datetime.now(timezone.utc)
    
    # Determine duration based on plan
    if body.plan_id == 'annual':
        days = 365
        plan_name = 'Annuel'
    else:  # monthly
        days = 30
        plan_name = 'Mensuel'
    
    expires_at = now + timedelta(days=days)
    
    # Update user with new subscription
    await db.users.update_one(
        {'user_id': user_id},
        {'$set': {
            'subscription': {
                'plan_id': body.plan_id,
                'plan_name': plan_name,
                'expires_at': expires_at.isoformat(),
                'status': 'active',
                'granted_by_admin': True,
                'granted_at': now.isoformat()
            }
        }}
    )
    
    logger.info(f"Subscription {body.plan_id} granted to user {user_id} until {expires_at.isoformat()}")
    return {
        'message': f'Abonnement {plan_name} accordé',
        'user_id': user_id,
        'expires_at': expires_at.isoformat()
    }

# ─── Admin: Referrals Management ───────────────────────────────────────────────

@api_router.get("/admin/referrals")
async def admin_list_referrals(request: Request):
    """List all referrals."""
    await require_admin(request)
    referrals = await db.referrals.find({}, {'_id': 0}).sort('created_at', -1).to_list(200)
    return referrals

@api_router.get("/admin/referrals/stats")
async def admin_referral_stats(request: Request):
    """Get referral statistics."""
    await require_admin(request)
    
    total_referrals = await db.referrals.count_documents({})
    converted_referrals = await db.referrals.count_documents({'status': 'converted'})
    pending_referrals = await db.referrals.count_documents({'status': 'pending'})
    
    # Top referrers
    pipeline = [
        {'$match': {'referral_count': {'$gt': 0}}},
        {'$sort': {'referral_count': -1}},
        {'$limit': 10},
        {'$project': {
            '_id': 0,
            'user_id': 1,
            'name': 1,
            'email': 1,
            'referral_count': 1,
            'free_months_earned': 1
        }}
    ]
    top_referrers = await db.users.aggregate(pipeline).to_list(10)
    
    return {
        'total_referrals': total_referrals,
        'converted': converted_referrals,
        'pending': pending_referrals,
        'conversion_rate': round(converted_referrals / total_referrals * 100, 1) if total_referrals > 0 else 0,
        'top_referrers': top_referrers
    }

@api_router.post("/admin/users/{user_id}/grant-free-months")
async def admin_grant_free_months(user_id: str, body: GrantFreeMonthRequest, request: Request):
    """Grant free months to a user (for professors, sponsors, etc.)."""
    await require_admin(request)
    
    user = await db.users.find_one({'user_id': user_id})
    if not user:
        raise HTTPException(404, "Utilisateur non trouvé")
    
    now = datetime.now(timezone.utc)
    
    # Calculate new subscription end date
    current_end = user.get('subscription_end_date')
    if current_end:
        if isinstance(current_end, str):
            current_end = datetime.fromisoformat(current_end.replace('Z', '+00:00'))
        if current_end.tzinfo is None:
            current_end = current_end.replace(tzinfo=timezone.utc)
    
    if not current_end or current_end < now:
        current_end = now
    
    new_end = current_end + timedelta(days=30 * body.months)
    
    # Update user
    await db.users.update_one(
        {'user_id': user_id},
        {
            '$inc': {'free_months_remaining': body.months},
            '$set': {
                'subscription_end_date': new_end,
                'subscription': {
                    'plan_id': 'free_grant',
                    'plan_name': f'{body.months} mois gratuit(s) - {body.reason}',
                    'expires_at': new_end.isoformat(),
                    'status': 'active',
                    'granted_by_admin': True,
                    'granted_at': now.isoformat(),
                    'reason': body.reason
                }
            }
        }
    )
    
    logger.info(f"Granted {body.months} free months to user {user_id} ({body.reason})")
    return {
        'message': f'{body.months} mois gratuit(s) accordé(s)',
        'user_id': user_id,
        'new_expires_at': new_end.isoformat(),
        'reason': body.reason
    }

@api_router.post("/admin/users/{user_id}/grant-lifetime")
async def admin_grant_lifetime(user_id: str, request: Request, reason: str = "professeur"):
    """Grant lifetime access to a user (for professors, sponsors, etc.)."""
    await require_admin(request)
    
    user = await db.users.find_one({'user_id': user_id})
    if not user:
        raise HTTPException(404, "Utilisateur non trouvé")
    
    now = datetime.now(timezone.utc)
    lifetime_end = now + timedelta(days=365 * 100)  # 100 years
    
    # Update user with lifetime subscription
    await db.users.update_one(
        {'user_id': user_id},
        {'$set': {
            'subscription_end_date': lifetime_end,
            'subscription': {
                'plan_id': 'lifetime',
                'plan_name': f'Accès à vie - {reason}',
                'expires_at': lifetime_end.isoformat(),
                'status': 'active',
                'is_lifetime': True,
                'granted_by_admin': True,
                'granted_at': now.isoformat(),
                'reason': reason
            }
        }}
    )
    
    logger.info(f"Granted lifetime access to user {user_id} ({reason})")
    return {
        'message': f'Accès à vie accordé ({reason})',
        'user_id': user_id,
        'is_lifetime': True
    }

# ─── Admin: Settings Management ────────────────────────────────────────────────

class StripeSettingsRequest(BaseModel):
    api_key: str
    webhook_secret: Optional[str] = None

class PricingSettingsRequest(BaseModel):
    trial_days: int = 3
    monthly: float = 9.99
    annual: float = 89.99

class ReferralSettingsRequest(BaseModel):
    enabled: bool = True
    referrer_months: int = 1
    referee_months: int = 1

@api_router.get("/admin/settings")
async def admin_get_settings(request: Request):
    """Get all platform settings."""
    await require_admin(request)
    
    # Get settings from database or return defaults
    settings = await db.settings.find_one({'id': 'platform_settings'}, {'_id': 0})
    if not settings:
        settings = {
            'id': 'platform_settings',
            'stripe_configured': False,
            'pricing': {
                'trial_days': 3,
                'monthly': 9.99,
                'annual': 89.99
            },
            'referral': {
                'enabled': True,
                'referrer_months': 1,
                'referee_months': 1
            }
        }
    
    # Check if Stripe is configured from environment
    stripe_key = os.environ.get('STRIPE_API_KEY', '')
    if stripe_key and stripe_key.startswith('sk_'):
        settings['stripe_configured'] = True
        settings['stripe_key_last4'] = stripe_key[-4:]
        settings['stripe_mode'] = 'live' if 'live' in stripe_key else 'test'
    
    return settings

@api_router.post("/admin/settings/stripe")
async def admin_save_stripe_settings(body: StripeSettingsRequest, request: Request):
    """Save Stripe API configuration."""
    await require_admin(request)
    
    # Validate the key format
    if not body.api_key.startswith('sk_'):
        raise HTTPException(400, "La clé API doit commencer par sk_test_ ou sk_live_")
    
    # Update the .env file
    env_path = '/app/backend/.env'
    
    # Read current .env content
    with open(env_path, 'r') as f:
        lines = f.readlines()
    
    # Update or add STRIPE_API_KEY
    key_found = False
    new_lines = []
    for line in lines:
        if line.startswith('STRIPE_API_KEY='):
            new_lines.append(f'STRIPE_API_KEY={body.api_key}\n')
            key_found = True
        else:
            new_lines.append(line)
    
    if not key_found:
        new_lines.append(f'STRIPE_API_KEY={body.api_key}\n')
    
    # Add webhook secret if provided
    if body.webhook_secret:
        webhook_found = False
        final_lines = []
        for line in new_lines:
            if line.startswith('STRIPE_WEBHOOK_SECRET='):
                final_lines.append(f'STRIPE_WEBHOOK_SECRET={body.webhook_secret}\n')
                webhook_found = True
            else:
                final_lines.append(line)
        if not webhook_found:
            final_lines.append(f'STRIPE_WEBHOOK_SECRET={body.webhook_secret}\n')
        new_lines = final_lines
    
    # Write updated .env
    with open(env_path, 'w') as f:
        f.writelines(new_lines)
    
    # Update settings in database
    await db.settings.update_one(
        {'id': 'platform_settings'},
        {'$set': {
            'stripe_configured': True,
            'stripe_key_last4': body.api_key[-4:],
            'stripe_mode': 'live' if 'live' in body.api_key else 'test',
            'updated_at': datetime.now(timezone.utc)
        }},
        upsert=True
    )
    
    logger.info(f"Stripe API key updated (mode: {'live' if 'live' in body.api_key else 'test'})")
    return {'message': 'Configuration Stripe enregistrée', 'mode': 'live' if 'live' in body.api_key else 'test'}

@api_router.post("/admin/settings/pricing")
async def admin_save_pricing_settings(body: PricingSettingsRequest, request: Request):
    """Save pricing configuration."""
    await require_admin(request)
    
    await db.settings.update_one(
        {'id': 'platform_settings'},
        {'$set': {
            'pricing': {
                'trial_days': body.trial_days,
                'monthly': body.monthly,
                'annual': body.annual
            },
            'updated_at': datetime.now(timezone.utc)
        }},
        upsert=True
    )
    
    logger.info(f"Pricing updated: trial={body.trial_days}d, monthly={body.monthly}€, annual={body.annual}€")
    return {'message': 'Tarifs enregistrés'}

@api_router.post("/admin/settings/referral")
async def admin_save_referral_settings(body: ReferralSettingsRequest, request: Request):
    """Save referral system configuration."""
    await require_admin(request)
    
    await db.settings.update_one(
        {'id': 'platform_settings'},
        {'$set': {
            'referral': {
                'enabled': body.enabled,
                'referrer_months': body.referrer_months,
                'referee_months': body.referee_months
            },
            'updated_at': datetime.now(timezone.utc)
        }},
        upsert=True
    )
    
    logger.info(f"Referral settings updated: enabled={body.enabled}, referrer={body.referrer_months}m, referee={body.referee_months}m")
    return {'message': 'Configuration parrainage enregistrée'}

# ─── Admin Panel: Settings Page ────────────────────────────────────────────────

@api_router.get("/admin-panel/settings")
async def admin_panel_settings(request: Request):
    """Serve the settings admin page."""
    return templates.TemplateResponse(
        "settings_new.html",
        {"request": request, "active_page": "settings"}
    )

# ─── Admin: Thematiques CRUD ───────────────────────────────────────────────────

# ─── Admin: Cursus CRUD (was Thematiques) ──────────────────────────────────────

@api_router.get("/admin/cursus")
async def admin_list_cursus(request: Request):
    """List all cursus (was thematiques)."""
    await require_admin(request)
    cursus_list = await db.cursus.find({}, {'_id': 0}).sort('order', 1).to_list(100)
    # Count courses for each cursus
    for c in cursus_list:
        c['course_count'] = await db.courses.count_documents({'cursus_id': c['id']})
    return cursus_list

# Keep old endpoint for compatibility
@api_router.get("/admin/thematiques")
async def admin_list_thematiques_compat(request: Request):
    """Compatibility endpoint - redirects to cursus."""
    return await admin_list_cursus(request)

@api_router.post("/admin/thematiques")
async def admin_create_thematique_compat(body: CursusCreate, request: Request):
    """Compatibility endpoint for web admin panel."""
    return await admin_create_cursus(body, request)

@api_router.put("/admin/thematiques/{cursus_id}")
async def admin_update_thematique_compat(cursus_id: str, body: CursusUpdate, request: Request):
    """Compatibility endpoint for web admin panel — saves description to cursus."""
    return await admin_update_cursus(cursus_id, body, request)

@api_router.delete("/admin/thematiques/{cursus_id}")
async def admin_delete_thematique_compat(cursus_id: str, request: Request):
    """Compatibility endpoint for web admin panel."""
    return await admin_delete_cursus(cursus_id, request)

@api_router.get("/cursus")
async def public_list_cursus():
    """Public endpoint to list active cursus.

    The `course_count` is recomputed live so it always reflects the current
    state of the `courses` collection (the value persisted on the cursus doc
    can drift after migrations like v15e that add/remove courses).
    """
    cursus_list = await db.cursus.find({'is_active': True}, {'_id': 0}).sort('order', 1).to_list(100)
    count_pipeline = await db.courses.aggregate([
        {'$match': {'is_active': {'$ne': False}}},
        {'$group': {'_id': '$cursus_id', 'n': {'$sum': 1}}},
    ]).to_list(50)
    count_map = {c['_id']: c['n'] for c in count_pipeline}
    for c in cursus_list:
        c['course_count'] = count_map.get(c['id'], 0)
    return cursus_list

@api_router.get("/cursus/{cursus_id}/scholars")
async def get_cursus_scholars(cursus_id: str):
    """Get all scholars teaching courses in a cursus."""
    # Find all courses in this cursus
    courses = await db.courses.find(
        {'$or': [{'cursus_id': cursus_id}, {'thematique_id': cursus_id}], 'is_active': {'$ne': False}},
        {'_id': 0, 'scholar_id': 1, 'scholar_name': 1, 'title': 1, 'id': 1}
    ).to_list(100)
    
    # Collect unique scholar IDs and their courses
    scholars_map = {}
    for c in courses:
        scholar_id = c.get('scholar_id')
        if scholar_id:
            if scholar_id not in scholars_map:
                scholars_map[scholar_id] = {
                    'id': scholar_id,
                    'name': c.get('scholar_name', ''),
                    'courses': []
                }
            scholars_map[scholar_id]['courses'].append({
                'id': c['id'],
                'title': clean_title(c.get('title', ''))
            })
    
    # Enrich with scholar details from scholars collection
    result = []
    for sid, data in scholars_map.items():
        scholar = await db.scholars.find_one({'id': sid}, {'_id': 0})
        if scholar:
            result.append({
                'id': sid,
                'name': scholar.get('name', data['name']),
                'title': scholar.get('title', ''),
                'bio': scholar.get('bio', ''),
                'photo': scholar.get('photo', ''),
                'courses_count': len(data['courses']),
                'courses': data['courses']
            })
        else:
            # Fallback if no scholar doc
            result.append({
                'id': sid,
                'name': data['name'],
                'title': '',
                'bio': '',
                'photo': '',
                'courses_count': len(data['courses']),
                'courses': data['courses']
            })
    
    return result

@api_router.get("/cursus/{cursus_id}/resources")
async def get_cursus_resources(cursus_id: str):
    """Get resources for a cursus (books, articles, links)."""
    # Check if there's a resources collection or field in cursus
    cursus = await db.cursus.find_one({'id': cursus_id}, {'_id': 0})
    if not cursus:
        raise HTTPException(404, "Cursus non trouvé")
    
    # For now, return empty list as resources are not yet implemented
    # This can be expanded later to include books, PDFs, external links
    resources = await db.resources.find({'cursus_id': cursus_id, 'is_active': True}, {'_id': 0}).to_list(50)
    return resources

@api_router.get("/thematiques")
async def public_list_thematiques_compat():
    """Compatibility endpoint."""
    return await public_list_cursus()

@api_router.post("/admin/cursus")
async def admin_create_cursus(body: CursusCreate, request: Request):
    await require_admin(request)
    cursus_id = f"cursus_{uuid.uuid4().hex[:8]}"
    last = await db.cursus.find_one(sort=[('order', -1)])
    next_order = (last.get('order', 0) + 1) if last else 1
    
    doc = {
        'id': cursus_id,
        'name': body.name,
        'description': body.description,
        'icon': body.icon,
        'order': body.order or next_order,
        'is_active': body.is_active,
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    await db.cursus.insert_one(doc)
    return {'message': 'Cursus créé', 'id': cursus_id, 'cursus': {k: v for k, v in doc.items() if k != '_id'}}

@api_router.put("/admin/cursus/{cursus_id}")
async def admin_update_cursus(cursus_id: str, body: CursusUpdate, request: Request):
    await require_admin(request)
    # Use exclude_unset to only update provided fields; allows explicit null to clear hero_title/hero_description
    raw = body.model_dump(exclude_unset=True)
    update = {}
    for k, v in raw.items():
        if v is not None:
            update[k] = v
        elif k in ('hero_title', 'hero_description'):
            # Allow explicit null to clear hero text
            update[k] = None
    if update:
        update['updated_at'] = datetime.now(timezone.utc).isoformat()
        # Lock against the on-startup cursus seed (see Migration v3) so the
        # admin's edits to description/subtitle/title survive redeploys.
        update['seed_locked'] = True
        result = await db.cursus.update_one({'id': cursus_id}, {'$set': update})
        if result.matched_count == 0:
            raise HTTPException(404, "Cursus non trouvé")
    return {'message': 'Cursus mis à jour', 'id': cursus_id}

@api_router.delete("/admin/cursus/{cursus_id}")
async def admin_delete_cursus(cursus_id: str, request: Request):
    await require_admin(request)
    # Check for linked courses
    course_count = await db.courses.count_documents({'cursus_id': cursus_id})
    if course_count > 0:
        raise HTTPException(400, f"Impossible de supprimer: {course_count} cours liés")
    result = await db.cursus.delete_one({'id': cursus_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Cursus non trouvé")
    return {'message': 'Cursus supprimé'}

@api_router.patch("/admin/cursus/{cursus_id}/toggle")
async def admin_toggle_cursus(cursus_id: str, request: Request):
    await require_admin(request)
    doc = await db.cursus.find_one({'id': cursus_id})
    if not doc:
        raise HTTPException(404, "Cursus non trouvé")
    new_status = not doc.get('is_active', False)
    await db.cursus.update_one({'id': cursus_id}, {'$set': {'is_active': new_status}})
    return {'id': cursus_id, 'is_active': new_status}

@api_router.post("/admin/cursus/bulk-toggle")
async def admin_bulk_toggle_cursus(body: BulkToggleRequest, request: Request):
    """Toggle multiple cursus at once."""
    await require_admin(request)
    result = await db.cursus.update_many(
        {'id': {'$in': body.ids}},
        {'$set': {'is_active': body.is_active, 'updated_at': datetime.now(timezone.utc).isoformat()}}
    )
    return {'message': f'{result.modified_count} cursus mis à jour', 'is_active': body.is_active}

# ─── Admin: Modules CRUD ───────────────────────────────────────────────────────

@api_router.get("/admin/modules")
async def admin_list_modules(request: Request, course_id: Optional[str] = None):
    """List all modules, optionally filtered by course_id."""
    await require_admin(request)
    query = {}
    if course_id:
        query['course_id'] = course_id
    modules = await db.modules.find(query, {'_id': 0}).sort('order', 1).to_list(500)
    # Count episodes for each module
    for m in modules:
        m['episode_count_actual'] = await db.audios.count_documents({'module_id': m['id']})
    return modules

@api_router.get("/modules")
async def public_list_modules(course_id: Optional[str] = None):
    """Public endpoint to list active modules."""
    query = {'is_active': True}
    if course_id:
        query['course_id'] = course_id
    modules = await db.modules.find(query, {'_id': 0}).sort('order', 1).to_list(500)
    return modules

@api_router.post("/admin/modules")
async def admin_create_module(body: ModuleCreate, request: Request):
    await require_admin(request)
    module_id = f"mod_{uuid.uuid4().hex[:8]}"
    # Get max order for this course
    last = await db.modules.find_one({'course_id': body.course_id}, sort=[('order', -1)])
    next_order = (last.get('order', 0) + 1) if last else 1
    
    doc = {
        'id': module_id,
        'name': body.name,
        'description': body.description,
        'course_id': body.course_id,
        'scholar_name': body.scholar_name,
        'order': body.order or next_order,
        'episode_count': body.episode_count,
        'is_active': body.is_active,
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    await db.modules.insert_one(doc)
    return {'message': 'Module créé', 'id': module_id, 'module': {k: v for k, v in doc.items() if k != '_id'}}

@api_router.put("/admin/modules/{module_id}")
async def admin_update_module(module_id: str, body: ModuleUpdate, request: Request):
    await require_admin(request)
    update = {k: v for k, v in body.model_dump().items() if v is not None}
    if update:
        update['updated_at'] = datetime.now(timezone.utc).isoformat()
        result = await db.modules.update_one({'id': module_id}, {'$set': update})
        if result.matched_count == 0:
            raise HTTPException(404, "Module non trouvé")
    return {'message': 'Module mis à jour', 'id': module_id}

@api_router.delete("/admin/modules/{module_id}")
async def admin_delete_module(module_id: str, request: Request):
    await require_admin(request)
    # Check for linked audios
    audio_count = await db.audios.count_documents({'module_id': module_id})
    if audio_count > 0:
        raise HTTPException(400, f"Impossible de supprimer: {audio_count} épisodes liés")
    result = await db.modules.delete_one({'id': module_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Module non trouvé")
    return {'message': 'Module supprimé'}

@api_router.patch("/admin/modules/{module_id}/toggle")
async def admin_toggle_module(module_id: str, request: Request):
    await require_admin(request)
    doc = await db.modules.find_one({'id': module_id})
    if not doc:
        raise HTTPException(404, "Module non trouvé")
    new_status = not doc.get('is_active', False)
    await db.modules.update_one({'id': module_id}, {'$set': {'is_active': new_status}})
    return {'id': module_id, 'is_active': new_status}

@api_router.post("/admin/modules/bulk-toggle")
async def admin_bulk_toggle_modules(body: BulkToggleRequest, request: Request):
    """Toggle multiple modules at once."""
    await require_admin(request)
    result = await db.modules.update_many(
        {'id': {'$in': body.ids}},
        {'$set': {'is_active': body.is_active, 'updated_at': datetime.now(timezone.utc).isoformat()}}
    )
    return {'message': f'{result.modified_count} modules mis à jour', 'is_active': body.is_active}

# ─── Admin: Courses Bulk Actions ───────────────────────────────────────────────

@api_router.post("/admin/courses/bulk-toggle")
async def admin_bulk_toggle_courses(body: BulkToggleRequest, request: Request):
    """Toggle multiple courses at once."""
    await require_admin(request)
    result = await db.courses.update_many(
        {'id': {'$in': body.ids}},
        {'$set': {'is_active': body.is_active, 'updated_at': datetime.now(timezone.utc).isoformat()}}
    )
    return {'message': f'{result.modified_count} cours mis à jour', 'is_active': body.is_active}

# ─── Admin: Bibliographies CRUD ────────────────────────────────────────────────

@api_router.get("/admin/bibliographies")
async def admin_list_bibliographies(request: Request):
    await require_admin(request)
    biblio = await db.bibliographies.find({}, {'_id': 0}).to_list(100)
    return biblio

@api_router.post("/admin/bibliographies")
async def admin_create_bibliography(request: Request):
    await require_admin(request)
    body = await request.json()
    biblio_id = f"biblio_{uuid.uuid4().hex[:8]}"
    doc = {
        'id': biblio_id,
        'title': body['title'],
        'thematique_id': body.get('thematique_id', ''),
        'content_fr': body.get('content_fr', ''),
        'content_en': body.get('content_en', ''),
        'is_active': True,
        'created_at': datetime.now(timezone.utc)
    }
    await db.bibliographies.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put("/admin/bibliographies/{biblio_id}")
async def admin_update_bibliography(biblio_id: str, request: Request):
    await require_admin(request)
    body = await request.json()
    update = {k: v for k, v in body.items() if k in ['title', 'thematique_id', 'content_fr', 'content_en', 'is_active']}
    update['updated_at'] = datetime.now(timezone.utc)
    result = await db.bibliographies.update_one({'id': biblio_id}, {'$set': update})
    if result.matched_count == 0:
        raise HTTPException(404, "Bibliographie non trouvée")
    return {'message': 'Bibliographie mise à jour', 'id': biblio_id}

@api_router.delete("/admin/bibliographies/{biblio_id}")
async def admin_delete_bibliography(biblio_id: str, request: Request):
    await require_admin(request)
    result = await db.bibliographies.delete_one({'id': biblio_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Bibliographie non trouvée")
    return {'message': 'Bibliographie supprimée'}

@api_router.post("/admin/bibliographies/standardize-titles")
async def admin_standardize_bibliography_titles(request: Request):
    """
    Uniformiser tous les titres des bibliographies au format :
    "Bibliographie — [Titre du Cours]"
    Uses local module_number (1-based per cursus) since sync is now positional.
    """
    await require_admin(request)
    import re
    
    letter_to_cursus = {
        'A': 'cursus-histoire',
        'B': 'cursus-theologie',
        'C': 'cursus-sciences-islamiques',
        'D': 'cursus-arts',
        'E': 'cursus-falsafa',
        'F': 'cursus-spiritualites',
        'G': 'cursus-pensees-non-islamiques',
    }
    
    # Pre-load courses per cursus
    cursus_courses_map = {}
    for letter, cursus_id in letter_to_cursus.items():
        courses = await db.courses.find({'cursus_id': cursus_id}, {'_id': 0}).to_list(50)
        cursus_courses_map[letter] = courses
    
    biblios = await db.bibliographies.find({'module_number': {'$exists': True}}, {'_id': 0}).to_list(100)
    
    updated_count = 0
    
    for biblio in biblios:
        module_num = biblio.get('module_number')
        cursus_letter = biblio.get('cursus_letter', '')
        cursus_id = biblio.get('cursus_id')
        
        if not module_num or not cursus_letter:
            continue
        
        updates = {}
        
        # Fix cursus_id if needed
        correct_cursus = letter_to_cursus.get(cursus_letter)
        if correct_cursus and cursus_id != correct_cursus:
            cursus_id = correct_cursus
            updates['cursus_id'] = cursus_id
        
        # module_number is now LOCAL (1-based per cursus)
        # So position = module_number directly
        courses = cursus_courses_map.get(cursus_letter, [])
        position = module_num  # Already 1-based local
        
        course_title = None
        if courses and 1 <= position <= len(courses):
            course = courses[position - 1]
            course_title = course.get('title', '')
            if biblio.get('course_id') != course['id']:
                updates['course_id'] = course['id']
        
        # Build title
        if course_title:
            clean_title = re.sub(r'^Cours\s*\d+\s*[:\-—]\s*', '', course_title)
            new_title = f"Bibliographie — {clean_title}"
        else:
            new_title = f"Bibliographie — Module {module_num}"
        
        if biblio.get('title') != new_title:
            updates['title'] = new_title
        
        if updates:
            await db.bibliographies.update_one(
                {'id': biblio['id']},
                {'$set': updates}
            )
            updated_count += 1
    
    return {'message': f'{updated_count} titre(s) mis à jour', 'updated': updated_count}

# ─── Admin: Masterclasses CRUD ─────────────────────────────────────────────────

@api_router.get("/admin/masterclasses")
async def admin_list_masterclasses(request: Request):
    await require_admin(request)
    mcs = await db.masterclasses.find({}, {'_id': 0}).to_list(100)
    return mcs

@api_router.post("/admin/masterclasses")
async def admin_create_masterclass(request: Request):
    await require_admin(request)
    body = await request.json()
    mc_id = f"mc_{uuid.uuid4().hex[:8]}"
    doc = {
        'id': mc_id,
        'title': body['title'],
        'description': body.get('description', ''),
        'thematique_id': body.get('thematique_id', ''),
        'scholar_id': body.get('scholar_id', ''),
        'scholar_name': body.get('scholar_name', ''),
        'date': body.get('date'),
        'duration': body.get('duration', 60),
        'price': body.get('price', 0),
        'price_type': 'free' if body.get('price', 0) == 0 else 'paid',
        'max_participants': body.get('max_participants', 100),
        'current_participants': 0,
        'thumbnail': body.get('thumbnail', ''),
        'video_url': body.get('video_url', ''),
        'registered_users': [],
        'is_active': True,
        'created_at': datetime.now(timezone.utc)
    }
    await db.masterclasses.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put("/admin/masterclasses/{mc_id}")
async def admin_update_masterclass(mc_id: str, request: Request):
    await require_admin(request)
    body = await request.json()
    allowed = ['title', 'description', 'thematique_id', 'scholar_id', 'scholar_name', 'date', 
               'duration', 'price', 'max_participants', 'thumbnail', 'video_url', 'is_active']
    update = {k: v for k, v in body.items() if k in allowed}
    if 'price' in update:
        update['price_type'] = 'free' if update['price'] == 0 else 'paid'
    update['updated_at'] = datetime.now(timezone.utc)
    result = await db.masterclasses.update_one({'id': mc_id}, {'$set': update})
    if result.matched_count == 0:
        raise HTTPException(404, "Masterclass non trouvée")
    return {'message': 'Masterclass mise à jour', 'id': mc_id}

@api_router.delete("/admin/masterclasses/{mc_id}")
async def admin_delete_masterclass(mc_id: str, request: Request):
    await require_admin(request)
    result = await db.masterclasses.delete_one({'id': mc_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Masterclass non trouvée")
    return {'message': 'Masterclass supprimée'}

@api_router.patch("/admin/masterclasses/{mc_id}/toggle")
async def admin_toggle_masterclass(mc_id: str, request: Request):
    await require_admin(request)
    doc = await db.masterclasses.find_one({'id': mc_id})
    if not doc:
        raise HTTPException(404, "Masterclass non trouvée")
    new_status = not doc.get('is_active', True)
    await db.masterclasses.update_one({'id': mc_id}, {'$set': {'is_active': new_status}})
    return {'id': mc_id, 'is_active': new_status}

# ─── Admin: Audio Categories CRUD ─────────────────────────────────────────────

@api_router.get("/admin/audio-categories")
async def admin_list_audio_categories(request: Request):
    """List all audio categories."""
    await require_admin(request)
    categories = await db.audio_categories.find({}, {'_id': 0}).to_list(100)
    return categories

@api_router.get("/audio-categories")
async def public_list_audio_categories():
    """Public endpoint to list active audio categories."""
    categories = await db.audio_categories.find({'is_active': True}, {'_id': 0}).to_list(100)
    return categories

@api_router.post("/admin/audio-categories")
async def admin_create_audio_category(body: AudioCategoryCreate, request: Request):
    """Create a new audio category."""
    await require_admin(request)
    cat_id = f"cat_{uuid.uuid4().hex[:8]}"
    doc = {
        'id': cat_id,
        **body.model_dump(),
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    await db.audio_categories.insert_one(doc)
    logger.info(f"Audio category '{body.name}' created with id {cat_id}")
    return {'message': 'Catégorie créée', 'id': cat_id, 'category': {**doc, '_id': None}}

@api_router.put("/admin/audio-categories/{cat_id}")
async def admin_update_audio_category(cat_id: str, body: AudioCategoryUpdate, request: Request):
    """Update an audio category."""
    await require_admin(request)
    update_data = {k: v for k, v in body.model_dump().items() if v is not None}
    if update_data:
        update_data['updated_at'] = datetime.now(timezone.utc).isoformat()
        result = await db.audio_categories.update_one({'id': cat_id}, {'$set': update_data})
        if result.matched_count == 0:
            raise HTTPException(404, "Catégorie non trouvée")
    doc = await db.audio_categories.find_one({'id': cat_id}, {'_id': 0})
    return {'message': 'Catégorie mise à jour', 'id': cat_id, 'category': doc}

@api_router.delete("/admin/audio-categories/{cat_id}")
async def admin_delete_audio_category(cat_id: str, request: Request):
    """Delete an audio category."""
    await require_admin(request)
    # Check if any audios use this category
    audios_count = await db.audios.count_documents({'category_id': cat_id})
    if audios_count > 0:
        raise HTTPException(400, f"Impossible de supprimer: {audios_count} audio(s) utilisent cette catégorie")
    result = await db.audio_categories.delete_one({'id': cat_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Catégorie non trouvée")
    return {'message': 'Catégorie supprimée'}

@api_router.patch("/admin/audio-categories/{cat_id}/toggle")
async def admin_toggle_audio_category(cat_id: str, request: Request):
    """Toggle audio category active status."""
    await require_admin(request)
    doc = await db.audio_categories.find_one({'id': cat_id})
    if not doc:
        raise HTTPException(404, "Catégorie non trouvée")
    new_status = not doc.get('is_active', True)
    await db.audio_categories.update_one({'id': cat_id}, {'$set': {'is_active': new_status}})
    return {'id': cat_id, 'is_active': new_status}

@api_router.get("/audios/by-category/{cat_id}")
async def get_audios_by_category(cat_id: str):
    """Get all audios in a specific category."""
    audios = await db.audios.find({'category_id': cat_id, 'is_active': True}, {'_id': 0}).to_list(100)
    for a in audios:
        a['stream_url'] = resolve_audio_url(a)
    return audios

# ─── Admin: Set Featured Course ─────────────────────────────────────────────────

@api_router.patch("/admin/courses/{course_id}/set-featured")
async def admin_set_featured_course(course_id: str, request: Request):
    """Set a course as the featured hero (unfeatures all cursus and other courses)."""
    await require_admin(request)
    await db.cursus.update_many({}, {'$set': {'is_featured': False}})
    await db.courses.update_many({}, {'$set': {'is_featured': False}})
    result = await db.courses.update_one(
        {'id': course_id},
        {'$set': {'is_featured': True}}
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Cours non trouvé")
    logger.info(f"Course {course_id} set as featured")
    return {'message': 'Cours mis en avant', 'id': course_id}

@api_router.delete("/admin/courses/featured")
async def admin_remove_featured_course(request: Request):
    """Remove featured status from all courses."""
    await require_admin(request)
    await db.courses.update_many({}, {'$set': {'is_featured': False}})
    return {'message': 'Aucun cours mis en avant'}

@api_router.patch("/admin/cursus/{cursus_id}/set-featured")
async def admin_set_featured_cursus(cursus_id: str, request: Request):
    """Set a cursus as the featured hero (unfeatures all courses and other cursus)."""
    await require_admin(request)
    await db.courses.update_many({}, {'$set': {'is_featured': False}})
    await db.cursus.update_many({}, {'$set': {'is_featured': False}})
    result = await db.cursus.update_one(
        {'id': cursus_id},
        {'$set': {'is_featured': True}}
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Cursus non trouvé")
    logger.info(f"Cursus {cursus_id} set as featured")
    return {'message': 'Cursus mis en avant', 'id': cursus_id}

@api_router.delete("/admin/cursus/featured")
async def admin_remove_featured_cursus(request: Request):
    """Remove featured status from all cursus."""
    await require_admin(request)
    await db.cursus.update_many({}, {'$set': {'is_featured': False}})
    return {'message': 'Aucun cursus mis en avant'}

# ─── Admin Panel Web Routes ────────────────────────────────────────────────────

ADMIN_TEMPLATES_DIR = ROOT_DIR / 'admin_templates'

@api_router.get("/admin-panel/login", response_class=HTMLResponse)
async def admin_panel_login():
    """Admin panel login page."""
    template_path = ADMIN_TEMPLATES_DIR / 'login.html'
    if not template_path.exists():
        raise HTTPException(404, "Template login non trouvé")
    return HTMLResponse(content=template_path.read_text(encoding='utf-8'))

@api_router.get("/admin-panel/", response_class=HTMLResponse)
async def admin_panel_dashboard(request: Request):
    """Admin panel dashboard page."""
    return templates.TemplateResponse(
        "dashboard_new.html",
        {"request": request, "active_page": "dashboard"}
    )

@api_router.get("/admin-panel/scholars", response_class=HTMLResponse)
async def admin_panel_scholars(request: Request):
    """Admin panel scholars page - redirect to professors."""
    return templates.TemplateResponse(
        "professors_new.html",
        {"request": request, "active_page": "professors"}
    )

@api_router.get("/admin-panel/professors", response_class=HTMLResponse)
async def admin_panel_professors(request: Request):
    """Admin panel professors page."""
    return templates.TemplateResponse(
        "professors_new.html",
        {"request": request, "active_page": "professors"}
    )

@api_router.get("/admin-panel/courses", response_class=HTMLResponse)
async def admin_panel_courses(request: Request):
    """Admin panel courses page."""
    return templates.TemplateResponse(
        "courses_new.html",
        {"request": request, "active_page": "courses"}
    )

@api_router.get("/admin-panel/tree", response_class=HTMLResponse)
async def admin_panel_tree(request: Request):
    """Hierarchical view: cursus → courses → episodes with inline editing."""
    return templates.TemplateResponse(
        "tree.html",
        {"request": request, "active_page": "tree"}
    )


@api_router.get("/admin-panel/episodes", response_class=HTMLResponse)
async def admin_panel_episodes(request: Request):
    """Admin panel flat episodes view — fast per-episode editing."""
    return templates.TemplateResponse(
        "episodes.html",
        {"request": request, "active_page": "episodes"}
    )

@api_router.get("/admin-panel/users", response_class=HTMLResponse)
async def admin_panel_users(request: Request):
    """Admin panel users page."""
    return templates.TemplateResponse(
        "users_new.html",
        {"request": request, "active_page": "users"}
    )

@api_router.get("/admin-panel/listening-stats", response_class=HTMLResponse)
async def admin_panel_listening_stats(request: Request):
    """Admin panel listening statistics page."""
    return templates.TemplateResponse(
        "listening-stats.html",
        {"request": request, "active_page": "listening-stats"}
    )

@api_router.get("/admin-panel/highlight", response_class=HTMLResponse)
async def admin_panel_highlight(request: Request):
    """Admin panel highlight configuration page."""
    return templates.TemplateResponse(
        "highlight.html",
        {"request": request, "active_page": "highlight"}
    )

@api_router.get("/admin-panel/audios", response_class=HTMLResponse)
async def admin_panel_audios(request: Request):
    """Admin panel audios page."""
    response = templates.TemplateResponse(
        "audios_new.html",
        {"request": request, "active_page": "audios"}
    )
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

@api_router.get("/admin-panel/articles", response_class=HTMLResponse)
async def admin_panel_articles():
    """Admin panel articles page."""
    return HTMLResponse(content=(ADMIN_TEMPLATES_DIR / 'dashboard.html').read_text(encoding='utf-8'))

@api_router.get("/admin-panel/preregistrations", response_class=HTMLResponse)
async def admin_panel_preregistrations(request: Request):
    """Admin panel pre-registrations page."""
    return templates.TemplateResponse(
        "preregistrations.html",
        {"request": request, "active_page": "preregistrations"}
    )

@api_router.get("/admin-panel/thematiques", response_class=HTMLResponse)
async def admin_panel_thematiques(request: Request):
    """Admin panel thematiques page - redirects to cursus."""
    return templates.TemplateResponse(
        "cursus_new.html",
        {"request": request, "active_page": "cursus"}
    )

@api_router.get("/admin-panel/cursus", response_class=HTMLResponse)
async def admin_panel_cursus(request: Request):
    """Admin panel cursus page."""
    return templates.TemplateResponse(
        "cursus_new.html",
        {"request": request, "active_page": "cursus"}
    )

@api_router.get("/admin-panel/modules", response_class=HTMLResponse)
async def admin_panel_modules(request: Request):
    """Admin panel modules page."""
    return templates.TemplateResponse(
        "modules_new.html",
        {"request": request, "active_page": "modules"}
    )

@api_router.get("/admin-panel/blog", response_class=HTMLResponse)
async def admin_panel_blog(request: Request):
    """Admin panel blog page."""
    response = templates.TemplateResponse(
        "blog.html",
        {"request": request, "active_page": "blog"}
    )
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    return response

@api_router.get("/admin-panel/audio-categories", response_class=HTMLResponse)
async def admin_panel_audio_categories():
    """Admin panel audio categories page."""
    template_path = ADMIN_TEMPLATES_DIR / 'audio-categories.html'
    if template_path.exists():
        return HTMLResponse(content=template_path.read_text(encoding='utf-8'))
    return HTMLResponse(content=(ADMIN_TEMPLATES_DIR / 'dashboard.html').read_text(encoding='utf-8'))

@api_router.get("/admin-panel/r2", response_class=HTMLResponse)
async def admin_panel_r2():
    """Admin panel R2 storage page."""
    template_path = ADMIN_TEMPLATES_DIR / 'r2.html'
    if template_path.exists():
        return HTMLResponse(content=template_path.read_text(encoding='utf-8'))
    return HTMLResponse(content=(ADMIN_TEMPLATES_DIR / 'dashboard.html').read_text(encoding='utf-8'))

@api_router.get("/admin-panel/r2-medias", response_class=HTMLResponse)
async def admin_panel_r2_medias(request: Request):
    """Admin panel R2 media auto-detection page (Phase 3)."""
    return templates.TemplateResponse(
        "r2-medias.html",
        {"request": request, "active_page": "r2-medias"}
    )


@api_router.get("/admin-panel/legal", response_class=HTMLResponse)
async def admin_panel_legal():
    """Admin panel legal pages editor."""
    template_path = ADMIN_TEMPLATES_DIR / 'legal.html'
    if template_path.exists():
        return HTMLResponse(content=template_path.read_text(encoding='utf-8'))
    return HTMLResponse(content=(ADMIN_TEMPLATES_DIR / 'dashboard.html').read_text(encoding='utf-8'))

# ─── Payment & Subscription Routes ─────────────────────────────────────────────

async def check_user_access(user_id: str, content_type: str = None, content_id: str = None) -> dict:
    """Check if user has access to content (subscription, purchase, or trial)."""
    user = await db.users.find_one({'user_id': user_id}, {'_id': 0})
    if not user:
        return {'has_access': False, 'reason': 'user_not_found'}
    
    # Admin or free_access users have full access
    if user.get('role') == 'admin' or user.get('free_access'):
        return {'has_access': True, 'reason': 'admin_or_free'}
    
    now = datetime.now(timezone.utc)
    
    # Check active free trial
    trial = user.get('trial')
    if trial and trial.get('expires_at'):
        expires_at = trial['expires_at']
        if isinstance(expires_at, str):
            expires_at = datetime.fromisoformat(expires_at.replace('Z', '+00:00'))
        if expires_at > now:
            return {'has_access': True, 'reason': 'free_trial', 'expires_at': expires_at.isoformat()}
    
    # Check active subscription
    subscription = user.get('subscription')
    if subscription and subscription.get('expires_at'):
        expires_at = subscription['expires_at']
        if isinstance(expires_at, str):
            expires_at = datetime.fromisoformat(expires_at.replace('Z', '+00:00'))
        if expires_at > now:
            return {'has_access': True, 'reason': 'active_subscription', 'expires_at': expires_at.isoformat()}
    
    # Check specific content purchases
    if content_type and content_id:
        purchases = user.get('purchases', [])
        for purchase in purchases:
            if purchase.get('content_type') == content_type and purchase.get('content_id') == content_id:
                expires_at = purchase.get('expires_at')
                if isinstance(expires_at, str):
                    expires_at = datetime.fromisoformat(expires_at.replace('Z', '+00:00'))
                if expires_at and expires_at > now:
                    return {'has_access': True, 'reason': 'content_purchase', 'expires_at': expires_at.isoformat()}
        
        # For courses, also check if parent cursus is purchased
        if content_type == 'course':
            course = await db.courses.find_one({'id': content_id}, {'thematique_id': 1})
            if course and course.get('thematique_id'):
                for purchase in purchases:
                    if purchase.get('content_type') == 'cursus' and purchase.get('content_id') == course['thematique_id']:
                        expires_at = purchase.get('expires_at')
                        if isinstance(expires_at, str):
                            expires_at = datetime.fromisoformat(expires_at.replace('Z', '+00:00'))
                        if expires_at and expires_at > now:
                            return {'has_access': True, 'reason': 'cursus_purchase', 'expires_at': expires_at.isoformat()}
    
    return {'has_access': False, 'reason': 'no_access'}

@api_router.get("/user/access")
async def get_user_access(request: Request, content_type: Optional[str] = None, content_id: Optional[str] = None):
    """Check user's access status."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    return await check_user_access(user['user_id'], content_type, content_id)

@api_router.delete("/user/delete-account")
async def delete_user_account(request: Request):
    """Delete the current user's account and all associated data."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    user_id = user['user_id']
    
    # Delete user data from various collections
    await db.user_favorites.delete_many({'user_id': user_id})
    await db.payment_transactions.delete_many({'user_id': user_id})
    
    # Delete the user account
    result = await db.users.delete_one({'user_id': user_id})
    
    if result.deleted_count == 0:
        raise HTTPException(404, "Utilisateur non trouvé")
    
    logger.info(f"User account deleted: {user_id}")
    return {'message': 'Compte supprimé avec succès'}

@api_router.get("/plans")
async def get_plans():
    """Get all active subscription and purchase plans."""
    plans = await db.plans.find({'is_active': True}, {'_id': 0}).to_list(100)
    if not plans:
        # Return fondateur plans (active now) + standard plans (inactive, for later)
        return [
            {'plan_id': 'fondateur_mensuel', 'name': 'Fondateur Mensuel', 'price': 7.00, 'duration_days': 30, 'type': 'subscription', 'description': '7 \u20ac/mois \u00b7 engagement 12 mois', 'is_fondateur': True, 'is_active': True, 'max_places': 200},
            {'plan_id': 'fondateur_annuel', 'name': 'Fondateur Annuel', 'price': 84.00, 'duration_days': 365, 'type': 'subscription', 'description': '84 \u20ac/an (soit 7 \u20ac/mois) \u00b7 paiement unique', 'is_fondateur': True, 'is_active': True, 'max_places': 200},
            {'plan_id': 'standard_mensuel', 'name': 'Standard Mensuel', 'price': 12.00, 'duration_days': 30, 'type': 'subscription', 'description': '12 \u20ac/mois', 'is_fondateur': False, 'is_active': False},
            {'plan_id': 'standard_annuel', 'name': 'Standard Annuel', 'price': 120.00, 'duration_days': 365, 'type': 'subscription', 'description': '120 \u20ac/an (soit 10 \u20ac/mois)', 'is_fondateur': False, 'is_active': False},
        ]
    return plans

@api_router.post("/checkout/create")
async def create_checkout_session(body: CheckoutRequest, request: Request):
    """Create a Stripe checkout session for subscription or one-time purchase."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    if not STRIPE_API_KEY:
        raise HTTPException(500, "Stripe non configure")
    
    # Determine what we're purchasing
    amount = 0.0
    product_name = ""
    metadata = {'user_id': user['user_id'], 'user_email': user.get('email', '')}
    duration_days = 0
    
    if body.plan_id:
        # Subscription plan
        plan = await db.plans.find_one({'plan_id': body.plan_id, 'is_active': True}, {'_id': 0})
        if not plan:
            # Check default plans
            default = DEFAULT_PLANS.get(body.plan_id)
            if default:
                plan = {'plan_id': body.plan_id, **default}
            else:
                raise HTTPException(404, "Plan non trouve")
        amount = float(plan['price'])
        product_name = plan['name']
        duration_days = plan.get('duration_days', 30)
        metadata['plan_id'] = body.plan_id
        metadata['purchase_type'] = 'subscription'
        metadata['duration_days'] = str(duration_days)
    
    elif body.course_id:
        # Course purchase (6 months access)
        course = await db.courses.find_one({'id': body.course_id}, {'_id': 0})
        if not course:
            raise HTTPException(404, "Cours non trouve")
        # Get course price from plans or use default
        course_plan = await db.plans.find_one({'plan_id': f'course_{body.course_id}'}, {'_id': 0})
        if course_plan:
            amount = float(course_plan['price'])
        else:
            amount = 19.99  # Default course price
        product_name = f"Cours: {course['title']}"
        duration_days = 180  # 6 months
        metadata['course_id'] = body.course_id
        metadata['purchase_type'] = 'course'
        metadata['duration_days'] = str(duration_days)
    
    elif body.cursus_id:
        # Cursus purchase (6 months access)
        cursus = await db.thematiques.find_one({'id': body.cursus_id}, {'_id': 0})
        if not cursus:
            raise HTTPException(404, "Cursus non trouve")
        # Get cursus price from plans or use default
        cursus_plan = await db.plans.find_one({'plan_id': f'cursus_{body.cursus_id}'}, {'_id': 0})
        if cursus_plan:
            amount = float(cursus_plan['price'])
        else:
            amount = 49.99  # Default cursus price
        product_name = f"Cursus: {cursus['name']}"
        duration_days = 180  # 6 months
        metadata['cursus_id'] = body.cursus_id
        metadata['purchase_type'] = 'cursus'
        metadata['duration_days'] = str(duration_days)
    
    else:
        raise HTTPException(400, "Veuillez specifier un plan, cours ou cursus")
    
    # Apply promo code if provided
    original_amount = amount
    promo_applied = None
    if body.promo_code:
        promo = await db.promo_codes.find_one({'code': body.promo_code.upper(), 'is_active': True}, {'_id': 0})
        if promo:
            # Validate promo code
            valid = True
            if promo.get('expires_at'):
                expires = promo['expires_at']
                if isinstance(expires, str):
                    expires = datetime.fromisoformat(expires.replace('Z', '+00:00'))
                if expires < datetime.now(timezone.utc):
                    valid = False
            if promo.get('max_uses') and promo.get('uses_count', 0) >= promo['max_uses']:
                valid = False
            applicable = promo.get('applicable_plans', [])
            if applicable and body.plan_id and body.plan_id not in applicable:
                valid = False
            
            if valid:
                if promo.get('discount_percent'):
                    discount = amount * (promo['discount_percent'] / 100)
                    amount = max(0, amount - discount)
                elif promo.get('discount_amount'):
                    amount = max(0, amount - promo['discount_amount'])
                
                promo_applied = promo['code']
                metadata['promo_code'] = promo['code']
                metadata['original_amount'] = str(original_amount)
                
                # Increment uses count
                await db.promo_codes.update_one({'code': promo['code']}, {'$inc': {'uses_count': 1}})
    
    # Build URLs
    success_url = f"{body.origin_url}/payment/success?session_id={{CHECKOUT_SESSION_ID}}"
    cancel_url = f"{body.origin_url}/payment/cancel"
    
    # Create Stripe checkout session
    host_url = str(request.base_url).rstrip('/')
    webhook_url = f"{host_url}/api/webhook/stripe"
    stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
    
    checkout_request = CheckoutSessionRequest(
        amount=amount,
        currency="eur",
        success_url=success_url,
        cancel_url=cancel_url,
        metadata=metadata
    )
    
    try:
        session = await stripe_checkout.create_checkout_session(checkout_request)
        
        # Store transaction
        transaction_id = f"txn_{uuid.uuid4().hex[:12]}"
        await db.payment_transactions.insert_one({
            'transaction_id': transaction_id,
            'session_id': session.session_id,
            'user_id': user['user_id'],
            'user_email': user.get('email', ''),
            'amount': amount,
            'currency': 'eur',
            'product_name': product_name,
            'metadata': metadata,
            'status': 'pending',
            'payment_status': 'initiated',
            'created_at': datetime.now(timezone.utc)
        })
        
        return {'url': session.url, 'session_id': session.session_id}
    
    except Exception as e:
        logger.error(f"Stripe checkout error: {e}")
        raise HTTPException(500, f"Erreur de paiement: {str(e)}")


# ════════════════════════════════════════════════════════════════════════════
# GIFT CARDS — purchase, redeem, scheduled delivery
# ════════════════════════════════════════════════════════════════════════════
from utils.gift_cards import (
    generate_code as _gift_generate_code,
    GiftCardPurchaseRequest, GiftCardRedeemRequest,
    GIFT_PLAN_PRICES, gift_email_html, purchaser_confirmation_html,
)


@api_router.post("/gift-cards/purchase")
async def gift_card_purchase(body: GiftCardPurchaseRequest, request: Request):
    """Create a pending gift card + Stripe Checkout session.

    Generic plan resolution: tries the hardcoded `GIFT_PLAN_PRICES` dict first
    (back-compat for `founder_monthly` / `founder_yearly`), then falls back to
    `db.plans` so any plan added later (admin or migration) becomes giftable
    automatically without code change.
    """
    if not STRIPE_API_KEY:
        raise HTTPException(500, "Stripe non configuré")

    plan = GIFT_PLAN_PRICES.get(body.plan_id)
    if not plan:
        # Generic fallback: look up the plan in db.plans
        db_plan = await db.plans.find_one({'plan_id': body.plan_id, 'is_active': True}, {'_id': 0})
        if not db_plan:
            raise HTTPException(400, f"Plan inconnu: {body.plan_id}")
        plan = {
            'amount': float(db_plan.get('price', 0)),
            'duration_days': int(db_plan.get('duration_days', 365)),
            'label': db_plan.get('name', body.plan_id),
        }
        if plan['amount'] <= 0:
            raise HTTPException(400, f"Plan {body.plan_id} sans tarif valide")
    # Validate deliver_at
    deliver_at = None
    if body.deliver_at:
        try:
            d = datetime.fromisoformat(body.deliver_at)
            if d.date() < datetime.now(timezone.utc).date():
                raise ValueError("date passée")
            deliver_at = d.date().isoformat()
        except Exception:
            raise HTTPException(400, "Date de livraison invalide (YYYY-MM-DD)")

    gift_id = f"gift_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)
    gift_doc = {
        'gift_id': gift_id,
        'plan_id': body.plan_id,
        'plan_label': plan['label'],
        'amount_paid': plan['amount'],
        'duration_days': plan['duration_days'],
        'currency': 'eur',
        'purchaser_name': body.purchaser_name,
        'purchaser_email': body.purchaser_email,
        'recipient_name': body.recipient_name,
        'recipient_email': body.recipient_email,
        'personal_message': body.personal_message or '',
        'status': 'pending',
        'code': None,
        'redeemed_by_user_id': None,
        'redeemed_at': None,
        'deliver_at': deliver_at,
        'delivered_at': None,
        'created_at': now,
        'expires_at': now + timedelta(days=365),
        'stripe_session_id': None,
    }
    await db.gift_cards.insert_one(gift_doc)

    # Stripe Checkout (one-shot payment)
    success_url = f"{body.origin_url}/cadeau/confirmation?session_id={{CHECKOUT_SESSION_ID}}"
    cancel_url = f"{body.origin_url}/cadeau"
    host_url = str(request.base_url).rstrip('/')
    webhook_url = f"{host_url}/api/webhook/stripe"
    stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
    checkout_request = CheckoutSessionRequest(
        amount=plan['amount'],
        currency='eur',
        success_url=success_url,
        cancel_url=cancel_url,
        metadata={
            'type': 'gift_card',
            'gift_id': gift_id,
            'plan_id': body.plan_id,
            'purchaser_email': body.purchaser_email,
            'recipient_email': body.recipient_email,
        },
    )
    try:
        session = await stripe_checkout.create_checkout_session(checkout_request)
        await db.gift_cards.update_one(
            {'gift_id': gift_id},
            {'$set': {'stripe_session_id': session.session_id}},
        )
        return {'gift_id': gift_id, 'url': session.url, 'session_id': session.session_id}
    except Exception as e:
        logger.error(f"Gift-card Stripe error: {e}")
        await db.gift_cards.delete_one({'gift_id': gift_id})
        raise HTTPException(500, f"Erreur de paiement: {e}")


async def _gift_card_finalize_after_payment(session_id: str) -> Optional[dict]:
    """Called from Stripe webhook AND from the success page polling.
    Idempotent: only generates the code+emails once.
    """
    gift = await db.gift_cards.find_one({'stripe_session_id': session_id}, {'_id': 0})
    if not gift:
        return None
    if gift.get('status') == 'paid' and gift.get('code'):
        return gift  # already finalized
    # Verify Stripe status
    try:
        host_url = os.environ.get('PUBLIC_BASE_URL', 'http://localhost:8001')
        stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=f"{host_url}/api/webhook/stripe")
        status = await stripe_checkout.get_checkout_status(session_id)
        if status.payment_status != 'paid':
            return None
    except Exception as e:
        logger.error(f"Gift finalize: status check failed: {e}")
        return None
    # Generate unique code (max 5 retries)
    for _ in range(5):
        code = _gift_generate_code()
        if not await db.gift_cards.find_one({'code': code}):
            break
    await db.gift_cards.update_one(
        {'gift_id': gift['gift_id']},
        {'$set': {'status': 'paid', 'code': code, 'paid_at': datetime.now(timezone.utc)}},
    )
    gift['code'] = code; gift['status'] = 'paid'
    # Send emails (or schedule)
    if not gift.get('deliver_at'):
        await _gift_card_send_to_recipient(gift)
    # Always send purchaser confirmation
    if is_email_configured():
        smtp_send_email(
            to_email=gift['purchaser_email'],
            to_name=gift['purchaser_name'],
            subject=f"Votre cadeau Sijill pour {gift['recipient_name']} est confirmé",
            html_content=purchaser_confirmation_html(
                purchaser_name=gift['purchaser_name'],
                recipient_name=gift['recipient_name'],
                plan_label=gift['plan_label'],
                deliver_at=gift.get('deliver_at'),
                code=code,
            ),
        )
    return gift


async def _gift_card_send_to_recipient(gift: dict):
    """Send the gift email to the recipient and mark delivered."""
    if not is_email_configured():
        logger.warning("Gift card recipient email skipped (SMTP not configured)")
        return
    site_url = os.environ.get('PUBLIC_SITE_URL', 'https://sijillproject.com')
    redeem_url = f"{site_url}/cadeau/recu?code={gift['code']}"
    html = gift_email_html(
        purchaser_name=gift['purchaser_name'],
        recipient_name=gift['recipient_name'],
        plan_label=gift['plan_label'],
        code=gift['code'],
        redeem_url=redeem_url,
        personal_message=gift.get('personal_message', ''),
    )
    smtp_send_email(
        to_email=gift['recipient_email'],
        to_name=gift['recipient_name'],
        subject=f"{gift['purchaser_name']} vous offre un abonnement Sijill",
        html_content=html,
    )
    await db.gift_cards.update_one(
        {'gift_id': gift['gift_id']},
        {'$set': {'delivered_at': datetime.now(timezone.utc)}},
    )


@api_router.get("/gift-cards/lookup/{code}")
async def gift_card_lookup(code: str):
    """Public preview of a gift card before redemption."""
    g = await db.gift_cards.find_one(
        {'code': code.upper()},
        {'_id': 0, 'plan_label': 1, 'purchaser_name': 1, 'recipient_name': 1,
         'personal_message': 1, 'status': 1, 'expires_at': 1, 'duration_days': 1},
    )
    if not g:
        raise HTTPException(404, "Code introuvable")
    if g['status'] == 'redeemed':
        raise HTTPException(410, "Ce code a déjà été utilisé")
    if g['status'] != 'paid':
        raise HTTPException(409, "Ce code n'est pas encore actif")
    exp = g.get('expires_at')
    if exp and exp < datetime.now(timezone.utc):
        raise HTTPException(410, "Ce code a expiré")
    return {
        'plan_label': g['plan_label'],
        'purchaser_name': g['purchaser_name'],
        'recipient_name': g['recipient_name'],
        'personal_message': g.get('personal_message', ''),
        'duration_days': g['duration_days'],
    }


@api_router.post("/gift-cards/redeem")
async def gift_card_redeem(body: GiftCardRedeemRequest, request: Request):
    """Apply a paid gift card to the authenticated user's subscription."""
    user = await get_current_user(request)
    g = await db.gift_cards.find_one({'code': body.code.strip().upper()})
    if not g:
        raise HTTPException(404, "Code introuvable")
    if g['status'] == 'redeemed':
        raise HTTPException(410, "Ce code a déjà été utilisé")
    if g['status'] != 'paid':
        raise HTTPException(409, "Ce code n'est pas encore activé (paiement non confirmé)")
    exp = g.get('expires_at')
    if exp and exp < datetime.now(timezone.utc):
        raise HTTPException(410, "Ce code a expiré")
    # Extend subscription
    user_doc = await db.users.find_one({'user_id': user['user_id']})
    now = datetime.now(timezone.utc)
    current_end = user_doc.get('subscription_end_date')
    if isinstance(current_end, str):
        try: current_end = datetime.fromisoformat(current_end.replace('Z', '+00:00'))
        except: current_end = None
    base = max(now, current_end) if current_end else now
    new_end = base + timedelta(days=g['duration_days'])
    await db.users.update_one(
        {'user_id': user['user_id']},
        {'$set': {
            'subscription_status': 'active',
            'subscription_plan': g['plan_id'],
            'subscription_end_date': new_end,
            'subscription_source': 'gift_card',
        }},
    )
    await db.gift_cards.update_one(
        {'gift_id': g['gift_id']},
        {'$set': {
            'status': 'redeemed',
            'redeemed_by_user_id': user['user_id'],
            'redeemed_by_email': user.get('email'),
            'redeemed_at': now,
        }},
    )
    return {
        'success': True,
        'subscription_end_date': new_end.isoformat(),
        'plan_label': g['plan_label'],
        'days_added': g['duration_days'],
    }


@api_router.get("/gift-cards/by-session/{session_id}")
async def gift_card_finalize_endpoint(session_id: str):
    """Used by the success page (post-Stripe redirect) to finalize the gift card.
    Idempotent: safe to call multiple times."""
    gift = await _gift_card_finalize_after_payment(session_id)
    if not gift:
        return {'status': 'pending'}
    return {
        'status': gift.get('status'),
        'plan_label': gift.get('plan_label'),
        'recipient_name': gift.get('recipient_name'),
        'recipient_email': gift.get('recipient_email'),
        'deliver_at': gift.get('deliver_at'),
        'code': gift.get('code'),
    }


@api_router.get("/admin/gift-cards")
async def admin_list_gift_cards(request: Request):
    await require_admin(request)
    docs = await db.gift_cards.find({}, {'_id': 0}).sort('created_at', -1).to_list(500)
    # Convert datetime to ISO for JSON
    for d in docs:
        for k in ('created_at', 'expires_at', 'paid_at', 'delivered_at', 'redeemed_at'):
            v = d.get(k)
            if hasattr(v, 'isoformat'):
                d[k] = v.isoformat()
    return docs


async def _process_scheduled_gift_deliveries():
    """Background task: deliver gift cards whose `deliver_at` is today."""
    try:
        today = datetime.now(timezone.utc).date().isoformat()
        async for gift in db.gift_cards.find({
            'status': 'paid',
            'deliver_at': {'$lte': today},
            'delivered_at': None,
        }, {'_id': 0}):
            await _gift_card_send_to_recipient(gift)
            logger.info(f"Gift card {gift['gift_id']} delivered to {gift['recipient_email']}")
    except Exception as e:
        logger.error(f"Scheduled gift delivery error: {e}")



@api_router.get("/checkout/status/{session_id}")
async def get_checkout_status(session_id: str, request: Request):
    """Get checkout session status and update user access if paid."""
    if not STRIPE_API_KEY:
        raise HTTPException(500, "Stripe non configure")
    
    # Get transaction
    transaction = await db.payment_transactions.find_one({'session_id': session_id}, {'_id': 0})
    if not transaction:
        raise HTTPException(404, "Transaction non trouvee")
    
    # If already processed, return cached status
    if transaction.get('payment_status') == 'paid':
        return {
            'status': 'complete',
            'payment_status': 'paid',
            'amount_total': int(transaction['amount'] * 100),
            'currency': transaction['currency'],
            'metadata': transaction['metadata']
        }
    
    # Check with Stripe
    host_url = str(request.base_url).rstrip('/')
    webhook_url = f"{host_url}/api/webhook/stripe"
    stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
    
    try:
        status = await stripe_checkout.get_checkout_status(session_id)
        
        # Update transaction
        await db.payment_transactions.update_one(
            {'session_id': session_id},
            {'$set': {'status': status.status, 'payment_status': status.payment_status, 'updated_at': datetime.now(timezone.utc)}}
        )
        
        # If paid, grant access
        if status.payment_status == 'paid' and transaction.get('payment_status') != 'paid':
            await grant_user_access(transaction)
        
        return {
            'status': status.status,
            'payment_status': status.payment_status,
            'amount_total': status.amount_total,
            'currency': status.currency,
            'metadata': status.metadata
        }
    
    except Exception as e:
        logger.error(f"Stripe status check error: {e}")
        raise HTTPException(500, f"Erreur: {str(e)}")

async def grant_user_access(transaction: dict):
    """Grant user access based on payment."""
    user_id = transaction['metadata'].get('user_id')
    if not user_id:
        return
    
    now = datetime.now(timezone.utc)
    duration_days = int(transaction['metadata'].get('duration_days', 30))
    expires_at = now + timedelta(days=duration_days)
    
    purchase_type = transaction['metadata'].get('purchase_type')
    
    # Get current user to check for referral
    user = await db.users.find_one({'user_id': user_id})
    
    if purchase_type == 'subscription':
        # Update subscription and subscription_end_date
        await db.users.update_one(
            {'user_id': user_id},
            {'$set': {
                'subscription': {
                    'plan_id': transaction['metadata'].get('plan_id'),
                    'started_at': now,
                    'expires_at': expires_at,
                    'transaction_id': transaction['transaction_id'],
                    'status': 'active'
                },
                'subscription_end_date': expires_at
            }}
        )
        
        # Send subscription confirmation email
        if user and is_email_configured():
            plan_name = transaction['metadata'].get('plan_name', 'Abonnement Sijill')
            amount = float(transaction.get('amount', 0))
            end_date = expires_at.strftime('%d/%m/%Y')
            
            send_subscription_confirmation(
                user_email=user.get('email', ''),
                user_name=user.get('name', 'Utilisateur'),
                plan_name=plan_name,
                amount=amount,
                end_date=end_date
            )
            logger.info(f"Subscription confirmation email sent to {user.get('email')}")
        
        # Check if this user was referred and convert the referral
        if user and user.get('referred_by'):
            referrer_id = user['referred_by']
            
            # Find the pending referral
            referral = await db.referrals.find_one({
                'referrer_id': referrer_id,
                'referee_id': user_id,
                'status': 'pending'
            })
            
            if referral:
                # Update referral status
                await db.referrals.update_one(
                    {'id': referral['id']},
                    {'$set': {
                        'status': 'converted',
                        'converted_at': now,
                        'referrer_rewarded': True,
                    }}
                )
                
                # Reward the referrer with 1 free month
                referrer = await db.users.find_one({'user_id': referrer_id})
                if referrer:
                    current_end = referrer.get('subscription_end_date')
                    if current_end:
                        if isinstance(current_end, str):
                            current_end = datetime.fromisoformat(current_end.replace('Z', '+00:00'))
                        if current_end.tzinfo is None:
                            current_end = current_end.replace(tzinfo=timezone.utc)
                    
                    if not current_end or current_end < now:
                        current_end = now
                    
                    new_end = current_end + timedelta(days=30)
                    
                    await db.users.update_one(
                        {'user_id': referrer_id},
                        {
                            '$inc': {
                                'referral_count': 1,
                                'free_months_earned': 1,
                            },
                            '$set': {'subscription_end_date': new_end}
                        }
                    )
                    
                    # Send referral conversion notification to referrer
                    if is_email_configured():
                        send_referral_conversion_notification(
                            referrer_email=referrer.get('email', ''),
                            referrer_name=referrer.get('name', 'Membre'),
                            referee_name=user.get('name', 'Un membre'),
                            free_months=1
                        )
                    
                    logger.info(f"Referrer {referrer_id} rewarded with 1 free month for referral conversion")
    
    elif purchase_type in ('course', 'cursus'):
        # Add to purchases
        content_type = purchase_type
        content_id = transaction['metadata'].get('course_id') or transaction['metadata'].get('cursus_id')
        
        purchase_record = {
            'content_type': content_type,
            'content_id': content_id,
            'purchased_at': now,
            'expires_at': expires_at,
            'transaction_id': transaction['transaction_id']
        }
        
        await db.users.update_one(
            {'user_id': user_id},
            {'$push': {'purchases': purchase_record}}
        )
    
    logger.info(f"Access granted to user {user_id}: {purchase_type}")

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    """Handle Stripe webhook events."""
    if not STRIPE_API_KEY:
        raise HTTPException(500, "Stripe non configure")
    
    body = await request.body()
    signature = request.headers.get("Stripe-Signature")
    
    host_url = str(request.base_url).rstrip('/')
    webhook_url = f"{host_url}/api/webhook/stripe"
    stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
    
    try:
        webhook_response = await stripe_checkout.handle_webhook(body, signature)
        
        if webhook_response.payment_status == 'paid':
            transaction = await db.payment_transactions.find_one({'session_id': webhook_response.session_id}, {'_id': 0})
            if transaction and transaction.get('payment_status') != 'paid':
                await db.payment_transactions.update_one(
                    {'session_id': webhook_response.session_id},
                    {'$set': {'status': 'complete', 'payment_status': 'paid', 'updated_at': datetime.now(timezone.utc)}}
                )
                await grant_user_access(transaction)
            # Gift card finalization (idempotent)
            try:
                await _gift_card_finalize_after_payment(webhook_response.session_id)
            except Exception as e:
                logger.error(f"Gift card webhook finalize error: {e}")
        
        return {'received': True}
    
    except Exception as e:
        logger.error(f"Webhook error: {e}")
        return {'received': True, 'error': str(e)}

# ─── Stripe Phase B — Subscriptions (12-month commitment) ─────────────────────
# Uses the official stripe SDK with `mode=subscription` and webhook signature
# verification via STRIPE_WEBHOOK_SECRET. Coexists with the legacy one-shot
# `/api/checkout/create` flow above (still used by gift cards).

@api_router.post("/subscription/checkout")
async def subscription_checkout(body: CheckoutRequest, request: Request):
    """Create a Stripe subscription Checkout session for a recurring plan.

    Required: body.plan_id ∈ {founder_monthly, founder_yearly, standard_monthly, standard_yearly}.
    Enforces a 12-month minimum commitment on monthly plans (yearly is upfront).
    """
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    if not STRIPE_API_KEY:
        raise HTTPException(500, "Stripe non configuré")
    if not body.plan_id:
        raise HTTPException(400, "plan_id requis")
    if body.plan_id not in stripe_subs.SUBSCRIPTION_PLANS:
        raise HTTPException(404, f"Plan {body.plan_id} inconnu (attendu: {list(stripe_subs.SUBSCRIPTION_PLANS.keys())})")

    origin = (body.origin_url or '').rstrip('/') or str(request.base_url).rstrip('/')
    try:
        result = await stripe_subs.create_subscription_checkout(
            db, user=user, plan_id=body.plan_id, origin_url=origin,
        )
        return result
    except Exception as e:
        logger.exception(f"subscription_checkout error: {e}")
        raise HTTPException(500, f"Erreur Stripe: {str(e)}")


@api_router.post("/stripe/webhook")
async def stripe_subscription_webhook(request: Request):
    """Verified Stripe webhook for Phase B subscription events.

    Validates the `Stripe-Signature` header against STRIPE_WEBHOOK_SECRET.
    Dispatches to handlers in utils.stripe_subscriptions.
    """
    payload = await request.body()
    signature = request.headers.get('Stripe-Signature', '')
    secret = os.environ.get('STRIPE_WEBHOOK_SECRET', '')
    if not secret:
        logger.error("STRIPE_WEBHOOK_SECRET not configured")
        raise HTTPException(503, "Webhook secret non configuré")

    try:
        event = await stripe_subs.construct_event(payload, signature)
    except ValueError as e:
        logger.warning(f"Stripe webhook invalid payload: {e}")
        raise HTTPException(400, "Invalid payload")
    except stripe_sdk.error.SignatureVerificationError as e:
        logger.warning(f"Stripe webhook signature mismatch: {e}")
        raise HTTPException(400, "Invalid signature")
    except Exception as e:
        logger.exception(f"Stripe webhook verification error: {e}")
        raise HTTPException(400, "Webhook verification failed")

    result = await stripe_subs.handle_event(db, event)
    return {'received': True, **result}


@api_router.post("/subscription/cancel")
async def subscription_cancel(request: Request):
    """Cancel the user's active subscription at period end.

    Strict policy: HTTP 400 if `paid_invoices_count < commitment_months`
    (12 by default) for monthly plans. Yearly upfront subscribers are
    exempt (their first invoice covers the full commitment).
    """
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    if not STRIPE_API_KEY:
        raise HTTPException(500, "Stripe non configuré")

    try:
        result = await stripe_subs.cancel_subscription_with_policy(db, user_id=user['user_id'])
        return {'ok': True, **result}
    except stripe_subs.CommitmentNotMet as exc:
        raise HTTPException(
            status_code=400,
            detail={
                'error': 'commitment_not_met',
                'message': str(exc),
                'paid_invoices_count': exc.paid,
                'required': exc.required,
                'commitment_min_end': exc.commitment_min_end.isoformat() if exc.commitment_min_end else None,
            },
        )
    except ValueError as exc:
        raise HTTPException(404, str(exc))
    except Exception as exc:
        logger.exception(f"subscription_cancel error: {exc}")
        raise HTTPException(500, f"Erreur Stripe: {str(exc)}")


@api_router.get("/subscription/status")
async def subscription_status(request: Request):
    """Return the user's current Stripe subscription state (commitment progress)."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    sub = await db.subscriptions.find_one(
        {'user_id': user['user_id']},
        {'_id': 0},
        sort=[('created_at', -1)],
    )
    if not sub:
        return {'has_subscription': False}
    # Serialise datetimes
    for k in ('started_at', 'current_period_end', 'commitment_min_end', 'canceled_at',
              'cancel_requested_at', 'last_payment_failed_at', 'created_at', 'updated_at'):
        v = sub.get(k)
        if isinstance(v, datetime):
            sub[k] = v.isoformat()
    return {'has_subscription': True, 'subscription': sub}

# ─── Admin: Plans CRUD ────────────────────────────────────────────────────────

@api_router.get("/admin/plans")
async def admin_get_plans(request: Request):
    """Get all plans (admin)."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin requis")
    plans = await db.plans.find({}, {'_id': 0}).to_list(100)
    return plans

@api_router.post("/admin/plans")
async def admin_create_plan(plan: PlanCreate, request: Request):
    """Create a new plan."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin requis")
    
    existing = await db.plans.find_one({'plan_id': plan.plan_id})
    if existing:
        raise HTTPException(400, "Ce plan existe deja")
    
    plan_doc = plan.model_dump()
    plan_doc['created_at'] = datetime.now(timezone.utc)
    await db.plans.insert_one(plan_doc)
    return {k: v for k, v in plan_doc.items() if k != '_id'}

@api_router.put("/admin/plans/{plan_id}")
async def admin_update_plan(plan_id: str, plan: PlanUpdate, request: Request):
    """Update a plan."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin requis")
    
    updates = {k: v for k, v in plan.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(400, "Aucune modification")
    
    updates['updated_at'] = datetime.now(timezone.utc)
    result = await db.plans.update_one({'plan_id': plan_id}, {'$set': updates})
    if result.matched_count == 0:
        raise HTTPException(404, "Plan non trouve")
    
    updated = await db.plans.find_one({'plan_id': plan_id}, {'_id': 0})
    return updated

@api_router.delete("/admin/plans/{plan_id}")
async def admin_delete_plan(plan_id: str, request: Request):
    """Delete a plan."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin requis")
    
    result = await db.plans.delete_one({'plan_id': plan_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Plan non trouve")
    return {'message': 'Plan supprime'}

@api_router.get("/admin/transactions")
async def admin_get_transactions(request: Request, limit: int = 50):
    """Get recent transactions (admin)."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin requis")
    transactions = await db.payment_transactions.find({}, {'_id': 0}).sort('created_at', -1).limit(limit).to_list(limit)
    return transactions

@api_router.get("/admin-panel/pricing", response_class=HTMLResponse)
async def admin_panel_pricing(request: Request):
    """Admin panel pricing management page."""
    return templates.TemplateResponse(
        "pricing_new.html",
        {"request": request, "active_page": "pricing"}
    )

@api_router.get("/admin-panel/commercial", response_class=HTMLResponse)
async def admin_panel_commercial(request: Request):
    """Unified Commercial page (Tarifs / Parrainage / Codes promo tabs)."""
    return templates.TemplateResponse(
        "commercial.html",
        {"request": request, "active_page": "commercial"}
    )

@api_router.get("/admin-panel/referrals", response_class=HTMLResponse)
async def admin_panel_referrals(request: Request):
    """Admin panel referrals management page."""
    return templates.TemplateResponse(
        "referrals.html",
        {"request": request, "active_page": "referrals"}
    )

# ─── Promo Codes & Free Trial ─────────────────────────────────────────────────

@api_router.post("/promo/validate")
async def validate_promo_code(code: str, plan_id: Optional[str] = None):
    """Validate a promo code and return discount info."""
    promo = await db.promo_codes.find_one({'code': code.upper(), 'is_active': True}, {'_id': 0})
    if not promo:
        raise HTTPException(404, "Code promo invalide")
    
    now = datetime.now(timezone.utc)
    
    # Check start date (if promo hasn't started yet)
    if promo.get('start_date'):
        start = promo['start_date']
        if isinstance(start, str):
            start = datetime.fromisoformat(start.replace('Z', '+00:00'))
        if start.tzinfo is None:
            start = start.replace(tzinfo=timezone.utc)
        if start > now:
            raise HTTPException(400, "Code promo pas encore valide")
    
    # Check expiration
    if promo.get('expires_at'):
        expires = promo['expires_at']
        if isinstance(expires, str):
            expires = datetime.fromisoformat(expires.replace('Z', '+00:00'))
        if expires.tzinfo is None:
            expires = expires.replace(tzinfo=timezone.utc)
        if expires < now:
            raise HTTPException(400, "Code promo expiré")
    
    # Check max uses
    if promo.get('max_uses') and promo.get('uses_count', 0) >= promo['max_uses']:
        raise HTTPException(400, "Code promo épuisé")
    
    # Check applicable plans
    applicable = promo.get('applicable_plans', [])
    if applicable and plan_id and plan_id not in applicable:
        raise HTTPException(400, "Code promo non applicable à ce plan")
    
    return {
        'valid': True,
        'code': promo['code'],
        'discount_percent': promo.get('discount_percent'),
        'discount_amount': promo.get('discount_amount'),
        'description': promo.get('description', ''),
        'start_date': promo.get('start_date'),
        'expires_at': promo.get('expires_at')
    }

@api_router.get("/admin/promo-codes")
async def admin_get_promo_codes(request: Request):
    """Get all promo codes (admin)."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin requis")
    promos = await db.promo_codes.find({}, {'_id': 0}).to_list(100)
    return promos

@api_router.post("/admin/promo-codes")
async def admin_create_promo_code(promo: PromoCodeCreate, request: Request):
    """Create a new promo code."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin requis")
    
    code = promo.code.upper().strip()
    existing = await db.promo_codes.find_one({'code': code})
    if existing:
        raise HTTPException(400, "Ce code existe deja")
    
    promo_doc = promo.model_dump()
    promo_doc['code'] = code
    promo_doc['uses_count'] = 0
    promo_doc['created_at'] = datetime.now(timezone.utc)
    
    await db.promo_codes.insert_one(promo_doc)
    return {k: v for k, v in promo_doc.items() if k != '_id'}

@api_router.put("/admin/promo-codes/{code}")
async def admin_update_promo_code(code: str, promo: PromoCodeUpdate, request: Request):
    """Update a promo code."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin requis")
    
    updates = {k: v for k, v in promo.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(400, "Aucune modification")
    
    updates['updated_at'] = datetime.now(timezone.utc)
    result = await db.promo_codes.update_one({'code': code.upper()}, {'$set': updates})
    if result.matched_count == 0:
        raise HTTPException(404, "Code promo non trouve")
    
    updated = await db.promo_codes.find_one({'code': code.upper()}, {'_id': 0})
    return updated

@api_router.delete("/admin/promo-codes/{code}")
async def admin_delete_promo_code(code: str, request: Request):
    """Delete a promo code."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin requis")
    
    result = await db.promo_codes.delete_one({'code': code.upper()})
    if result.deleted_count == 0:
        raise HTTPException(404, "Code promo non trouve")
    return {'message': 'Code promo supprime'}

@api_router.post("/trial/start")
async def start_free_trial(body: StartTrialRequest, request: Request):
    """Start a free trial for a subscription plan."""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(401, "Authentification requise")
    
    # Check if user already had a trial
    user_doc = await db.users.find_one({'user_id': user['user_id']})
    if user_doc.get('had_trial'):
        raise HTTPException(400, "Vous avez déjà utilisé votre essai gratuit")
    
    # Special 3-day trial for new users
    if body.plan_id == 'trial_3days':
        trial_days = 3
    else:
        # Get plan with trial
        plan = await db.plans.find_one({'plan_id': body.plan_id, 'is_active': True}, {'_id': 0})
        if not plan:
            # Check default plans
            default = DEFAULT_PLANS.get(body.plan_id)
            if default:
                plan = {'plan_id': body.plan_id, 'trial_days': 7, **default}
            else:
                raise HTTPException(404, "Plan non trouvé")
        
        trial_days = plan.get('trial_days', 7)
    
    if trial_days <= 0:
        raise HTTPException(400, "Ce plan n'offre pas d'essai gratuit")
    
    now = datetime.now(timezone.utc)
    trial_expires = now + timedelta(days=trial_days)
    
    # Grant trial access
    await db.users.update_one(
        {'user_id': user['user_id']},
        {'$set': {
            'had_trial': True,
            'trial': {
                'plan_id': body.plan_id,
                'started_at': now.isoformat(),
                'expires_at': trial_expires.isoformat()
            }
        }}
    )
    
    logger.info(f"Trial {trial_days} days started for user {user['user_id']} until {trial_expires.isoformat()}")
    
    return {
        'success': True,
        'trial_days': trial_days,
        'expires_at': trial_expires.isoformat(),
        'message': f'Essai gratuit de {trial_days} jours activé !'
    }

@api_router.get("/admin-panel/promos", response_class=HTMLResponse)
async def admin_panel_promos(request: Request):
    """Admin panel promo codes management page."""
    return templates.TemplateResponse(
        "promos_new.html",
        {"request": request, "active_page": "promos"}
    )

# ─── Legal Pages API ──────────────────────────────────────────────────────────

DEFAULT_LEGAL = {
    "privacy": {
        "title": "Politique de confidentialité",
        "content": """**Dernière mise à jour : Février 2025**

## 1. Introduction

Sijill s'engage à protéger la confidentialité de vos données personnelles.

## 2. Données collectées

- Données d'identification : nom, prénom, adresse email
- Données d'utilisation : historique d'écoute, progression, favoris
- Données techniques : type d'appareil, système d'exploitation

## 3. Utilisation des données

Vos données sont utilisées pour personnaliser votre expérience d'apprentissage.

## 4. Vos droits

Conformément au RGPD, vous avez le droit d'accéder, rectifier et supprimer vos données.

## 5. Contact

Email : privacy@sijill.com"""
    },
    "terms": {
        "title": "Conditions d'utilisation",
        "content": """**Dernière mise à jour : Février 2025**

## 1. Acceptation des conditions

En utilisant l'application Sijill, vous acceptez les présentes conditions.

## 2. Description du service

Sijill est une plateforme d'apprentissage dédiée aux études islamiques.

## 3. Propriété intellectuelle

Tous les contenus sont protégés par le droit d'auteur.

## 4. Contact

Email : contact@sijillproject.com"""
    }
}

@api_router.get("/legal/{page_type}")
async def get_legal_page(page_type: str):
    """Get legal page content (privacy or terms)"""
    if page_type not in ["privacy", "terms"]:
        raise HTTPException(400, "Invalid page type")
    
    # Try to get from database
    doc = await db.legal_pages.find_one({"type": page_type}, {"_id": 0})
    if doc:
        return {"title": doc.get("title", ""), "content": doc.get("content", "")}
    
    # Return default
    return DEFAULT_LEGAL.get(page_type, DEFAULT_LEGAL["privacy"])

@api_router.put("/admin/legal/{page_type}")
async def update_legal_page(page_type: str, request: Request):
    """Update legal page content (admin only)"""
    if page_type not in ["privacy", "terms"]:
        raise HTTPException(400, "Invalid page type")
    
    body = await request.json()
    title = body.get("title", "")
    content = body.get("content", "")
    
    await db.legal_pages.update_one(
        {"type": page_type},
        {"$set": {"type": page_type, "title": title, "content": content, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )
    
    return {"success": True, "message": f"Page '{page_type}' mise à jour"}

# ─── Health Check ──────────────────────────────────────────────────────────────

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "service": "HikmabyLM API"}

# ─── Download Endpoint for Website ZIP ─────────────────────────────────────────

@api_router.get("/download/website-zip")
async def download_website_zip():
    """Download the website-hostinger.zip file"""
    zip_path = Path(__file__).parent.parent / "website-hostinger.zip"
    if zip_path.exists():
        return FileResponse(
            path=str(zip_path),
            filename="website-hostinger.zip",
            media_type="application/zip"
        )
    return {"error": "File not found"}

# ─── Trial Expiration Check & Emails ───────────────────────────────────────────

async def check_and_send_trial_expiration_emails():
    """Check for expiring trials and send notification emails."""
    now = datetime.now(timezone.utc)
    
    # Find users with trials expiring in 1 day (reminder) or already expired
    # Look for trials that haven't had an email sent yet
    
    try:
        # Users with trial expiring in ~1 day (send reminder)
        one_day_from_now = now + timedelta(days=1)
        users_expiring_soon = await db.users.find({
            'trial.expires_at': {'$exists': True},
            'subscription': {'$exists': False},  # No active subscription
            'trial_reminder_sent': {'$ne': True}
        }).to_list(100)
        
        for user in users_expiring_soon:
            try:
                trial_expires = user.get('trial', {}).get('expires_at')
                if not trial_expires:
                    continue
                    
                # Parse the expiration date
                if isinstance(trial_expires, str):
                    expires_dt = datetime.fromisoformat(trial_expires.replace('Z', '+00:00'))
                else:
                    expires_dt = trial_expires
                
                # Make it timezone aware if needed
                if expires_dt.tzinfo is None:
                    expires_dt = expires_dt.replace(tzinfo=timezone.utc)
                
                time_remaining = expires_dt - now
                days_remaining = time_remaining.days
                
                # Send reminder 1 day before or on expiration day
                if 0 <= days_remaining <= 1:
                    result = send_trial_expiration_email(
                        user_email=user['email'],
                        user_name=user.get('name', 'Cher utilisateur'),
                        days_remaining=days_remaining
                    )
                    if result.get('success'):
                        await db.users.update_one(
                            {'user_id': user['user_id']},
                            {'$set': {'trial_reminder_sent': True}}
                        )
                        logger.info(f"Trial reminder sent to {user['email']} ({days_remaining} days remaining)")
                
                # Send expiration notice if trial has ended
                elif days_remaining < 0 and not user.get('trial_expired_sent'):
                    result = send_trial_expiration_email(
                        user_email=user['email'],
                        user_name=user.get('name', 'Cher utilisateur'),
                        days_remaining=0
                    )
                    if result.get('success'):
                        await db.users.update_one(
                            {'user_id': user['user_id']},
                            {'$set': {'trial_expired_sent': True}}
                        )
                        logger.info(f"Trial expired notice sent to {user['email']}")
                        
            except Exception as e:
                logger.error(f"Error processing trial email for {user.get('email')}: {e}")
                
    except Exception as e:
        logger.error(f"Error in trial expiration check: {e}")

@api_router.post("/admin/send-trial-emails")
async def admin_send_trial_emails(request: Request):
    """Manually trigger trial expiration emails (admin only)."""
    user = await get_current_user(request)
    if not user or user.get('role') != 'admin':
        raise HTTPException(403, "Admin access required")
    
    await check_and_send_trial_expiration_emails()
    return {"success": True, "message": "Trial expiration emails processed"}

# ─── App Setup ────────────────────────────────────────────────────────────────

app.include_router(api_router)

# Mount static website files at /api/site (accessible via /api/site/)
WEBSITE_REACT_DIR = Path(__file__).parent.parent / "website-react" / "dist"
WEBSITE_DIR = Path(__file__).parent.parent / "website"

# ─── robots.txt ──────────────────────────────────────────────────────────────
@app.get("/api/site/robots.txt")
async def robots_txt():
    content = """User-agent: *
Allow: /

Sitemap: https://sijillproject.com/sitemap.xml
"""
    return Response(content=content, media_type="text/plain")

# ─── sitemap.xml ─────────────────────────────────────────────────────────────
@app.get("/api/site/sitemap.xml")
async def sitemap_xml():
    base = "https://sijillproject.com"
    static_pages = [
        ("", "1.0", "weekly"),
        ("/cursus", "0.9", "weekly"),
        ("/catalogue", "0.9", "weekly"),
        ("/blog", "0.9", "daily"),
        ("/a-propos", "0.7", "monthly"),
        ("/mentions-legales", "0.3", "yearly"),
        ("/politique-de-confidentialite", "0.3", "yearly"),
        ("/conditions-utilisation", "0.3", "yearly"),
    ]

    articles = await db.blog_articles.find({}, {"_id": 0, "id": 1, "synced_at": 1}).to_list(500)

    urls = []
    for path, priority, freq in static_pages:
        urls.append(f"""  <url>
    <loc>{base}{path}</loc>
    <changefreq>{freq}</changefreq>
    <priority>{priority}</priority>
  </url>""")

    for a in articles:
        urls.append(f"""  <url>
    <loc>{base}/blog/{a['id']}</loc>
    <changefreq>monthly</changefreq>
    <priority>0.8</priority>
  </url>""")

    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
{chr(10).join(urls)}
</urlset>"""
    return Response(content=xml, media_type="application/xml")

# ─── Server-side OG meta injection for blog articles (social crawlers) ───────
async def inject_og_meta(index_html: str, article_id: str) -> str:
    """Inject Open Graph meta tags into the HTML for blog article pages."""
    article = await db.blog_articles.find_one({"id": article_id}, {"_id": 0})
    if not article:
        return index_html

    base = "https://sijillproject.com"
    title = f"{article.get('title', '')} — Sijill Times #{article.get('number', '')}"
    desc = article.get('seo_description', article.get('hook', ''))[:200]
    url = f"{base}/blog/{article_id}"
    image = f"{base}/api/blog/image/{article_id}"
    tags_meta = "".join(f'\n    <meta property="article:tag" content="{t}" />' for t in (article.get('tags') or []))

    og_tags = f"""
    <meta property="og:type" content="article" />
    <meta property="og:title" content="{title}" />
    <meta property="og:description" content="{desc}" />
    <meta property="og:url" content="{url}" />
    <meta property="og:image" content="{image}" />
    <meta property="og:site_name" content="Sijill Project" />
    <meta property="og:locale" content="fr_FR" />
    <meta name="twitter:card" content="summary_large_image" />
    <meta name="twitter:title" content="{title}" />
    <meta name="twitter:description" content="{desc}" />
    <meta name="twitter:image" content="{image}" />{tags_meta}
    <meta name="description" content="{desc}" />
    <meta name="keywords" content="{', '.join(article.get('tags') or [])}" />
    <title>{title}</title>
    <script type="application/ld+json">
    {{
      "@context": "https://schema.org",
      "@type": "Article",
      "headline": "{article.get('title', '')}",
      "description": "{desc}",
      "url": "{url}",
      "image": "{image}",
      "author": {{"@type": "Organization", "name": "Sijill Project", "url": "{base}"}},
      "publisher": {{"@type": "Organization", "name": "Sijill Project", "url": "{base}"}},
      "keywords": "{', '.join(article.get('tags') or [])}",
      "inLanguage": "fr"
    }}
    </script>"""

    return index_html.replace("</head>", f"{og_tags}\n  </head>", 1)


async def inject_og_meta_course(index_html: str, course_id: str) -> str:
    """Inject OG meta + JSON-LD (Article + BreadcrumbList) for course pages.

    Crawlers (Google, social) read HTML synchronously and don't execute the
    React SPA. We inject the schema server-side from the course doc.
    """
    course = await db.courses.find_one({"id": course_id}, {"_id": 0})
    if not course:
        return index_html

    base = "https://sijillproject.com"
    raw_title = (course.get('title') or course.get('name') or '').strip()
    # Strip leading "Cours N : " prefix if present (matches frontend behavior).
    import re as _re
    clean_title = _re.sub(r'^Cours\s+\d+\s*:\s*', '', raw_title)
    title = f"{clean_title} — Sijill Project"
    desc = (course.get('summary') or course.get('description') or '').strip()
    desc = (desc[:197] + '…') if len(desc) > 200 else desc
    desc = desc.replace('"', '\\"').replace('\n', ' ')
    url = f"{base}/cours/{course_id}"
    image = course.get('thumbnail') or f"{base}/api/site/favicon.svg"
    scholar = course.get('scholar_name') or 'Sijill Project'

    # Escape for HTML attributes / JSON
    def _esc(s: str) -> str:
        return (s or '').replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')

    def _jesc(s: str) -> str:
        return (s or '').replace('\\', '\\\\').replace('"', '\\"').replace('\n', ' ')

    cursus_letter = (course.get('cursus_id') or '').replace('cursus-', '').upper()[:1] or ''
    breadcrumb_items = [
        ('Accueil', base + '/'),
        ('Catalogue', base + '/catalogue'),
    ]
    if cursus_letter:
        breadcrumb_items.append((f'Cursus {cursus_letter}', f'{base}/cursus'))
    breadcrumb_items.append((clean_title, url))
    breadcrumb_list = ",".join(
        f'{{"@type":"ListItem","position":{i+1},"name":"{_jesc(n)}","item":"{u}"}}'
        for i, (n, u) in enumerate(breadcrumb_items)
    )

    og_tags = f"""
    <meta property="og:type" content="article" />
    <meta property="og:title" content="{_esc(title)}" />
    <meta property="og:description" content="{_esc(desc)}" />
    <meta property="og:url" content="{url}" />
    <meta property="og:image" content="{_esc(image)}" />
    <meta property="og:site_name" content="Sijill Project" />
    <meta property="og:locale" content="fr_FR" />
    <meta name="twitter:card" content="summary_large_image" />
    <meta name="twitter:title" content="{_esc(title)}" />
    <meta name="twitter:description" content="{_esc(desc)}" />
    <meta name="twitter:image" content="{_esc(image)}" />
    <script type="application/ld+json">
    {{
      "@context": "https://schema.org",
      "@type": "Article",
      "headline": "{_jesc(clean_title)}",
      "description": "{_jesc(desc)}",
      "url": "{url}",
      "image": "{_jesc(image)}",
      "author": {{"@type": "Person", "name": "{_jesc(scholar)}"}},
      "publisher": {{"@type": "Organization", "name": "Sijill Project", "url": "{base}", "logo": {{"@type": "ImageObject", "url": "{base}/api/site/favicon.svg"}}}},
      "about": "{_jesc(course.get('topic') or 'Sciences islamiques')}",
      "inLanguage": "fr",
      "isAccessibleForFree": false
    }}
    </script>
    <script type="application/ld+json">
    {{
      "@context": "https://schema.org",
      "@type": "BreadcrumbList",
      "itemListElement": [{breadcrumb_list}]
    }}
    </script>"""

    # Replace existing title/canonical/description rather than appending duplicates.
    out = index_html
    out = _re.sub(r'<title>.*?</title>', f'<title>{_esc(title)}</title>', out, count=1, flags=_re.DOTALL)
    out = _re.sub(r'<link\s+rel="canonical"[^>]*?/>', f'<link rel="canonical" href="{url}" />', out, count=1)
    out = _re.sub(r'<meta\s+name="description"[^>]*?/>', f'<meta name="description" content="{_esc(desc)}" />', out, count=1)
    return out.replace("</head>", f"{og_tags}\n  </head>", 1)

# ─── WebApp (Expo Web) served at /webapp/* ─────────────────────────────────
WEBAPP_DIR = Path(__file__).parent.parent / "webapp" / "dist"
if WEBAPP_DIR.exists():
    app.mount("/webapp", StaticFiles(directory=str(WEBAPP_DIR), html=True), name="webapp")
    logger.info(f"WebApp mounted from {WEBAPP_DIR}")

# SPA catch-all: serve index.html for all non-file /api/site/* routes
@app.get("/api/site/{full_path:path}")
async def serve_website_spa(full_path: str):
    """Serve the React SPA. Static assets are served by StaticFiles mount below."""
    if WEBSITE_REACT_DIR.exists():
        file_path = WEBSITE_REACT_DIR / full_path
        if file_path.is_file():
            return FileResponse(str(file_path))
        index_path = WEBSITE_REACT_DIR / "index.html"
        if index_path.exists():
            if full_path.startswith("blog/") and full_path != "blog":
                article_id = full_path.replace("blog/", "", 1)
                html_content = index_path.read_text()
                injected = await inject_og_meta(html_content, article_id)
                return Response(content=injected, media_type="text/html")
            if full_path.startswith("cours/") and full_path != "cours":
                course_id = full_path.replace("cours/", "", 1).split("/", 1)[0]
                html_content = index_path.read_text()
                injected = await inject_og_meta_course(html_content, course_id)
                return Response(content=injected, media_type="text/html")
            return FileResponse(str(index_path))
    raise HTTPException(404, "Website not found")

# Mount React website assets
if WEBSITE_REACT_DIR.exists():
    app.mount("/api/site-assets", StaticFiles(directory=str(WEBSITE_REACT_DIR)), name="website-assets")

# Mount backend static files at /api/static
STATIC_DIR = Path(__file__).parent / "static"
if STATIC_DIR.exists():
    app.mount("/api/static", StaticFiles(directory=str(STATIC_DIR), html=True), name="static")

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup():
    await seed_data()
    # Check for trial expirations on startup (non-blocking)
    import asyncio
    asyncio.create_task(check_and_send_trial_expiration_emails())

    # Stripe Phase B — provision subscription catalog (idempotent via lookup_key).
    # Runs in background to avoid blocking startup if Stripe is slow / unreachable.
    async def _provision_stripe():
        try:
            if STRIPE_API_KEY and STRIPE_API_KEY.startswith('sk_'):
                summary = await stripe_subs.provision_catalog(db)
                logger.info(f"Stripe Phase B catalog ready: {summary}")
            else:
                logger.info("Stripe Phase B: STRIPE_API_KEY missing → catalog provisioning skipped")
        except Exception as exc:
            logger.exception(f"Stripe Phase B provisioning failed: {exc}")
    asyncio.create_task(_provision_stripe())

    # Gift cards scheduled delivery — check every hour
    async def _gift_delivery_loop():
        while True:
            await _process_scheduled_gift_deliveries()
            await asyncio.sleep(3600)
    asyncio.create_task(_gift_delivery_loop())

@app.on_event("shutdown")
async def shutdown():
    client.close()
